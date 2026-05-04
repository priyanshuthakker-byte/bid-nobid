"""
Submission Document Generator
Generates all standard Nascent tender submission documents:
1. Cover Letter (Form 1 equivalent)
2. Non-Blacklisting Declaration (notarised format)
3. Financial Standing Undertaking
4. Turnover Certificate Format (for CA to fill)
5. Employee Strength Declaration (Form 6 equivalent)
6. MII / Make in India Declaration
7. MSME Exemption Letter (for EMD waiver)
8. Pre-Bid Query Letter (formal format)
9. Custom annexure filler (fills Nascent data into RFP-specific forms)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Any, List
import json

PROFILE_PATH = Path(__file__).parent / "nascent_profile.json"

# ── Load Nascent profile ─────────────────────────────────────

def load_nascent() -> Dict:
    defaults = {
        "company": {
            "name": "Nascent Info Technologies Pvt. Ltd.",
            "cin": "U72200GJ2006PTC048723",
            "incorporated": "23 June 2006",
            "years_in_operation": 19,
            "address": "A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad - 380015, Gujarat",
            "pan": "AACCN3670J",
            "gstin": "24AACCN3670J1ZG",
            "msme_udyam": "UDYAM-GJ-01-0007420",
            "email": "nascent.tender@nascentinfo.com",
            "phone": "+91-79-40200400",
            "website": "www.nascentinfo.com",
            "signatory_name": "Hitesh Patel",
            "signatory_designation": "Chief Administrative Officer",
            "prepared_by": "Parthav Thakkar",
            "prepared_designation": "Bid Executive",
        },
        "finance": {
            "fy2223_cr": 16.36,
            "fy2324_cr": 16.36,
            "fy2425_cr": 18.83,
            "avg_turnover_last_3_fy": 17.18,
            "net_worth_cr": 26.09,
            "bank_name": "State Bank of India",
            "bank_branch": "SG Highway Branch, Ahmedabad",
            "account_no": "[Account Number]",
        },
        "employees": {
            "total_confirmed": 67,
            "gis_staff": 11,
            "it_dev_staff": 21,
        },
        "certifications": {
            "cmmi_level": "V2.0 Level 3",
            "cmmi_valid": "19-Dec-2026",
            "iso9001_cert": "ISO 9001:2015",
            "iso27001_cert": "ISO/IEC 27001:2022",
            "iso20000_cert": "ISO/IEC 20000-1:2018",
            "iso_valid": "08-Sep-2028",
        }
    }
    try:
        if PROFILE_PATH.exists():
            p = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
            # Deep merge
            for key in defaults:
                if key in p:
                    defaults[key].update(p[key])
    except Exception:
        pass
    return defaults


# ── Doc helpers ──────────────────────────────────────────────

def _setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    return doc


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.bold = True
    r.font.underline = (level == 1)


def _para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = bold
    return p


def _line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("_" * 80)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _signature_block(doc, n):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Authorised Signatory\n{n['company']['signatory_name']}\n{n['company']['signatory_designation']}\n{n['company']['name']}")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    doc.add_paragraph(f"Date: {date.today().strftime('%d-%B-%Y')}")
    doc.add_paragraph(f"Place: Ahmedabad, Gujarat")
    doc.add_paragraph()
    _para(doc, "Company Seal:", bold=True)
    doc.add_paragraph()
    doc.add_paragraph()


def _tender_ref_block(doc, tender: Dict):
    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    _para(doc, f"Tender / Bid No.: {tender.get('tender_no', '[Tender Number]')}")
    _para(doc, f"Tender Name: {tender.get('tender_name', tender.get('org_name', '[Tender Name]'))}")
    _para(doc, f"Organization: {tender.get('org_name', '[Organization Name]')}")
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════
# DOCUMENT 1 — COVER LETTER
# ════════════════════════════════════════════════════════════

def generate_cover_letter(tender: Dict, output_path: str):
    n = load_nascent()
    doc = _setup_doc()

    # Letterhead
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(n["company"]["name"])
    r.font.bold = True; r.font.size = Pt(13)
    _para(doc, n["company"]["address"]).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para(doc, f"Email: {n['company']['email']} | Ph: {n['company']['phone']}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _line(doc)
    doc.add_paragraph()

    _heading(doc, "COVER LETTER / BID SUBMISSION LETTER")
    doc.add_paragraph()

    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph()
    _para(doc, "To,")
    _para(doc, f"The Authorized Officer,")
    _para(doc, f"{tender.get('org_name', '[Organization Name]')},")
    _para(doc, f"[Address of Organization]")
    doc.add_paragraph()

    _para(doc, f"Sub: Submission of Bid for {tender.get('tender_name', '[Tender Name]')}")
    _para(doc, f"Ref: Tender / Bid No.: {tender.get('tender_no', '[Tender No.]')}")
    doc.add_paragraph()
    _para(doc, "Dear Sir / Madam,")
    doc.add_paragraph()

    _para(doc,
        f"We, {n['company']['name']}, a Private Limited Company incorporated under the Companies Act, 1956 "
        f"on {n['company']['incorporated']} (CIN: {n['company']['cin']}), having our registered office at "
        f"{n['company']['address']}, hereby submit our bid for the above-referenced tender."
    )
    doc.add_paragraph()
    _para(doc,
        "We have carefully read and examined the entire tender document including all clauses, "
        "schedules, annexures, and any corrigendum/addendum issued thereafter. We confirm our "
        "unconditional acceptance of all the terms and conditions stipulated therein."
    )
    doc.add_paragraph()
    _para(doc, "We hereby declare and confirm that:", bold=True)

    declarations = [
        f"1. We are a registered Indian company with PAN: {n['company']['pan']} and GSTIN: {n['company']['gstin']}.",
        f"2. We are a registered MSME under Udyam Registration No. {n['company']['msme_udyam']} (lifetime validity).",
        f"3. We have been in existence and providing IT/GIS services for {n['company']['years_in_operation']} years.",
        "4. We are not blacklisted or debarred by any Central / State Government, Ministry, Department, PSU, or Autonomous Body.",
        "5. We are not under any liquidation, court receivership, or similar proceedings.",
        "6. All information and documents submitted with this bid are true, accurate, and complete.",
        "7. We have not paid, and will not pay, any commissions or gratification to any person in connection with this bid.",
        "8. We confirm compliance with all applicable labour laws, GST regulations, and statutory requirements.",
    ]
    for d in declarations:
        _para(doc, d, indent=True)

    doc.add_paragraph()
    _para(doc,
        "We understand that the organization reserves the right to accept or reject any or all bids "
        "without assigning any reason whatsoever, and that our bid shall remain valid for "
        f"{tender.get('bid_validity', '180')} days from the date of bid submission."
    )
    doc.add_paragraph()
    _para(doc, "Thanking you,")
    _para(doc, "Yours faithfully,")
    _signature_block(doc, n)

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 2 — NON-BLACKLISTING DECLARATION
# ════════════════════════════════════════════════════════════

def generate_non_blacklisting(tender: Dict, output_path: str):
    n = load_nascent()
    doc = _setup_doc()

    _heading(doc, "SELF-DECLARATION — NON-BLACKLISTING / NON-DEBARMENT")
    _heading(doc, "(On Non-Judicial Stamp Paper of Rs. 100/-)", level=2)
    doc.add_paragraph()

    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph()
    _para(doc, f"Tender / Bid No.: {tender.get('tender_no', '[Tender No.]')}")
    _para(doc, f"Tender Name: {tender.get('tender_name', '[Tender Name]')}")
    doc.add_paragraph()

    _para(doc, "TO WHOMSOEVER IT MAY CONCERN", bold=True)
    doc.add_paragraph()

    _para(doc,
        f"I, {n['company']['signatory_name']}, {n['company']['signatory_designation']} of "
        f"{n['company']['name']} (CIN: {n['company']['cin']}), having registered office at "
        f"{n['company']['address']}, do hereby solemnly affirm and declare as under:"
    )
    doc.add_paragraph()

    paras = [
        f"1. That {n['company']['name']} (hereinafter referred to as 'the Company') has NOT been blacklisted "
        "by any Central Government Ministry, Department, Public Sector Undertaking (PSU), Public Sector Bank (PSB), "
        "State Government, Union Territory Administration, or any Government Autonomous Body in India, "
        f"as on the date of this declaration i.e. {date.today().strftime('%d %B %Y')}.",

        "2. That the Company has NOT been debarred from participation in public procurement by any "
        "Central / State / UT Government authority, Ministry, Department, PSU, or PSB for any reason "
        "whatsoever, including but not limited to breach of contract, corrupt practices, fraudulent "
        "misrepresentation, or any conviction under the Prevention of Corruption Act, 1988.",

        "3. That the Company is not currently involved in any legal proceedings that have resulted "
        "in disqualification from participating in Government tenders.",

        "4. That the Company is not under liquidation, court receivership, or similar proceedings, "
        "and has not been declared bankrupt or insolvent.",

        "5. That this declaration is made in good faith and all information provided herein is true "
        "and accurate to the best of my knowledge and belief.",

        "6. I am duly authorised to make this declaration on behalf of the Company.",
    ]
    for para in paras:
        _para(doc, para, indent=True)
    doc.add_paragraph()

    _para(doc,
        "I undertake that if the above declaration is found to be false at any point of time, "
        "the Company shall be liable for legal action and disqualification from this tender/contract."
    )
    doc.add_paragraph()
    _para(doc, "DEPONENT", bold=True)
    _signature_block(doc, n)

    _para(doc, "NOTARISATION:", bold=True)
    _para(doc, "Sworn and signed before me on _____________ at Ahmedabad.")
    doc.add_paragraph()
    _para(doc, "Notary Public: _______________________________")
    _para(doc, "Registration No.: ____________________________")
    _para(doc, "Stamp: ")

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 3 — TURNOVER CERTIFICATE FORMAT (for CA)
# ════════════════════════════════════════════════════════════

def generate_turnover_certificate(tender: Dict, output_path: str):
    n = load_nascent()
    fin = n["finance"]
    doc = _setup_doc()

    _heading(doc, "TURNOVER CERTIFICATE")
    _heading(doc, "(To be issued by Statutory Chartered Accountant / Cost Accountant on Letterhead)", level=2)
    doc.add_paragraph()

    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph()
    _para(doc, "To,")
    _para(doc, f"The Authorized Officer,")
    _para(doc, f"{tender.get('org_name', '[Organization Name]')}")
    doc.add_paragraph()

    _para(doc, f"Sub: Certificate of Annual Turnover for {n['company']['name']}")
    _para(doc, f"Ref: Tender / Bid No.: {tender.get('tender_no', '[Tender No.]')}")
    doc.add_paragraph()

    _para(doc,
        f"This is to certify that M/s {n['company']['name']} (CIN: {n['company']['cin']}), "
        f"having its registered office at {n['company']['address']}, PAN: {n['company']['pan']}, "
        "is a client of our firm and we have audited their financial statements."
    )
    doc.add_paragraph()

    _para(doc, "The annual turnover from IT Solution / Software Application Development & Implementation "
         "for the last three audited financial years is as under:", bold=True)
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Financial Year", "Total Turnover (Rs. Lakhs)", "IT Solution / Software Dev Turnover (Rs. Lakhs)"]
    rows_data = [
        ["FY 2022-23", f"{fin['fy2223_cr']*100:.2f}", "[IT-specific — to be filled by CA]"],
        ["FY 2023-24", f"{fin['fy2324_cr']*100:.2f}", "[IT-specific — to be filled by CA]"],
        ["FY 2024-25", f"{fin['fy2425_cr']*100:.2f}", "[IT-specific — to be filled by CA]"],
        ["Average (3 Years)", f"{fin['avg_turnover_last_3_fy']*100:.2f}", "[Average IT-specific — to be filled by CA]"],
        ["Net Worth", f"{fin['net_worth_cr']*100:.2f}", "—"],
    ]
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for ri, row_data in enumerate(rows_data, 1):
        for ci, val in enumerate(row_data):
            table.rows[ri].cells[ci].text = val

    doc.add_paragraph()
    _para(doc,
        "We further certify that the above figures have been derived from the audited financial statements "
        "of the company and are true and correct to the best of our knowledge and belief."
    )
    doc.add_paragraph()

    _para(doc, "CA Firm Name: _________________________________")
    _para(doc, "CA Registration No.: __________________________")
    _para(doc, "CA Name: _____________________________________")
    _para(doc, "Membership No.: ______________________________")
    _para(doc, "Seal & Signature:")
    _para(doc, f"Place: Ahmedabad | Date: {date.today().strftime('%d-%B-%Y')}")

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 4 — EMPLOYEE STRENGTH DECLARATION
# ════════════════════════════════════════════════════════════

def generate_employee_declaration(tender: Dict, output_path: str):
    n = load_nascent()
    emp = n["employees"]
    doc = _setup_doc()

    _heading(doc, "EMPLOYEE STRENGTH DECLARATION")
    _heading(doc, "(Self-Declaration on Company Letterhead — duly signed by HR / Head of Organization)", level=2)
    doc.add_paragraph()

    _tender_ref_block(doc, tender)

    _para(doc,
        f"I, {n['company']['signatory_name']}, {n['company']['signatory_designation']} of "
        f"{n['company']['name']}, do hereby declare and confirm that as on the date of this "
        f"declaration ({date.today().strftime('%d %B %Y')}), the Company has the following "
        "technically qualified personnel on its payroll:"
    )
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    rows_data = [
        ["Category", "Number", "Minimum Qualification"],
        ["GIS Specialists / GIS Developers", str(emp["gis_staff"]), "B.E./B.Tech/M.Sc (CS/IT/Geography/Geomatics)"],
        ["IT / Software Developers", str(emp["it_dev_staff"]), "B.E./B.Tech/MCA/M.Sc (CS/IT)"],
        ["QA / Testing Engineers", "[count]", "B.E./B.Tech/MCA"],
        ["Project Managers / Business Analysts", "[count]", "B.E./B.Tech/MBA"],
        ["TOTAL Technically Qualified Staff", str(emp["total_confirmed"]), "As above categories"],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            if ri == 0 or ci == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    _para(doc,
        "All the above personnel are on the company's regular payroll and are covered under "
        "Employee Provident Fund (EPF). EPF Combined Challan for the latest month is enclosed "
        "as supporting evidence."
    )
    doc.add_paragraph()
    _para(doc,
        "I declare that the information provided above is true and correct to the best of my knowledge."
    )
    doc.add_paragraph()

    _para(doc, "HR / Head of Organization:", bold=True)
    _para(doc, f"Name: {n['company']['signatory_name']}")
    _para(doc, f"Designation: {n['company']['signatory_designation']}")
    _para(doc, f"Organization: {n['company']['name']}")
    _para(doc, f"Date: {date.today().strftime('%d-%B-%Y')}")
    _para(doc, "Seal & Signature:")

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 5 — MSME EMD EXEMPTION LETTER
# ════════════════════════════════════════════════════════════

def generate_msme_emd_exemption(tender: Dict, output_path: str):
    n = load_nascent()
    doc = _setup_doc()

    # Letterhead
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(n["company"]["name"]).bold = True

    _heading(doc, "REQUEST FOR EMD EXEMPTION UNDER MSME PROCUREMENT POLICY")
    doc.add_paragraph()

    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph()
    _para(doc, "To,")
    _para(doc, "The Authorized Officer,")
    _para(doc, f"{tender.get('org_name', '[Organization Name]')},")
    doc.add_paragraph()

    _para(doc, f"Sub: Request for EMD Exemption — {tender.get('tender_name', '[Tender Name]')}")
    _para(doc, f"Ref: Tender / Bid No.: {tender.get('tender_no', '[Tender No.]')}")
    doc.add_paragraph()

    _para(doc, "Dear Sir / Madam,")
    doc.add_paragraph()

    _para(doc,
        f"We, {n['company']['name']}, are a registered Micro, Small and Medium Enterprise (MSME) "
        f"under the MSMED Act, 2006, vide Udyam Registration Number: {n['company']['msme_udyam']} "
        "(Lifetime Validity), and are a Service Provider as per the said registration."
    )
    doc.add_paragraph()

    _para(doc,
        "In terms of the Public Procurement Policy for Micro and Small Enterprises Order, 2012 "
        "(as amended in 2018 and clarified vide DoE Office Memorandum dated November 2020), "
        "Micro and Small Enterprises registered as Service Providers are EXEMPT from payment of "
        "Earnest Money Deposit (EMD) in Government tenders."
    )
    doc.add_paragraph()

    _para(doc,
        f"We hereby request that our Udyam Registration Certificate (No. {n['company']['msme_udyam']}) "
        "enclosed herewith be accepted in lieu of EMD Demand Draft / Bank Guarantee for this bid, "
        f"as per the above-cited policy provisions. The EMD amount mentioned in the bid is "
        f"Rs. {tender.get('emd', '[EMD Amount]')}."
    )
    doc.add_paragraph()

    _para(doc,
        "Policy References:\n"
        "1. Public Procurement Policy for MSMEs Order, 2012 (Ministry of MSME)\n"
        "2. DoE Office Memorandum F.No.6/18/2019-PPD dated 16.11.2020\n"
        "3. GeM General Terms & Conditions — EMD Exemption for MSE Service Providers",
        indent=True
    )
    doc.add_paragraph()

    _para(doc, "Enclosure: Copy of Udyam Registration Certificate")
    doc.add_paragraph()
    _signature_block(doc, n)

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 6 — FINANCIAL STANDING UNDERTAKING
# ════════════════════════════════════════════════════════════

def generate_financial_standing(tender: Dict, output_path: str):
    n = load_nascent()
    doc = _setup_doc()

    _heading(doc, "UNDERTAKING — FINANCIAL STANDING")
    doc.add_paragraph()
    _tender_ref_block(doc, tender)

    _para(doc,
        f"I, {n['company']['signatory_name']}, {n['company']['signatory_designation']} of "
        f"{n['company']['name']} (CIN: {n['company']['cin']}), do hereby solemnly affirm and declare:"
    )
    doc.add_paragraph()

    for stmt in [
        f"1. {n['company']['name']} is NOT under liquidation, court receivership, or similar proceedings.",
        "2. The Company has NOT been declared bankrupt or insolvent by any court or authority.",
        "3. The Company's accounts are not frozen or under any attachment by any government authority.",
        f"4. The Company has a positive Net Worth of Rs. {n['finance']['net_worth_cr']} Crores as per audited financial statements.",
        "5. The Company has sufficient financial capacity to execute the project if awarded.",
        "6. All statutory payments including GST, PF, PT, and IT are current and up to date.",
    ]:
        _para(doc, stmt, indent=True)

    doc.add_paragraph()
    _para(doc, "I confirm that the above information is true and accurate.")
    doc.add_paragraph()
    _signature_block(doc, n)

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# DOCUMENT 7 — MII DECLARATION
# ════════════════════════════════════════════════════════════

def generate_mii_declaration(tender: Dict, output_path: str):
    n = load_nascent()
    doc = _setup_doc()

    _heading(doc, "MAKE IN INDIA (MII) DECLARATION")
    doc.add_paragraph()
    _tender_ref_block(doc, tender)

    _para(doc,
        f"I, {n['company']['signatory_name']}, {n['company']['signatory_designation']} of "
        f"{n['company']['name']}, do hereby declare and certify that:"
    )
    doc.add_paragraph()

    for stmt in [
        f"1. {n['company']['name']} is an Indian Company incorporated under the Companies Act, 1956/2013, "
        f"having CIN: {n['company']['cin']}.",
        "2. All IT solutions, software applications, GIS applications, mobile applications, and web portals "
        "developed and delivered by the Company are designed, developed, and maintained entirely within India "
        "by Indian professionals.",
        "3. The Company's development teams are based in India. No offshore development or outsourcing to "
        "non-Indian entities is involved.",
        "4. All services to be provided under this contract, if awarded, will be performed from India by "
        "Indian nationals employed by the Company.",
        "5. The Company complies with the 'Preference to Make in India' policy as per DPIIT/MeitY guidelines "
        "and subsequent Orders/OM issued in this regard.",
        "6. The Company is NOT a company from a country sharing land border with India (as per GFR Rule 144(xi)).",
    ]:
        _para(doc, stmt, indent=True)

    doc.add_paragraph()
    _signature_block(doc, n)

    doc.save(output_path)


# ════════════════════════════════════════════════════════════
# MASTER GENERATOR — generates all docs as a package
# ════════════════════════════════════════════════════════════

def generate_submission_package(tender: Dict, output_dir: str) -> List[Dict]:
    """
    Generate complete submission document package for a tender.
    Returns list of {filename, title, description} for each generated file.
    """
    from pathlib import Path
    out = Path(output_dir)
    out.mkdir(exist_ok=True, parents=True)

    safe_no = __import__("re").sub(r'[^\w\-]', '_', tender.get("tender_no", "Tender"))[:40]
    generated = []

    docs = [
        ("cover_letter", f"01_Cover_Letter_{safe_no}.docx",
         "Cover Letter", "Bid submission cover letter on company letterhead",
         generate_cover_letter),
        ("non_blacklisting", f"02_Non_Blacklisting_Declaration_{safe_no}.docx",
         "Non-Blacklisting Declaration", "Notarised self-declaration on Rs.100 stamp paper",
         generate_non_blacklisting),
        ("turnover_cert", f"03_Turnover_Certificate_Format_{safe_no}.docx",
         "Turnover Certificate (CA Format)", "Format to be filled and signed by Statutory CA",
         generate_turnover_certificate),
        ("employee_decl", f"04_Employee_Strength_Declaration_{safe_no}.docx",
         "Employee Strength Declaration", "Self-declaration of technical staff count",
         generate_employee_declaration),
        ("msme_emd", f"05_MSME_EMD_Exemption_Letter_{safe_no}.docx",
         "MSME EMD Exemption Letter", "Request letter citing MSME Policy 2012 for EMD waiver",
         generate_msme_emd_exemption),
        ("financial_standing", f"06_Financial_Standing_Undertaking_{safe_no}.docx",
         "Financial Standing Undertaking", "Declaration that company is not under liquidation",
         generate_financial_standing),
        ("mii_decl", f"07_MII_Declaration_{safe_no}.docx",
         "Make in India Declaration", "MII compliance declaration",
         generate_mii_declaration),
    ]

    for doc_key, filename, title, description, generator_fn in docs:
        try:
            filepath = str(out / filename)
            generator_fn(tender, filepath)
            generated.append({
                "doc_key": doc_key,
                "filename": filename,
                "title": title,
                "description": description,
                "status": "generated",
                "path": filepath,
            })
        except Exception as e:
            generated.append({
                "doc_key": doc_key,
                "filename": filename,
                "title": title,
                "description": description,
                "status": "error",
                "error": str(e),
            })

    return generated


def merge_docs_to_pdf(file_paths: List[str], output_path: str) -> bool:
    """
    Merge multiple Word docs into a single PDF.
    Requires LibreOffice on the server (available on Render).
    Falls back to returning False if not available.
    """
    try:
        import subprocess
        import tempfile, os
        from pathlib import Path

        # Convert each DOCX to PDF first
        pdf_paths = []
        for fp in file_paths:
            fp = Path(fp)
            if not fp.exists():
                continue
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", str(fp.parent), str(fp)],
                capture_output=True, timeout=30
            )
            pdf = fp.with_suffix(".pdf")
            if pdf.exists():
                pdf_paths.append(str(pdf))

        if not pdf_paths:
            return False

        # Merge PDFs using pypdf
        try:
            from pypdf import PdfWriter
            writer = PdfWriter()
            for pdf_path in pdf_paths:
                from pypdf import PdfReader
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    writer.add_page(page)
            with open(output_path, "wb") as f:
                writer.write(f)
            return True
        except Exception:
            pass

        return False
    except Exception:
        return False
