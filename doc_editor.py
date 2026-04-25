"""
Doc Editor — docx ↔ HTML preview + chatbot-assisted edits + version history.

Flow:
  1. Frontend requests /tender/{id}/doc-html → we render current docx as HTML.
  2. User edits HTML in contenteditable pane OR types chat instruction.
  3. Chat instruction: /tender/{id}/ai-edit → Gemini rewrites text sections.
  4. Save: /tender/{id}/doc-save → HTML → docx (overwrites, keeps version snapshot).
  5. Version list: /tender/{id}/doc-versions.
  6. Restore: /tender/{id}/doc-restore/{version}.

Dependencies:
  - python-docx (already present)
  - mammoth (added to requirements for rich docx→HTML)
  - falls back to python-docx direct extraction if mammoth unavailable.
"""

import re
import json
import shutil
import hashlib
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from html.parser import HTMLParser

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────
# docx → HTML
# ─────────────────────────────────────────────────────────

def docx_to_html(docx_path: Path) -> str:
    """Prefer mammoth for styled HTML, fallback to python-docx plain extraction."""
    try:
        import mammoth
        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
            html = result.value
            return _wrap_html(html)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"mammoth failed: {e} — falling back to python-docx")

    try:
        from docx import Document
        doc = Document(str(docx_path))
        parts = []
        for block in _iter_block_items(doc):
            parts.append(_block_to_html(block))
        return _wrap_html("\n".join(parts))
    except Exception as e:
        logger.error(f"docx_to_html fallback failed: {e}")
        return _wrap_html(f"<p>⚠ Could not render document: {e}</p>")


def _iter_block_items(doc):
    """Yield paragraphs + tables in source order."""
    from docx.document import Document as _D
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def _block_to_html(block) -> str:
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    if isinstance(block, Paragraph):
        text = _escape(block.text)
        if not text.strip():
            return "<p>&nbsp;</p>"
        style = (block.style.name or "").lower() if block.style else ""
        if style.startswith("heading 1"):
            return f"<h1>{text}</h1>"
        if style.startswith("heading 2"):
            return f"<h2>{text}</h2>"
        if style.startswith("heading 3"):
            return f"<h3>{text}</h3>"
        if style.startswith("heading"):
            return f"<h4>{text}</h4>"
        if style.startswith("list"):
            return f"<li>{text}</li>"
        return f"<p>{text}</p>"
    if isinstance(block, Table):
        rows = []
        for row in block.rows:
            cells = "".join(f"<td>{_escape(c.text)}</td>" for c in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        return f"<table class='docx-table' border='1'>{''.join(rows)}</table>"
    return ""


def _escape(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _wrap_html(body: str) -> str:
    return (
        '<div class="docx-body" style="font-family:Calibri,Arial,sans-serif;'
        'font-size:11pt;line-height:1.5;padding:24px;background:#fff;color:#1f2937;">'
        + body +
        "</div>"
    )


# ─────────────────────────────────────────────────────────
# HTML → docx
# ─────────────────────────────────────────────────────────

class _HtmlDocxBuilder(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.stack: List[str] = []
        self.current_text = []
        self.in_table = False
        self.table_rows: List[List[str]] = []
        self.current_row: List[str] = []
        self.cell_text: List[str] = []
        self.bold = False
        self.italic = False
        self.current_para = None

    def _flush_para(self, style: Optional[str] = None):
        text = "".join(self.current_text).strip()
        if text:
            p = self.doc.add_paragraph()
            if style:
                try:
                    p.style = style
                except Exception:
                    pass
            run = p.add_run(text)
            run.bold = self.bold
            run.italic = self.italic
        self.current_text = []

    def handle_starttag(self, tag, attrs):
        t = tag.lower()
        if t == "br":
            self.current_text.append("\n")
            return
        if self.in_table:
            if t == "tr":
                self.current_row = []
            elif t in ("td", "th"):
                self.cell_text = []
            return
        if t == "table":
            self.in_table = True
            self.table_rows = []
            return
        if t in ("b", "strong"):
            self.bold = True
        elif t in ("i", "em"):
            self.italic = True
        self.stack.append(t)

    def handle_endtag(self, tag):
        t = tag.lower()
        if self.in_table:
            if t in ("td", "th"):
                self.current_row.append("".join(self.cell_text).strip())
                self.cell_text = []
            elif t == "tr":
                self.table_rows.append(self.current_row)
                self.current_row = []
            elif t == "table":
                self._flush_table()
                self.in_table = False
            return
        if t in ("p", "div"):
            self._flush_para()
        elif t in ("h1", "h2", "h3", "h4"):
            style = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3", "h4": "Heading 4"}[t]
            self._flush_para(style=style)
        elif t == "li":
            self._flush_para(style="List Bullet")
        elif t in ("b", "strong"):
            self.bold = False
        elif t in ("i", "em"):
            self.italic = False
        if self.stack and self.stack[-1] == t:
            self.stack.pop()

    def handle_data(self, data):
        if self.in_table:
            self.cell_text.append(data)
        else:
            self.current_text.append(data)

    def _flush_table(self):
        if not self.table_rows:
            return
        cols = max(len(r) for r in self.table_rows)
        table = self.doc.add_table(rows=len(self.table_rows), cols=cols)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for i, row in enumerate(self.table_rows):
            for j in range(cols):
                table.cell(i, j).text = row[j] if j < len(row) else ""
        self.table_rows = []

    def close(self):
        self._flush_para()
        super().close()


def html_to_docx(html: str, out_path: Path, title: str = "Document") -> Path:
    from docx import Document
    doc = Document()
    body_match = re.search(r'<body[^>]*>(.*)</body>', html, re.S | re.I)
    body = body_match.group(1) if body_match else html
    builder = _HtmlDocxBuilder(doc)
    try:
        builder.feed(body)
        builder.close()
    except Exception as e:
        logger.warning(f"HTML parse hiccup: {e}")
    doc.save(str(out_path))
    return out_path


# ─────────────────────────────────────────────────────────
# Version history
# ─────────────────────────────────────────────────────────

def versions_dir_for(docx_path: Path) -> Path:
    vd = docx_path.parent / ".versions" / docx_path.stem
    vd.mkdir(parents=True, exist_ok=True)
    return vd


def snapshot_version(docx_path: Path, note: str = "") -> Dict:
    """Copy current docx to versions dir with timestamp."""
    if not docx_path.exists():
        return {"error": "source file missing"}
    vd = versions_dir_for(docx_path)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = vd / f"{ts}.docx"
    shutil.copy2(docx_path, dest)
    meta = {
        "ts": ts,
        "created": datetime.now().isoformat(),
        "note": note,
        "size": dest.stat().st_size,
        "sha256": _sha256(dest),
    }
    (vd / f"{ts}.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return meta


def list_versions(docx_path: Path) -> List[Dict]:
    vd = versions_dir_for(docx_path)
    out = []
    for meta_file in sorted(vd.glob("*.json"), reverse=True)[:50]:
        try:
            m = json.loads(meta_file.read_text(encoding="utf-8"))
            m["version_id"] = meta_file.stem
            out.append(m)
        except Exception:
            continue
    return out


def restore_version(docx_path: Path, version_id: str) -> Dict:
    vd = versions_dir_for(docx_path)
    src = vd / f"{version_id}.docx"
    if not src.exists():
        return {"error": "version not found"}
    snapshot_version(docx_path, note=f"auto-snapshot before restore of {version_id}")
    shutil.copy2(src, docx_path)
    return {"restored": version_id, "target": docx_path.name}


def _sha256(p: Path) -> str:
    h = hashlib.sha256()
    with open(p, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()[:16]


# ─────────────────────────────────────────────────────────
# AI chat-edit
# ─────────────────────────────────────────────────────────

AI_EDIT_PROMPT = """You are editing a bid-submission Word document on behalf of the bidder.

USER INSTRUCTION:
{instruction}

CURRENT DOCUMENT CONTENT (HTML, keep the tags):
{html}

Task:
- Apply the user's instruction faithfully.
- If user says "this is wrong" or "change X to Y", make that edit precisely.
- Keep the HTML structure (<p>, <h1>-<h4>, <table>, <tr>, <td>, <ul>, <li>, <b>, <i>).
- Do NOT add commentary, explanations, or markdown code fences.
- Do NOT invent facts — if instruction is vague, make minimal plausible change.
- If instruction targets a specific section only, leave the rest IDENTICAL.
- Preserve existing figures, numbers, dates unless instruction explicitly changes them.

Return ONLY the complete edited HTML — nothing else.
"""


def ai_edit_html(html: str, instruction: str) -> Dict:
    """Send current HTML + instruction to Gemini; return edited HTML."""
    try:
        from ai_analyzer import _call, get_all_api_keys, load_config
    except Exception as e:
        return {"error": f"analyzer unavailable: {e}"}

    keys = get_all_api_keys()
    if not keys:
        return {"error": "No Gemini API key configured."}
    cfg = load_config()
    groq_key = cfg.get("groq_api_key", cfg.get("groq_key", ""))

    prompt = AI_EDIT_PROMPT.format(instruction=instruction.strip()[:2000],
                                    html=html[:60000])
    try:
        raw = _call(prompt, keys[0], keys, groq_key, max_tokens=8192)
    except Exception as e:
        return {"error": f"AI call failed: {str(e)[:200]}"}

    raw = raw.strip()
    raw = re.sub(r'^```(?:html)?\s*', '', raw, flags=re.MULTILINE)
    raw = re.sub(r'\s*```$', '', raw, flags=re.MULTILINE)
    if "<" not in raw or ">" not in raw:
        return {"error": "AI returned non-HTML content", "raw_preview": raw[:200]}
    return {"html": raw, "chars": len(raw)}


# ─────────────────────────────────────────────────────────
# Compliance matrix + risk scoring (bonus modules)
# ─────────────────────────────────────────────────────────

def build_compliance_matrix(tender: Dict) -> List[Dict]:
    """Map every PQ/TQ clause → coverage status + evidence."""
    matrix = []
    for item in tender.get("pq_criteria", []):
        matrix.append({
            "category": "PQ",
            "clause_ref": item.get("clause_ref", "—"),
            "criteria": item.get("criteria", ""),
            "required_docs": item.get("documents_required", ""),
            "status": item.get("nascent_status", "Review"),
            "color": item.get("nascent_color", "BLUE"),
            "evidence": item.get("evidence_projects", "—"),
            "remark": item.get("nascent_remark", ""),
            "raises_query": item.get("raises_query", "NO"),
        })
    for item in tender.get("tq_criteria", []):
        matrix.append({
            "category": "TQ",
            "clause_ref": item.get("clause_ref", "—"),
            "criteria": item.get("criteria", ""),
            "required_docs": item.get("documents_required", ""),
            "status": item.get("nascent_status", "Review"),
            "color": item.get("nascent_color", "BLUE"),
            "evidence": "—",
            "marks": item.get("max_marks", ""),
            "nascent_score": item.get("nascent_score", ""),
            "remark": item.get("nascent_remark", ""),
            "raises_query": item.get("raises_query", "NO"),
        })
    return matrix


def compute_risk_score(tender: Dict) -> Dict:
    """
    Lightweight deterministic risk scoring.
    Higher score = riskier. Ranges 0-100.
    """
    score = 0
    factors = []

    pq = tender.get("pq_criteria", [])
    red_pq = sum(1 for p in pq if p.get("nascent_color") == "RED")
    amber_pq = sum(1 for p in pq if p.get("nascent_color") == "AMBER")
    if red_pq:
        score += min(40, red_pq * 12)
        factors.append(f"{red_pq} red PQ clause(s) — hard disqualifier risk")
    if amber_pq:
        score += min(20, amber_pq * 5)
        factors.append(f"{amber_pq} amber PQ clause(s) — borderline")

    tq = tender.get("tq_criteria", [])
    red_tq = sum(1 for t in tq if t.get("nascent_color") == "RED")
    if red_tq:
        score += min(15, red_tq * 5)
        factors.append(f"{red_tq} red TQ clause(s) — marks loss")

    if tender.get("hard_disqualifiers"):
        dq = tender["hard_disqualifiers"]
        if isinstance(dq, list) and dq:
            score += 20
            factors.append(f"{len(dq)} hard disqualifier(s)")

    if tender.get("has_corrigendum"):
        score += 3
        factors.append("Corrigendum issued — verify latest deadlines")

    ld = str(tender.get("ld_rate", "") or "").lower()
    if "%" in ld:
        try:
            pct = float(re.search(r"(\d+(?:\.\d+)?)", ld).group(1))
            if pct >= 10:
                score += 8
                factors.append(f"High LD penalty: {ld}")
            elif pct >= 5:
                score += 4
                factors.append(f"LD penalty: {ld}")
        except Exception:
            pass

    pbg = str(tender.get("pbg_details", "") or tender.get("performance_security", "") or "")
    pbg_pct = re.search(r"(\d+(?:\.\d+)?)\s*%", pbg)
    if pbg_pct and float(pbg_pct.group(1)) >= 10:
        score += 5
        factors.append(f"High PBG: {pbg_pct.group(0)}")

    score = min(100, score)
    if score >= 60:
        grade, color = "HIGH", "#ef4444"
    elif score >= 30:
        grade, color = "MEDIUM", "#f59e0b"
    else:
        grade, color = "LOW", "#10b981"

    return {
        "score": score,
        "grade": grade,
        "color": color,
        "factors": factors or ["No major risk factors detected"],
        "verdict_align": tender.get("verdict", "—"),
    }
