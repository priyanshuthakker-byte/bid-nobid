"""
Technical Proposal Generator v1.0
Nascent Info Technologies Bid/No-Bid System

Generates a fully designed, professional Technical Proposal Word document
automatically after AI analysis of any tender. Uses Nascent's project database
to match relevant experience to the tender domain automatically.
"""

import json
import re
from pathlib import Path
from datetime import date
from typing import Optional
import io

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

BASE_DIR = Path(__file__).parent

# ─────────────────────────────────────────────────────────────
# NASCENT STATIC PROFILE  (also loaded from nascent_profile.json)
# ─────────────────────────────────────────────────────────────
NASCENT = {
    "name": "Nascent Info Technologies Pvt. Ltd.",
    "short": "Nascent",
    "cin": "U72200GJ2006PTC048723",
    "pan": "AACCN3670J",
    "gstin": "24AACCN3670J1ZG",
    "msme": "UDYAM-GJ-01-0007420",
    "address": "A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad – 380015, Gujarat",
    "phone": "+91-79-40200400",
    "email": "nascent.tender@nascentinfo.com",
    "web": "www.nascentinfo.com",
    "signatory": "Hitesh Patel",
    "designation": "Chief Administrative Officer",
    "incorporated": "23 June 2006",
    "years": "19",
    "turnover_2223": "Rs. 16.36 Cr",
    "turnover_2324": "Rs. 16.36 Cr",
    "turnover_2425": "Rs. 18.83 Cr",
    "avg_turnover": "Rs. 17.18 Cr",
    "net_worth": "Rs. 26.09 Cr",
    "employees": "67",
    "gis_staff": "11",
    "it_staff": "21",
    "cmmi": "CMMI V2.0 Level 3 (valid till 19-Dec-2026)",
    "iso": "ISO 9001:2015, ISO/IEC 27001:2022, ISO/IEC 20000-1:2018 (till Sep-2028)",
    "bank": "State Bank of India, SG Highway Branch, Ahmedabad",
}

# All Nascent projects — used for smart matching
ALL_PROJECTS = [
    {"name": "AMC GIS – Ahmedabad Municipal Corporation", "value": "Rs. 10.55 Cr", "status": "Ongoing AMC", "domain": ["gis", "mobile", "municipal", "ulb", "ahmedabad", "amc", "smart city"], "scope": "Web GIS, property mapping, utilities, ward management, mobile field app, ongoing AMC"},
    {"name": "JuMC GIS – Junagadh Municipal Corporation", "value": "Rs. 9.78 Cr", "status": "Ongoing", "domain": ["gis", "municipal", "ulb", "junagadh", "survey"], "scope": "Web GIS, mobile survey, GIS layer management, ongoing AMC"},
    {"name": "BMC GIS Mobile App – Bhavnagar Municipal Corporation", "value": "Rs. 4.2 Cr", "status": "Completed", "domain": ["gis", "mobile", "android", "ios", "app", "municipal", "ulb", "field", "geo-tagging"], "scope": "Android + iOS GIS Mobile App, field survey, geo-tagging, offline mode"},
    {"name": "VMC GIS + ERP – Vadodara Municipal Corporation", "value": "Rs. 20.5 Cr", "status": "Completed", "domain": ["gis", "erp", "municipal", "ulb", "vadodara"], "scope": "Web GIS + ERP modules, property, utilities, large scale ULB"},
    {"name": "KVIC Mobile GIS – Khadi Village Industries Commission", "value": "Rs. 5.15 Cr", "status": "Completed", "domain": ["gis", "mobile", "android", "ios", "field", "geo-tagging", "offline", "central", "psu"], "scope": "PAN India mobile GIS, geo-tagging, offline sync for field workers"},
    {"name": "PCSCL Smart City – Pimpri-Chinchwad", "value": "Rs. 61.19 Cr", "status": "Ongoing", "domain": ["smart city", "gis", "erp", "scada", "iot", "municipal"], "scope": "Smart City GIS, ERP, SCADA integration, large scale"},
    {"name": "AMC Heritage App (AHA) – Ahmedabad Smart City", "value": "Rs. 3.8 Cr", "status": "Completed", "domain": ["mobile", "android", "ios", "ar", "augmented reality", "heritage", "smart city", "ahmedabad"], "scope": "Heritage mobile app with AR, QR, location-based experiences"},
    {"name": "CEICED – CEICED Government Portal", "value": "Rs. 2.1 Cr", "status": "Completed", "domain": ["portal", "web", "government", "egovernance", "central"], "scope": "eGovernance web portal, data management, reporting"},
    {"name": "NSO FOD Survey – National Statistics Office", "value": "Rs. 8.4 Cr", "status": "Completed", "domain": ["survey", "mobile", "android", "field", "gis", "central", "psu", "government", "data collection"], "scope": "Mobile field data collection, survey management, GIS integration for NSO"},
    {"name": "NP Lalganj GIS – Nagar Panchayat Lalganj", "value": "Rs. 1.2 Cr", "status": "Completed", "domain": ["gis", "municipal", "ulb", "property", "survey", "small"], "scope": "GIS mapping, property survey, urban local body"},
    {"name": "Digital Road Map – Gujarat Tourism", "value": "Rs. 1.8 Cr", "status": "Completed", "domain": ["gis", "mobile", "tourism", "map", "navigation"], "scope": "Interactive GIS road map, tourism navigation mobile app"},
    {"name": "ERP – Government Medical College", "value": "Rs. 4.5 Cr", "status": "Completed", "domain": ["erp", "healthcare", "hospital", "government", "education"], "scope": "Hospital ERP, patient management, billing, inventory"},
]

# Domain keywords → matching weights
DOMAIN_MAP = {
    "gis": ["gis", "geospatial", "spatial", "map", "mapping", "geoinformatics", "geodata", "citylayers", "geoserver", "postgis", "qgis"],
    "mobile": ["mobile", "android", "ios", "app", "application", "smartphone", "tablet"],
    "ulb": ["municipal", "municipality", "ulb", "corporation", "urban local body", "nagar", "panchayat", "ward"],
    "smart_city": ["smart city", "smart cities mission", "scm", "iot", "sensor", "dashboard"],
    "survey": ["survey", "field", "geo-tagging", "geotagging", "data collection", "enumeration"],
    "erp": ["erp", "enterprise resource", "hr", "payroll", "accounting", "finance", "inventory"],
    "egovernance": ["egovernance", "e-governance", "portal", "citizen", "service", "government"],
    "central": ["central government", "ministry", "psu", "public sector", "national"],
}


def load_profile() -> dict:
    """Load nascent_profile.json if available, merge with static data."""
    profile_file = BASE_DIR / "nascent_profile.json"
    if profile_file.exists():
        try:
            data = json.loads(profile_file.read_text(encoding="utf-8"))
            merged = {**NASCENT, **data}
            return merged
        except Exception:
            pass
    return NASCENT.copy()


def match_projects(tender_data: dict, top_n: int = 6) -> list:
    """Score all Nascent projects against tender keywords and return top N."""
    brief = (tender_data.get("brief", "") + " " +
             tender_data.get("scope_summary", "") + " " +
             tender_data.get("tender_name", "") + " " +
             tender_data.get("org_name", "")).lower()

    scored = []
    for proj in ALL_PROJECTS:
        score = 0
        for kw in proj["domain"]:
            if kw in brief:
                score += 10
        # Boost if scope items mention project domain
        for item in tender_data.get("scope_items", []):
            for kw in proj["domain"]:
                if kw in item.lower():
                    score += 3
        if score > 0:
            scored.append((score, proj))

    scored.sort(key=lambda x: x[0], reverse=True)
    result = [p for _, p in scored[:top_n]]

    # Always include at least 3 top-value projects even if no keyword match
    if len(result) < 3:
        for proj in ALL_PROJECTS:
            if proj not in result:
                result.append(proj)
            if len(result) >= 5:
                break

    return result


def generate_approach(tender_data: dict) -> dict:
    """Generate technical approach based on tender domain."""
    brief = (tender_data.get("brief", "") + " " +
             tender_data.get("scope_summary", "")).lower()

    approach = {
        "understanding": "",
        "architecture": [],
        "tech_stack": [],
        "methodology": "",
        "team_size": "7",
        "timeline_months": "5",
    }

    # GIS + Mobile
    if any(k in brief for k in ["gis", "mobile app", "android", "ios", "geo"]):
        approach["understanding"] = (
            "The client requires a GIS-enabled Mobile Application integrating with existing "
            "geospatial infrastructure to enable field data collection, spatial visualization, "
            "offline operation, and seamless sync with the central GIS database."
        )
        approach["architecture"] = [
            "Mobile Layer: Android/iOS Native App with offline GIS engine",
            "API Layer: Spring Boot REST APIs on client servers",
            "Data Layer: Existing GeoServer + PostGIS (zero replacement)"
        ]
        approach["tech_stack"] = [
            ["Android App", "Android Native (Java/Kotlin)", "Best offline map + GPS performance"],
            ["iOS App", "iOS Native (Swift) / Flutter", "Feature parity with Android"],
            ["GIS SDK", "MapLibre GL + OpenLayers", "OGC compliant, works with GeoServer WMS/WFS"],
            ["GIS Backend", "Existing GeoServer (client)", "Direct reuse — zero replacement"],
            ["Offline Maps", "MBTiles + OSMDroid", "Ward-level offline packs"],
            ["Backend APIs", "Java Spring Boot REST", "Thin API on existing infrastructure"],
            ["Database", "Existing PostGIS (client)", "All field data to existing spatial DB"],
            ["Auth", "OAuth 2.0 + JWT + LDAP", "SSO for all client staff"],
            ["Notifications", "Firebase Cloud Messaging", "Real-time push for field staff"],
        ]
        approach["timeline_months"] = "5"
        approach["team_size"] = "7"

    # Smart City / ERP
    elif any(k in brief for k in ["smart city", "erp", "enterprise", "scada"]):
        approach["understanding"] = (
            "The client requires an integrated Smart City / Enterprise platform covering "
            "multiple operational domains with real-time dashboards, inter-department "
            "data sharing, and citizen-facing service delivery."
        )
        approach["tech_stack"] = [
            ["Frontend", "React.js + Progressive Web App", "Responsive, citizen-friendly UI"],
            ["Backend", "Java Spring Boot microservices", "Scalable, independently deployable modules"],
            ["Database", "PostgreSQL + Redis cache", "Relational data + high-performance caching"],
            ["Integration", "REST APIs + Apache Kafka", "Real-time data streams between modules"],
            ["Dashboard", "Custom BI + Apache Superset", "Executive and operational dashboards"],
            ["Mobile", "Android + iOS apps", "Field staff and citizen mobile access"],
            ["Security", "OAuth 2.0 + SSL + WAF", "Government-grade security"],
        ]
        approach["timeline_months"] = "8"
        approach["team_size"] = "12"

    # eGovernance Portal
    elif any(k in brief for k in ["portal", "egovernance", "citizen", "e-governance"]):
        approach["understanding"] = (
            "The client requires a citizen-centric eGovernance portal enabling online "
            "service delivery, application processing, status tracking, and data "
            "management with secure government-grade infrastructure."
        )
        approach["tech_stack"] = [
            ["Frontend", "React.js + Next.js SSR", "Fast, SEO-friendly citizen portal"],
            ["Backend", "Python FastAPI + Django", "Rapid development, robust REST APIs"],
            ["Database", "PostgreSQL + Elasticsearch", "Structured data + full-text search"],
            ["File Storage", "MinIO / NFS", "Document management on government infrastructure"],
            ["Auth", "OpenID Connect + Aadhaar OTP", "Secure citizen authentication"],
            ["Integration", "REST + SFTP + API Gateway", "Connect with existing government systems"],
        ]
        approach["timeline_months"] = "6"
        approach["team_size"] = "8"

    # Survey / Data Collection
    elif any(k in brief for k in ["survey", "data collection", "enumeration", "field"]):
        approach["understanding"] = (
            "The client requires a field data collection and survey management system "
            "enabling large-scale enumeration with offline capability, geo-tagging, "
            "photo capture, and real-time dashboard monitoring."
        )
        approach["tech_stack"] = [
            ["Mobile App", "Android (Java/Kotlin)", "Offline data collection, GPS, camera"],
            ["Offline DB", "SQLite + Room DB", "Store 50,000+ survey records offline"],
            ["Sync Engine", "REST APIs + conflict resolution", "Smart background sync"],
            ["Backend", "Python FastAPI", "High-throughput survey submission processing"],
            ["Dashboard", "React.js + Chart.js", "Real-time survey progress monitoring"],
            ["GIS Layer", "PostGIS + GeoServer", "Geo-tagged survey visualization"],
        ]
        approach["timeline_months"] = "4"
        approach["team_size"] = "6"

    # Default fallback
    else:
        approach["understanding"] = (
            "The client requires a comprehensive IT solution covering system design, "
            "development, implementation, and maintenance with focus on performance, "
            "security, and ease of use for government stakeholders."
        )
        approach["tech_stack"] = [
            ["Frontend", "React.js / Angular", "Modern, responsive UI"],
            ["Backend", "Python FastAPI / Java Spring Boot", "Robust REST API layer"],
            ["Database", "PostgreSQL", "Reliable relational database"],
            ["Mobile", "Android + iOS (Flutter)", "Cross-platform mobile access"],
            ["Security", "OAuth 2.0 + SSL + WAF", "Government-grade security standards"],
            ["DevOps", "Docker + CI/CD pipeline", "Automated build, test, deploy"],
        ]
        approach["timeline_months"] = "6"
        approach["team_size"] = "8"

    approach["methodology"] = (
        "Nascent follows an Agile-Waterfall hybrid methodology: requirements and design "
        "phases use structured waterfall for thorough documentation, while development "
        "uses 2-week sprints with regular client demos and feedback integration. "
        "This approach reduces risk while maintaining flexibility for scope refinements."
    )

    return approach


def generate_technical_proposal(tender_data: dict, output_path: str) -> dict:
    """
    Main entry point. Generates a complete Technical Proposal .docx file.

    Args:
        tender_data: dict from AI analysis (tender_name, org_name, scope_items, pq_criteria, etc.)
        output_path: where to save the .docx

    Returns:
        dict with status, path, pages, matched_projects count
    """
    if not DOCX_AVAILABLE:
        return {"status": "error", "message": "python-docx not installed. Run: pip install python-docx"}

    try:
        profile = load_profile()
        matched_projects = match_projects(tender_data)
        approach = generate_approach(tender_data)

        doc = Document()

        # ── Page setup ────────────────────────────────────────
        from docx.shared import Mm
        section = doc.sections[0]
        section.page_width = Mm(210)   # A4
        section.page_height = Mm(297)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)

        tender_name = tender_data.get("tender_name", tender_data.get("brief", "Tender"))
        org_name = tender_data.get("org_name", "")
        tender_no = tender_data.get("tender_no", tender_data.get("ref_no", ""))
        today = date.today().strftime("%d %B %Y")

        # ─────────────────────────────────────────────
        # COVER PAGE
        # ─────────────────────────────────────────────
        _add_cover_page(doc, profile, tender_name, org_name, tender_no, today)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # TABLE OF CONTENTS
        # ─────────────────────────────────────────────
        _add_toc(doc)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 1. EXECUTIVE SUMMARY
        # ─────────────────────────────────────────────
        _add_heading(doc, "1.  Executive Summary", 1)
        verdict = tender_data.get("overall_verdict", {}).get("verdict", "CONDITIONAL BID")
        _add_verdict_box(doc, verdict, tender_data)
        doc.add_paragraph()
        _add_para(doc,
            f"{profile['name']} is pleased to present this Technical Proposal for {tender_name}. "
            f"With {profile['years']} years of experience in IT, GIS, and Mobile Application development "
            f"exclusively for Government and Urban Local Bodies, Nascent is technically aligned and "
            f"operationally ready to deliver this project successfully."
        )
        _add_para(doc,
            f"Nascent is Gujarat-based (Ahmedabad), CMMI V2.0 Level 3 certified, ISO 9001/27001/20000 "
            f"certified, and MSME registered. Our team of {profile['employees']} professionals includes "
            f"{profile['gis_staff']} GIS specialists and {profile['it_staff']} IT developers. "
            f"Average annual turnover: {profile['avg_turnover']}. Net Worth: {profile['net_worth']}."
        )

        # Key stats table
        _add_stat_row(doc, [
            (profile["years"], "Years Experience"),
            (f"{len(matched_projects)}+", "Relevant Projects"),
            (profile["avg_turnover"], "Avg Turnover"),
            (profile["employees"], "IT Professionals"),
        ])
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 2. COMPANY PROFILE
        # ─────────────────────────────────────────────
        _add_heading(doc, "2.  Company Profile & Credentials", 1)
        company_rows = [
            ("Company Name", profile["name"]),
            ("CIN", profile["cin"]),
            ("Incorporated", f"{profile['incorporated']} ({profile['years']} years in operation)"),
            ("PAN", profile["pan"]),
            ("GSTIN", profile["gstin"]),
            ("MSME Registration", f"{profile['msme']} (Lifetime validity)"),
            ("Total Employees", f"{profile['employees']} ({profile['gis_staff']} GIS, {profile['it_staff']} IT/Dev, rest QA/PM/BA)"),
            ("Avg. Annual Turnover", f"{profile['avg_turnover']} (FY 2022-23 to 2024-25)"),
            ("Net Worth", profile["net_worth"]),
            ("CMMI", profile["cmmi"]),
            ("ISO Certifications", profile["iso"]),
            ("Address", profile["address"]),
            ("Contact", f"{profile['phone']}  |  {profile['email']}"),
        ]
        _add_kv_table(doc, company_rows)

        doc.add_paragraph()
        _add_heading(doc, "Quality Certifications", 2)
        _add_cert_table(doc)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 3. RELEVANT EXPERIENCE
        # ─────────────────────────────────────────────
        _add_heading(doc, "3.  Relevant Experience", 1)
        _add_para(doc,
            f"Nascent has delivered {len(ALL_PROJECTS)}+ major IT/GIS projects for Government clients. "
            f"The {len(matched_projects)} most relevant projects for this tender are highlighted below."
        )
        doc.add_paragraph()
        _add_experience_table(doc, matched_projects)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 4. TECHNICAL APPROACH
        # ─────────────────────────────────────────────
        _add_heading(doc, "4.  Technical Approach and Methodology", 1)

        _add_heading(doc, "4.1  Understanding of Requirement", 2)
        _add_para(doc, approach["understanding"])

        # Scope items from AI analysis
        scope_items = tender_data.get("scope_items", [])
        if scope_items:
            _add_heading(doc, "4.2  Scope of Work Identified", 2)
            for item in scope_items[:15]:
                _add_bullet(doc, str(item))

        _add_heading(doc, "4.3  Methodology", 2)
        _add_para(doc, approach["methodology"])

        _add_heading(doc, "4.4  Technology Stack", 2)
        if approach["tech_stack"]:
            _add_tech_table(doc, approach["tech_stack"])

        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 5. IMPLEMENTATION TIMELINE
        # ─────────────────────────────────────────────
        _add_heading(doc, "5.  Project Implementation Timeline", 1)
        months = int(approach["timeline_months"])
        _add_timeline_table(doc, months)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 6. TEAM STRUCTURE
        # ─────────────────────────────────────────────
        _add_heading(doc, "6.  Proposed Team Structure", 1)
        _add_team_table(doc, approach)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 7. QUALITY ASSURANCE
        # ─────────────────────────────────────────────
        _add_heading(doc, "7.  Quality Assurance & Testing", 1)
        _add_qa_table(doc)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 8. PQ COMPLIANCE SUMMARY
        # ─────────────────────────────────────────────
        pq_criteria = tender_data.get("pq_criteria", [])
        if pq_criteria:
            _add_heading(doc, "8.  PQ Criteria Compliance", 1)
            _add_pq_table(doc, pq_criteria)
            doc.add_page_break()

        # ─────────────────────────────────────────────
        # 9. WHY NASCENT
        # ─────────────────────────────────────────────
        _add_heading(doc, "9.  Why Nascent", 1)
        differentiators = [
            ("Gujarat-based (Ahmedabad)", "100 km from Surat / 0 km from Ahmedabad. On-site support within hours. Deep understanding of Gujarat Government processes and procurement norms."),
            ("Proven Track Record", f"Delivered {len(ALL_PROJECTS)}+ Government IT/GIS projects. All on time, all in production. Strong reference clients across Gujarat and India."),
            ("CMMI V2.0 Level 3", "Highest process maturity certification for IT companies. Ensures disciplined, repeatable, measurable software delivery."),
            ("ISO Triple Certification", "ISO 9001 (Quality) + ISO 27001 (Security) + ISO 20000 (ITSM) — all active, all valid. Government-grade assurance."),
            ("MSME Registered", f"Udyam: {profile['msme']}. EMD exempt in Government tenders. Gujarat state-based MSME preference applicable."),
            ("Financial Stability", f"Net Worth {profile['net_worth']}. Average turnover {profile['avg_turnover']}. 19 years profitable operation. Zero project abandonment risk."),
            ("Domain Expertise", f"{profile['gis_staff']} GIS specialists + {profile['it_staff']} IT developers with hands-on GeoServer, PostGIS, Android, iOS, Spring Boot experience."),
            ("AMC Proven", "Currently maintaining 3+ live Government GIS systems. Proven ability to provide long-term support, not just delivery."),
        ]
        _add_differentiator_table(doc, differentiators)
        doc.add_page_break()

        # ─────────────────────────────────────────────
        # 10. DECLARATION
        # ─────────────────────────────────────────────
        _add_heading(doc, "10.  Declaration & Authorisation", 1)
        _add_declaration(doc, profile, tender_no, today)

        # ── Save ──────────────────────────────────────
        doc.save(output_path)

        return {
            "status": "success",
            "path": output_path,
            "matched_projects": len(matched_projects),
            "sections": 10,
            "timeline_months": approach["timeline_months"],
        }

    except Exception as e:
        import traceback
        return {"status": "error", "message": str(e), "trace": traceback.format_exc()}


# ═════════════════════════════════════════════════════════════
# DOCX HELPER FUNCTIONS
# ═════════════════════════════════════════════════════════════

def _rgb(r, g, b):
    return RGBColor(r, g, b)

NAVY  = _rgb(31, 56, 100)
BLUE  = _rgb(46, 117, 182)
WHITE = _rgb(255, 255, 255)
GRAY  = _rgb(238, 242, 247)
GREEN = _rgb(55, 86, 35)
AMBER = _rgb(127, 96, 0)
RED   = _rgb(192, 0, 0)
DNAVY = _rgb(13, 27, 55)


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _cell_para(cell, text: str, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.paragraphs[0].alignment = align
    # Cell padding
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top','bottom','left','right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '80')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _add_heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.font.name = "Arial"
        run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)


def _add_para(doc, text: str, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_bullet(doc, text: str):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)


def _add_cover_page(doc, profile, tender_name, org_name, tender_no, today):
    """Dark professional cover page."""
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(profile["name"].upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(4)

    # Tag line
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("GIS  •  Smart City  •  Mobile Applications  •  eGovernance  •  ERP")
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLUE
    r2.font.name = "Arial"
    p2.paragraph_format.space_after = Pt(20)

    # Divider
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("━" * 55)
    r3.font.color.rgb = NAVY
    r3.font.name = "Arial"
    p3.paragraph_format.space_after = Pt(20)

    # TECHNICAL PROPOSAL title
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("TECHNICAL PROPOSAL")
    rt.bold = True
    rt.font.size = Pt(28)
    rt.font.color.rgb = NAVY
    rt.font.name = "Arial"
    pt.paragraph_format.space_after = Pt(16)

    # Tender name
    ptn = doc.add_paragraph()
    ptn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtn = ptn.add_run("For")
    rtn.font.size = Pt(11)
    rtn.font.color.rgb = BLUE
    rtn.font.name = "Arial"
    ptn.paragraph_format.space_after = Pt(4)

    ptn2 = doc.add_paragraph()
    ptn2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtn2 = ptn2.add_run(tender_name[:120] + ("..." if len(tender_name) > 120 else ""))
    rtn2.bold = True
    rtn2.font.size = Pt(14)
    rtn2.font.color.rgb = NAVY
    rtn2.font.name = "Arial"
    ptn2.paragraph_format.space_after = Pt(20)

    # Submitted to / by table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    c1 = tbl.cell(0, 0)
    c2 = tbl.cell(0, 1)
    _set_cell_bg(c1, "EBF3FB")
    _set_cell_bg(c2, "1F3864")

    # Left cell
    c1.paragraphs[0].clear()
    c1.add_paragraph("")
    for line, bold, color in [
        ("SUBMITTED TO", True, NAVY),
        (org_name or "Client Organisation", True, NAVY),
        (f"Tender No.: {tender_no}" if tender_no else "", False, BLUE),
    ]:
        if line:
            p = c1.add_paragraph()
            r = p.add_run(line)
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(11 if bold else 10)
            r.font.color.rgb = color
            p.paragraph_format.space_after = Pt(4)

    # Right cell
    c2.paragraphs[0].clear()
    c2.add_paragraph("")
    for line, bold, color in [
        ("SUBMITTED BY", True, _rgb(157, 195, 230)),
        (profile["name"], True, WHITE),
        (profile["address"][:60], False, _rgb(222, 235, 247)),
        (f"CMMI V2.0 L3  •  ISO 9001/27001/20000  •  MSME", True, _rgb(189, 215, 238)),
        (f"Date: {today}", False, WHITE),
    ]:
        p = c2.add_paragraph()
        r = p.add_run(line)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(4)

    # Credential badges
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    bt = doc.add_table(rows=1, cols=5)
    bt.style = 'Table Grid'
    badges = [("CMMI V2.0\nLevel 3", "1D4E8C"), ("ISO 9001\n:2015", "2E75B6"),
              ("ISO 27001\n:2022", "0070C0"), ("ISO 20000\n-1:2018", "005A9E"), ("MSME\nUdyam", "1F3864")]
    for i, (text, bg) in enumerate(badges):
        c = bt.cell(0, i)
        _set_cell_bg(c, bg)
        _cell_para(c, text, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_toc(doc):
    _add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1.  Executive Summary", "2.  Company Profile & Credentials",
        "3.  Relevant Experience", "4.  Technical Approach and Methodology",
        "5.  Project Implementation Timeline", "6.  Proposed Team Structure",
        "7.  Quality Assurance & Testing", "8.  PQ Criteria Compliance",
        "9.  Why Nascent", "10. Declaration & Authorisation"
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"  {item}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(5)


def _add_verdict_box(doc, verdict: str, tender_data: dict):
    """Coloured verdict summary box."""
    if "NO-BID" in verdict or "NO BID" in verdict:
        bg, tc = "FCE4D6", "C00000"
    elif "CONDITIONAL" in verdict:
        bg, tc = "FFF2CC", "7F6000"
    else:
        bg, tc = "E2EFDA", "375623"

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.cell(0, 0)
    _set_cell_bg(c, bg)
    c.paragraphs[0].clear()
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"VERDICT:  {verdict}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(tc)
    r.font.name = "Arial"
    sub = c.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    counts = tender_data.get("overall_verdict", {})
    sub_text = f"{counts.get('green',0)} Met  ·  {counts.get('amber',0)} Conditional  ·  {counts.get('red',0)} Not Met"
    sr = sub.add_run(sub_text)
    sr.font.name = "Arial"
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor.from_string(tc)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_stat_row(doc, stats: list):
    """4-column coloured stats row."""
    tbl = doc.add_table(rows=1, cols=len(stats))
    tbl.style = 'Table Grid'
    colors = ["1D4E8C", "2E75B6", "0070C0", "1F3864"]
    for i, (val, label) in enumerate(stats):
        c = tbl.cell(0, i)
        _set_cell_bg(c, colors[i % len(colors)])
        c.paragraphs[0].clear()
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(val))
        r.bold = True; r.font.size = Pt(18); r.font.color.rgb = WHITE; r.font.name = "Arial"
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9); r2.font.color.rgb = _rgb(222, 235, 247); r2.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_kv_table(doc, rows: list):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        _set_cell_bg(tbl.cell(i, 0), "EEF2F7")
        _cell_para(tbl.cell(i, 0), k, bold=True, color=NAVY, size=10)
        _cell_para(tbl.cell(i, 1), v, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_cert_table(doc):
    certs = [
        ("CMMI V2.0 Level 3", "Process Maturity\nValid: 19-Dec-2026", "1D4E8C"),
        ("ISO 9001:2015", "Quality Management\nValid: Sep-2028", "2E75B6"),
        ("ISO 27001:2022", "Information Security\nValid: Sep-2028", "0070C0"),
        ("ISO 20000-1:2018", "IT Service Mgmt\nValid: Sep-2028", "005A9E"),
        ("MSME Udyam", "Govt Registered\nLifetime Validity", "1F3864"),
    ]
    tbl = doc.add_table(rows=1, cols=len(certs))
    tbl.style = 'Table Grid'
    for i, (name, detail, bg) in enumerate(certs):
        c = tbl.cell(0, i)
        _set_cell_bg(c, bg)
        c.paragraphs[0].clear()
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE; r.font.name = "Arial"
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(detail); r2.font.size = Pt(8); r2.font.color.rgb = _rgb(222, 235, 247); r2.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_experience_table(doc, projects: list):
    headers = ["Client / Project", "Scope Delivered", "Value", "Status", "Relevance"]
    col_w = [Inches(1.5), Inches(2.0), Inches(0.9), Inches(0.9), Inches(2.0)]
    tbl = doc.add_table(rows=1 + len(projects), cols=5)
    tbl.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    highlight_colors = {"Completed": "E2EFDA", "Ongoing": "EBF3FB", "Ongoing AMC": "EBF3FB"}
    for row_i, proj in enumerate(projects):
        ri = row_i + 1
        bg = "EBF3FB" if row_i % 2 == 0 else "F5F8FF"
        status_bg = highlight_colors.get(proj.get("status", ""), bg)
        _cell_para(tbl.cell(ri, 0), proj["name"], bold=True, color=NAVY, size=10)
        _cell_para(tbl.cell(ri, 1), proj["scope"], size=9)
        _cell_para(tbl.cell(ri, 2), proj["value"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(tbl.cell(ri, 3), status_bg)
        _cell_para(tbl.cell(ri, 3), proj["status"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(tbl.cell(ri, 4), ", ".join(proj["domain"][:4]), size=9, color=BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_tech_table(doc, stack: list):
    headers = ["Component", "Technology", "Justification"]
    tbl = doc.add_table(rows=1 + len(stack), cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (comp, tech, why) in enumerate(stack):
        _set_cell_bg(tbl.cell(ri+1, 0), "EEF2F7")
        _cell_para(tbl.cell(ri+1, 0), comp, bold=True, color=NAVY, size=10)
        _set_cell_bg(tbl.cell(ri+1, 1), "F0F6FF")
        _cell_para(tbl.cell(ri+1, 1), tech, size=10)
        _cell_para(tbl.cell(ri+1, 2), why, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_timeline_table(doc, months: int):
    phases = [
        ("Phase 1", "Requirements Analysis & SRS", "SRS, Architecture Doc"),
        ("Phase 2", "UI/UX Design & Prototype", "Wireframes, Prototype"),
        ("Phase 3", "Development & Integration", "Beta Application"),
        ("Phase 4", "Testing & UAT", "Test Reports, UAT Sign-off"),
        ("Phase 5", "Deployment & Training", "Live System, Manuals"),
        ("AMC", "Ongoing Maintenance & Support", "Monthly MIS, SLA Compliance"),
    ]
    headers = ["Phase", "Activities", "Deliverables"] + [f"M{i+1}" for i in range(min(months, 5))] + ["AMC"]
    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(phases), cols=col_count)
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    phase_colors = ["1D4E8C", "2E75B6", "0070C0", "375623", "375623", "7F6000"]
    phase_month_map = {0: [0], 1: [0,1], 2: [1,2,3], 3: [3,4], 4: [4], 5: [-1]}
    for ri, (phase, act, deliv) in enumerate(phases):
        col_bg = phase_colors[ri]
        _set_cell_bg(tbl.cell(ri+1, 0), col_bg)
        _cell_para(tbl.cell(ri+1, 0), phase, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(tbl.cell(ri+1, 1), act, size=9)
        _cell_para(tbl.cell(ri+1, 2), deliv, size=9, color=NAVY)
        for mi in range(min(months, 5)):
            c = tbl.cell(ri+1, 3+mi)
            if mi in phase_month_map.get(ri, []):
                _set_cell_bg(c, col_bg)
                _cell_para(c, "●", bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        # AMC column
        amc_c = tbl.cell(ri+1, 3+min(months,5))
        if ri == 5:
            _set_cell_bg(amc_c, col_bg)
            _cell_para(amc_c, "●", bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_team_table(doc, approach: dict):
    team = [
        ("Project Manager", f"10+ yrs", "Overall delivery, client coordination, milestone tracking, SLA oversight"),
        ("Technical Architect / Lead", "8+ yrs", "Architecture design, technology decisions, code review, integration oversight"),
        ("Senior Developer 1", "5+ yrs", "Core module development, API integration, performance optimisation"),
        ("Senior Developer 2", "5+ yrs", "Frontend / Mobile development, UI/UX implementation"),
        ("GIS Specialist", "6+ yrs GIS", "GeoServer, PostGIS, spatial queries, GIS layer management"),
        ("QA Engineer", "5+ yrs", "Test planning, automated testing, performance testing, UAT coordination"),
        ("AMC Support Engineer", "3+ yrs", "Post go-live support, bug triage, enhancements, SLA monitoring"),
    ]
    tbl = doc.add_table(rows=1+len(team), cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Role", "Experience", "Responsibilities"]):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (role, exp, resp) in enumerate(team):
        _set_cell_bg(tbl.cell(ri+1, 0), "EEF2F7")
        _cell_para(tbl.cell(ri+1, 0), role, bold=True, color=NAVY, size=10)
        _cell_para(tbl.cell(ri+1, 1), exp, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(tbl.cell(ri+1, 2), resp, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_qa_table(doc):
    qa_data = [
        ("Unit Testing", "TDD — all business logic tested before integration", "JUnit / Pytest — 80%+ code coverage"),
        ("Integration Testing", "API-level integration + third-party connector tests", "Postman + Newman in CI pipeline"),
        ("Device / Browser Testing", "15+ Android devices, 5+ iOS devices, major browsers", "BrowserStack + physical device lab"),
        ("Performance Testing", "Load testing, stress testing, database query optimisation", "Apache JMeter + profiling tools"),
        ("Security Testing", "OWASP Top 10, penetration testing, auth validation", "OWASP ZAP + manual review"),
        ("User Acceptance Testing", "Structured UAT with client stakeholders and end users", "Test scenarios, sign-off checklist"),
    ]
    tbl = doc.add_table(rows=1+len(qa_data), cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Testing Type", "Methodology", "Tools"]):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (tp, meth, tools) in enumerate(qa_data):
        _set_cell_bg(tbl.cell(ri+1, 0), "EEF2F7")
        _cell_para(tbl.cell(ri+1, 0), tp, bold=True, color=NAVY, size=10)
        _cell_para(tbl.cell(ri+1, 1), meth, size=9)
        _cell_para(tbl.cell(ri+1, 2), tools, size=9, color=BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_pq_table(doc, pq_criteria: list):
    headers = ["Sl.", "Clause", "Criteria", "Status", "Nascent Remark"]
    tbl = doc.add_table(rows=1+len(pq_criteria), cols=5)
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        _set_cell_bg(c, "1F3864")
        _cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    status_colors = {"GREEN": ("E2EFDA", "375623"), "AMBER": ("FFF2CC", "7F6000"), "RED": ("FCE4D6", "C00000")}
    for ri, pq in enumerate(pq_criteria):
        color_key = pq.get("nascent_color", "AMBER")
        bg, tc = status_colors.get(color_key, ("EBF3FB", "2E75B6"))
        _cell_para(tbl.cell(ri+1, 0), str(pq.get("sl_no", ri+1)), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(tbl.cell(ri+1, 1), pq.get("clause_ref", ""), size=9)
        _cell_para(tbl.cell(ri+1, 2), pq.get("criteria", ""), size=9)
        _set_cell_bg(tbl.cell(ri+1, 3), bg)
        _cell_para(tbl.cell(ri+1, 3), pq.get("nascent_status", ""), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor.from_string(tc))
        _cell_para(tbl.cell(ri+1, 4), pq.get("nascent_remark", ""), size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_differentiator_table(doc, diffs: list):
    tbl = doc.add_table(rows=(len(diffs)+1)//2, cols=2)
    tbl.style = 'Table Grid'
    colors = ["1D4E8C", "2E75B6", "0070C0", "005A9E", "1F3864", "0D3B66", "1A4060", "0A2040"]
    for i, (title, desc) in enumerate(diffs):
        r, c = divmod(i, 2)
        cell = tbl.cell(r, c)
        _set_cell_bg(cell, colors[i % len(colors)])
        cell.paragraphs[0].clear()
        tp = cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tr = tp.add_run(title)
        tr.bold = True; tr.font.size = Pt(11); tr.font.color.rgb = WHITE; tr.font.name = "Arial"
        dp = cell.add_paragraph()
        dr = dp.add_run(desc)
        dr.font.size = Pt(9); dr.font.color.rgb = _rgb(222, 235, 247); dr.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_declaration(doc, profile: dict, tender_no: str, today: str):
    decl = (
        f"We, {profile['name']} (CIN: {profile['cin']}), do hereby solemnly declare:\n\n"
        f"1. All information, facts, and claims in this Technical Proposal are true, accurate, and complete.\n"
        f"2. We have thoroughly studied the tender document and unconditionally accept all terms and conditions.\n"
        f"3. We possess the technical capability, qualified team, and financial resources to deliver this project.\n"
        f"4. Our company has NOT been blacklisted by any Central/State Government, PSU, or ULB in India.\n"
        f"5. This proposal shall remain valid for 120 days from bid submission date."
    )
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.cell(0, 0)
    _set_cell_bg(c, "EBF3FB")
    c.paragraphs[0].clear()
    for line in decl.split("\n"):
        if line:
            p = c.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    _add_para(doc, f"For {profile['name']}", bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    _add_para(doc, "_" * 45)
    _add_para(doc, profile["signatory"], bold=True)
    _add_para(doc, profile["designation"])
    _add_para(doc, profile["name"], bold=True)
    _add_para(doc, f"Date: {today}                    Place: Ahmedabad, Gujarat")
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    _add_para(doc, "Company Seal:", bold=True)
