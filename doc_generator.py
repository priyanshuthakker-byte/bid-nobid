"""
BidDocGenerator v4 - Clean rewrite matching KELTRON format
No emojis. Plain status text. Human language.
FIXED: verdict color now reads from data["verdict_color"] fallback
FIXED: scope_items renders correctly for both dict and string formats (v7 AI output)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, Any, List
import re as _re

C = {
    "dark_blue": "0A1F3D",   # deep navy
    "mid_blue":  "1A5276",   # rich blue
    "light_blue":"C8DDEF",
    "label_col": "D0E8F7",
    "alt_row":   "F0F7FD",
    "white":     "FFFFFF",
    "green_bg":  "E6F4EA",
    "green_text":"1E5631",
    "amber_bg":  "FFF8E1",
    "amber_text":"7A5800",
    "red_bg":    "FDECE8",
    "red_text":  "9B1C1C",
    "blue_bg":   "EBF4FB",
    "blue_text": "0A1F3D",
    "orange":    "C2560B",
    "gray":      "6B7280",
    "dark":      "1A1A2E",
    "gold":      "C9A800",
    "navy_mid":  "1B3F6E",
}

STATUS_STYLE = {
    "GREEN": ("green_bg", "green_text"),
    "AMBER": ("amber_bg", "amber_text"),
    "RED":   ("red_bg",   "red_text"),
    "BLUE":  ("blue_bg",  "blue_text"),
}

def hex_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_borders(cell, color="9DC3E6", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for b in ["top", "left", "bottom", "right"]:
        el = OxmlElement("w:" + b)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color.lstrip("#"))
        tcB.append(el)
    tcPr.append(tcB)

def set_table_borders(table, color="2E75B6"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblB = OxmlElement("w:tblBorders")
    for b in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement("w:" + b)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color.lstrip("#"))
        tblB.append(el)
    tblPr.append(tblB)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)

def add_run(para, text, bold=False, size=10, color=None, italic=False):
    r = para.add_run(str(text) if text else "")
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        c = hex_rgb(color) if isinstance(color, str) else color
        r.font.color.rgb = RGBColor(*c)
    return r

def cell_write(cell, text, bold=False, size=9, color=None, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, pad=5):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(pad)
    add_run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def strip_emojis(text):
    if not text:
        return text
    return _re.sub(
        r"[\U00010000-\U0010ffff"
        r"\u2600-\u26FF\u2700-\u27BF"
        r"\u2300-\u23FF\u2B50-\u2BFF"
        r"\u2100-\u214F\u0080-\u00FF"
        r"\u2000-\u206F]",
        "", text
    ).strip()

def clean_status(s):
    s = strip_emojis(str(s))
    for old, new in [
        ("MEETS", "Met"), ("MET", "Met"),
        ("DOES NOT MEET", "Not Met"), ("NOT MET", "Not Met"),
        ("Critical", "Not Met"), ("CRITICAL", "Not Met"),
        ("CONDITIONAL", "Conditional"), ("Pending", "Conditional"), ("PENDING", "Conditional"),
        ("REVIEW", "Review"),
    ]:
        if old.lower() in s.lower():
            return new
    if "Met" in s: return "Met"
    if "Not" in s: return "Not Met"
    if "Cond" in s: return "Conditional"
    return "Review"

def status_color(status_text):
    s = status_text.lower()
    if "not met" in s or "critical" in s: return "RED"
    if "conditional" in s or "pending" in s: return "AMBER"
    if "met" in s: return "GREEN"
    return "BLUE"


def field_value(v):
    if isinstance(v, dict):
        return str(v.get("value", "—") or "—")
    if v is None:
        return "—"
    return str(v)


def scope_item_text(item) -> str:
    if isinstance(item, dict):
        parts = []
        sec_no = item.get("section_no", "")
        title = item.get("section_title", "") or item.get("title", "")
        prose = item.get("prose", "") or item.get("description", "")
        phase = item.get("phase", "")
        tech = item.get("tech_specified", "")
        deliverables = item.get("deliverables", [])
        if sec_no and title:
            parts.append(f"[{sec_no}] {title}")
        elif title:
            parts.append(title)
        if phase and phase not in ("—", ""):
            parts.append(f"Phase: {phase}")
        if prose:
            parts.append(prose)
        if tech and tech not in ("—", ""):
            parts.append(f"Technology: {tech}")
        if deliverables:
            dl = [str(d) for d in deliverables if d]
            if dl:
                parts.append("Deliverables: " + " | ".join(dl))
        return " — ".join(parts) if parts else str(item)
    return strip_emojis(str(item)) if item else ""


class BidDocGenerator:

    def generate(self, data: Dict[str, Any], output_path: str):
        self.doc = Document()
        self._setup_page()
        self._header_block(data)
        self._section_snapshot(data)
        self._section_pq(data)
        self._section_scope(data)
        self._section_payment(data)
        self._section_recommendation(data)
        self._section_notes(data)
        self._section_authorization()
        self._footer(data)
        self.doc.save(output_path)

    def _setup_page(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.left_margin = sec.right_margin = Cm(1.8)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        self.doc.styles["Normal"].font.name = "Calibri"
        self.doc.styles["Normal"].font.size = Pt(10)
        hdr = sec.header
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("Nascent Info Technologies Pvt. Ltd.  |  BID / NO-BID ANALYSIS  |  CONFIDENTIAL")
        hr.font.name = "Calibri"
        hr.font.size = Pt(7)
        hr.font.color.rgb = RGBColor(*hex_rgb("6B7280"))
        hr.font.italic = True
        ftr = sec.footer
        fp = ftr.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("Page ")
        fr.font.size = Pt(7); fr.font.name = "Calibri"; fr.font.color.rgb = RGBColor(*hex_rgb("6B7280"))
        fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
        fp.runs[-1]._r.append(fld)
        instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
        run2 = OxmlElement("w:r"); run2.append(instrText)
        fp._p.append(run2)
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run3 = OxmlElement("w:r"); run3.append(fld2)
        fp._p.append(run3)
        fr2 = fp.add_run("  |  Nascent Info Technologies Pvt. Ltd.  |  For Internal Use Only")
        fr2.font.size = Pt(7); fr2.font.name = "Calibri"; fr2.font.color.rgb = RGBColor(*hex_rgb("6B7280")); fr2.font.italic = True

    def _sec_heading(self, number, title, source_note=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        rn = p.add_run("  " + number + "  ")
        rn.font.name = "Calibri"
        rn.font.size = Pt(11)
        rn.font.bold = True
        rn.font.color.rgb = RGBColor(*hex_rgb("C9A800"))
        r = p.add_run("  " + title.upper() + "  ")
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), C["dark_blue"])
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
        p_line = self.doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(4)
        pPr2 = p_line._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), "C9A800")
        pBdr.append(bot); pPr2.append(pBdr)
        if source_note:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(4)
            add_run(p2, source_note, size=8, italic=True, color=C["navy_mid"])

    def _header_block(self, data):
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color=C["dark_blue"])

        c0 = table.rows[0].cells[0]
        c0.width = Cm(10)
        set_bg(c0, C["dark_blue"])
        set_borders(c0, color="FFFFFF", size=6)
        p = c0.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        add_run(p, "Nascent Info Technologies Pvt. Ltd.", bold=True, size=11, color="FFFFFF")
        p2 = c0.add_paragraph()
        add_run(p2, "A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad 380015", size=8, color="BDD7EE")
        p3 = c0.add_paragraph()
        add_run(p3, "www.nascentinfo.com | nascent.tender@nascentinfo.com", size=8, color="BDD7EE")
        p4 = c0.add_paragraph()
        p4.paragraph_format.space_after = Pt(4)
        add_run(p4, "MSME | CMMI L3 | ISO 9001 | ISO 27001 | ISO 20000", size=8, color="FFF2CC")

        c1 = table.rows[0].cells[1]
        c1.width = Cm(15.5)
        set_bg(c1, C["navy_mid"])
        set_borders(c1, color="FFFFFF", size=6)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        add_run(p, "BID / NO-BID ANALYSIS REPORT", bold=True, size=18, color="FFFFFF")

        tender_title = strip_emojis(field_value(data.get("tender_name", data.get("org_name", ""))))[:100]
        p2 = c1.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        add_run(p2, tender_title, bold=False, size=9, color="C8DDEF")

        verdict_data = data.get("overall_verdict", {})
        verdict = strip_emojis(verdict_data.get("verdict", "PENDING REVIEW"))
        vcolor = verdict_data.get("color") or data.get("verdict_color", "BLUE")
        v_bg  = {"GREEN": "1E5631", "AMBER": "7A5800", "RED": "9B1C1C", "BLUE": "1A5276"}.get(vcolor, "1A5276")
        v_txt = {"GREEN": "E6F4EA", "AMBER": "FFF8E1", "RED": "FDECE8", "BLUE": "EBF4FB"}.get(vcolor, "EBF4FB")

        p3 = c1.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(6)
        p3.paragraph_format.space_after = Pt(2)
        pPr3 = p3._p.get_or_add_pPr()
        shd3 = OxmlElement("w:shd"); shd3.set(qn("w:fill"), v_bg); shd3.set(qn("w:val"), "clear")
        pPr3.append(shd3)
        add_run(p3, "  " + verdict + "  ", bold=True, size=15, color=v_txt)

        p4 = c1.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_after = Pt(4)
        add_run(p4, "Prepared: " + datetime.now().strftime("%d-%b-%Y") + "  |  Nascent Bid Team",
                size=8, color="C8DDEF")

        self.doc.add_paragraph()

    def _section_snapshot(self, data):
        self._sec_heading("1", "Tender Overview")
        fields = [
            ("Tender No.",               field_value(data.get("tender_no", "—"))),
            ("Tender ID",                field_value(data.get("tender_id", "—"))),
            ("T247 ID",                  field_value(data.get("t247_id", "—"))),
            ("Portal / Website",         field_value(data.get("portal", "—"))),
            ("Organization / Department",field_value(data.get("org_name", "—"))),
            ("Tender Name",              field_value(data.get("tender_name", "—"))),
            ("Form of Contract",         field_value(data.get("tender_type", "—"))),
            ("Bid Submission Start Date",field_value(data.get("bid_start_date", "—"))),
            ("Bid Submission End Date",  field_value(data.get("bid_submission_date", "—"))),
            ("Bid Opening Date",         field_value(data.get("bid_opening_date", "—"))),
            ("Mode of Selection",        field_value(data.get("mode_of_selection", "—"))),
            ("Pre-Bid Meeting",          field_value(data.get("prebid_meeting", "Not specified"))),
            ("Pre-Bid Query Deadline",   field_value(data.get("prebid_query_date", "Not specified"))),
            ("Estimated Cost",           field_value(data.get("estimated_cost", "Not specified"))),
            ("Tender Fee",               field_value(data.get("tender_fee", "Not specified"))),
            ("EMD",                      field_value(data.get("emd", "Not specified"))),
            ("EMD Exemption",            field_value(data.get("emd_exemption", "—"))),
            ("Performance Bank Guarantee",field_value(data.get("performance_security", "As per tender"))),
            ("Bid Validity",             field_value(data.get("bid_validity", "—"))),
            ("Period of Work",           field_value(data.get("contract_period", "—"))),
            ("Post-Implementation Support", field_value(data.get("post_implementation", "—"))),
            ("Project Location",         field_value(data.get("location", "—"))),
            ("Contact",                  field_value(data.get("contact", "—"))),
            ("JV / Consortium Allowed",  field_value(data.get("jv_allowed", "Not specified"))),
        ]

        table = self.doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color=C["mid_blue"])

        highlight_keys = ["Bid Submission End Date", "EMD", "Tender Fee", "JV / Consortium Allowed"]
        for idx, (key, val) in enumerate(fields):
            row = table.add_row()
            row.height = Cm(0.75)
            bg_l = "D6E4F0" if idx % 2 == 0 else "E8F0F8"
            bg_v = C["white"] if idx % 2 == 0 else C["alt_row"]

            c0 = row.cells[0]
            c0.width = Cm(7)
            set_bg(c0, bg_l)
            set_borders(c0)
            cell_write(c0, key, bold=True, size=9, color=C["dark_blue"])

            c1 = row.cells[1]
            c1.width = Cm(18.5)
            set_bg(c1, bg_v)
            set_borders(c1)
            hl = key in highlight_keys and val not in ["—", "Not specified", ""]
            cell_write(c1, str(val), bold=hl, size=9, color=(C["orange"] if hl else None))

        self.doc.add_paragraph()

    def _criteria_table(self, criteria, headers, col_w):
        if not criteria:
            p = self.doc.add_paragraph()
            add_run(p, "No criteria extracted — refer tender document.", italic=True, size=9)
            return

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color=C["mid_blue"])

        hrow = table.rows[0]
        repeat_header(hrow)
        for cell, hdr in zip(hrow.cells, headers):
            set_bg(cell, C["dark_blue"])
            set_borders(cell, color="FFFFFF", size=4)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, hdr, bold=True, size=9, color="FFFFFF")

        for ri, item in enumerate(criteria):
            row = table.add_row()
            bg = C["white"] if ri % 2 == 0 else C["alt_row"]

            raw_status = item.get("nascent_status", "Review")
            sym = clean_status(raw_status)
            sc = item.get("nascent_color") or status_color(sym)
            s_bg, s_tc = STATUS_STYLE.get(sc, ("blue_bg", "blue_text"))

            c = row.cells[0]
            set_bg(c, C["label_col"])
            set_borders(c)
            cell_write(c, str(item.get("sl_no", ri + 1)), bold=True, size=9,
                       color=C["dark_blue"], align=WD_ALIGN_PARAGRAPH.CENTER)

            c = row.cells[1]
            set_bg(c, bg)
            set_borders(c)
            ref = item.get("clause_ref", "—")
            page = item.get("page_no", "")
            ref_text = f"{ref}\n{page}" if page and page != "—" else ref
            cell_write(c, ref_text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

            c = row.cells[2]
            set_bg(c, bg)
            set_borders(c)
            cell_write(c, item.get("criteria", ""), size=9)

            c = row.cells[3]
            set_bg(c, bg)
            set_borders(c)
            cell_write(c, item.get("details", "") or item.get("documents_required", ""), size=9)

            c = row.cells[4]
            set_bg(c, C[s_bg])
            set_borders(c)
            cell_write(c, sym, bold=True, size=8, color=C[s_tc],
                       align=WD_ALIGN_PARAGRAPH.CENTER)

            c = row.cells[5]
            set_bg(c, C[s_bg] if sc != "BLUE" else bg)
            set_borders(c)
            remark = strip_emojis(item.get("nascent_remark", ""))
            calc = item.get("calculation_shown", "")
            if calc and calc != "—":
                remark = remark + ("\n" if remark else "") + f"Calc: {calc}"
            cell_write(c, remark, size=8)

    def _section_pq(self, data):
        self._sec_heading("2", "Pre-Qualification (PQ) Criteria",
                          "Criteria reproduced word-for-word from tender | Nascent status checked against company profile")
        headers = [
            "Sr.",
            "Clause No.\nPage No.",
            "Eligibility Criteria\n(word-for-word from tender)",
            "Supporting Documents\nRequired",
            "Nascent\nStatus",
            "Remarks / Action Required",
        ]
        col_w = [Cm(0.9), Cm(2.3), Cm(8.5), Cm(5.0), Cm(2.0), Cm(6.8)]
        self._criteria_table(data.get("pq_criteria", []), headers, col_w)

        tq = data.get("tq_criteria", [])
        if tq:
            self._sec_heading("2B", "Technical Qualification (TQ) Criteria",
                              "Marking scheme and slab calculations")
            tq_headers = [
                "Sr.",
                "Clause\nPage",
                "Criteria\n(word-for-word)",
                "Evaluation Criteria\n(Marks / Slabs)",
                "Score",
                "Remarks / Calculation",
            ]
            mapped = []
            for item in tq:
                if not isinstance(item, dict):
                    continue
                mapped.append({
                    "sl_no": item.get("sl_no", ""),
                    "clause_ref": item.get("clause_ref", "—"),
                    "page_no": item.get("page_no", "—"),
                    "criteria": item.get("criteria", ""),
                    "details": item.get("eval_criteria", "") or item.get("details", ""),
                    "nascent_status": item.get("nascent_status", "Review"),
                    "nascent_color": item.get("nascent_color", "BLUE"),
                    "nascent_remark": (
                        f"Score: {item.get('nascent_score','?')}/{item.get('max_marks','?')}\n"
                        + strip_emojis(str(item.get("slab_calculation", "") or ""))
                        + ("\n" + strip_emojis(str(item.get("nascent_remark", "") or "")) if item.get("nascent_remark") else "")
                    ),
                    "calculation_shown": "",
                })
            self._criteria_table(mapped, tq_headers, col_w)

        self.doc.add_paragraph()

    def _section_scope(self, data):
        self._sec_heading("3", "Scope of Work",
                          "Source: Tender document — key deliverables and phases")

        scope_sections = data.get("scope_sections", [])
        scope_items = data.get("scope_items", [])
        items_to_render = scope_sections if scope_sections else scope_items

        bg = data.get("scope_background", "")
        if bg and isinstance(bg, str) and len(bg) > 10:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_run(p, strip_emojis(bg), size=9, italic=True)

        if not items_to_render:
            p = self.doc.add_paragraph()
            add_run(p, "Refer to tender document for scope of work.", italic=True, size=9)
            self.doc.add_paragraph()
            return

        for item in items_to_render:
            text = scope_item_text(item)
            if not text:
                continue
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, "• ", size=9)
            add_run(p, text, size=9)

        integrations = data.get("key_integrations", [])
        if integrations:
            p = self.doc.add_paragraph()
            add_run(p, "Key Integrations:", bold=True, size=9)
            for intg in integrations:
                if not isinstance(intg, dict):
                    continue
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.3)
                p2.paragraph_format.space_after = Pt(2)
                txt = f"{intg.get('system','')} ({intg.get('type','')}) — {intg.get('purpose','')}"
                add_run(p2, "• " + txt, size=8)

        self.doc.add_paragraph()

    def _section_payment(self, data):
        self._sec_heading("4", "Payment Terms & Timeline",
                          "Note: Detailed payment schedule from tender document.")

        sched = data.get("payment_schedule", [])
        items = data.get("payment_terms", [])

        if sched:
            table = self.doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table, color=C["mid_blue"])
            hrow = table.rows[0]
            for cell, hdr in zip(hrow.cells, ["Milestone", "Activity / Trigger", "Timeline", "%", "Phase"]):
                set_bg(cell, C["dark_blue"])
                set_borders(cell, color="FFFFFF")
                cell_write(cell, hdr, bold=True, size=9, color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for ri, row_data in enumerate(sched):
                if not isinstance(row_data, dict):
                    continue
                row = table.add_row()
                bg = C["white"] if ri % 2 == 0 else C["alt_row"]
                vals = [
                    row_data.get("milestone_name", row_data.get("milestone_no", "")),
                    row_data.get("trigger_activity", ""),
                    row_data.get("timeline", ""),
                    row_data.get("payment_percent", ""),
                    row_data.get("phase", ""),
                ]
                for cell, val in zip(row.cells, vals):
                    set_bg(cell, bg)
                    set_borders(cell)
                    cell_write(cell, str(val), size=9)
        elif items:
            for item in items:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(3)
                add_run(p, "• ", size=9)
                add_run(p, strip_emojis(str(item)), size=9)
        else:
            fallback = [
                "Period of work: " + field_value(data.get("contract_period", "As per tender")),
                "EMD: " + field_value(data.get("emd", "As per tender")),
                "Performance Bank Guarantee: " + field_value(data.get("performance_security", "As per tender")),
                "Payment schedule: Not explicitly defined — refer tender document.",
                "Penalty / LD clause: Refer tender document for applicable clauses.",
            ]
            for item in fallback:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(3)
                add_run(p, "• ", size=9)
                add_run(p, strip_emojis(str(item)), size=9)

        penalties = data.get("penalty_clauses", [])
        if penalties:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            add_run(p, "Penalty / Risk Clauses:", bold=True, size=9, color=C["red_text"])
            for pen in penalties:
                if not isinstance(pen, dict):
                    continue
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.2)
                p2.paragraph_format.space_after = Pt(2)
                txt = f"{pen.get('type','')} — {pen.get('condition','')} | Penalty: {pen.get('penalty','')} | Cap: {pen.get('max_cap','—')}"
                add_run(p2, "• " + strip_emojis(txt), size=8, color=C["red_text"])

        self.doc.add_paragraph()

    def _section_recommendation(self, data):
        self._sec_heading("5", "Bid / No-Bid Recommendation")

        verdict_data = data.get("overall_verdict", {})
        green_count = verdict_data.get("green", 0)
        amber_count = verdict_data.get("amber", 0)
        red_count   = verdict_data.get("red", 0)

        stats = [
            (str(green_count), "Criteria Met",  "green_bg", "green_text"),
            (str(amber_count), "Conditional",   "amber_bg", "amber_text"),
            (str(red_count),   "Not Met",        "red_bg",   "red_text"),
        ]
        table = self.doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color=C["mid_blue"])
        for cell, (count, label, bg, tc) in zip(table.rows[0].cells, stats):
            set_bg(cell, C[bg])
            set_borders(cell, color="FFFFFF", size=6)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            add_run(p, count, bold=True, size=18, color=C[tc])
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(4)
            add_run(p2, label, size=9, color=C[tc])

        self.doc.add_paragraph()

        vcolor = verdict_data.get("color") or data.get("verdict_color", "BLUE")
        s_bg, s_tc = STATUS_STYLE.get(vcolor, ("blue_bg", "blue_text"))

        verdict = strip_emojis(verdict_data.get("verdict", "PENDING REVIEW"))
        reason  = strip_emojis(verdict_data.get("reason", ""))

        v_fill = {"GREEN":"1E5631","AMBER":"7A5800","RED":"9B1C1C","BLUE":"1A5276"}.get(vcolor,"1A5276")
        v_fc   = {"GREEN":"E6F4EA","AMBER":"FFF8E1","RED":"FDECE8","BLUE":"EBF4FB"}.get(vcolor,"EBF4FB")
        pv = self.doc.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv.paragraph_format.space_before = Pt(8)
        pv.paragraph_format.space_after  = Pt(4)
        pPrv = pv._p.get_or_add_pPr()
        shdv = OxmlElement("w:shd"); shdv.set(qn("w:fill"), v_fill); shdv.set(qn("w:val"), "clear")
        pPrv.append(shdv)
        add_run(pv, "   " + verdict + "   ", bold=True, size=22, color=v_fc)

        try:
            from nascent_checker import load_profile
            p = load_profile()
            fin = p.get("finance", {})
            emp = p.get("employees", {})
            certs = p.get("certifications", {})
            turnover_str = f"Rs. {fin.get('avg_turnover_last_3_fy', 17.18):.2f} Cr (avg last 3 FY). Net worth Rs. {fin.get('net_worth_cr', 26.09):.2f} Cr."
            emp_str = f"{emp.get('total_confirmed', 67)} employees — {emp.get('gis_staff', 11)} GIS, {emp.get('it_dev_staff', 21)} IT/Dev."
            cmmi = certs.get("cmmi", {})
            cmmi_str = f"CMMI {cmmi.get('version', 'V2.0')} Level {cmmi.get('level', 3)} — valid till {cmmi.get('valid_to', '19-Dec-2026')}."
        except Exception:
            turnover_str = "Rs. 17.18 Cr (avg last 3 FY). Net worth Rs. 26.09 Cr."
            emp_str = "67 employees — 11 GIS, 21 IT/Dev, plus QA, PM, BA teams."
            cmmi_str = "CMMI V2.0 Level 3 — valid till 19-Dec-2026."

        proj_matches = data.get("project_matches", [])
        proj_str = ""
        if proj_matches:
            proj_str = " | ".join([
                f"{pm.get('matching_project','')}: {pm.get('relevance','')[:60]}"
                for pm in proj_matches[:3] if isinstance(pm, dict)
            ])

        strengths = data.get("key_strengths", [])
        strengths_str = " | ".join([strip_emojis(str(s))[:80] for s in strengths[:3]]) if strengths else ""

        assessment_rows = [
            ("Financial Eligibility",   turnover_str),
            ("Company Registration",    "Private Limited Company since 2006. 19 years in operation."),
            ("CMMI Certification",      cmmi_str),
            ("ISO Certifications",      "ISO 9001:2015, ISO 27001:2022, ISO 20000-1:2018 — all valid till Sep-2028."),
            ("GIS / IT Experience",     "9 major projects — AMC, JuMC, VMC, KVIC, PCSCL, TCGL."),
            ("Mobile GIS Experience",   "KVIC mobile GIS (PAN India), BMC mobile app, AMC Heritage App."),
            ("Employee Strength",       emp_str),
            ("MSME Status",             "UDYAM-GJ-01-0007420 — eligible for EMD exemption."),
            ("PQ Criteria Summary",     f"Met: {green_count}  Conditional: {amber_count}  Not Met: {red_count}"),
        ]
        if proj_str:
            assessment_rows.append(("Matching Projects", proj_str))
        if strengths_str:
            assessment_rows.append(("Key Strengths", strengths_str))
        assessment_rows.append(("FINAL RECOMMENDATION", verdict + ("\n" + reason if reason else "")))

        tbl2 = self.doc.add_table(rows=0, cols=2)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl2, color=C["mid_blue"])
        for ri, (key, val) in enumerate(assessment_rows):
            row = tbl2.add_row()
            is_final = (key == "FINAL RECOMMENDATION")
            c0 = row.cells[0]
            c0.width = Cm(8)
            set_bg(c0, C["dark_blue"] if is_final else C["label_col"])
            set_borders(c0)
            cell_write(c0, key, bold=True,
                       size=10 if is_final else 9,
                       color=("FFFFFF" if is_final else C["dark_blue"]))
            c1 = row.cells[1]
            c1.width = Cm(17.5)
            set_bg(c1, C[s_bg] if is_final else (C["white"] if ri % 2 == 0 else C["alt_row"]))
            set_borders(c1)
            cell_write(c1, val, bold=is_final,
                       size=11 if is_final else 9,
                       color=(C[s_tc] if is_final else None))

        action_items = data.get("action_items", [])
        if action_items:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            add_run(p, "Action Items:", bold=True, size=10, color=C["dark_blue"])
            for ai in action_items[:8]:
                if not isinstance(ai, dict):
                    continue
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.2)
                p2.paragraph_format.space_after = Pt(2)
                priority = ai.get("priority", "")
                action = strip_emojis(ai.get("action", ""))
                by = ai.get("target_date", "")
                clr = C["red_text"] if priority == "URGENT" else C["amber_text"] if priority == "HIGH" else C["dark_blue"]
                add_run(p2, f"[{priority}] ", bold=True, size=9, color=clr)
                add_run(p2, action, size=9)
                if by:
                    add_run(p2, f" — by {by}", size=8, italic=True)

        self.doc.add_paragraph()

    def _section_notes(self, data):
        self._sec_heading("6", "Notes & Action Items",
                          "Items to be resolved before committing to bid.")
        action_items = []

        for item in data.get("pq_criteria", []):
            sc = item.get("nascent_color", "BLUE")
            if sc in ["RED", "AMBER"]:
                priority = "URGENT" if sc == "RED" else "ACTION"
                crit   = strip_emojis(item.get("criteria", ""))[:80]
                remark = strip_emojis(item.get("nascent_remark", ""))[:200]
                action_items.append({
                    "priority": priority,
                    "text":   crit,
                    "detail": remark,
                    "color":  C["red_text"] if sc == "RED" else C["amber_text"],
                    "bg":     "red_bg"   if sc == "RED" else "amber_bg",
                })
