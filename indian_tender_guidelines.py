"""
Indian Tender Guidelines Engine v1.0
Nascent Info Technologies Bid/No-Bid System

Complete database of ALL Indian government tender procurement rules.
Cross-checks every RFP clause against:
  1. Legal compliance (what they CAN and CANNOT ask)
  2. Nascent's specific profile (where we meet / fall short)
  3. Pre-bid query generation (what to challenge / clarify)

Sources:
  - General Financial Rules (GFR) 2017, Chapter 6
  - MSME Public Procurement Policy Order 2012
  - DPIIT Startup India Procurement Guidelines
  - CVC Guidelines on Transparency in Procurement
  - Make in India / Public Procurement Order 2017 (Rule 144 GFR)
  - GeM Guidelines and Prohibited Clauses
  - Ministry of Finance Manual for Procurement of Goods 2022
  - Gujarat State Procurement Rules

Author: Built for Nascent Info Technologies Pvt. Ltd.
"""

import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

# ─────────────────────────────────────────────────────────────────────────────
# NASCENT PROFILE (what we have / don't have)
# ─────────────────────────────────────────────────────────────────────────────

NASCENT = {
    "name":            "Nascent Info Technologies Pvt. Ltd.",
    "type":            "Private Limited",
    "msme":            True,
    "msme_reg":        "UDYAM-GJ-01-0007420",
    "msme_category":   "Small Enterprise",
    "pan":             "AACCN3670J",
    "gstin":           "24AACCN3670J1ZG",
    "incorporated":    "23 June 2006",
    "years_in_ops":    19,
    "state":           "Gujarat",
    "hq":              "Ahmedabad",
    "branch_offices":  [],   # No branches in other states
    "employees":       67,
    "turnover_cr": {
        "fy2223": 16.36,
        "fy2324": 16.36,
        "fy2425": 18.83,
        "avg_3yr": 17.18,
        "avg_2yr": 17.60,
    },
    "net_worth_cr":    26.09,
    "certifications": {
        "cmmi":    "CMMI V2.0 Level 3 (valid till 19-Dec-2026)",
        "iso9001": "ISO 9001:2015 (valid till 08-Sep-2028)",
        "iso27001":"ISO/IEC 27001:2022 (valid till 08-Sep-2028)",
        "iso20000":"ISO/IEC 20000-1:2018 (valid till 08-Sep-2028)",
        "cert_in": None,     # NOT held
        "stqc":    None,     # NOT held
        "nsic":    None,     # NOT registered (MSME Udyam instead)
    },
    "domains": [
        "GIS", "Mobile Applications", "Smart City", "eGovernance",
        "Web Portals", "ERP", "Survey & Mapping", "IT Services",
        "Data Analytics", "AMC / Maintenance"
    ],
    "oem_partnerships": [],  # No SAP, Oracle, Microsoft OEM
    "blacklisted":     False,
    "govt_debarred":   False,
    "make_in_india":   True,   # Software developed in India
    "startup":         False,  # 19 years old
}


# ─────────────────────────────────────────────────────────────────────────────
# COMPLETE GUIDELINES DATABASE
# ─────────────────────────────────────────────────────────────────────────────

GUIDELINES = {

    # ═══════════════════════════════════════════════════════════════════
    # 1. EMD / BID SECURITY
    # ═══════════════════════════════════════════════════════════════════
    "emd": {
        "title": "Earnest Money Deposit (EMD) / Bid Security",
        "rules": [
            {
                "id": "EMD-01",
                "rule": "GFR 2017 Rule 170 + MSME Policy 2012",
                "what_law_says": "EMD must be between 2% and 5% of estimated tender value only. Cannot exceed 5%.",
                "nascent_right": "As a registered MSME (UDYAM-GJ-01-0007420), Nascent is FULLY EXEMPT from paying EMD. This is mandatory — not at department's discretion.",
                "challenge_if": "EMD exceeds 5% of value | MSME exemption not mentioned | Department denies MSME exemption",
                "pre_bid_query": "The tender prescribes EMD of Rs. {emd_amount}. As per MSME Public Procurement Policy Order 2012 read with GFR 2017 Rule 170, Micro and Small Enterprises registered under Udyam Registration are fully exempt from payment of Earnest Money Deposit. Nascent Info Technologies Pvt. Ltd. is registered as a Small Enterprise vide UDYAM-GJ-01-0007420. We request the authority to: (a) Confirm that Nascent is exempt from EMD payment, and (b) Amend the bid document to explicitly state the MSME EMD exemption clause.",
                "legal_cite": "MSME Public Procurement Policy for MSEs Order 2012, Section 11 | GFR 2017 Rule 170 | Ministry of MSME Circular dated 26.03.2012",
                "severity": "HIGH",
            },
            {
                "id": "EMD-02",
                "rule": "GFR 2017 Rule 170",
                "what_law_says": "EMD can be in form of DD, Banker's Cheque, or Bank Guarantee only.",
                "nascent_right": "MSME exempt — no form required",
                "challenge_if": "Department asks for Cash EMD | Asks for FDR instead of BG | Asks for EMD from MSME despite exemption",
                "pre_bid_query": "Clause {clause_no} requires EMD in cash or FDR form. GFR 2017 Rule 170 permits EMD only in the form of Demand Draft, Banker's Cheque, or Bank Guarantee. Please clarify and amend accordingly.",
                "legal_cite": "GFR 2017 Rule 170(i)",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 2. PERFORMANCE SECURITY / BANK GUARANTEE
    # ═══════════════════════════════════════════════════════════════════
    "performance_security": {
        "title": "Performance Security / Bank Guarantee",
        "rules": [
            {
                "id": "PBG-01",
                "rule": "GFR 2017 Rule 171",
                "what_law_says": "Performance Security is 3% to 5% of contract value. NOT more than 10% in any case.",
                "nascent_right": "MSMEs are NOT exempt from Performance Security (unlike EMD). Must pay if awarded contract.",
                "challenge_if": "PBG > 10% of contract value | PBG validity unreasonably long (beyond warranty + 60 days)",
                "pre_bid_query": "Clause {clause_no} prescribes Performance Bank Guarantee of {pbg_pct}% of contract value. As per GFR 2017 Rule 171, Performance Security should be 3-5% of contract value. A PBG of {pbg_pct}% appears disproportionate and creates unnecessary financial burden. We request the authority to revise this to 5% (standard GFR rate) or justify the higher amount as per CVC guidelines.",
                "legal_cite": "GFR 2017 Rule 171 | CVC Circular No. 07/02/10 dated 04.02.2010",
                "severity": "MEDIUM",
            },
            {
                "id": "PBG-02",
                "rule": "GFR 2017 Rule 171",
                "what_law_says": "PBG must be returned within 60 days after completion of all contractual obligations including warranty.",
                "nascent_right": "Can insist on timely PBG release",
                "challenge_if": "No clear PBG release clause | Retention longer than 60 days post warranty",
                "pre_bid_query": "The tender document does not specify the timeline for release of Performance Bank Guarantee after project completion. As per GFR 2017 Rule 171, PBG must be returned within 60 days after completion of all contractual obligations including warranty period. We request inclusion of this clause explicitly.",
                "legal_cite": "GFR 2017 Rule 171(ii)",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 3. TURNOVER / FINANCIAL ELIGIBILITY
    # ═══════════════════════════════════════════════════════════════════
    "turnover": {
        "title": "Annual Turnover / Financial Eligibility Criteria",
        "rules": [
            {
                "id": "TRN-01",
                "rule": "MSME Policy 2012 + Ministry of MSME Circular 10.03.2016",
                "what_law_says": "MSMEs are exempt from minimum turnover criteria. No minimum turnover requirement for MSEs in public procurement.",
                "nascent_right": "As MSME, Nascent can request waiver of turnover criteria. However with Rs.17.18 Cr avg turnover, Nascent likely meets most criteria anyway.",
                "challenge_if": "Turnover > 3× estimated project value | Turnover mentioned without MSME exemption clause",
                "pre_bid_query": "Clause {clause_no} prescribes minimum average annual turnover of Rs. {turnover_cr} Cr for the last 3 financial years. As per Ministry of MSME Policy Circular No. 1(2)(1)/2016-MA dated 10.03.2016 and Public Procurement Policy for MSEs Order 2012, Micro and Small Enterprises are exempt from turnover criteria in government procurement. Nascent Info Technologies Pvt. Ltd. is registered as a Small Enterprise vide UDYAM-GJ-01-0007420 and requests: (a) Confirmation of MSME exemption from turnover criteria, OR (b) In the event turnover criteria is maintained, that it be reduced to a maximum of 3× the estimated project value as per established procurement norms.",
                "legal_cite": "MSME Policy Circular 1(2)(1)/2016-MA dated 10.03.2016 | GFR 2017 Rule 173",
                "severity": "HIGH",
            },
            {
                "id": "TRN-02",
                "rule": "General procurement norms + CVC guidelines",
                "what_law_says": "Turnover requirement must be proportionate to project size. Industry standard is maximum 3× estimated project value for average annual turnover.",
                "nascent_right": "Nascent turnover Rs.17.18 Cr avg. Can bid on projects up to ~Rs.5-6 Cr comfortably. For larger projects, raise pre-bid query.",
                "challenge_if": "Turnover requirement is more than 3× project value",
                "pre_bid_query": "The minimum turnover requirement of Rs. {turnover_cr} Cr appears disproportionate to the estimated project value of Rs. {project_value_cr} Cr (ratio: {ratio}×). Standard procurement practice limits turnover criteria to 3× the estimated project value to ensure fair competition. We request the authority to revise turnover criteria to Rs. {suggested_cr} Cr (3× project value) to encourage competitive bidding as per CVC guidelines on fair procurement.",
                "legal_cite": "CVC Circular on Transparency in Procurement | GFR 2017 Rule 173(xi)",
                "severity": "HIGH",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 4. EXPERIENCE / PRIOR WORK CRITERIA
    # ═══════════════════════════════════════════════════════════════════
    "experience": {
        "title": "Prior Experience / Similar Work Criteria",
        "rules": [
            {
                "id": "EXP-01",
                "rule": "MSME Policy 2012 + MSME Circular 10.03.2016",
                "what_law_says": "MSMEs are exempt from prior experience criteria subject to meeting quality and technical specifications.",
                "nascent_right": "Nascent has 19 years experience and 9+ completed municipal/GIS projects. Usually meets experience criteria. But can claim exemption if needed.",
                "challenge_if": "Experience criteria not met for specific domain | Only central/state govt experience accepted | Only projects in specific state accepted",
                "pre_bid_query": "Clause {clause_no} requires prior experience of similar work. We note that as per MSME Public Procurement Policy Order 2012, Micro and Small Enterprises are exempt from prior experience criteria subject to meeting technical specifications. Nascent Info Technologies Pvt. Ltd. (UDYAM-GJ-01-0007420) requests confirmation of this exemption. Additionally, we note that restricting experience to a specific state/department is contrary to procurement norms. Experience from any government department/PSU across India should be acceptable.",
                "legal_cite": "MSME Policy Circular 1(2)(1)/2016-MA | GFR 2017 Rule 173",
                "severity": "HIGH",
            },
            {
                "id": "EXP-02",
                "rule": "CVC Guidelines + GFR 2017 Rule 173",
                "what_law_says": "Experience requirement must be in similar/related work, not identical. Cannot mandate experience only from a specific department/state/country.",
                "nascent_right": "Nascent has AMC GIS, VMC GIS, SMC GIS, KVIC Mobile GIS, Tourism Portal, eGov projects across multiple clients.",
                "challenge_if": "Experience required only from municipal corporations | Only from Gujarat | Only from specific department | Only from projects >100 Cr (unreasonable bar)",
                "pre_bid_query": "Clause {clause_no} restricts eligibility to bidders having experience specifically from {restriction}. As per GFR 2017 Rule 173 and CVC guidelines on fair procurement, experience in similar works from any Central/State Government, PSU, Autonomous Body, or reputed Private Organisation should be acceptable. Restricting to {restriction} unduly limits competition. We request this clause to be amended to accept experience from similar nature of work from any government or reputed private organisation.",
                "legal_cite": "GFR 2017 Rule 173(xi) | CVC Circular 07/02/10",
                "severity": "HIGH",
            },
            {
                "id": "EXP-03",
                "rule": "GeM Prohibited Clauses",
                "what_law_says": "Cannot ask for experience from a specific single organization only. Cannot mandate foreign/export experience.",
                "nascent_right": "All Nascent experience is domestic — valid",
                "challenge_if": "Only multinational company experience accepted | Only foreign project experience | Only experience with specific named client",
                "pre_bid_query": "Clause {clause_no} requires experience exclusively from {specific_org}. As per GeM guidelines and GFR 2017 Rule 173, mandating experience from a specific organisation is prohibited as it restricts competition to a single vendor. We request this clause be amended to accept experience from any similar organisation.",
                "legal_cite": "GeM Prohibited Clauses List | GFR 2017 Rule 173",
                "severity": "HIGH",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 5. LOCAL OFFICE / PRESENCE REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════
    "local_office": {
        "title": "Local Office / State Presence Requirements",
        "rules": [
            {
                "id": "LOC-01",
                "rule": "CVC Guidelines + Competition Commission norms",
                "what_law_says": "Mandatory requirement of local/state office before bid submission is restrictive and against open competition principles. Office can be set up after award.",
                "nascent_right": "Nascent HQ in Ahmedabad. No branches. If office required in another state — raise query. Feasible to open post-award.",
                "challenge_if": "Local office required before bid | Registered office in specific city/state | Proof of local presence at bid stage",
                "pre_bid_query": "Clause {clause_no} requires the bidder to have a registered/local office in {city/state} as a pre-condition for bidding. This requirement is restrictive and discourages competition. As per established procurement principles and CVC guidelines, physical presence should not be a pre-bid eligibility condition. A successful bidder can establish necessary local presence within a reasonable time post award. We request this clause be modified to: 'The successful bidder shall establish operational presence in {city/state} within 30 days of award of contract.'",
                "legal_cite": "CVC Circular on Fair Procurement | GFR 2017 Rule 173(xi)",
                "severity": "CONDITIONAL",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 6. CERTIFICATIONS REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════
    "certifications": {
        "title": "Certification Requirements",
        "rules": [
            {
                "id": "CERT-01",
                "rule": "GFR 2017 Rule 173 + GeM Prohibited Clauses",
                "what_law_says": "Cannot mandate foreign/international certifications when equivalent Indian standards exist. ISO certifications from NABCB-accredited bodies are fully valid.",
                "nascent_right": "Nascent has ISO 9001, 27001, 20000 + CMMI L3 — covers most certification requirements.",
                "challenge_if": "Foreign certifications mandated exclusively | ISO not accepted, only proprietary certification required",
                "pre_bid_query": "Clause {clause_no} mandates {foreign_cert}. As per GeM guidelines and GFR 2017 Rule 173, mandating foreign/international certifications when equivalent Indian standards exist is not permitted. Nascent Info Technologies holds ISO 9001:2015, ISO/IEC 27001:2022, ISO/IEC 20000-1:2018, and CMMI V2.0 Level 3 — all internationally recognised and NABCB-accredited certifications. We request confirmation that these certifications satisfy the clause requirement.",
                "legal_cite": "GeM Prohibited Clauses Item 9 | GFR 2017 Rule 173",
                "severity": "MEDIUM",
            },
            {
                "id": "CERT-02",
                "rule": "CERT-In Empanelment requirement",
                "what_law_says": "CERT-In empanelment is mandatory for cybersecurity audits specifically. For IT development work, CERT-In is not a blanket requirement.",
                "nascent_right": "Nascent does NOT have CERT-In empanelment. This is a genuine gap. If CERT-In is mandated, must raise pre-bid query to clarify scope.",
                "challenge_if": "CERT-In mandated for non-security work | CERT-In required for general software development",
                "pre_bid_query": "Clause {clause_no} requires the bidder to be empanelled with CERT-In. We note that CERT-In empanelment is specifically required for: (a) IS/IT security audit, (b) Vulnerability Assessment and Penetration Testing (VAPT), and (c) Information security consulting services. For {project_scope}, which primarily involves {work_description}, CERT-In empanelment does not appear to be a mandatory statutory requirement. We request the authority to: (a) Clarify whether CERT-In empanelment is a mandatory eligibility condition or a preferred qualification, and (b) If mandatory, specify the exact CERT-In services panel under which empanelment is required.",
                "legal_cite": "IT Act 2000 Section 70B | CERT-In Guidelines | MeitY Advisory",
                "severity": "HIGH",
                "nascent_gap": True,
            },
            {
                "id": "CERT-03",
                "rule": "STQC / Testing certification",
                "what_law_says": "STQC certification may be required for certain software products. Not universally applicable.",
                "nascent_right": "Nascent does NOT have STQC certification. Must raise pre-bid query.",
                "challenge_if": "STQC mandated for custom software development (not product)",
                "pre_bid_query": "Clause {clause_no} requires STQC certification. STQC testing and certification is typically applicable for standardised software products, not custom-developed software solutions. Since this engagement involves custom software development specifically for {org_name}'s requirements, STQC product certification is not applicable. We request the authority to clarify: (a) Is STQC certification mandatory for custom-developed software in this project? (b) If yes, which specific STQC panel/service is required?",
                "legal_cite": "MeitY STQC Guidelines | GFR 2017 Rule 173",
                "severity": "MEDIUM",
                "nascent_gap": True,
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 7. OEM / PARTNERSHIP REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════
    "oem_requirements": {
        "title": "OEM Authorization / Partnership Requirements",
        "rules": [
            {
                "id": "OEM-01",
                "rule": "GFR 2017 Rule 173 + CVC Guidelines",
                "what_law_says": "Cannot mandate OEM authorization unless the specific product is truly proprietary and no alternative exists. For custom IT development, OEM partnership should not be mandatory.",
                "nascent_right": "Nascent is not SAP/Oracle/Microsoft partner. If these are mandated, must raise pre-bid query.",
                "challenge_if": "SAP certified partner required | Oracle authorized partner required | Microsoft Gold Partner required",
                "pre_bid_query": "Clause {clause_no} mandates that the bidder must be a {oem_name} authorized partner/reseller. The scope of this project involves {work_description}. If the intent is to use {oem_name} software/platform, we request clarification on: (a) Is {oem_name} the only acceptable platform, or will open-source/alternative platforms meeting functional requirements be considered? (b) If {oem_name} is mandated, is it specified anywhere in the functional/technical requirements? Mandating a specific OEM partnership without functional justification restricts competition and may violate GFR 2017 Rule 173(xi) which prohibits restrictive specifications. We request the clause be modified to specify functional/technical requirements rather than specific product mandates.",
                "legal_cite": "GFR 2017 Rule 173(xi) | CVC Circular 07/02/10 | GeM Prohibited Clauses",
                "severity": "HIGH",
                "nascent_gap": True,
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 8. TENDER FEE / BID PARTICIPATION FEE
    # ═══════════════════════════════════════════════════════════════════
    "tender_fee": {
        "title": "Tender Fee / Document Fee",
        "rules": [
            {
                "id": "FEE-01",
                "rule": "MSME Policy 2012 | GeM Guidelines",
                "what_law_says": "Tender documents must be provided FREE OF COST to MSMEs. No tender fee can be charged from MSEs.",
                "nascent_right": "Nascent as MSME — should get tender documents free. If fee charged, claim exemption.",
                "challenge_if": "Tender fee charged to MSME | Document fee demanded from MSME | Non-refundable registration fee",
                "pre_bid_query": "The tender prescribes a document fee of Rs. {tender_fee}. As per MSME Public Procurement Policy Order 2012, Micro and Small Enterprises registered under Udyam Registration are entitled to tender documents free of cost. Nascent Info Technologies Pvt. Ltd. (UDYAM-GJ-01-0007420) is a registered Small Enterprise and requests waiver of tender document fee as per statutory entitlement.",
                "legal_cite": "MSME Policy 2012 Section 11 | GeM Prohibited Clause 14",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 9. BID VALIDITY
    # ═══════════════════════════════════════════════════════════════════
    "bid_validity": {
        "title": "Bid Validity Period",
        "rules": [
            {
                "id": "VAL-01",
                "rule": "GFR 2017 Rule 168",
                "what_law_says": "Bid validity period should be reasonable — generally 90 to 180 days. Anything beyond 180 days should be justified by complexity.",
                "nascent_right": "If validity >180 days, raise query — long validity locks up resources without guarantee of award.",
                "challenge_if": "Bid validity > 180 days | No extension compensation offered | Validity of 360 days or more",
                "pre_bid_query": "Clause {clause_no} requires bid validity of {days} days. For IT software development projects, a validity period beyond 180 days is unreasonable as it blocks significant bidder resources without certainty of contract award. We request: (a) Bid validity be reduced to 120 days (standard for IT projects), or (b) In case validity >180 days is maintained, provision for price escalation at prevailing rates be included for the extended period.",
                "legal_cite": "GFR 2017 Rule 168 | Standard Bid Document Clause",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 10. MAKE IN INDIA / LOCAL CONTENT
    # ═══════════════════════════════════════════════════════════════════
    "make_in_india": {
        "title": "Make in India / Local Content Requirements",
        "rules": [
            {
                "id": "MII-01",
                "rule": "Public Procurement (PP) Order 2017 | GFR Rule 144(xi)",
                "what_law_says": "Government mandates preference to 'Class I Local Supplier' (50%+ local content) and 'Class II Local Supplier' (20-50% local content). Foreign suppliers from border-sharing countries cannot participate without registration.",
                "nascent_right": "Nascent qualifies as Class I Local Supplier — all software developed in India by Indian engineers.",
                "challenge_if": "MII compliance not mentioned | Local content percentage not specified | Foreign vendors from border countries not restricted",
                "pre_bid_query": "The tender document does not specify Make in India compliance requirements. As per Public Procurement (Preference to Make in India) Order 2017 and GFR Rule 144(xi), procurement entities must specify the local content requirement and give preference to Class I Local Suppliers. Nascent Info Technologies Pvt. Ltd. qualifies as a Class I Local Supplier with 100% local content in its software solutions developed entirely in India. We request: (a) Explicit mention of MII compliance requirement in the bid document, and (b) Confirmation that local suppliers will receive applicable purchase preference.",
                "legal_cite": "PP Order 2017 | GFR 2017 Rule 144(xi) | DoE OM F.6/18/2019-PPD",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 11. PAYMENT TERMS
    # ═══════════════════════════════════════════════════════════════════
    "payment_terms": {
        "title": "Payment Terms and Conditions",
        "rules": [
            {
                "id": "PAY-01",
                "rule": "GFR 2017 Rule 172 | MSMED Act Section 15-23",
                "what_law_says": "Payment must be made within 45 days of acceptance/delivery. MSMED Act Section 15 mandates payment to MSMEs within 45 days; delay attracts compound interest at 3× RBI bank rate.",
                "nascent_right": "As MSME, legally entitled to payment within 45 days. Late payment attracts mandatory interest.",
                "challenge_if": "Payment terms >45 days | No milestone-based payment | 100% payment only after final handover | Retention money held beyond warranty",
                "pre_bid_query": "Clause {clause_no} specifies payment within {payment_days} days. As per MSMED Act 2006 Sections 15-23 and GFR 2017 Rule 172, payment to Micro and Small Enterprises must be made within 45 days of acceptance of goods/services. Delay beyond 45 days attracts compound interest at three times the RBI bank rate. We request: (a) Payment terms be revised to 45 days from invoice/acceptance, and (b) Milestone-based payment schedule be included for a project of this size to ensure adequate cash flow.",
                "legal_cite": "MSMED Act 2006 Sections 15-23 | GFR 2017 Rule 172 | TReDS Framework",
                "severity": "MEDIUM",
            },
            {
                "id": "PAY-02",
                "rule": "Standard IT project payment norms",
                "what_law_says": "IT projects should have milestone-based payments aligned to deliverables. Common split: 10% on signing, 30% on design, 30% on development, 20% on go-live, 10% on stability/warranty.",
                "nascent_right": "Milestone payments protect Nascent's cash flow for multi-year projects.",
                "challenge_if": "Single payment only at end | >30% retention held | No advance payment for mobilisation",
                "pre_bid_query": "The payment terms do not specify a milestone-based payment schedule. For IT software development projects of this nature (estimated value Rs. {project_value_cr} Cr, duration {months} months), a milestone-based payment schedule is essential for: (a) Ensuring continuous project momentum, (b) Managing vendor cash flow, and (c) Aligning payment with delivered value. We propose the following payment schedule: 10% on project kickoff/contract signing, 20% on SRS approval, 30% on development completion/UAT, 30% on go-live, 10% on 3-month post go-live stability. We request the authority to incorporate a milestone-based payment structure.",
                "legal_cite": "Standard IT Procurement Norms | GFR 2017 Rule 172",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 12. CONTRACT PERIOD / PROJECT TIMELINE
    # ═══════════════════════════════════════════════════════════════════
    "contract_period": {
        "title": "Contract Period and Timeline",
        "rules": [
            {
                "id": "TIME-01",
                "rule": "Standard procurement norms",
                "what_law_says": "Contract period must be realistic and achievable. Unreasonably short deadlines create execution risk and invite low-quality work.",
                "nascent_right": "Nascent can complete most GIS/portal projects in 9-18 months depending on scope.",
                "challenge_if": "Timeline < 6 months for complex portal | No timeline given | Timeline not commensurate with scope",
                "pre_bid_query": "The project timeline of {months} months for the defined scope appears insufficient considering: (a) Requirements gathering and SRS approval typically takes 4-6 weeks, (b) Design and development of {modules} modules requires {dev_weeks} weeks, (c) UAT, security testing, and go-live preparation needs 4-6 weeks, (d) Data migration and training requires {migration_weeks} weeks. The total realistic timeline for quality delivery is {suggested_months} months. We request extension of contract period to {suggested_months} months or clarification on which modules are in scope for the initial phase.",
                "legal_cite": "Standard IT project management norms",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 13. LIQUIDATED DAMAGES (LD)
    # ═══════════════════════════════════════════════════════════════════
    "liquidated_damages": {
        "title": "Liquidated Damages / Penalty Clauses",
        "rules": [
            {
                "id": "LD-01",
                "rule": "GFR 2017 Rule 173 + Standard contract norms",
                "what_law_says": "Liquidated Damages should be reasonable — typically 0.5% per week, capped at 10% of contract value. Anything above 10% is considered a penalty (not LD) and is legally challengeable.",
                "nascent_right": "Nascent can challenge excessive LD clauses in pre-bid.",
                "challenge_if": "LD >0.5% per week | LD cap >10% | No cap on LD | LD applicable from day 1 without cure period",
                "pre_bid_query": "Clause {clause_no} prescribes Liquidated Damages of {ld_pct}% per {period}, with no cap mentioned. As per GFR 2017 Rule 173 and established contract law, Liquidated Damages for IT projects should be: (a) 0.5% per week of delay, (b) Capped at maximum 10% of contract value, (c) Applied only after a reasonable cure period of 15-30 days. An uncapped LD clause or LD exceeding 10% constitutes a penalty clause, which is not enforceable under Indian Contract Act 1872 Section 74. We request revision to standard LD terms.",
                "legal_cite": "GFR 2017 Rule 173 | Indian Contract Act 1872 Section 74 | FIDIC Standard Terms",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 14. IP / OWNERSHIP RIGHTS
    # ═══════════════════════════════════════════════════════════════════
    "ip_rights": {
        "title": "Intellectual Property / Source Code Ownership",
        "rules": [
            {
                "id": "IP-01",
                "rule": "IT Act 2000 | Contract law + Standard IT procurement",
                "what_law_says": "Government clients typically claim ownership of all deliverables. Source code escrow is increasingly required. However, reuse of Nascent's existing IP frameworks should be compensated.",
                "nascent_right": "Nascent should clarify: (a) Who owns custom-developed code, (b) Whether Nascent's existing frameworks/libraries are licensed or transferred, (c) Source code escrow arrangements.",
                "challenge_if": "Blanket IP transfer with no compensation | Existing library/framework code demanded | No source code escrow option",
                "pre_bid_query": "Clause {clause_no} transfers all intellectual property to {org_name} without qualification. We seek clarification on: (a) Does IP transfer include only project-specific custom code or also Nascent's existing proprietary frameworks, libraries, and tools used in development? (b) If existing frameworks are to be transferred, what is the mechanism for valuation and compensation? (c) Would the authority accept a Source Code Escrow arrangement as an alternative to outright transfer, ensuring government access to source code while Nascent retains IP? Clarification on these points is essential for accurate cost estimation.",
                "legal_cite": "IT Act 2000 | Copyright Act 1957 | Standard IT contract practices",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 15. CONSORTIUM / JOINT VENTURE
    # ═══════════════════════════════════════════════════════════════════
    "consortium": {
        "title": "Consortium / Joint Venture / Sub-contracting",
        "rules": [
            {
                "id": "CON-01",
                "rule": "GFR 2017 Rule 173 | Standard procurement norms",
                "what_law_says": "Consortium/JV is permissible if allowed in bid document. Lead partner must meet majority (usually 60%) of eligibility criteria. Sub-contracting must be disclosed upfront.",
                "nascent_right": "If a tender is too large for Nascent alone (turnover/experience gap), can propose consortium with a larger partner.",
                "challenge_if": "Consortium not permitted | Lead partner must meet 100% criteria | Sub-contracting prohibited",
                "pre_bid_query": "The tender document does not permit consortium bidding. Given the large scope of this project (estimated value Rs. {value_cr} Cr), restricting participation to single firms may: (a) Reduce competition, (b) Exclude capable MSMEs who could partner with larger firms. We request the authority to: (a) Allow consortium/JV bidding with lead partner meeting 60% of financial criteria and individual partners meeting their respective technical criteria, OR (b) Clarify whether sub-contracting of specific modules is permitted with prior disclosure.",
                "legal_cite": "GFR 2017 Rule 173 | Standard Procurement Manual Chapter 4",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 16. SCOPE CLARIFICATIONS
    # ═══════════════════════════════════════════════════════════════════
    "scope": {
        "title": "Scope of Work Clarifications",
        "rules": [
            {
                "id": "SCO-01",
                "rule": "Standard pre-bid practice",
                "what_law_says": "Bidders have right to seek clarifications on scope ambiguities. All clarifications published as corrigendum apply equally to all bidders.",
                "nascent_right": "Always ask about data migration, server infrastructure, training, AMC scope",
                "challenge_if": "Scope is vague | Existing data not specified | Hardware supply mixed with software | Training not quantified",
                "pre_bid_query": "With respect to the scope of work, we seek the following clarifications: (a) Data Migration: Is migration of existing data ({existing_data}) in scope? If yes, what is the estimated volume and format? (b) Infrastructure: Is server/hosting infrastructure to be provided by {org_name} or included in Nascent's scope? (c) Training: How many users are to be trained and at which locations? (d) AMC: Is Annual Maintenance Contract separate from development, and if so, what is its duration and scope? (e) Integration: Are API integrations with {integration_systems} in scope? If yes, will {org_name} provide API documentation and sandbox environment?",
                "legal_cite": "Standard pre-bid practice | GFR 2017 Rule 168 (Bidding Documents)",
                "severity": "MEDIUM",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 17. EMPLOYEE / MANPOWER REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════
    "manpower": {
        "title": "Employee Strength / Manpower Requirements",
        "rules": [
            {
                "id": "MAN-01",
                "rule": "CVC Guidelines + GFR 2017 Rule 173",
                "what_law_says": "Minimum employee requirement must be proportionate to project needs. Requiring 200+ employees for a 5 Cr IT project is disproportionate.",
                "nascent_right": "Nascent has 67 employees. If >100 employees required, raise pre-bid query.",
                "challenge_if": "Employee count >100 for small project | Permanent employees only (not contractual) | Employees must be in same state",
                "pre_bid_query": "Clause {clause_no} requires minimum {emp_count} full-time employees as an eligibility criterion. For a project of this scope and value, we request clarification: (a) What is the technical justification for requiring {emp_count} employees for this engagement? (b) The effective deployment for this project would require {project_team} professionals. Basing eligibility on total company headcount rather than project-specific capability is not standard procurement practice. We request the criterion be modified to: 'The bidder shall deploy a minimum project team of {project_team} qualified professionals for this engagement.'",
                "legal_cite": "GFR 2017 Rule 173(xi) | CVC Circular on Fair Procurement",
                "severity": "CONDITIONAL",
                "nascent_gap": True,  # If >100 required
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 18. DATA SECURITY / PRIVACY REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════
    "data_security": {
        "title": "Data Security and Privacy Requirements",
        "rules": [
            {
                "id": "SEC-01",
                "rule": "IT Act 2000 | DPDP Act 2023 | ISO 27001",
                "what_law_says": "Data protection requirements must comply with DPDP Act 2023. ISO 27001 certification is the standard for information security management.",
                "nascent_right": "Nascent has ISO/IEC 27001:2022 — covers information security requirements.",
                "challenge_if": "CERT-In mandated for non-security work | NDA terms unreasonably broad | Data localisation not specified",
                "pre_bid_query": "The tender specifies data security requirements. Nascent Info Technologies holds ISO/IEC 27001:2022 certification (valid till Sep-2028) which covers comprehensive information security management. We request confirmation that ISO 27001 certification satisfies the data security requirements of Clause {clause_no}. Additionally, we note that the Digital Personal Data Protection Act 2023 (DPDP Act) imposes obligations on both the Data Processor (Nascent) and Data Fiduciary ({org_name}). We request the bid document be updated to include standard DPDP Act compliance obligations for both parties.",
                "legal_cite": "IT Act 2000 Section 43A | DPDP Act 2023 | ISO/IEC 27001:2022",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 19. BLACKLISTING / DEBARMENT
    # ═══════════════════════════════════════════════════════════════════
    "blacklisting": {
        "title": "Non-Blacklisting / Non-Debarment Declaration",
        "rules": [
            {
                "id": "BLK-01",
                "rule": "GFR 2017 Rule 151 | Standard bid conditions",
                "what_law_says": "Bidder must not be blacklisted/debarred by any government. Declaration required on Rs.100 stamp paper.",
                "nascent_right": "Nascent is clean — not blacklisted anywhere. Must provide declaration.",
                "challenge_if": "Blacklisting certificate required from all state governments separately (unreasonable) | Notarised declaration required (expensive)",
                "pre_bid_query": "Clause {clause_no} requires non-blacklisting certificates from all state governments. This is operationally impractical as there is no centralised registry of blacklisted vendors across all states. As per GFR 2017 Rule 151 and standard procurement practice, a self-declaration on company letterhead (on Rs.100 stamp paper) that the bidder is not blacklisted/debarred by any government authority is sufficient and standard practice. We request amendment of this clause to accept a self-declaration in the prescribed format.",
                "legal_cite": "GFR 2017 Rule 151 | Standard Bid Conditions",
                "severity": "LOW",
            },
        ]
    },

    # ═══════════════════════════════════════════════════════════════════
    # 20. BID DOCUMENT / CORRIGENDUM TIMELINES
    # ═══════════════════════════════════════════════════════════════════
    "timelines": {
        "title": "Bid Timeline and Corrigendum Rules",
        "rules": [
            {
                "id": "TIM-01",
                "rule": "GFR 2017 Rule 161",
                "what_law_says": "Minimum 3 weeks (21 days) must be given from publication of tender to bid submission for domestic tenders. For complex projects, 4-8 weeks is standard.",
                "nascent_right": "If insufficient time is given, raise pre-bid query for extension.",
                "challenge_if": "Less than 21 days from NIT to bid submission | Pre-bid query deadline too close to NIT publication",
                "pre_bid_query": "The bid submission deadline of {submission_date} allows only {days_given} days from the date of NIT publication ({nit_date}). As per GFR 2017 Rule 161, a minimum of 21 days must be provided for domestic tenders. For a complex IT project of this nature, adequate time is needed for: (a) Downloading and reviewing the complete bid document, (b) Site visit/pre-bid meeting participation, (c) Preparation of technical proposal and supporting documents. We request extension of bid submission deadline by at least {additional_days} days to ensure fair competition.",
                "legal_cite": "GFR 2017 Rule 161(iii) | Standard Bid Document Clause",
                "severity": "MEDIUM",
            },
        ]
    },

}


# ─────────────────────────────────────────────────────────────────────────────
# AI PROMPT FOR COMPLETE PRE-BID ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

def build_prebid_analysis_prompt(tender_text: str, ai_analysis: dict) -> str:
    """
    Build a comprehensive AI prompt that:
    1. Reads the full RFP text
    2. Cross-checks every clause against Indian tender guidelines
    3. Identifies where Nascent meets / falls short / can challenge
    4. Generates complete pre-bid query letter with exact citations
    """

    nascent_summary = f"""
NASCENT INFO TECHNOLOGIES PVT. LTD. — COMPLETE LEGAL PROFILE:
- MSME Registration: UDYAM-GJ-01-0007420 (Small Enterprise) — FULLY EXEMPT FROM EMD & TENDER FEE
- Turnover: FY22-23: Rs.16.36Cr | FY23-24: Rs.16.36Cr | FY24-25: Rs.18.83Cr | Avg 3yr: Rs.17.18Cr
- Net Worth: Rs.26.09 Cr (positive)
- Employees: 67 (HQ Ahmedabad, no branch offices)
- Certifications HELD: CMMI V2.0 L3 (till Dec-2026) | ISO 9001:2015 | ISO 27001:2022 | ISO 20000-1:2018
- Certifications NOT HELD: CERT-In | STQC | NSIC | SAP Partner | Oracle Partner | Microsoft Gold
- Make in India: YES — all software developed in India
- Blacklisted: NO
- States with office: Gujarat only
- Years in operation: 19 years
"""

    guidelines_summary = """
INDIAN TENDER LEGAL GUIDELINES — APPLY THESE TO EVERY CLAUSE:

1. EMD RULES (GFR Rule 170 + MSME Policy 2012):
   - EMD maximum 2-5% of tender value. Cannot exceed 5%.
   - MSMEs (Udyam registered) are FULLY EXEMPT from EMD — mandatory, not discretionary.
   - EMD forms: DD / Banker's Cheque / Bank Guarantee only.

2. PERFORMANCE SECURITY (GFR Rule 171):
   - PBG = 3-5% of contract value. Max 10%.
   - MSMEs are NOT exempt from PBG (only EMD exempt).
   - PBG must be released within 60 days of warranty completion.

3. TURNOVER CRITERIA (MSME Circular 10.03.2016 + GFR Rule 173):
   - MSMEs exempt from minimum turnover criteria.
   - Even for non-MSMEs, turnover should not exceed 3× estimated project value.
   - Experience accepted from any Central/State Govt/PSU/Autonomous Body — not restricted by state/department.

4. PRIOR EXPERIENCE (MSME Policy 2012 + GFR Rule 173):
   - MSMEs exempt from prior experience criteria.
   - Cannot mandate experience from specific state/org/department alone.
   - Cannot mandate foreign experience.
   - Cannot mandate experience with specific named client.

5. LOCAL OFFICE (CVC Guidelines):
   - Local office as PRE-BID eligibility condition is challengeable.
   - Can be required POST-AWARD only (within 30-60 days of award).

6. CERTIFICATIONS (GeM Prohibited Clauses + GFR Rule 173):
   - Cannot mandate foreign certifications when Indian equivalents exist.
   - ISO 9001/27001/20000 = internationally recognised, NABCB-accredited.
   - CERT-In: only for security audit/VAPT/IS consulting — not general IT dev.
   - STQC: only for standardised software products — not custom development.

7. OEM REQUIREMENTS (CVC + GFR Rule 173):
   - Cannot mandate specific OEM partnership without functional justification.
   - Must specify functional requirements, not product brand names.
   - SAP/Oracle/Microsoft partner mandates are challengeable for custom dev.

8. TENDER FEE (MSME Policy 2012 + GeM):
   - MSMEs get tender documents FREE OF COST.

9. PAYMENT TERMS (GFR Rule 172 + MSMED Act 2006):
   - Payment within 45 days mandatory for MSMEs.
   - Late payment attracts compound interest at 3× RBI bank rate.
   - Milestone-based payment is standard for IT projects.

10. LIQUIDATED DAMAGES (GFR Rule 173 + Contract Act):
    - LD standard: 0.5% per week of delay.
    - LD must be capped at 10% of contract value.
    - Above 10% = penalty clause, not enforceable under Contract Act Section 74.

11. BID TIMELINE (GFR Rule 161):
    - Minimum 21 days from NIT to bid submission.
    - Complex IT projects: 4-8 weeks standard.

12. MAKE IN INDIA (PP Order 2017 + GFR Rule 144):
    - Class I Local Supplier (50%+ local content) gets purchase preference.
    - Nascent = Class I (100% locally developed software).
"""

    prompt = f"""
You are a senior government tender compliance expert working for Nascent Info Technologies Pvt. Ltd. 
Your task is to:
1. Read the complete RFP/tender document text below
2. Identify EVERY clause that is either illegal, challengeable, or where Nascent has a gap
3. Cross-check each clause against Indian tender guidelines
4. Generate a professional, complete Pre-Bid Query Letter

{nascent_summary}

{guidelines_summary}

TENDER DOCUMENT TEXT:
{tender_text[:30000]}

EXISTING AI ANALYSIS FINDINGS:
{json.dumps(ai_analysis, indent=2)[:3000]}

TASK — Generate a complete Pre-Bid Query Letter with the following structure:

OUTPUT FORMAT (strict JSON):
{{
  "clause_violations": [
    {{
      "clause_no": "Clause number or section",
      "issue_type": "EMD|TURNOVER|EXPERIENCE|LOCAL_OFFICE|CERTIFICATION|OEM|PAYMENT|LD|TIMELINE|SCOPE|IP|OTHER",
      "severity": "HIGH|MEDIUM|LOW",
      "what_rfp_says": "Exact clause content",
      "what_law_says": "Applicable law/rule with citation",
      "nascent_status": "MET|NOT_MET|CONDITIONAL|EXEMPT_AS_MSME",
      "action": "RAISE_PREBID_QUERY|CLAIM_MSME_EXEMPTION|WATCH|ACCEPT",
      "query_text": "Complete professional query text to include in pre-bid letter"
    }}
  ],
  "nascent_gaps": [
    {{
      "gap": "What Nascent lacks",
      "clause_ref": "Relevant clause",
      "severity": "DISQUALIFYING|CONDITIONAL|MANAGEABLE",
      "mitigation": "How to address this gap — query/workaround/consortium"
    }}
  ],
  "prebid_letter": {{
    "salutation": "To, The [Authority], [Org Name], [Address]",
    "subject": "Pre-bid Queries for Tender No. [XXX] — [Project Name]",
    "opening": "Opening paragraph",
    "queries": [
      {{
        "query_no": "Q1",
        "clause_ref": "Clause reference",
        "query": "Complete formal query text with legal citations",
        "legal_basis": "GFR Rule / MSME Policy / etc."
      }}
    ],
    "msme_exemption_para": "Standard para claiming MSME exemptions for EMD, tender fee, turnover, experience",
    "closing": "Closing paragraph",
    "signature_block": "Parthav Thakkar, Bid Executive | Nascent Info Technologies Pvt. Ltd."
  }},
  "overall_recommendation": {{
    "verdict": "BID|NO-BID|CONDITIONAL",
    "critical_gaps": ["List of disqualifying gaps if any"],
    "queries_must_be_raised": ["List of most important queries"],
    "confidence_if_queries_resolved": "HIGH|MEDIUM|LOW"
  }}
}}

RULES FOR QUERY GENERATION:
- Every query must cite the specific GFR Rule / MSME Policy circular / GeM guidelines
- Queries must be professional, factual, and non-confrontational
- Each query must have a specific ask (confirm / amend / clarify)
- Group related queries together
- The MSME exemption claim must appear prominently
- Include Nascent's registration numbers, certifications, and turnover facts where relevant
- Do NOT generate generic queries — every query must reference actual RFP clause content
- Minimum 5 queries, maximum 15 queries

Generate the complete output now.
"""
    return prompt


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def analyze_rfp_against_guidelines(
    tender_text: str,
    ai_analysis: dict,
    tender_data: dict,
) -> dict:
    """
    Cross-check tender RFP against all Indian guidelines.
    Returns: clause violations, Nascent gaps, pre-bid letter.
    """
    from ai_analyzer import load_config, call_gemini_with_keys

    prompt = build_prebid_analysis_prompt(tender_text, ai_analysis)
    cfg    = load_config()
    key    = cfg.get("gemini_api_key", "")
    keys   = cfg.get("gemini_api_keys", [])

    try:
        raw = call_gemini_with_keys(prompt, key, keys)
        # Extract JSON from response
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if json_match:
            result = json.loads(json_match.group())
            return result
        return {"error": "Could not parse AI response", "raw": raw[:500]}
    except Exception as e:
        return {"error": str(e)}


def generate_prebid_letter_docx(
    analysis_result: dict,
    tender_data: dict,
    out_path: str,
) -> dict:
    """
    Generate a professional pre-bid query letter as DOCX
    using the analysis result from analyze_rfp_against_guidelines().
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from datetime import date

        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2)
        section.bottom_margin = Cm(2)

        NAVY  = RGBColor(31, 56, 100)
        BLUE  = RGBColor(46, 117, 182)
        RED   = RGBColor(192, 0, 0)
        GRAY  = RGBColor(80, 80, 80)
        WHITE = RGBColor(255, 255, 255)

        def para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
            p = doc.add_paragraph()
            r = p.add_run(text or "")
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(size)
            if color: r.font.color.rgb = color
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            return p

        def heading(text, level=1):
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                run.font.color.rgb = NAVY
                run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)

        # ── LETTERHEAD ───────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NASCENT INFO TECHNOLOGIES PVT. LTD.")
        r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY; r.font.name = "Arial"

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad – 380015, Gujarat")
        r2.font.size = Pt(9); r2.font.color.rgb = GRAY; r2.font.name = "Arial"

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("Ph: +91-79-40200400  |  nascent.tender@nascentinfo.com  |  www.nascentinfo.com")
        r3.font.size = Pt(9); r3.font.color.rgb = GRAY; r3.font.name = "Arial"

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run("UDYAM-GJ-01-0007420 (MSME – Small Enterprise)  |  CMMI V2.0 L3  |  ISO 9001 / 27001 / 20000")
        r4.font.size = Pt(9); r4.font.color.rgb = BLUE; r4.font.name = "Arial"

        # Horizontal line
        phr = doc.add_paragraph()
        phr.paragraph_format.space_after = Pt(4)
        pPr = phr._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F3864')
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph()

        # ── DATE & REF ───────────────────────────────────────────────
        para(f"Date: {date.today().strftime('%d %B %Y')}")
        doc.add_paragraph()

        letter = analysis_result.get("prebid_letter", {})
        queries = analysis_result.get("clause_violations", [])

        # ── TO ADDRESS ───────────────────────────────────────────────
        para("To,")
        org   = tender_data.get("org_name", "[Organisation Name]")
        ref   = tender_data.get("ref_no",   "[Tender Reference Number]")
        brief = tender_data.get("brief",    "[Project Name]")[:80]
        para(letter.get("salutation", f"The Authorised Officer,\n{org}"))

        doc.add_paragraph()

        # ── SUBJECT ──────────────────────────────────────────────────
        subj = letter.get("subject", f"Pre-bid Queries — Tender No. {ref} — {brief}")
        para(f"Sub: {subj}", bold=True)
        para(f"Ref: Tender No.: {ref}  |  Pre-bid Query Submission", bold=True)

        doc.add_paragraph()
        para("Dear Sir / Madam,")
        doc.add_paragraph()

        # ── OPENING ──────────────────────────────────────────────────
        opening = letter.get("opening",
            f"We, Nascent Info Technologies Pvt. Ltd., a CMMI V2.0 Level 3 and ISO 9001/27001/20000 certified IT company registered as a Small Enterprise vide UDYAM-GJ-01-0007420, are interested in participating in the above-referenced tender. After careful review of the bid document, we wish to raise the following queries and seek clarifications before bid submission.")
        para(opening)
        doc.add_paragraph()

        # ── MSME EXEMPTION BLOCK ─────────────────────────────────────
        heading("MSME Statutory Entitlements", level=2)
        msme_para = letter.get("msme_exemption_para",
            "As a Micro and Small Enterprise registered under Udyam Registration (UDYAM-GJ-01-0007420), Nascent Info Technologies is entitled to the following benefits under MSME Public Procurement Policy Order 2012 and GFR 2017:\n(a) Full exemption from Earnest Money Deposit (EMD)\n(b) Tender documents free of cost\n(c) Relaxation from minimum turnover criteria\n(d) Relaxation from prior experience criteria\nWe hereby formally claim these statutory entitlements and request the authority to confirm the same and update the bid evaluation process accordingly.")
        para(msme_para)
        doc.add_paragraph()

        # ── QUERIES ──────────────────────────────────────────────────
        heading("Pre-bid Queries", level=2)

        # Use AI-generated queries from prebid_letter
        ai_queries = letter.get("queries", [])

        # Also add queries from clause_violations
        violation_queries = [
            {
                "query_no": f"Q{i+len(ai_queries)+1}",
                "clause_ref": v.get("clause_no", "General"),
                "query": v.get("query_text", ""),
                "legal_basis": v.get("what_law_says", ""),
            }
            for i, v in enumerate(queries)
            if v.get("action") in ["RAISE_PREBID_QUERY"] and v.get("query_text")
        ]

        all_queries = ai_queries + violation_queries

        if not all_queries:
            # Fallback: generate standard queries from guidelines
            all_queries = _generate_standard_queries(tender_data, analysis_result)

        for q in all_queries[:15]:
            q_no = q.get("query_no", "Q?")
            clause = q.get("clause_ref", "")
            query_text = q.get("query", "") or q.get("query_text", "")
            legal = q.get("legal_basis", "") or q.get("what_law_says", "")

            if not query_text:
                continue

            # Query heading
            p = doc.add_paragraph()
            r = p.add_run(f"{q_no}. Clause: {clause}")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)

            # Query text
            para(query_text)

            if legal:
                p_legal = doc.add_paragraph()
                r_legal = p_legal.add_run(f"Legal Basis: {legal}")
                r_legal.font.name = "Arial"
                r_legal.font.size = Pt(9)
                r_legal.font.color.rgb = BLUE
                r_legal.italic = True
                p_legal.paragraph_format.space_after = Pt(8)

        doc.add_paragraph()

        # ── NASCENT GAPS SECTION ─────────────────────────────────────
        gaps = analysis_result.get("nascent_gaps", [])
        if gaps:
            heading("Areas Requiring Your Clarification / Nascent's Position", level=2)
            for gap in gaps[:5]:
                para(f"• {gap.get('gap','')}", size=10)
                if gap.get("mitigation"):
                    para(f"  Our position: {gap.get('mitigation','')}", size=10, color=GRAY)

        doc.add_paragraph()

        # ── CLOSING ──────────────────────────────────────────────────
        closing = letter.get("closing",
            "We request the authority to provide written responses to the above queries via the official pre-bid meeting or corrigendum, and to share the revised bid document if any amendments are made. We assure our commitment to participating in this procurement with full compliance to all applicable laws and with the highest standards of quality and professionalism.")
        para(closing)
        doc.add_paragraph()

        para("Thanking you,")
        para("Yours faithfully,")

        doc.add_paragraph().paragraph_format.space_after = Pt(30)
        para("Parthav Thakkar", bold=True)
        para("Bid Executive")
        para("Nascent Info Technologies Pvt. Ltd.")
        para(f"Date: {date.today().strftime('%d %B %Y')}")

        doc.add_paragraph()
        para("Enclosures:", bold=True)
        para("1. Copy of Udyam Registration Certificate (UDYAM-GJ-01-0007420)")
        para("2. Copy of Certificate of Incorporation")
        para("3. CMMI V2.0 Level 3 Certificate")
        para("4. ISO 9001/27001/20000 Certificates")

        doc.save(out_path)
        return {
            "status": "success",
            "path": out_path,
            "query_count": len(all_queries),
            "gaps_found": len(gaps),
        }

    except Exception as e:
        return {"status": "error", "message": str(e)}


def _generate_standard_queries(tender_data: dict, analysis: dict) -> list:
    """Generate standard queries when AI doesn't produce them."""
    queries = []
    q = 1

    # Always: MSME EMD claim
    queries.append({
        "query_no": f"Q{q}",
        "clause_ref": "EMD / Bid Security",
        "query": f"As a registered Small Enterprise vide UDYAM-GJ-01-0007420, Nascent Info Technologies Pvt. Ltd. is entitled to full exemption from EMD as per MSME Public Procurement Policy Order 2012. We request confirmation that Nascent is exempt from depositing EMD and that this exemption is reflected in the bid evaluation process.",
        "legal_basis": "MSME Policy 2012 | GFR 2017 Rule 170",
    }); q += 1

    # Turnover relaxation
    queries.append({
        "query_no": f"Q{q}",
        "clause_ref": "Financial Eligibility / Turnover",
        "query": "As a Micro/Small Enterprise, Nascent is entitled to relaxation from minimum turnover criteria as per MSME Circular 1(2)(1)/2016-MA dated 10.03.2016. We request: (a) Confirmation of MSME turnover relaxation, OR (b) If turnover criteria is maintained, confirmation that Rs.17.18 Cr average annual turnover (FY22-25) qualifies for this project.",
        "legal_basis": "MSME Policy Circular 10.03.2016 | GFR Rule 173",
    }); q += 1

    # Experience relaxation
    queries.append({
        "query_no": f"Q{q}",
        "clause_ref": "Technical Eligibility / Experience",
        "query": "As an MSME, Nascent is eligible for relaxation from prior experience criteria as per MSME Policy 2012. However, Nascent has 19 years of experience and has completed 9+ GIS/eGovernance/Mobile App projects for AMC, VMC, JuMC, BMC, KVIC, Tourism Corp Gujarat, and PCSCL. We request confirmation that experience from any government/PSU/autonomous body across India is acceptable, not restricted to a specific state or department.",
        "legal_basis": "MSME Policy 2012 | GFR Rule 173(xi)",
    }); q += 1

    # Scope clarification
    queries.append({
        "query_no": f"Q{q}",
        "clause_ref": "Scope of Work",
        "query": "We seek the following scope clarifications to accurately price the bid: (a) Is existing data migration in scope? If yes, what is the estimated volume/format? (b) Is server/hosting infrastructure to be provided by the authority or quoted by the bidder? (c) How many users are to be trained and at which locations? (d) Is Annual Maintenance Contract (AMC) post-warranty period part of this tender scope?",
        "legal_basis": "GFR 2017 Rule 168 (Bidding Documents)",
    }); q += 1

    # Payment terms
    queries.append({
        "query_no": f"Q{q}",
        "clause_ref": "Payment Terms",
        "query": "We request inclusion of milestone-based payment schedule to ensure smooth project execution. Suggested structure: 10% on project kickoff, 20% on SRS/design approval, 30% on development completion, 30% on UAT/go-live, 10% post 3-month stability period. We further request confirmation that payment will be processed within 45 days of invoice submission as mandated by MSMED Act 2006.",
        "legal_basis": "MSMED Act 2006 Section 15 | GFR 2017 Rule 172",
    }); q += 1

    return queries


def get_all_guidelines_summary() -> dict:
    """Return all guidelines in structured format for UI display."""
    return {
        "total_categories": len(GUIDELINES),
        "categories": {
            cat_id: {
                "title": cat_data["title"],
                "rule_count": len(cat_data["rules"]),
                "rules": [
                    {
                        "id": r["id"],
                        "law": r["rule"],
                        "summary": r["what_law_says"][:100],
                        "nascent_right": r["nascent_right"][:100],
                        "severity": r.get("severity", "MEDIUM"),
                        "nascent_gap": r.get("nascent_gap", False),
                    }
                    for r in cat_data["rules"]
                ]
            }
            for cat_id, cat_data in GUIDELINES.items()
        }
    }
