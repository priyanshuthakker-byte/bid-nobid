# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
NASCENT PROJECT DATA SCANNER v1.0
  Scans H:/Nascent/backup vt - external hard disk
  Extracts ALL project data -> Excel + CSV for Google Sheet
  Runs OFFLINE - no internet, no API key needed
  RESUME-SAFE: saves progress every 10 folders

HOW TO RUN:
  1. Make sure H: drive is connected
  2. Double-click this file  OR  run: python nascent_scanner.py
  3. Leave it running — can take hours for 100GB
  4. Output files appear on Desktop when done
  5. Import CSV to Google Sheets → File > Import

If interrupted: just run again — it resumes from where it stopped.
"""

import os, re, json, csv, sys, time, traceback
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, List, Tuple

# ══ CONFIG ════════════════════════════════════════════════════════════════════
ROOT_DIR       = r"H:\Nascent\backup vt"
DESKTOP        = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT_CSV     = os.path.join(DESKTOP, "nascent_projects.csv")
OUTPUT_XLSX    = os.path.join(DESKTOP, "nascent_projects.xlsx")
PROGRESS_FILE  = os.path.join(DESKTOP, "scanner_progress.json")
LOG_FILE       = os.path.join(DESKTOP, "scanner_log.txt")

MAX_PAGES      = 6       # pages per PDF
MAX_CHARS      = 15000   # chars per document
MAX_DOCS       = 6       # docs to read per project folder
MAX_DEPTH      = 6       # folder depth limit

# ══ COLUMNS (Google Sheet schema) ════════════════════════════════════════════
COLUMNS = [
    "Project_ID",
    "Folder_Name",
    "Client_Name",
    "Client_Type",            # Central Govt / State Govt / PSU / ULB / SPV / Private
    "State",
    "Project_Name",
    "Scope_Summary",          # 2-3 line description
    "Work_Type",              # Implementation / Survey / AMC / Consulting / Development
    "Primary_Technology",     # GIS / Mobile App / Web Portal / ERP / eGovernance / Analytics
    "Tech_Stack_Details",     # QGIS, Spring Boot, React, PostgreSQL etc.
    "Sector_Tags",            # Smart City / Land Records / Tourism / ULB etc.
    "Our_Role",               # Lead / Consortium Lead / Consortium Member / Sub
    "Contract_Value_Cr",      # Numeric
    "Contract_Value_Raw",     # Original text
    "Duration_Months",
    "Start_Date",
    "End_Date",
    "Project_Status",         # Completed / Ongoing / Unknown
    "Manpower_Count",
    "Key_Roles_Deployed",
    "Survey_Area",            # e.g. 250 Sq.Km
    "Tender_WO_Number",
    "Has_Work_Order",         # Yes/No
    "Has_Completion_Cert",    # Yes/No
    "Has_LOA",                # Yes/No
    "Has_BOQ",                # Yes/No
    "Has_NIT_RFP",            # Yes/No
    "Similar_Work_Category",  # For eligibility matching in new tenders
    "Relevant_For",           # BOQ / PreBid / TQ / Profile / TechProposal
    "Total_Files_Found",
    "Key_Documents",
    "Notes",
    "Folder_Path",
    "Scan_Date",
]

# ══ TECHNOLOGY & SECTOR KEYWORDS ═════════════════════════════════════════════
TECH_KW = {
    "GIS/Geospatial": ["gis", "geographic information", "geospatial", "arcgis", "qgis", "geoserver",
                        "postgis", "shapefile", "raster", "vector", "citylayers", "mapinfo",
                        "openstreet", "leaflet", "openlayers", "mapbox", "geojson", "kml", "wms", "wfs"],
    "Mobile App":     ["mobile app", "android", "ios", "flutter", "react native", "mobile application",
                        "smartphone", "tablet app", "play store", "app store"],
    "Web Portal":     ["web portal", "web application", "web-based", "online portal", "website",
                        "web platform", "web interface", "react", "angular", "django", "spring boot"],
    "ERP":            ["erp", "enterprise resource", "sap", "oracle erp", "odoo", "tally integration"],
    "eGovernance":    ["e-governance", "egov", "e-gov", "digital governance", "government portal",
                        "citizen service", "public service delivery", "nic", "diksha"],
    "Data Analytics": ["data analytics", "dashboard", "business intelligence", "power bi", "tableau",
                        "data platform", "analytics portal", "reporting system", "mis"],
    "Survey/Mapping": ["survey", "field survey", "gps survey", "drone survey", "lidar", "uav survey",
                        "topographic", "photogrammetry", "ground truth", "total station", "dgps", "rtk",
                        "field data collection", "land parcels"],
    "IoT/Smart City": ["iot", "sensor", "scada", "smart city", "surveillance", "cctv", "iccc",
                        "command control centre", "smart infrastructure", "smart grid"],
    "Cloud/Infra":    ["cloud hosting", "aws", "azure", "google cloud", "kubernetes", "docker",
                        "microservice", "server setup", "datacenter"],
}

SECTOR_KW = {
    "Smart City":         ["smart city", "smart cities mission", "iccc", "intelligent city"],
    "Urban Local Body":   ["municipal corporation", "nagar palika", "nagar nigam", "ulb", "urban local body",
                            "city corporation", "mc ", "nmc", "amc", "vmc", "smc", "bmc", "pmc", "pcmc", "rmc"],
    "Land Records":       ["land record", "cadastral", "land demarcation", "property survey", "revenue survey",
                            "bhulekh", "bhu-naksha", "land parcel", "dlr", "cadastral mapping"],
    "Tourism":            ["tourism", "tourist", "heritage", "travel platform", "hospitality", "tcgl", "gujrat tourism"],
    "Water/Sanitation":   ["water supply", "sewerage", "sanitation", "drainage", "jal", "amrut", "swachh"],
    "Health":             ["health", "hospital", "medical", "clinic", "nhm", "ayushman"],
    "Education":          ["education", "school", "college", "university", "diksha", "vidya"],
    "Agriculture":        ["agriculture", "kisan", "farmer", "crop", "irrigation", "agri"],
    "Transport/Traffic":  ["transport", "traffic", "road", "highway", "metro", "railway", "vehicle tracking"],
    "Power/Utility":      ["electricity", "power", "energy", "utility mapping", "discoms", "feeder"],
    "KVIC/Industry":      ["kvic", "msme", "khadi", "village industry", "industrial estate"],
    "eGovernance/IT":     ["e-governance", "government it", "nic", "meity", "digital india", "common service"],
}

CLIENT_TYPE_KW = {
    "Central PSU":   ["coal india", "ongc", "bhel", "ntpc", "gail", "iocl", "bpcl", "hpcl", "sail", "nmdc",
                       "nhpc", "rites", "ircon", "nbcc", "tcil", "kvic", "nsic", "wapcos", "npci", "uidai"],
    "Central Govt":  ["ministry of", "government of india", "central government", "nic ", "meity", "niti aayog",
                       "nhm", "isro", "drdo", "doit", "mha", "moe", "mof", "dpiit"],
    "State PSU":     ["gujarat state", "maharashtra state", "state road", "state electricity", "state transport",
                       "vidhut", "discoms", "genco", "transco", "sldc", "gsfc", "gmdc", "gidc"],
    "ULB":           ["municipal corporation", "nagar palika", "nagar nigam", "town planning",
                       "urban development authority", "bmc", "amc", "vmc", "smc", "rmc", "pmc", "pcmc",
                       "nmmc", "kdmc", "ulb", "corporation"],
    "Smart City SPV":["smart city", "smartcity", "spv", "special purpose vehicle"],
    "State Govt":    ["state government", "government of gujarat", "government of maharashtra",
                       "government of rajasthan", "collector office", "district", "taluka panchayat",
                       "gram panchayat", "state dept", "secretariat"],
    "Private":       ["pvt. ltd.", "pvt ltd", "private limited", "ltd.", " llp", "group", "industries"],
}

WORK_TYPE_KW = {
    "Implementation":  ["implementation", "development", "develop", "design and development",
                         "system development", "software development", "portal development", "creation of"],
    "Survey":          ["survey", "field survey", "gis mapping", "land survey", "drone survey",
                         "cadastral survey", "property survey"],
    "AMC/O&M":         ["amc", "annual maintenance", "operation and maintenance", "o&m",
                         "support and maintenance", "camc", "comprehensive maintenance"],
    "Consulting/DPR":  ["consulting", "consultancy", "advisory", "study", "feasibility",
                         "dpr", "project report", "assessment"],
    "Training":        ["training", "capacity building", "workshop", "handholding"],
    "Supply":          ["supply", "procurement", "purchase", "supply of"],
    "Data Digitization":["digitization", "digitisation", "data entry", "data conversion", "scanning"],
}

DOC_SCORE = {
    "nit": 10, "rfp": 10, "tender_doc": 9, "notice_inviting": 10, "bid_doc": 9,
    "work_order": 9, "_wo_": 9, "_wo.": 9, "purchase_order": 8, "loa": 9,
    "letter_of_award": 9, "letter_of_acceptance": 9, "loi": 8,
    "completion": 8, "completion_cert": 9, "pac": 7, "final_acceptance": 7,
    "agreement": 7, "contract": 7, "mou": 6,
    "boq": 6, "price_schedule": 6, "bill_of_quantities": 6,
    "scope": 5, "sow": 5, "technical_proposal": 5, "project_report": 5,
    "payment": 3, "invoice": 2, "minutes": 2, "photo": 1, "image": 0,
}

INDIAN_STATES = [
    "gujarat", "maharashtra", "rajasthan", "uttar pradesh", "madhya pradesh",
    "karnataka", "tamil nadu", "andhra pradesh", "telangana", "kerala",
    "odisha", "orissa", "west bengal", "bihar", "jharkhand", "haryana",
    "punjab", "himachal pradesh", "uttarakhand", "goa", "chhattisgarh",
    "assam", "tripura", "meghalaya", "nagaland", "manipur", "mizoram",
    "arunachal pradesh", "sikkim", "delhi", "nct of delhi", "jammu",
    "kashmir", "ladakh", "chandigarh", "puducherry", "andaman", "lakshadweep",
]

# ══ LOGGING ═══════════════════════════════════════════════════════════════════
def log(msg: str, level: str = "INFO"):
    ts = datetime.now().strftime("%H:%M:%S")
    line = f"[{ts}] [{level:5s}] {msg}"
    print(line)
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass

# ══ PROGRESS SAVE/LOAD ════════════════════════════════════════════════════════
def load_progress() -> Dict:
    if os.path.exists(PROGRESS_FILE):
        try:
            with open(PROGRESS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"processed": [], "rows": [], "total": 0}

def save_progress(progress: Dict):
    try:
        with open(PROGRESS_FILE, "w", encoding="utf-8") as f:
            json.dump(progress, f, indent=2, default=str)
    except Exception as e:
        log(f"Progress save error: {e}", "WARN")

# ══ FILE READERS ══════════════════════════════════════════════════════════════
def read_pdf(path: str) -> str:
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= MAX_PAGES: break
                try:
                    t = page.extract_text() or ""
                    text += t + "\n"
                    if len(text) >= MAX_CHARS: break
                except Exception: continue
        return text[:MAX_CHARS]
    except Exception:
        return ""

def read_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        text = ""
        for para in doc.paragraphs[:250]:
            text += para.text + "\n"
            if len(text) >= MAX_CHARS: break
        for table in doc.tables[:8]:
            for row in table.rows:
                text += " | ".join(c.text for c in row.cells) + "\n"
                if len(text) >= MAX_CHARS: break
        return text[:MAX_CHARS]
    except Exception:
        return ""

def read_xlsx(path: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        text = ""
        for sh in wb.sheetnames[:3]:
            ws = wb[sh]
            text += f"\n--- {sh} ---\n"
            for row in ws.iter_rows(max_row=60, values_only=True):
                r = " | ".join(str(c) for c in row if c is not None)
                if r.strip(): text += r + "\n"
                if len(text) >= MAX_CHARS: break
        return text[:MAX_CHARS]
    except Exception:
        return ""

def read_file(path: str) -> str:
    try:
        size = os.path.getsize(path)
        if size < 200 or size > 60 * 1024 * 1024: return ""
        ext = Path(path).suffix.lower()
        if ext == ".pdf":    return read_pdf(path)
        if ext == ".docx":   return read_docx(path)
        if ext in (".xlsx", ".xls"): return read_xlsx(path)
        if ext in (".txt", ".csv"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(MAX_CHARS)
        return ""
    except Exception:
        return ""

def doc_score(filename: str) -> int:
    name = filename.lower().replace(" ", "_").replace("-", "_")
    score = 0
    for kw, pts in DOC_SCORE.items():
        if kw in name:
            score = max(score, pts)
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":  score += 1
    if ext == ".docx": score += 1
    return score

# ══ EXTRACTORS ════════════════════════════════════════════════════════════════
def ex_value(text: str) -> Tuple[Optional[float], str]:
    patterns = [
        r"(?:rs\.?|inr|₹)\s*([0-9,]+(?:\.[0-9]+)?)\s*(?:crore|cr\.?|crores?)",
        r"\b([0-9]+(?:\.[0-9]+)?)\s*(?:crore|cr\.?|crores?)\b",
        r"(?:rs\.?|inr|₹)\s*([0-9]{1,3}(?:,[0-9]{2,3}){2,}(?:\.[0-9]+)?)",
        r"(?:contract|total|project|tender|estimated|bid)\s+(?:value|amount|cost)[:\s]+(?:rs\.?|inr|₹)?\s*([0-9,]+(?:\.[0-9]+)?)",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text[:5000], re.IGNORECASE):
            raw = m.group(1).replace(",", "").strip()
            try:
                v = float(raw)
                if v > 1_00_00_000:   return round(v/1_00_00_000, 2), f"Rs.{v:,.0f}"
                if v > 1_00_000:      return round(v/100, 2), f"Rs.{v:,.0f} (Lakh)"
                if 0.01 < v <= 500:   return round(v, 2), f"Rs.{v} Cr"
            except Exception: continue
    return None, ""

def ex_dates(text: str) -> Tuple[str, str]:
    pats = [
        r'\b(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b',
        r'\b(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{2,4})\b',
        r'\b((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\-\s]\d{4})\b',
    ]
    found = []
    for pat in pats:
        found.extend(re.findall(pat, text[:4000], re.IGNORECASE))
    found = list(dict.fromkeys(found))[:8]
    return (found[0] if found else ""), (found[-1] if len(found) > 1 else "")

def ex_client(text: str, folder: str) -> Tuple[str, str]:
    combined = (text[:3000] + " " + folder).lower()
    for ctype, kws in CLIENT_TYPE_KW.items():
        for kw in kws:
            if kw in combined:
                idx = combined.find(kw)
                snippet = text[max(0,idx-30):idx+120] if idx < len(text) else folder
                lines = [l.strip() for l in snippet.split('\n') if len(l.strip()) > 3]
                name = lines[0][:80] if lines else kw.title()
                return name, ctype
    return folder[:60], "Unknown"

def ex_tech(text: str, folder: str) -> Tuple[str, str, str]:
    combined = (text + " " + folder).lower()
    techs, details, sectors = [], [], []
    for tech, kws in TECH_KW.items():
        hits = [kw for kw in kws if kw in combined]
        if hits:
            techs.append(tech)
            details.extend(hits[:2])
    for sector, kws in SECTOR_KW.items():
        if any(kw in combined for kw in kws):
            sectors.append(sector)
    primary = techs[0] if techs else "IT Services"
    return primary, ", ".join(list(dict.fromkeys(details))[:7]), ", ".join(sectors[:4])

def ex_work_type(text: str, folder: str) -> str:
    combined = (text[:2000] + " " + folder).lower()
    for wt, kws in WORK_TYPE_KW.items():
        if any(kw in combined for kw in kws):
            return wt
    return "Implementation"

def ex_manpower(text: str) -> Tuple[int, str]:
    pats = [
        r'(\d+)\s*(?:nos?\.?\s*)?(?:professionals?|engineers?|staff|employees?|personnel|resources?|team\s*members?)',
        r'(?:team\s+(?:of|size(?:\s+of)?))\s*(\d+)',
        r'(\d+)\s*(?:full[\s-]?time|dedicated)\s*(?:resources?|staff|engineers?)',
    ]
    count = 0
    for pat in pats:
        for m in re.finditer(pat, text.lower()):
            try:
                n = int(m.group(1))
                if 1 <= n <= 1000: count = max(count, n)
            except Exception: pass
    roles = ["project manager","gis expert","developer","analyst","surveyor","engineer","coordinator","architect","team lead"]
    found_roles = [r.title() for r in roles if r in text.lower()]
    return count, ", ".join(found_roles[:5])

def ex_tender_no(text: str) -> str:
    pats = [
        r'(?:tender\s*no\.?|nit\s*no\.?|rfp\s*no\.?|ref\s*no\.?)[:\s]+([A-Z0-9/_\-\.]{5,50})',
        r'(?:work\s+order\s*no\.?|wo\s*no\.?|p\.?\s*o\.?\s*no\.?)[:\s]+([A-Z0-9/_\-\.]{5,50})',
        r'(?:contract\s*no\.?|agreement\s*no\.?)[:\s]+([A-Z0-9/_\-\.]{5,50})',
    ]
    for pat in pats:
        m = re.search(pat, text[:3000], re.IGNORECASE)
        if m: return m.group(1).strip()[:50]
    return ""

def ex_area(text: str) -> str:
    pats = [
        r'(\d+(?:,\d+)?(?:\.\d+)?)\s*(?:sq\.?\s*km|square\s*km|sqkm|km\s*2)',
        r'(\d+(?:,\d+)?(?:\.\d+)?)\s*(?:hectare|ha)\b',
        r'(\d+(?:,\d+)?(?:\.\d+)?)\s*(?:acre|acres)\b',
    ]
    for pat in pats:
        m = re.search(pat, text.lower())
        if m: return m.group(0)[:30]
    return ""

def ex_scope(text: str) -> str:
    scope_pats = [
        r'scope\s+of\s+work\s*[:\-]\s*(.{80,400}?)(?:\n\n|\Z)',
        r'scope\s+of\s+(?:the\s+)?(?:project|services?)\s*[:\-]\s*(.{80,400}?)(?:\n\n|\Z)',
        r'objective\s*[:\-]\s*(.{60,300}?)(?:\n|\Z)',
        r'brief\s+(?:about|description)\s*[:\-]\s*(.{60,300}?)(?:\n|\Z)',
    ]
    for pat in scope_pats:
        m = re.search(pat, text[:4000], re.IGNORECASE | re.DOTALL)
        if m:
            s = re.sub(r'\s+', ' ', m.group(1)).strip()
            if len(s) > 40: return s[:280]
    lines = [l.strip() for l in text[:1500].split('\n') if len(l.strip()) > 50]
    return (' '.join(lines[:3]))[:280] if lines else ""

def ex_state(text: str, folder: str) -> str:
    combined = (text[:2000] + " " + folder).lower()
    for st in INDIAN_STATES:
        if st in combined: return st.title()
    return ""

def ex_role(text: str) -> str:
    tl = text.lower()
    if any(k in tl for k in ["lead partner","lead firm","lead bidder","prime bidder","consortium leader","lead member"]):
        return "Consortium Lead"
    if any(k in tl for k in ["consortium member","associate member","jv member","partner firm","member firm"]):
        return "Consortium Member"
    if any(k in tl for k in ["sub-contractor","subcontractor","sub contractor","sub-vendor"]):
        return "Sub-contractor"
    return "Lead"

def ex_project_name(text: str, folder: str) -> str:
    pats = [
        r'(?:project\s+(?:name|title)|name\s+of\s+(?:the\s+)?(?:project|work|assignment))\s*[:\-]\s*([^\n]{15,120})',
        r'(?:supply[,\s]+implementation|design[,\s]+development|development\s+of)\s+([^\n]{15,100})',
        r'(?:rfp|nit)\s+for\s+([^\n]{15,100})',
    ]
    for pat in pats:
        m = re.search(pat, text[:2000], re.IGNORECASE)
        if m:
            name = re.sub(r'\s+', ' ', m.group(1)).strip()
            if len(name) > 10: return name[:100]
    return folder[:80]

def ex_status(text: str, has_completion: bool) -> str:
    if has_completion: return "Completed"
    tl = text.lower()
    if any(k in tl for k in ["ongoing", "in progress", "under execution", "work in progress"]): return "Ongoing"
    if any(k in tl for k in ["completed", "successfully completed", "work completed", "project completed"]): return "Completed"
    if any(k in tl for k in ["terminated", "abandoned", "foreclosed"]): return "Terminated"
    return "Unknown"

# ══ FOLDER ANALYSIS ═══════════════════════════════════════════════════════════
DOC_EXTS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.pptx', '.ppt'}
SKIP_DIRS = {'$recycle.bin', 'system volume information', '__macosx', 'thumbs',
             '.git', '.svn', 'node_modules', '__pycache__', 'temp', 'tmp'}

def get_all_files(folder: str) -> List[str]:
    files = []
    try:
        for root, dirs, fnames in os.walk(folder):
            dirs[:] = [d for d in dirs if d.lower() not in SKIP_DIRS and not d.startswith('.')]
            for fn in fnames:
                if Path(fn).suffix.lower() in DOC_EXTS:
                    files.append(os.path.join(root, fn))
    except PermissionError:
        pass
    return files

def has_doc_files(folder: str) -> bool:
    try:
        for fn in os.listdir(folder):
            if Path(fn).suffix.lower() in DOC_EXTS:
                return True
    except Exception:
        pass
    return False

def find_project_folders(root: str) -> List[Tuple[str, int]]:
    """Smart project folder detection."""
    log(f"Scanning folder tree... (this may take a few minutes)")
    candidates = []

    for dirpath, dirnames, filenames in os.walk(root):
        # Depth check
        rel = dirpath.replace(root, "").replace("\\","/").replace("//","/")
        depth = rel.count("/")
        if depth == 0: continue           # skip root itself
        if depth > MAX_DEPTH:
            dirnames.clear(); continue

        # Skip system/junk dirs
        dirname = os.path.basename(dirpath)
        if dirname.lower() in SKIP_DIRS or dirname.startswith('.'):
            dirnames[:] = []; continue

        dirnames[:] = [d for d in dirnames if d.lower() not in SKIP_DIRS and not d.startswith('.')]

        # Count doc files directly in this folder
        doc_count = sum(1 for f in filenames if Path(f).suffix.lower() in DOC_EXTS)
        subdir_count = len(dirnames)

        # A project folder if: has docs directly, OR is a leaf with any content
        if doc_count >= 1:
            candidates.append((dirpath, depth))
        elif subdir_count >= 1 and depth <= 2:
            # Could be an org folder (Year / Client); its sub-folders are projects
            pass

    # De-duplicate: if parent already captured, don't also add its children separately
    # Sort by path length (shorter = higher up = parent)
    candidates.sort(key=lambda x: len(x[0]))
    final = []
    captured_parents = set()

    for path, depth in candidates:
        # Check if any ancestor is already a captured project folder at depth <=3
        parent = str(Path(path).parent)
        grandparent = str(Path(parent).parent)
        if parent in captured_parents or grandparent in captured_parents:
            # Already inside a captured project — skip (files belong to parent project)
            continue
        final.append((path, depth))
        if depth <= 3:
            captured_parents.add(path)

    log(f"Found {len(final)} project folders to process")
    return final

def check_doc_types(filenames: List[str]) -> Dict[str, bool]:
    r = {"has_work_order": False, "has_completion_cert": False, "has_loa": False, "has_boq": False, "has_nit": False}
    for f in filenames:
        fl = f.lower().replace(" ", "_")
        if any(k in fl for k in ["work_order","_wo_","_wo.","purchase_order","po_","supply_order"]): r["has_work_order"] = True
        if any(k in fl for k in ["completion","handover","pac","final_acceptance","taking_over"]): r["has_completion_cert"] = True
        if any(k in fl for k in ["loa","loi","letter_of_award","letter_of_acceptance","acceptance_letter"]): r["has_loa"] = True
        if any(k in fl for k in ["boq","bill_of_quantities","price_schedule","price_bid","financial_bid"]): r["has_boq"] = True
        if any(k in fl for k in ["nit","rfp","tender_doc","bid_doc","notice_inviting","request_for"]): r["has_nit"] = True
    return r

# ══ PROCESS ONE FOLDER ════════════════════════════════════════════════════════
def process_folder(folder_path: str, depth: int, project_id: str) -> Optional[Dict]:
    folder_name = os.path.basename(folder_path)
    all_files = get_all_files(folder_path)
    total_files = len(all_files)
    if total_files == 0: return None

    # Score & rank documents
    scored = sorted([(f, doc_score(os.path.basename(f))) for f in all_files], key=lambda x: -x[1])
    filenames_only = [os.path.basename(f) for f in all_files]
    key_docs = [os.path.basename(f) for f, _ in scored[:5]]

    # Read top documents
    combined_text = ""
    main_source = ""
    for filepath, score in scored[:MAX_DOCS]:
        text = read_file(filepath)
        if len(text) > 100:
            if not main_source: main_source = os.path.basename(filepath)
            combined_text += f"\n\n=== {os.path.basename(filepath)} ===\n{text}"
            if len(combined_text) > MAX_CHARS * 2: break

    if not combined_text.strip():
        for filepath, _ in scored[:5]:
            text = read_file(filepath)
            if len(text) > 50:
                combined_text = text
                main_source = os.path.basename(filepath)
                break

    txt = combined_text
    fname = folder_name

    # Extract all fields
    value_cr, value_raw    = ex_value(txt + " " + fname)
    start_d, end_d         = ex_dates(txt)
    client_name, ctype     = ex_client(txt, fname)
    primary_tech, tech_stk, sectors = ex_tech(txt, fname)
    work_type              = ex_work_type(txt, fname)
    manpower, key_roles    = ex_manpower(txt)
    tender_no              = ex_tender_no(txt)
    area                   = ex_area(txt)
    scope                  = ex_scope(txt)
    state                  = ex_state(txt, fname)
    our_role               = ex_role(txt)
    project_name           = ex_project_name(txt, fname)
    doc_types              = check_doc_types(filenames_only)
    status                 = ex_status(txt, doc_types["has_completion_cert"])

    # Relevant for
    rel_tags = ["Profile"]
    if doc_types["has_completion_cert"]: rel_tags += ["TQ", "TechProposal"]
    if doc_types["has_boq"]:            rel_tags.append("BOQ")
    if value_cr and value_cr > 0.5:     rel_tags.append("PreBid")
    rel_tags.append("BidNoBid")
    relevant_for = ", ".join(rel_tags)

    # Similar work category
    sim_cats = []
    for tech in primary_tech.split("/"):
        t = tech.strip()
        if "GIS" in t or "Geo" in t:       sim_cats.append("GIS/Geospatial Implementation")
        if "Mobile" in t:                   sim_cats.append("Mobile App Development")
        if "Web" in t or "Portal" in t:     sim_cats.append("Web Portal Development")
        if "Survey" in t or "Mapping" in t: sim_cats.append("Survey & Field Mapping")
        if "ERP" in t:                      sim_cats.append("ERP Implementation")
        if "eGov" in t or "Gov" in t:       sim_cats.append("eGovernance Project")
    for sec in sectors.split(","):
        s = sec.strip()
        if "ULB" in s or "Urban" in s: sim_cats.append("Municipal IT Project")
        if "Smart City" in s:           sim_cats.append("Smart City SPV Project")
        if "Land" in s:                 sim_cats.append("Land Records/Cadastral")
    similar_work = ", ".join(list(dict.fromkeys(sim_cats))[:4]) or primary_tech

    return {
        "Project_ID":          project_id,
        "Folder_Name":         fname,
        "Client_Name":         client_name,
        "Client_Type":         ctype,
        "State":               state,
        "Project_Name":        project_name,
        "Scope_Summary":       scope,
        "Work_Type":           work_type,
        "Primary_Technology":  primary_tech,
        "Tech_Stack_Details":  tech_stk,
        "Sector_Tags":         sectors,
        "Our_Role":            our_role,
        "Contract_Value_Cr":   value_cr if value_cr else "",
        "Contract_Value_Raw":  value_raw,
        "Duration_Months":     "",
        "Start_Date":          start_d,
        "End_Date":            end_d,
        "Project_Status":      status,
        "Manpower_Count":      manpower if manpower else "",
        "Key_Roles_Deployed":  key_roles,
        "Survey_Area":         area,
        "Tender_WO_Number":    tender_no,
        "Has_Work_Order":      "Yes" if doc_types["has_work_order"] else "No",
        "Has_Completion_Cert": "Yes" if doc_types["has_completion_cert"] else "No",
        "Has_LOA":             "Yes" if doc_types["has_loa"] else "No",
        "Has_BOQ":             "Yes" if doc_types["has_boq"] else "No",
        "Has_NIT_RFP":         "Yes" if doc_types["has_nit"] else "No",
        "Similar_Work_Category": similar_work,
        "Relevant_For":        relevant_for,
        "Total_Files_Found":   total_files,
        "Key_Documents":       " | ".join(key_docs[:3]),
        "Notes":               "",
        "Folder_Path":         folder_path,
        "Scan_Date":           datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

# ══ OUTPUT: EXCEL ══════════════════════════════════════════════════════════════
def save_excel(rows: List[Dict]):
    try:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Nascent Projects"

        hdr_fill = PatternFill(start_color="0D1B2A", end_color="0D1B2A", fill_type="solid")
        hdr_font = Font(color="FFFFFF", bold=True, size=11, name="Calibri")
        wrap = Alignment(wrap_text=True, vertical="top")
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # Headers
        for ci, col in enumerate(COLUMNS, 1):
            c = ws.cell(row=1, column=ci, value=col.replace("_", " "))
            c.fill = hdr_fill; c.font = hdr_font; c.alignment = center

        # Status colors
        STATUS_COLORS = {"Completed": "D1FAE5", "Ongoing": "DBEAFE", "Unknown": "F3F4F6", "Terminated": "FEE2E2"}
        CTYPE_COLORS  = {"ULB": "EDE9FE", "Central Govt": "DBEAFE", "State Govt": "D1FAE5",
                          "Central PSU": "FEF3C7", "Smart City SPV": "FCE7F3", "Private": "F3F4F6"}

        col_widths = {
            "Scope_Summary": 42, "Folder_Path": 52, "Tech_Stack_Details": 32,
            "Project_Name": 38, "Key_Documents": 38, "Client_Name": 28,
            "Similar_Work_Category": 32, "Sector_Tags": 22, "Relevant_For": 24,
        }

        for ri, row in enumerate(rows, 2):
            for ci, col in enumerate(COLUMNS, 1):
                val = row.get(col, "")
                c = ws.cell(row=ri, column=ci, value=val)
                c.alignment = wrap
                if col == "Project_Status":
                    color = STATUS_COLORS.get(str(val), "F9FAFB")
                    c.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                if col == "Client_Type":
                    color = CTYPE_COLORS.get(str(val), "FFFFFF")
                    c.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                if col == "Contract_Value_Cr" and val:
                    try:
                        c.number_format = "0.00"
                    except Exception: pass

        for ci, col in enumerate(COLUMNS, 1):
            letter = openpyxl.utils.get_column_letter(ci)
            ws.column_dimensions[letter].width = col_widths.get(col, 16)

        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

        # Summary sheet
        ws2 = wb.create_sheet("Summary")
        ws2["A1"] = "Nascent Project Scan Summary"
        ws2["A1"].font = Font(bold=True, size=14)
        ws2["A3"] = "Total Projects Found:"; ws2["B3"] = len(rows)
        ws2["A4"] = "Completed:";            ws2["B4"] = sum(1 for r in rows if r.get("Project_Status")=="Completed")
        ws2["A5"] = "Ongoing:";              ws2["B5"] = sum(1 for r in rows if r.get("Project_Status")=="Ongoing")
        ws2["A6"] = "Unknown Status:";       ws2["B6"] = sum(1 for r in rows if r.get("Project_Status")=="Unknown")
        ws2["A8"] = "With Work Order:";      ws2["B8"] = sum(1 for r in rows if r.get("Has_Work_Order")=="Yes")
        ws2["A9"] = "With Completion Cert:"; ws2["B9"] = sum(1 for r in rows if r.get("Has_Completion_Cert")=="Yes")
        ws2["A10"]= "With BOQ:";             ws2["B10"]= sum(1 for r in rows if r.get("Has_BOQ")=="Yes")
        ws2["A12"]= "Scan Date:";            ws2["B12"]= datetime.now().strftime("%Y-%m-%d %H:%M")
        ws2["A13"]= "Root Folder:";          ws2["B13"]= ROOT_DIR

        # Tech breakdown
        from collections import Counter
        tech_counts = Counter(r.get("Primary_Technology","Unknown") for r in rows)
        ws2["A15"] = "Technology Breakdown"; ws2["A15"].font = Font(bold=True)
        for i, (tech, cnt) in enumerate(tech_counts.most_common(), 16):
            ws2[f"A{i}"] = tech; ws2[f"B{i}"] = cnt

        wb.save(OUTPUT_XLSX)
        log(f"Excel saved → {OUTPUT_XLSX}")
    except Exception as e:
        log(f"Excel error: {e}", "ERROR")
        traceback.print_exc()

# ══ OUTPUT: CSV ═══════════════════════════════════════════════════════════════
def save_csv(rows: List[Dict]):
    try:
        with open(OUTPUT_CSV, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=COLUMNS, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)
        log(f"CSV saved → {OUTPUT_CSV}")
    except Exception as e:
        log(f"CSV error: {e}", "ERROR")

# ══ MAIN ══════════════════════════════════════════════════════════════════════
def main():
    print("\n" + "="*65)
    print("  NASCENT PROJECT DATA SCANNER v1.0")
    print("  Scans your 100GB project archive — completely OFFLINE")
    print("  Progress saved every 10 folders — safe to interrupt & resume")
    print("="*65 + "\n")

    # Check drive
    if not os.path.exists(ROOT_DIR):
        log(f"ERROR: Cannot find: {ROOT_DIR}", "ERROR")
        log("Make sure H: drive (external hard disk) is connected!", "ERROR")
        input("\nPress Enter to exit...")
        sys.exit(1)

    log(f"Root: {ROOT_DIR}")
    log(f"Output: {DESKTOP}")

    # Load progress
    progress = load_progress()
    already_done = set(progress.get("processed", []))
    all_rows = progress.get("rows", [])
    log(f"Previously processed: {len(already_done)} folders, {len(all_rows)} rows")

    # Scan for project folders
    project_folders = find_project_folders(ROOT_DIR)
    to_process = [(p, d) for p, d in project_folders if p not in already_done]
    log(f"Remaining to process: {len(to_process)} folders\n")

    if not to_process:
        log("All done! Saving final output...")
        save_excel(all_rows)
        save_csv(all_rows)
        print(f"\n✅ COMPLETE — {len(all_rows)} projects extracted")
        print(f"📊 Excel: {OUTPUT_XLSX}")
        print(f"📄 CSV:   {OUTPUT_CSV}")
        input("\nPress Enter to exit...")
        return

    # Process
    t_start = time.time()
    processed = 0
    errors = 0

    for idx, (folder_path, depth) in enumerate(to_process, 1):
        fname = os.path.basename(folder_path)
        elapsed = time.time() - t_start
        if idx > 1:
            avg = elapsed / (idx - 1)
            eta = avg * (len(to_process) - idx + 1)
            eta_str = f"| ETA {int(eta//60)}m{int(eta%60)}s"
        else:
            eta_str = ""

        log(f"[{idx:4d}/{len(to_process)}] {fname[:50]:<50} (d={depth}) {eta_str}")

        try:
            pid = f"PROJ-{(len(already_done) + idx):04d}"
            row = process_folder(folder_path, depth, pid)
            if row:
                all_rows.append(row)
                processed += 1
                log(f"         ✓ {row['Project_Status']} | {row['Primary_Technology']} | "
                    f"₹{row['Contract_Value_Cr']} Cr | {row['Client_Type']}")
            else:
                log(f"         – No documents found, skipped")
        except Exception as e:
            log(f"         ✗ Error: {e}", "ERROR")
            errors += 1

        already_done.add(folder_path)

        # Save every 10
        if idx % 10 == 0 or idx == len(to_process):
            progress["processed"] = list(already_done)
            progress["rows"] = all_rows
            progress["total"] = len(all_rows)
            save_progress(progress)
            save_csv(all_rows)
            log(f"         💾 Progress saved. Rows so far: {len(all_rows)}")

    # Final save
    log("\nFinal output...")
    save_excel(all_rows)
    save_csv(all_rows)

    elapsed = int(time.time() - t_start)
    print("\n" + "="*65)
    print(f"  ✅ COMPLETE in {elapsed//60}m {elapsed%60}s")
    print(f"  Projects extracted : {processed}")
    print(f"  Errors             : {errors}")
    print(f"  Total rows         : {len(all_rows)}")
    print(f"\n  📊 Excel : {OUTPUT_XLSX}")
    print(f"  📄 CSV   : {OUTPUT_CSV}")
    print(f"\n  NEXT: Import CSV to Google Sheets → File > Import > Upload")
    print("="*65 + "\n")
    input("Press Enter to exit...")

if __name__ == "__main__":
    main()
