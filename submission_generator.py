"""
Submission Package Generator
Generates all bid submission documents auto-filled from Nascent Profile + tender data
FIXED: Profile key access now reads nested structure (company.cin, company.pan etc.)
"""
import json, re, zipfile
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

BASE_DIR = Path(__file__).parent
PROFILE_PATH = BASE_DIR / "nascent_profile.json"

# ─── Nascent defaults (fallback if profile not saved) ───
DEFAULT_PROFILE = {
    "name": "Nascent Info Technologies Pvt. Ltd.",
    "cin": "U72200GJ2006PTC048723",
    "pan": "AACCN3670J",
    "gstin": "24AACCN3670J1ZG",
    "msme": "UDYAM-GJ-01-0007420",
    "signatory": "Hitesh Patel",
    "signatory_designation": "Chief Administrative Officer (CAO)",
    "md": "Maulik Bhagat",
    "address": "Ahmedabad, Gujarat - 380 015",
    "turnover_fy1": "18.83", "turnover_fy1_label": "2024-25",
    "turnover_fy2": "16.36", "turnover_fy2_label": "2023-24",
    "turnover_fy3": "16.36", "turnover_fy3_label": "2022-23",
    "net_worth": "26.09",
    "ca_firm": "Anuj J. Sharedalal & Co.",
    "cmmi": "CMMI V2.0 Level 3",
    "cmmi_valid": "19-Dec-2026",
    "iso_9001": "08-Sep-2028",
    "iso_27001": "08-Sep-2028",
    "iso_20000": "08-Sep-2028",
    "total_employees": "67",
    "it_employees": "21",
}


def load_profile() -> dict:
    """
    FIXED: Reads nested nascent_profile.json structure:
      company.name, company.cin, finance.net_worth_cr, etc.
    Falls back to DEFAULT_PROFILE for any missing key.
    """
    raw = {}
    for path in [Path("/tmp/bid-nobid/nascent_profile.json"), PROFILE_PATH]:
        if path.exists():
            try:
                raw = json.loads(path.read_text(encoding="utf-8"))
                break
            except Exception:
                continue

    if not raw:
        return DEFAULT_PROFILE

    company = raw.get("company", {})
    finance = raw.get("finance", {})
    certs = raw.get("certifications", {})
    employees = raw.get("employees", {})

    # Map nested keys → flat profile dict for document generation
    p = dict(DEFAULT_PROFILE)  # start with defaults
    p["name"] = company.get("name", p["name"])
    p["cin"] = company.get("cin", p["cin"])
    p["pan"] = company.get("pan", p["pan"])
    p["gstin"] = company.get("gstin", p["gstin"])
    p["msme"] = company.get("udyam", p["msme"])
    p["address"] = company.get("address", p["address"])
    p["md"] = company.get("md", p["md"])
    p["signatory"] = company.get("signatory", p["signatory"])
    p["signatory_designation"] = company.get("signatory_designation", p["signatory_designation"])

    # Finance
    turnover_by_year = finance.get("turnover_by_year", {})
    if turnover_by_year:
        sorted_fy = sorted(turnover_by_year.keys(), reverse=True)
        if len(sorted_fy) >= 1:
            p["turnover_fy1_label"] = sorted_fy[0]
            p["turnover_fy1"] = str(turnover_by_year[sorted_fy[0]])
        if len(sorted_fy) >= 2:
            p["turnover_fy2_label"] = sorted_fy[1]
            p["turnover_fy2"] = str(turnover_by_year[sorted_fy[1]])
        if len(sorted_fy) >= 3:
            p["turnover_fy3_label"] = sorted_fy[2]
            p["turnover_fy3"] = str(turnover_by_year[sorted_fy[2]])
    p["net_worth"] = str(finance.get("net_worth_cr", p["net_worth"]))
    p["ca_firm"] = finance.get("ca_firm", p["ca_firm"])

    # Certifications
    cmmi = certs.get("cmmi", {})
    if cmmi:
        p["cmmi"] = f"CMMI {cmmi.get('version', 'V2.0')} Level {cmmi.get('level', 3)}"
        p["cmmi_valid"] = cmmi.get("valid_to", p["cmmi_valid"])
    p["iso_9001"] = certs.get("iso_9001", {}).get("valid_to", p["iso_9001"])
    p["iso_27001"] = certs.get("iso_27001", {}).get("valid_to", p["iso_27001"])
    p["iso_20000"] = certs.get("iso_20000", {}).get("valid_to", p["iso_20000"])

    # Employees
    p["total_employees"] = str(employees.get("total_confirmed", p["total_employees"]))
    p["it_employees"] = str(employees.get("it_dev_staff", p["it_employees"]))

    return p


def today_str():
    return datetime.now().strftime("%d %B %Y")


def _safe_field(tender: dict, *keys, default="—"):
    """Get field from tender, unwrapping {value, clause_ref, page_no} dicts."""
    for key in keys:
        v = tender.get(key)
        if v is not None:
            if isinstance(v, dict):
                v = v.get("value", "")
            if v and str(v).strip() not in ("", "—", "null", "None"):
                return str(v).strip()
    return default


# ─── Word doc helpers ───
def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def h(doc, text, size=13, bold=True, center=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.bold = bold
    r.font.size = Pt(size)
    if color:
        rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
        r.font.color.rgb = RGBColor(*rgb)
    return p


def line(doc, text, size=11, bold=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    return p


def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def sign_block(doc, p_data: dict):
    """Standard Nascent signatory block."""
    blank(doc, 2)
    line(doc, "For Nascent Info Technologies Pvt. Ltd.", bold=True)
    blank(doc)
    blank(doc)
    line(doc, "Authorised Signatory")
    line(doc, f"Name: {p_data.get('signatory', 'Hitesh Patel')}")
    line(doc, f"Designation: {p_data.get('signatory_designation', 'Chief Administrative Officer (CAO)')}")
    line(doc, f"Date: {today_str()}")
    line(doc, f"Place: Ahmedabad")


# ─────────────────────────────────────────────────────────────────
# DOCUMENT GENERATORS
# ─────────────────────────────────────────────────────────────────

def gen_cover_letter(tender: dict, p_data: dict) -> Document:
    doc = new_doc()
    org = _safe_field(tender, "org_name")
    tender_name = _safe_field(tender, "tender_name", "brief")
    tender_no = _safe_field(tender, "tender_no", "ref_no")
    deadline = _safe_field(tender, "bid_submission_date", "deadline")
    location = _safe_field(tender, "location")

    h(doc, p_data["name"], size=14, center=True, color="1F497D")
    line(doc, f"A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad 380015", size=10)
    line(doc, f"CIN: {p_data['cin']} | GSTIN: {p_data['gstin']} | PAN: {p_data['pan']}", size=9)
    blank(doc)

    line(doc, f"Date: {today_str()}")
    blank(doc)
    line(doc, "To,", bold=True)
    line(doc, f"The Tender Inviting Authority,")
    line(doc, org)
    blank(doc)

    line(doc, f"Subject: Submission of Bid — {tender_name}", bold=True)
    line(doc, f"Reference: Tender No. {tender_no}", bold=False)
    blank(doc)

    line(doc, "Respected Sir/Madam,")
    blank(doc)
    line(doc, (
        f"We, {p_data['name']}, are pleased to submit our bid in response to the above-referenced tender. "
        f"We have carefully read and understood all the terms, conditions, and requirements mentioned in the "
        f"Request for Proposal / Notice Inviting Tender."
    ))
    blank(doc)
    line(doc, (
        f"We confirm that our bid is valid for the period specified in the tender document and that all "
        f"information furnished herein is true and accurate to the best of our knowledge. We understand that "
        f"submission of false or misleading information shall result in disqualification."
    ))
    blank(doc)
    line(doc, "We hereby declare that:")
    for decl in [
        "We are not blacklisted or debarred by any Government department, PSU, or autonomous body in India.",
        "We have not been convicted of any offence under the Prevention of Corruption Act or any other applicable law.",
        "We accept all terms and conditions of the tender document without any deviation.",
        "All documents submitted along with this bid are genuine and duly authenticated.",
        f"Our bid shall remain valid until the date specified in the tender document.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(decl).font.size = Pt(11)
    blank(doc)
    line(doc, "We look forward to a favourable consideration of our bid.")
    sign_block(doc, p_data)
    return doc


def gen_emd_exemption_letter(tender: dict, p_data: dict) -> Document:
    doc = new_doc()
    org = _safe_field(tender, "org_name")
    tender_no = _safe_field(tender, "tender_no", "ref_no")
    tender_name = _safe_field(tender, "tender_name", "brief")
    emd_amount = _safe_field(tender, "emd")

    h(doc, "APPLICATION FOR EMD EXEMPTION UNDER MSME PROCUREMENT POLICY", size=13, center=True, color="1F497D")
    blank(doc)

    line(doc, f"Date: {today_str()}")
    blank(doc)
    line(doc, "To,", bold=True)
    line(doc, "The Tender Inviting Authority,")
    line(doc, org)
    blank(doc)

    line(doc, f"Subject: Request for EMD Exemption — Tender No. {tender_no}", bold=True)
    blank(doc)

    line(doc, "Respected Sir/Madam,")
    blank(doc)
    line(doc, (
        f"We, {p_data['name']}, hereby apply for exemption from payment of Earnest Money Deposit (EMD) "
        f"of {emd_amount} specified in the above tender."
    ))
    blank(doc)
    line(doc, "GROUNDS FOR EXEMPTION:", bold=True)
    line(doc, (
        f"We are a registered Micro, Small, and Medium Enterprise (MSME) under the MSMED Act, 2006. "
        f"Our Udyam Registration Number is {p_data['msme']}. As per the Public Procurement Policy for "
        f"Micro and Small Enterprises (MSEs) Order, 2012 issued by the Ministry of MSME, Government of India, "
        f"and as amended by the Department of Expenditure Office Memorandum F.No.6/18/2019-PPD dated "
        f"16 November 2020, Micro and Small Enterprises are exempted from payment of Earnest Money Deposit "
        f"in Government procurement."
    ))
    blank(doc)
    line(doc, "We enclose the following documents in support of this application:", bold=True)
    for doc_item in [
        f"Udyam Registration Certificate — Reg. No. {p_data['msme']} (self-attested copy)",
        "PAN Card — self-attested copy",
        "GST Registration Certificate — self-attested copy",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(doc_item).font.size = Pt(11)
    blank(doc)
    line(doc, (
        "We request you to kindly accept our Udyam Registration Certificate in lieu of EMD and grant "
        "exemption as per the applicable policy."
    ))
    sign_block(doc, p_data)
    return doc


def gen_non_blacklisting_declaration(tender: dict, p_data: dict) -> Document:
    doc = new_doc()
    org = _safe_field(tender, "org_name")
    tender_no = _safe_field(tender, "tender_no", "ref_no")

    h(doc, "DECLARATION OF NON-BLACKLISTING / NON-DEBARMENT", size=13, center=True, color="1F497D")
    blank(doc)

    line(doc, f"Date: {today_str()}")
    blank(doc)
    line(doc, "To,", bold=True)
    line(doc, "The Tender Inviting Authority,")
    line(doc, org)
    blank(doc)
    line(doc, f"Ref: Tender No. {tender_no}", bold=True)
    blank(doc)
    line(doc, "DECLARATION", bold=True, size=12)
    blank(doc)
    line(doc, f"I/We, {p_data['name']} (CIN: {p_data['cin']}), hereby solemnly declare that:")
    blank(doc)

    declarations = [
        f"Our company has NOT been blacklisted, debarred, or banned by any Central Government Ministry, "
        f"State Government Department, Public Sector Undertaking, or any other Government body in India "
        f"as on the date of this declaration.",
        "No criminal proceedings or cases under the Prevention of Corruption Act, 1988 or under any "
        "other relevant statute are pending against our company or its directors.",
        "Our company has not been declared insolvent or bankrupt by any court or tribunal in India.",
        "Our company does not have any conflict of interest that could potentially impair the fair "
        "execution of the contract, if awarded.",
        "All information provided in this bid is true, correct, and complete to the best of our "
        "knowledge and belief.",
    ]
    for i, decl in enumerate(declarations, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{i}. {decl}").font.size = Pt(11)

    blank(doc)
    line(doc, (
        "We understand that furnishing of false information or documents in this regard shall result in "
        "immediate disqualification and legal action as per applicable laws."
    ))
    blank(doc)
    line(doc, "This declaration is made in good faith and is binding on our company.", bold=True)
    sign_block(doc, p_data)
    blank(doc)
    line(doc, "Place: Ahmedabad, Gujarat")
    line(doc, "Stamp / Seal of Company: ___________________________")
    return doc


def gen_turnover_certificate(tender: dict, p_data: dict) -> Document:
    doc = new_doc()

    h(doc, "CERTIFICATE OF ANNUAL TURNOVER", size=13, center=True, color="1F497D")
    h(doc, "(On CA Firm Letterhead)", size=10, center=True, bold=False)
    blank(doc)

    line(doc, f"Date: {today_str()}")
    blank(doc)
    line(doc, "To Whomsoever It May Concern")
    blank(doc)
    line(doc, (
        f"This is to certify that we have audited the financial statements of "
        f"{p_data['name']} (CIN: {p_data['cin']}, PAN: {p_data['pan']}) "
        f"for the financial years mentioned below."
    ))
    blank(doc)
    line(doc, "The annual turnover from IT/ITeS services for the last three financial years is as follows:", bold=True)
    blank(doc)

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Financial Year"
    hdr[1].text = "Annual Turnover (Rs. Crore)"
    hdr[2].text = "Nature of Business"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    fy_data = [
        (p_data.get("turnover_fy1_label", "2024-25"), p_data.get("turnover_fy1", "18.83")),
        (p_data.get("turnover_fy2_label", "2023-24"), p_data.get("turnover_fy2", "16.36")),
        (p_data.get("turnover_fy3_label", "2022-23"), p_data.get("turnover_fy3", "16.36")),
    ]
    for fy_label, fy_val in fy_data:
        row = table.add_row().cells
        row[0].text = fy_label
        row[1].text = f"Rs. {fy_val} Crore"
        row[2].text = "IT / ITeS Services"

    blank(doc)
    line(doc, (
        f"The above figures are based on audited financial statements and are true and correct to the "
        f"best of our knowledge. The entire turnover of the company is from IT and ITeS domain services "
        f"including GIS, Software Development, and e-Governance solutions."
    ))
    blank(doc)
    line(doc, f"Net Worth as on 31st March 2025: Rs. {p_data.get('net_worth', '26.09')} Crore", bold=True)
    blank(doc, 2)
    line(doc, f"For {p_data.get('ca_firm', 'Chartered Accountants')}", bold=True)
    blank(doc, 2)
    line(doc, "Chartered Accountant")
    line(doc, "Membership No.: _______________")
    line(doc, f"Date: {today_str()}")
    line(doc, "Place: Ahmedabad")
    return doc


def gen_employee_certificate(tender: dict, p_data: dict) -> Document:
    doc = new_doc()

    h(doc, "CERTIFICATE OF EMPLOYEE STRENGTH", size=13, center=True, color="1F497D")
    h(doc, "(On Company Letterhead)", size=10, center=True, bold=False)
    blank(doc)
    line(doc, f"Date: {today_str()}")
    blank(doc)
    line(doc, "To Whomsoever It May Concern")
    blank(doc)
    line(doc, (
        f"This is to certify that {p_data['name']} (CIN: {p_data['cin']}) currently employs "
        f"{p_data.get('total_employees', '67')} full-time employees on its payroll as on the date of "
        f"this certificate."
    ))
    blank(doc)
    line(doc, "Employee Breakdown:", bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Number of Employees"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    categories = [
        ("GIS Specialists", "11"),
        ("IT / Software Developers", p_data.get("it_employees", "21")),
        ("QA / Testing Engineers", "8"),
        ("Project Managers / BAs", "7"),
        ("Support / Admin / HR", "20"),
        (f"Total (All IT / ITeS)", p_data.get("total_employees", "67")),
    ]
    for cat, count in categories:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = count

    blank(doc)
    line(doc, (
        f"We hereby certify that all the above employees are full-time employees engaged in "
        f"IT / ITeS domain activities. Their EPF/ESIC contributions are being duly paid as per "
        f"applicable regulations."
    ))
    sign_block(doc, p_data)
    return doc


def gen_experience_summary(tender: dict, p_data: dict) -> Document:
    """Summary sheet of Nascent's relevant project experience."""
    doc = new_doc()
    tender_name = _safe_field(tender, "tender_name", "brief")

    h(doc, "STATEMENT OF RELEVANT PROJECT EXPERIENCE", size=13, center=True, color="1F497D")
    h(doc, f"Submitted for: {tender_name[:80]}", size=10, center=True, bold=False)
    blank(doc)

    line(doc, (
        f"The following is a summary of relevant completed and ongoing projects executed by "
        f"{p_data['name']}:"
    ))
    blank(doc)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ["Sr.", "Client / Project", "Value (Rs. Cr)", "Duration", "Status", "Key Technologies"]
    for i, hdr_text in enumerate(headers):
        table.rows[0].cells[i].text = hdr_text
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    projects = [
        ("AMC GIS", "Ahmedabad Municipal Corporation", "10.55", "2020–Ongoing", "Ongoing", "GIS, Web, Mobile, AMC"),
        ("JuMC GIS", "Junagadh Municipal Corporation", "9.78", "2021–Ongoing", "Ongoing", "GIS, Survey, Web, ULB"),
        ("BMC GIS", "Bhavnagar Municipal Corporation", "4.20", "2019–2021", "Completed", "Android, iOS, GeoServer"),
        ("VMC GIS+ERP", "Vadodara Municipal Corporation", "20.50", "2018–2022", "Completed", "GIS, ERP, ULB (Consortium)"),
        ("KVIC Geo Portal", "KVIC Central PSU", "5.15", "2020–2022", "Completed", "Mobile GIS, Geo-tagging, PAN India"),
        ("PCSCL Smart City", "Pimpri-Chinchwad Smart City", "61.19", "2019–Ongoing", "Ongoing", "Smart City, GIS, ERP, IoT"),
        ("TCGL Tourism", "Tourism Corp Gujarat", "9.31", "2021–2023", "Completed", "Web, Mobile, Tourism, GIS"),
        ("NSO Survey", "National Statistics Office", "8.40", "2022–2023", "Completed", "Mobile Survey, PAN India, GoI"),
        ("CEICED", "Electrical Inspector Gujarat", "3.59", "2023–Ongoing", "Ongoing", "eGov, Web, Mobile"),
        ("AMC Heritage App", "Ahmedabad Smart City", "3.80", "2022–2023", "Completed", "AR, Mobile, Tourism"),
    ]
    for i, (short, client, val, dur, status, tech) in enumerate(projects, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = f"{short}\n{client}"
        row[2].text = val
        row[3].text = dur
        row[4].text = status
        row[5].text = tech

    blank(doc)
    line(doc, "Note: Work Orders and Completion / Performance Certificates for all projects are available upon request.", size=9)
    sign_block(doc, p_data)
    return doc


# ─────────────────────────────────────────────────────────────────
# PACKAGE GENERATOR
# ─────────────────────────────────────────────────────────────────

def generate_submission_package(tender: dict, output_dir) -> dict:
    """
    Generate full submission package for the given tender.
    Returns dict with pkg_dir, zip_file, files list.
    """
    if not DOCX_OK:
        return {"error": "python-docx not installed"}

    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True, parents=True)

    p_data = load_profile()
    tender_no = _safe_field(tender, "tender_no", "ref_no", "t247_id")
    safe_no = re.sub(r'[^\w\-]', '_', str(tender_no))[:30]
    pkg_dir = output_dir / f"SubmissionPackage_{safe_no}"
    pkg_dir.mkdir(exist_ok=True, parents=True)

    generated = []
    errors = []

    docs_to_generate = [
        ("01_CoverLetter",          gen_cover_letter),
        ("02_EMD_Exemption_Letter", gen_emd_exemption_letter),
        ("03_NonBlacklisting_Decl", gen_non_blacklisting_declaration),
        ("04_Turnover_Certificate", gen_turnover_certificate),
        ("05_Employee_Certificate", gen_employee_certificate),
        ("06_Experience_Summary",   gen_experience_summary),
    ]

    for fname, gen_fn in docs_to_generate:
        try:
            doc = gen_fn(tender, p_data)
            fpath = pkg_dir / f"{fname}.docx"
            doc.save(str(fpath))
            generated.append({"name": fname.replace("_", " "), "filename": fpath.name})
        except Exception as e:
            errors.append(f"{fname}: {str(e)}")

    # Create ZIP
    zip_name = f"SubmissionPackage_{safe_no}.zip"
    zip_path = output_dir / zip_name
    try:
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            for fpath in sorted(pkg_dir.glob("*.docx")):
                zf.write(str(fpath), fpath.name)
        zip_ok = True
    except Exception as e:
        errors.append(f"ZIP: {str(e)}")
        zip_ok = False

    return {
        "status": "success" if generated else "error",
        "pkg_dir": str(pkg_dir),
        "zip_file": zip_name if zip_ok else None,
        "files": generated,
        "errors": errors,
        "profile_used": p_data.get("name", "Unknown"),
    }
