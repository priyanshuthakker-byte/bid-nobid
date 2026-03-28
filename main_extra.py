"""
PASTE THIS ENTIRE BLOCK AT THE BOTTOM OF main.py
These are the missing/new endpoints for the updated frontend.
"""

# ── SKIP TENDER ────────────────────────────────────────────────
@app.post("/tender/{t247_id}/skip")
async def skip_tender(t247_id: str, data: dict = Body(default={})):
    """Mark tender as skipped/not interested."""
    db = load_db()
    t = db["tenders"].get(t247_id, {"t247_id": t247_id})
    t["status"] = "Not Interested"
    t["skip_reason"] = data.get("reason", "Not interested")
    t["skipped_at"] = datetime.now().isoformat()
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "skipped", "t247_id": t247_id}


@app.get("/skipped-tenders")
async def get_skipped():
    """Return all skipped tenders."""
    db = load_db()
    skipped = [
        t for t in db["tenders"].values()
        if t.get("status") == "Not Interested"
    ]
    return {"tenders": skipped}


@app.post("/tender/{t247_id}/restore")
async def restore_tender(t247_id: str):
    """Restore a skipped tender back to active."""
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    if t.get("status") == "Not Interested":
        t["status"] = "Identified"
        t.pop("skip_reason", None)
        t.pop("skipped_at", None)
        db["tenders"][t247_id] = t
        save_db(db)
    return {"status": "restored", "t247_id": t247_id}


# ── FAVOURITE ──────────────────────────────────────────────────
@app.post("/tender/{t247_id}/favourite")
async def toggle_favourite(t247_id: str, data: dict = Body(default={})):
    """Toggle favourite flag on a tender."""
    db = load_db()
    t = db["tenders"].get(t247_id, {"t247_id": t247_id})
    current = t.get("favourite", False)
    t["favourite"] = not current
    db["tenders"][t247_id] = t
    save_db(db)
    return {"favourite": t["favourite"], "t247_id": t247_id}


# ── DOCUMENT VAULT ─────────────────────────────────────────────
VAULT_DIR = BASE_DIR / "data" / "vault"
VAULT_DIR.mkdir(exist_ok=True, parents=True)

VAULT_DOCS = [
    {"id": "pan_card", "name": "PAN Card", "category": "Company"},
    {"id": "cin_cert", "name": "CIN Certificate / MOA", "category": "Company"},
    {"id": "gst_cert", "name": "GST Certificate", "category": "Company"},
    {"id": "msme_cert", "name": "MSME / UDYAM Certificate", "category": "Company"},
    {"id": "poa_doc", "name": "Power of Attorney (Hitesh Patel)", "category": "Company"},
    {"id": "cmmi_cert", "name": "CMMI Level 3 Certificate", "category": "Certification"},
    {"id": "iso9001_cert", "name": "ISO 9001:2015 Certificate", "category": "Certification"},
    {"id": "iso27001_cert", "name": "ISO 27001:2022 Certificate", "category": "Certification"},
    {"id": "iso20000_cert", "name": "ISO 20000-1:2018 Certificate", "category": "Certification"},
    {"id": "audited_fy2223", "name": "Audited Accounts FY 2022-23", "category": "Finance"},
    {"id": "audited_fy2324", "name": "Audited Accounts FY 2023-24", "category": "Finance"},
    {"id": "audited_fy2425", "name": "Audited Accounts FY 2024-25", "category": "Finance"},
    {"id": "net_worth_cert", "name": "Net Worth Certificate (CA Signed)", "category": "Finance"},
    {"id": "solvency_cert", "name": "Solvency Certificate", "category": "Finance"},
    {"id": "blacklisting_dec", "name": "Non-Blacklisting Declaration Template", "category": "Declaration"},
    {"id": "mii_dec", "name": "Make in India Declaration Template", "category": "Declaration"},
    {"id": "integrity_pact", "name": "Integrity Pact Template", "category": "Declaration"},
    {"id": "amc_gis_cc", "name": "Completion Certificate — AMC GIS", "category": "Experience"},
    {"id": "pcscl_po", "name": "Purchase Order — PCSCL Smart City", "category": "Experience"},
    {"id": "kvic_cc", "name": "Completion Certificate — KVIC Geo Portal", "category": "Experience"},
    {"id": "tcgl_cc", "name": "Completion Certificate — TCGL Tourism", "category": "Experience"},
    {"id": "vmc_cc", "name": "Completion Certificate — VMC GIS+ERP", "category": "Experience"},
]


@app.get("/vault")
async def get_vault():
    """Return vault document list with upload status."""
    docs = []
    for doc in VAULT_DOCS:
        # Check if file exists
        existing = list(VAULT_DIR.glob(f"{doc['id']}.*"))
        docs.append({
            **doc,
            "uploaded": len(existing) > 0,
            "filename": existing[0].name if existing else None,
            "size_kb": round(existing[0].stat().st_size / 1024) if existing else 0,
        })
    return {"documents": docs}


@app.post("/vault/upload/{doc_id}")
async def upload_vault_doc(doc_id: str, file: UploadFile = File(...)):
    """Upload a document to the vault."""
    # Validate doc_id
    valid_ids = {d["id"] for d in VAULT_DOCS}
    if doc_id not in valid_ids:
        raise HTTPException(400, f"Unknown document ID: {doc_id}")

    # Remove old files for this doc_id
    for old in VAULT_DIR.glob(f"{doc_id}.*"):
        old.unlink(missing_ok=True)

    # Save new file
    ext = Path(file.filename or "file.pdf").suffix.lower() or ".pdf"
    dest = VAULT_DIR / f"{doc_id}{ext}"
    dest.write_bytes(await file.read())

    # Sync to Drive
    try:
        save_to_drive(dest)
    except Exception:
        pass

    return {
        "status": "uploaded",
        "doc_id": doc_id,
        "filename": dest.name,
        "size_kb": round(dest.stat().st_size / 1024),
    }


@app.delete("/vault/{doc_id}")
async def delete_vault_doc(doc_id: str):
    """Delete a vault document."""
    deleted = []
    for f in VAULT_DIR.glob(f"{doc_id}.*"):
        deleted.append(f.name)
        f.unlink(missing_ok=True)
    return {"deleted": deleted}


@app.get("/vault/download/{doc_id}")
async def download_vault_doc(doc_id: str):
    """Download a vault document."""
    files = list(VAULT_DIR.glob(f"{doc_id}.*"))
    if not files:
        raise HTTPException(404, f"Document {doc_id} not uploaded yet")
    f = files[0]
    media_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
    }
    media = media_map.get(f.suffix.lower(), "application/octet-stream")
    return FileResponse(str(f), filename=f.name, media_type=media)


# ── T247 AUTO-DOWNLOAD (HTTP approach — works on Render) ───────
@app.post("/tender/{t247_id}/auto-download")
async def auto_download_tender_http(t247_id: str, data: dict = Body(default={})):
    """
    Try to auto-download tender documents from T247 using HTTP session.
    Requires T247 credentials saved in config.
    """
    config = load_config()
    username = config.get("t247_username", "")
    password = config.get("t247_password", "")

    if not username or not password:
        return {
            "status": "no_credentials",
            "message": "T247 credentials not saved. Go to Settings → T247 Credentials and save your login.",
        }

    try:
        from downloader import download_via_http
        result = await download_via_http(t247_id, username, password)
        return result
    except ImportError:
        # Fallback — just return the T247 URL for manual download
        t247_url = f"https://www.tender247.com/keyword=~&category=~&country=IN#TenderDetail/{t247_id}"
        return {
            "status": "manual",
            "message": "Auto-download not available. Click the link to open T247.",
            "t247_url": t247_url,
            "t247_id": t247_id,
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e),
        }


# ── TEST T247 CONNECTION ────────────────────────────────────────
@app.get("/test-t247")
async def test_t247():
    """Test T247 connection with saved credentials."""
    config = load_config()
    username = config.get("t247_username", "")
    if not username:
        return {"status": "error", "message": "No T247 credentials saved"}
    return {
        "status": "ok",
        "username": username,
        "message": "Credentials saved. Auto-download requires T247 session.",
    }


# ── TENDER WORKSPACE (all-in-one tender data) ──────────────────
@app.get("/workspace/{t247_id}")
async def get_workspace(t247_id: str):
    """
    Return complete tender workspace data:
    tender details + analysis + checklist + prebid queries
    """
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, f"Tender {t247_id} not found")

    # Get checklist
    checklist = tender.get("doc_checklist", [])

    # Get prebid queries
    prebid = tender.get("prebid_queries", [])

    # Find report file
    reports = []
    for f in OUTPUT_DIR.glob(f"*{t247_id}*.docx"):
        reports.append({
            "filename": f.name,
            "size_kb": round(f.stat().st_size / 1024),
            "created": datetime.fromtimestamp(f.stat().st_mtime).strftime("%d-%b-%Y %H:%M"),
        })

    return {
        "tender": tender,
        "checklist": checklist,
        "prebid_queries": prebid,
        "reports": reports,
        "vault_docs": []  # Will be populated by separate vault call
    }


# ── GENERATE SUBMISSION DOCUMENTS ─────────────────────────────
@app.post("/generate-docs/{t247_id}")
async def generate_submission_docs(t247_id: str):
    """Generate all submission documents for a tender."""
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found. Analyse it first.")

    try:
        from submission_generator import SubmissionGenerator
        gen = SubmissionGenerator()
        result = gen.generate_all(tender)
        return {
            "status": "success",
            "files": result.get("files", []),
            "zip_url": result.get("zip_url"),
            "count": len(result.get("files", [])),
        }
    except ImportError:
        # submission_generator.py not uploaded yet — return what we have
        # Fall back to doc_generator if available
        try:
            from doc_generator import BidDocGenerator
            gen = BidDocGenerator()
            import re
            safe_no = re.sub(r'[^\w\-]', '_', tender.get("tender_no", t247_id))[:50]
            filename = f"BidNoBid_{safe_no}.docx"
            gen.generate(tender, str(OUTPUT_DIR / filename))
            return {
                "status": "partial",
                "files": [{"name": "Bid-No-Bid Report", "filename": filename}],
                "message": "submission_generator.py not found — generated bid/no-bid report only"
            }
        except Exception as e2:
            raise HTTPException(500, f"Document generation failed: {str(e2)}")
    except Exception as e:
        raise HTTPException(500, f"Generation error: {str(e)}")


# ── HEAD HANDLER (fixes Render health check 405) ───────────────
@app.head("/")
async def head_root():
    return {}


@app.head("/health")
async def head_health():
    return {}


# ── AI QUOTA STATUS ────────────────────────────────────────────
@app.get("/api-quota-status")
async def api_quota_status():
    """Return API key configuration status."""
    from ai_analyzer import get_all_api_keys, load_config
    keys = get_all_api_keys()
    cfg = load_config()
    groq_key = cfg.get("groq_api_key", "")
    return {
        "status": "ok" if keys else "no_keys",
        "gemini_keys": len(keys),
        "groq_configured": bool(groq_key),
        "total_keys": len(keys) + (1 if groq_key else 0),
        "message": f"{len(keys)} Gemini key(s) + {'Groq backup' if groq_key else 'no Groq backup'}",
    }


# ── TENDER REANALYSE FROM DRIVE ────────────────────────────────
@app.post("/tender/{t247_id}/reanalyse")
async def reanalyse_tender(t247_id: str):
    """Re-run AI analysis on previously uploaded tender (from saved text in DB)."""
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    saved_text = tender.get("raw_text", "")
    if not saved_text:
        raise HTTPException(400, "No saved document text. Upload files and analyse again.")
    # Re-run analysis
    from ai_analyzer import analyze_with_gemini, merge_results
    prebid_passed = tender.get("prebid_passed", False)
    ai_result = analyze_with_gemini(saved_text, prebid_passed)
    if "error" not in ai_result:
        updated = merge_results(tender, ai_result, prebid_passed)
        updated["bid_no_bid_done"] = True
        updated["analysed_at"] = datetime.now().isoformat()
        save_tender(t247_id, updated)
        return {"status": "success", "verdict": updated.get("verdict", "REVIEW")}
    return {"status": "error", "error": ai_result.get("error", "Unknown error")}


# ── GENERATE PREBID LETTER ─────────────────────────────────────
@app.post("/tender/{t247_id}/generate-prebid-letter")
async def gen_prebid_letter(t247_id: str):
    """Generate a pre-bid query letter as Word doc."""
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    queries = tender.get("prebid_queries", [])
    if not queries:
        queries = tender.get("prebid_queries_list", [])
    if not queries:
        raise HTTPException(400, "No pre-bid queries found. Analyse tender first.")

    try:
        from doc_generator import BidDocGenerator
        gen = BidDocGenerator()
        import re
        safe_no = re.sub(r'[^\w\-]', '_', tender.get("tender_no", t247_id))[:40]
        filename = f"PreBid_{safe_no}.docx"
        gen.generate_prebid_letter(tender, queries, str(OUTPUT_DIR / filename))
        return {
            "status": "success",
            "filename": filename,
            "query_count": len(queries),
            "download_url": f"/download/{filename}",
        }
    except AttributeError:
        # doc_generator doesn't have generate_prebid_letter method
        # Create a simple one
        from docx import Document
        import re
        doc = Document()
        profile_company = "Nascent Info Technologies Pvt. Ltd."
        doc.add_heading("Pre-Bid Queries", 0)
        doc.add_paragraph(f"Tender: {tender.get('tender_name', 'N/A')}")
        doc.add_paragraph(f"Tender No: {tender.get('tender_no', 'N/A')}")
        doc.add_paragraph(f"Organization: {tender.get('org_name', 'N/A')}")
        doc.add_paragraph("")
        doc.add_paragraph("Respected Sir/Madam,")
        doc.add_paragraph(
            f"We, {profile_company}, wish to participate in the above tender. "
            f"We request clarifications on the following points:"
        )
        doc.add_paragraph("")
        for i, q in enumerate(queries, 1):
            clause = q.get("clause", "") if isinstance(q, dict) else ""
            query_text = q.get("query", q) if isinstance(q, dict) else str(q)
            doc.add_paragraph(f"Q{i}. {clause}: {query_text}")
        doc.add_paragraph("")
        doc.add_paragraph("Thanking you,")
        doc.add_paragraph(profile_company)
        safe_no = re.sub(r'[^\w\-]', '_', tender.get("tender_no", t247_id))[:40]
        filename = f"PreBid_{safe_no}.docx"
        doc.save(str(OUTPUT_DIR / filename))
        return {
            "status": "success",
            "filename": filename,
            "query_count": len(queries),
            "download_url": f"/download/{filename}",
        }
    except Exception as e:
        raise HTTPException(500, f"Letter generation failed: {str(e)}")


# ── CHECKLIST ITEM TOGGLE ──────────────────────────────────────
@app.post("/checklist/{t247_id}/item")
async def toggle_checklist_item(t247_id: str, data: dict = Body(...)):
    """Toggle a checklist item done/not done."""
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    items = t.get("doc_checklist", [])
    item_id = data.get("id", "")
    done = data.get("done", False)
    for item in items:
        if item.get("id") == item_id or item.get("label") == item_id:
            item["done"] = done
            break
    t["doc_checklist"] = items
    # Recalculate completion %
    if items:
        t["checklist_pct"] = round(100 * sum(1 for i in items if i.get("done")) / len(items))
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved", "item_id": item_id, "done": done}


# ── SAVE T247 CREDENTIALS TO CONFIG ───────────────────────────
# (The /config POST endpoint in main.py already handles this,
#  but make sure these fields are accepted)
# The main.py /config POST needs to handle t247_username and t247_password:
# Add this inside the existing update_config_route function:
#   if "t247_username" in data:
#       config["t247_username"] = data["t247_username"]
#   if "t247_password" in data and data["t247_password"]:
#       config["t247_password"] = data["t247_password"]
