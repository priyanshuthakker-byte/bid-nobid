"""
Letterhead Manager
- Upload Nascent letterhead (Word file) — stored as template
- Apply letterhead to any generated document
- If no letterhead uploaded, use designed placeholder
- Tracks which documents need letterhead vs plain
"""

import json, io, shutil
from pathlib import Path
from typing import Optional
from datetime import datetime

BASE_DIR = Path(__file__).parent
LETTERHEAD_DIR = BASE_DIR / "letterhead"
LETTERHEAD_DIR.mkdir(exist_ok=True, parents=True)
LETTERHEAD_FILE = LETTERHEAD_DIR / "nascent_letterhead.docx"
LETTERHEAD_META = LETTERHEAD_DIR / "letterhead_meta.json"


def save_letterhead(file_bytes: bytes, filename: str) -> dict:
    """Save uploaded letterhead file."""
    LETTERHEAD_DIR.mkdir(exist_ok=True, parents=True)
    LETTERHEAD_FILE.write_bytes(file_bytes)
    meta = {
        "filename": filename,
        "uploaded_at": datetime.now().isoformat(),
        "size_kb": round(len(file_bytes) / 1024, 1),
    }
    LETTERHEAD_META.write_text(json.dumps(meta, indent=2))
    return meta


def get_letterhead_meta() -> Optional[dict]:
    """Get metadata about stored letterhead."""
    if LETTERHEAD_META.exists():
        try:
            return json.loads(LETTERHEAD_META.read_text())
        except Exception:
            pass
    return None


def has_letterhead() -> bool:
    return LETTERHEAD_FILE.exists() and LETTERHEAD_FILE.stat().st_size > 1000


def apply_letterhead_to_doc(content_doc_path: str, output_path: str) -> bool:
    """
    Prepend Nascent letterhead to a content document.
    Strategy:
    - If letterhead is available: copy letterhead, append content doc's body to it
    - If not available: add a styled placeholder header to the document
    Returns True if letterhead was applied, False if placeholder used.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        content_doc = Document(content_doc_path)

        if has_letterhead():
            # Load letterhead template
            lh_doc = Document(str(LETTERHEAD_FILE))

            # Append content doc's body elements to letterhead doc
            # Add a page break after letterhead header if it doesn't have one
            lh_body = lh_doc.element.body
            content_body = content_doc.element.body

            # Copy all content paragraphs/tables into letterhead doc
            for element in content_body:
                # Skip the final sectPr (section properties)
                if element.tag.endswith('}sectPr'):
                    continue
                lh_body.append(copy.deepcopy(element))

            lh_doc.save(output_path)
            return True

        else:
            # Add placeholder letterhead header to content doc
            _add_placeholder_letterhead(content_doc)
            content_doc.save(output_path)
            return False

    except Exception as e:
        print(f"Letterhead apply failed: {e}")
        # Just copy the file as-is
        shutil.copy2(content_doc_path, output_path)
        return False


def _add_placeholder_letterhead(doc):
    """Add a designed placeholder letterhead at the top of the document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Table as DocxTable

    # Insert at beginning — we need to add paragraphs before existing content
    # We'll use the approach of inserting XML before the first element
    body = doc.element.body
    first_element = body[0] if len(body) > 0 else None

    # Create header table
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsmap

    header_xml = '''<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
        <w:tblPr>
            <w:tblW w:w="9360" w:type="dxa"/>
            <w:tblBorders>
                <w:bottom w:val="single" w:sz="6" w:color="1F3864"/>
            </w:tblBorders>
        </w:tblPr>
        <w:tblGrid><w:gridCol w:w="4680"/><w:gridCol w:w="4680"/></w:tblGrid>
        <w:tr>
            <w:tc>
                <w:tcPr><w:tcW w:w="4680" w:type="dxa"/></w:tcPr>
                <w:p>
                    <w:pPr><w:jc w:val="left"/></w:pPr>
                    <w:r><w:rPr><w:b/><w:color w:val="1F3864"/><w:sz w:val="26"/></w:rPr>
                        <w:t>Nascent Info Technologies Pvt. Ltd.</w:t>
                    </w:r>
                </w:p>
                <w:p>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>A-805, Shapath IV, SG Highway, Prahlad Nagar</w:t></w:r>
                </w:p>
                <w:p>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>Ahmedabad - 380015, Gujarat, India</w:t></w:r>
                </w:p>
                <w:p>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>Ph: +91-79-40200400</w:t></w:r>
                </w:p>
            </w:tc>
            <w:tc>
                <w:tcPr><w:tcW w:w="4680" w:type="dxa"/></w:tcPr>
                <w:p>
                    <w:pPr><w:jc w:val="right"/></w:pPr>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>Email: nascent.tender@nascentinfo.com</w:t></w:r>
                </w:p>
                <w:p>
                    <w:pPr><w:jc w:val="right"/></w:pPr>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>Web: www.nascentinfo.com</w:t></w:r>
                </w:p>
                <w:p>
                    <w:pPr><w:jc w:val="right"/></w:pPr>
                    <w:r><w:rPr><w:color w:val="444444"/><w:sz w:val="18"/></w:rPr>
                        <w:t>PAN: AACCN3670J | GSTIN: 24AACCN3670J1ZG</w:t></w:r>
                </w:p>
                <w:p>
                    <w:pPr><w:jc w:val="right"/></w:pPr>
                    <w:r><w:rPr><w:color w:val="1F3864"/><w:sz w:val="16"/></w:rPr>
                        <w:t>CMMI L3 | ISO 9001 | ISO 27001 | ISO 20000 | MSME</w:t></w:r>
                </w:p>
            </w:tc>
        </w:tr>
    </w:tbl>'''

    try:
        header_tbl = parse_xml(header_xml)
        if first_element is not None:
            body.insert(0, header_tbl)
        else:
            body.append(header_tbl)
        # Add spacing paragraph
        from lxml import etree
        spacer = parse_xml('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>')
        body.insert(1, spacer)
    except Exception as e:
        print(f"Placeholder header insert failed: {e}")


def create_letterhead_doc(content_paragraphs: list, output_path: str,
                          use_letterhead: bool = True) -> bool:
    """
    Create a new document with letterhead applied.
    content_paragraphs: list of (text, bold, size, align) tuples
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Setup page
    sec = doc.sections[0]
    from docx.shared import Cm
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    for text, bold, size, align in content_paragraphs:
        p = doc.add_paragraph()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(str(text))
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.bold = bold

    # Save to temp, then apply letterhead
    import tempfile
    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    if use_letterhead:
        result = apply_letterhead_to_doc(tmp, output_path)
    else:
        shutil.copy2(tmp, output_path)
        result = False

    Path(tmp).unlink(missing_ok=True)
    return result
