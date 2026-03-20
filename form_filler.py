"""
Form Filler
- AI reads RFP-specific forms (Form 1, Form 4, Annexure-H, etc.)
- Extracts field names and what data goes where
- Fills Nascent data automatically
- Generates completed Word document
- For fields AI can't fill (blank boxes, signatures) — marks with [FILL: description]
"""

import json, re
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime, date

BASE_DIR = Path(__file__).parent
PROFILE_PATH = BASE_DIR / "nascent_profile.json"


def load_nascent_data() -> Dict:
    """Load all Nascent data that can be auto-filled into forms."""
    try:
        if PROFILE_PATH.exists():
            p = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
        else:
            p = {}
    except Exception:
        p = {}

    company = p.get("company", {})
    finance = p.get("finance", {})
    certs = p.get("certifications", {})
    employees = p.get("employees", {})

    # All fillable Nascent data — used for form filling
    return {
        # Company basics
        "company_name": company.get("name", "Nascent Info Technologies Pvt. Ltd."),
        "company_name_short": "Nascent Info Technologies Pvt. Ltd.",
        "cin": company.get("cin", "U72200GJ2006PTC048723"),
        "incorporation_date": company.get("incorporated", "23 June 2006"),
        "years_operation": str(company.get("years_in_operation", 19)),
        "address": company.get("address", "A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad - 380015, Gujarat"),
        "city": "Ahmedabad",
        "state": "Gujarat",
        "pin": "380015",
        "pan": company.get("pan", "AACCN3670J"),
        "gstin": company.get("gstin", "24AACCN3670J1ZG"),
        "email": company.get("email", "nascent.tender@nascentinfo.com"),
        "phone": company.get("phone", "+91-79-40200400"),
        "website": company.get("website", "www.nascentinfo.com"),
        "msme_udyam": company.get("msme_udyam", "UDYAM-GJ-01-0007420"),
        "signatory_name": company.get("signatory_name", "Hitesh Patel"),
        "signatory_designation": company.get("signatory_designation", "Chief Administrative Officer"),
        "prepared_by": company.get("prepared_by", "Parthav Thakkar"),
        "prepared_designation": company.get("prepared_designation", "Bid Executive"),

        # Financial data
        "turnover_fy2223": str(finance.get("fy2223_cr", 16.36)) + " Cr",
        "turnover_fy2324": str(finance.get("fy2324_cr", 16.36)) + " Cr",
        "turnover_fy2425": str(finance.get("fy2425_cr", 18.83)) + " Cr",
        "avg_turnover_3yr": str(finance.get("avg_turnover_last_3_fy", 17.18)) + " Cr",
        "net_worth": str(finance.get("net_worth_cr", 26.09)) + " Cr",
        "bank_name": finance.get("bank_name", "State Bank of India"),
        "bank_branch": finance.get("bank_branch", "SG Highway Branch, Ahmedabad"),

        # Certifications
        "cmmi_level": certs.get("cmmi_level", "V2.0 Level 3"),
        "cmmi_valid": certs.get("cmmi_valid", "19-Dec-2026"),
        "iso9001": "ISO 9001:2015",
        "iso27001": "ISO/IEC 27001:2022",
        "iso20000": "ISO/IEC 20000-1:2018",
        "iso_valid": certs.get("iso_valid", "08-Sep-2028"),

        # Employees
        "total_employees": str(employees.get("total_confirmed", 67)),
        "gis_staff": str(employees.get("gis_staff", 11)),
        "it_dev_staff": str(employees.get("it_dev_staff", 21)),

        # Dates
        "today_date": date.today().strftime("%d %B %Y"),
        "today_date_short": date.today().strftime("%d-%m-%Y"),
        "today_year": str(date.today().year),
    }


def build_form_fill_prompt(form_text: str, form_name: str, nascent_data: Dict,
                            tender: Dict) -> str:
    """Build prompt for AI to fill a specific form."""
    return f"""You are filling in a Government tender form on behalf of Nascent Info Technologies Pvt. Ltd.

FORM NAME: {form_name}

NASCENT DATA (use this to fill form fields):
{json.dumps(nascent_data, indent=2)}

TENDER DETAILS:
- Tender No: {tender.get('tender_no', '')}
- Organization: {tender.get('org_name', '')}
- Tender Name: {tender.get('tender_name', '')}
- Date: {date.today().strftime('%d %B %Y')}

FORM CONTENT:
{form_text[:8000]}

Instructions:
1. Identify every field/blank in this form
2. Fill each field with appropriate Nascent data
3. For fields you can fill with certainty — fill them
4. For fields requiring original documents (original work order, original certificate) — write [ATTACH: document name]
5. For fields requiring manual input (specific project name for this bid, specific amount) — write [FILL: description of what to enter]
6. For signature fields — write [SIGN: who signs this]
7. For date fields — use today's date unless context suggests otherwise

Return ONLY valid JSON. No markdown.

{{
  "form_name": "{form_name}",
  "form_type": "cover_letter / experience / turnover / employee / declaration / annexure / other",
  "fields": [
    {{
      "field_label": "Name of the bidder",
      "field_value": "Nascent Info Technologies Pvt. Ltd.",
      "auto_filled": true,
      "fill_note": ""
    }},
    {{
      "field_label": "Project value of qualifying project",
      "field_value": "[FILL: Enter value of qualifying project from your portfolio]",
      "auto_filled": false,
      "fill_note": "Select from: PCSCL Rs.61.19 Cr, VMC Rs.20.5 Cr, JuMC Rs.9.78 Cr, etc."
    }}
  ],
  "filled_form_text": "The complete form text with all fields filled in, ready to copy-paste",
  "manual_items_count": 0,
  "auto_filled_count": 0,
  "notes": "Any important notes about this form"
}}"""


def fill_form_with_ai(form_text: str, form_name: str, tender: Dict) -> Dict:
    """Use AI to fill a form with Nascent data."""
    from ai_analyzer import get_all_api_keys, call_gemini, call_groq, get_groq_key, clean_json

    nascent_data = load_nascent_data()
    prompt = build_form_fill_prompt(form_text, form_name, nascent_data, tender)

    all_keys = get_all_api_keys()
    response_text = ""

    for key in all_keys:
        try:
            response_text = call_gemini(prompt, key)
            result = clean_json(response_text)
            return result
        except json.JSONDecodeError:
            try:
                import re as _re
                m = _re.search(r'\{.*\}', response_text, _re.DOTALL)
                if m:
                    return json.loads(m.group(0))
            except Exception:
                pass
        except Exception as e:
            err = str(e)
            if "quota" in err.lower() or "429" in err:
                continue
            break

    # Groq fallback
    groq_key = get_groq_key()
    if groq_key:
        try:
            response_text = call_groq(prompt, groq_key)
            return clean_json(response_text)
        except Exception:
            pass

    return {"error": "AI unavailable. Fill form manually using Nascent data."}


def generate_filled_form_doc(form_result: Dict, output_path: str,
                              tender: Dict, use_letterhead: bool = True):
    """Generate a Word document with the filled form."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    form_name = form_result.get("form_name", "Form")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(form_name.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.bold = True

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Tender: {tender.get('tender_no', '')} | {tender.get('org_name', '')}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # If we have filled form text, use it
    filled_text = form_result.get("filled_form_text", "")
    if filled_text:
        for line in filled_text.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            # Highlight unfilled fields
            if "[FILL:" in line or "[ATTACH:" in line or "[SIGN:" in line:
                r = p.add_run(line)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Red for manual items
            else:
                r = p.add_run(line)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
    else:
        # Fallback: show field table
        fields = form_result.get("fields", [])
        if fields:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for cell, hdr in zip(table.rows[0].cells,
                                  ["Field", "Value", "Note"]):
                cell.text = hdr
                cell.paragraphs[0].runs[0].bold = True
            for field in fields:
                row = table.add_row()
                row.cells[0].text = field.get("field_label", "")
                val = field.get("field_value", "")
                row.cells[1].text = val
                row.cells[2].text = field.get("fill_note", "")
                if not field.get("auto_filled"):
                    for cell in row.cells:
                        for run in cell.paragraphs[0].runs:
                            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Summary of manual items
    manual_count = form_result.get("manual_items_count", 0)
    if manual_count > 0:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(f"Note: {manual_count} field(s) need manual input (shown in red above)")
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r.font.bold = True

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()
    for line in [
        "Authorised Signatory",
        form_result.get("signatory", "Hitesh Patel"),
        "Chief Administrative Officer",
        "Nascent Info Technologies Pvt. Ltd.",
        f"Date: {date.today().strftime('%d %B %Y')}",
        "Place: Ahmedabad",
        "",
        "Company Seal:",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Calibri"
        r.font.size = Pt(11)

    # Apply letterhead
    import tempfile
    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    if use_letterhead:
        from letterhead_manager import apply_letterhead_to_doc
        apply_letterhead_to_doc(tmp, output_path)
    else:
        import shutil
        shutil.copy2(tmp, output_path)

    Path(tmp).unlink(missing_ok=True)


def extract_forms_from_rfp(rfp_text: str, tender: Dict) -> List[Dict]:
    """
    AI reads the RFP and extracts list of all forms/annexures required.
    Returns list of {form_name, form_ref, required_section, description}.
    """
    from ai_analyzer import get_all_api_keys, call_gemini, clean_json

    prompt = f"""Read this tender document and extract ALL forms, annexures, and declarations required from bidders.

TENDER: {tender.get('tender_no', '')} | {tender.get('org_name', '')}

DOCUMENT TEXT:
{rfp_text[:12000]}

Return ONLY valid JSON. List every form/annexure/declaration asked from bidder.

{{
  "forms": [
    {{
      "form_ref": "Form 1",
      "form_name": "Cover Letter / Bid Submission Letter",
      "required_in": "Technical Bid",
      "letterhead_required": true,
      "stamp_paper": null,
      "notary_required": false,
      "description": "what this form is for",
      "nascent_can_auto_fill": true,
      "manual_inputs_needed": ["project reference number", "specific date"],
      "clause_ref": "Annexure-III Form 1"
    }}
  ],
  "total_forms": 0,
  "stamp_papers_needed": [
    {{"purpose": "Non-blacklisting declaration", "amount": "Rs. 100", "clause": "Form 7"}}
  ]
}}"""

    all_keys = get_all_api_keys()
    for key in all_keys:
        try:
            response_text = call_gemini(prompt, key)
            return clean_json(response_text)
        except Exception as e:
            err = str(e)
            if "quota" in err.lower() or "429" in err:
                continue
            break

    return {"forms": [], "error": "AI unavailable for form extraction"}
