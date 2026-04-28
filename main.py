"""
Bid/No-Bid Automation v6.2 — Production-ready
FIXES APPLIED:
- CRITICAL: Constants moved above startup event (NameError on Render fixed)
- CRITICAL: CORS locked to ALLOWED_ORIGIN env var (was wildcard *)
- CRITICAL: Path traversal fix in /download endpoint
- CRITICAL: Admin token guard on /config, /upload-db, /sync-drive
- CRITICAL: Config keys read from env vars first, disk write skipped on Render
- WARN: Deprecated @app.on_event replaced with lifespan
- WARN: Threading lock added to DB read/write
- WARN: File size limit on uploads (50MB)
- WARN: Temp dir cleanup on startup
"""

import zipfile, tempfile, shutil, json, re, os, base64, hashlib
import threading
import asyncio
import uuid
from pathlib import Path
from datetime import datetime, date, timedelta
from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Request, BackgroundTasks, Depends
from typing import List
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy import text
from extractor import TenderExtractor, read_document
from doc_generator import BidDocGenerator
from nascent_checker import NascentChecker
from ai_analyzer import analyze_with_gemini, merge_results, load_config, save_config, get_all_api_keys, call_gemini
try:
    from ai_analyzer import analyze_with_gemini_parallel
    PARALLEL_ANALYST_AVAILABLE = True
except ImportError:
    PARALLEL_ANALYST_AVAILABLE = False
    analyze_with_gemini_parallel = None
try:
    from core.api_pool import get_pool, get_slots, refresh_pool
    API_POOL_AVAILABLE = True
except ImportError:
    API_POOL_AVAILABLE = False
try:
    import doc_editor
    DOC_EDITOR_AVAILABLE = True
except ImportError:
    DOC_EDITOR_AVAILABLE = False
    doc_editor = None
from excel_processor import process_excel
from prebid_generator import generate_prebid_queries
from chatbot import process_message, load_history
from gdrive_sync import init_drive, save_to_drive, load_from_drive, is_available as drive_available
from tracker import (get_deadline_alerts, get_pipeline_stats,
                     get_win_loss_stats, generate_doc_checklist,
                     PIPELINE_STAGES, STAGE_COLORS)
from core.config import settings
from core.database import engine, Base, get_db_session
from core.models import User, WorkItem, TenderRecord, TenderSource, IngestedTender, ClauseEvidence
from core.auth import hash_password, verify_password, create_access_token, decode_token
from core.worker import work_queue
from core.ingestion import registry as ingestion_registry, JsonApiSource, CpppFeedSource, StatePortalTableSource
try:
    from submission_generator import generate_submission_package
    SUBMISSION_GEN_AVAILABLE = True
except Exception:
    SUBMISSION_GEN_AVAILABLE = False
    def generate_submission_package(tender, output_dir):
        return {"error": "submission_generator not available"}
try:
    from pdf_merger import merge_submission_package, get_doc_order_preview
    PDF_MERGE_AVAILABLE = True
except Exception:
    PDF_MERGE_AVAILABLE = False
    def merge_submission_package(*a, **k):
        return {"status": "error", "errors": ["pdf_merger not available"]}
    def get_doc_order_preview(*a, **k):
        return []

try:
    from boq_engine import extract_boq_from_scope, calculate_boq_totals, get_boq_constants
    BOQ_AVAILABLE = True
except ImportError:
    BOQ_AVAILABLE = False

try:
    from technical_proposal_generator import generate_technical_proposal as _gen_tech_proposal
    TECH_PROPOSAL_AVAILABLE = True
except Exception:
    TECH_PROPOSAL_AVAILABLE = False
    _gen_tech_proposal = None
    def extract_boq_from_scope(t): return []
    def calculate_boq_totals(i, m=15, g=18):
        return {"items": i, "base_total": 0, "margin_amount": 0, "subtotal": 0, "gst_amount": 0, "grand_total": 0}
    def get_boq_constants():
        return {"categories": ["Manpower","Software / Licenses","Hardware","Cloud / Hosting","Training","AMC / Support","Miscellaneous"],
                "unit_types": ["Months","Nos","Lumpsum","Per Year","Sq.Km","Per Day"],
                "manpower_roles": ["Project Manager","GIS Developer","Software Developer","QA Engineer"]}

# ── FIX 1: Constants defined HERE — above everything that uses them ──────────
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "data"
TEMP_DIR = BASE_DIR / "temp"
OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
TEMP_DIR.mkdir(exist_ok=True, parents=True)
DB_FILE = OUTPUT_DIR / "tenders_db.json"
DRAFTS_DIR = OUTPUT_DIR / "drafts"
LATEST_EXCEL_FILE = OUTPUT_DIR / "latest_tenders_import.xlsx"
DRAFTS_DIR.mkdir(exist_ok=True, parents=True)

# ── FIX 2: Threading lock for safe concurrent DB access ─────────────────────
_db_lock = threading.RLock()
db_lock = _db_lock
import time as _time
_last_drive_restore = 0.0  # rate-limit Drive calls in load_db to once per minute

# ── Background Job Store — file-based (survives OOM restarts) ──────────────
JOBS_DIR = OUTPUT_DIR / "jobs"
JOBS_DIR.mkdir(exist_ok=True, parents=True)
_jobs_lock = threading.Lock()
_t247_sync_lock = threading.Lock()
_t247_sync_stop = threading.Event()
_t247_sync_state = {
    "last_run_at": "",
    "last_status": "idle",
    "last_message": "",
    "last_total": 0,
    "last_added": 0,
    "last_updated": 0,
}
_digest_lock = threading.Lock()
_digest_state = {"last_generated_at": "", "last_generated_date": "", "status": "idle", "error": ""}

def _job_file(job_id: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_\-]", "", job_id)
    return JOBS_DIR / f"{safe}.json"

def _set_job(job_id: str, **kwargs):
    with _jobs_lock:
        jf = _job_file(job_id)
        try:
            existing = json.loads(jf.read_text()) if jf.exists() else {}
        except Exception:
            existing = {}
        # Merge dicts instead of overwriting for streaming fields
        for merge_key in ("segments", "seg_log"):
            if merge_key in kwargs:
                existing[merge_key] = {**existing.get(merge_key, {}), **kwargs.pop(merge_key)}
        existing.update(kwargs)
        # Don't persist huge base64 doc inline — store separately
        doc_b64 = existing.pop("doc_b64", None)
        try:
            jf.write_text(json.dumps(existing))
        except Exception:
            pass
        if doc_b64:
            try:
                (JOBS_DIR / f"{re.sub(r'[^a-zA-Z0-9_\\-]', '', job_id)}.b64").write_text(doc_b64)
            except Exception:
                pass

def _get_job(job_id: str) -> dict:
    with _jobs_lock:
        jf = _job_file(job_id)
        if not jf.exists():
            return {}
        try:
            data = json.loads(jf.read_text())
            # Re-attach b64 if result is being fetched and file exists
            b64f = JOBS_DIR / f"{re.sub(r'[^a-zA-Z0-9_\\-]', '', job_id)}.b64"
            if data.get("status") == "done" and b64f.exists():
                if data.get("result"):
                    data["result"]["doc_b64"] = b64f.read_text()
            return data
        except Exception:
            return {}



# ── FIX 3: Admin token for sensitive endpoints ───────────────────────────────
ADMIN_TOKEN = os.environ.get("ADMIN_TOKEN", "")

def check_admin(request: Request):
    """Raises 403 if ADMIN_TOKEN env var is set and request doesn't provide it."""
    if not ADMIN_TOKEN:
        return  # No token configured → open (backward compat for local dev)
    token = request.headers.get("X-Admin-Token", "") or request.query_params.get("token", "")
    if token != ADMIN_TOKEN:
        raise HTTPException(403, "Admin token required. Set X-Admin-Token header.")

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB
auth_scheme = HTTPBearer(auto_error=False)


def _current_user(credentials: HTTPAuthorizationCredentials | None) -> dict:
    if not credentials:
        return {}
    payload = decode_token(credentials.credentials)
    return payload or {}


def _require_role(credentials: HTTPAuthorizationCredentials | None, allowed_roles: set[str]) -> dict:
    user = _current_user(credentials)
    if not user:
        raise HTTPException(401, "Authentication required")
    if user.get("role") not in allowed_roles:
        raise HTTPException(403, "Insufficient role")
    return user


def _handle_tender_scoring(payload: dict) -> dict:
    """Initial worker handler; can evolve into full ML scoring."""
    risk_score = float(payload.get("risk_score", 0) or 0)
    margin_score = float(payload.get("margin_score", 0) or 0)
    fit_score = float(payload.get("fit_score", 0) or 0)
    win_probability = max(0.0, min(100.0, (margin_score * 0.35 + fit_score * 0.45 + (100 - risk_score) * 0.20)))
    return {
        "win_probability": round(win_probability, 2),
        "recommended": "BID" if win_probability >= 65 else "CONDITIONAL" if win_probability >= 45 else "NO-BID",
    }


def _extract_clause_candidates(text_blob: str) -> list[dict]:
    text_blob = text_blob or ""
    patterns = {
        "emd": r"(?:emd|earnest money deposit)[^.\n]{0,140}",
        "turnover": r"(?:turnover|annual turnover)[^.\n]{0,180}",
        "experience": r"(?:experience|work order|completion certificate)[^.\n]{0,200}",
        "security": r"(?:performance security|pbg|bank guarantee)[^.\n]{0,180}",
        "timeline": r"(?:timeline|completion period|contract period|delivery period)[^.\n]{0,180}",
    }
    found = []
    for ctype, pattern in patterns.items():
        for m in re.finditer(pattern, text_blob, flags=re.IGNORECASE):
            snippet = m.group(0).strip()
            if len(snippet) < 12:
                continue
            found.append({
                "clause_type": ctype,
                "clause_text": snippet[:500],
                "evidence_text": snippet[:500],
                "confidence": 0.55,
            })
            if len(found) >= 100:
                return found
    return found


def _handle_ingestion_sync(payload: dict) -> dict:
    if not settings.database_enabled:
        raise RuntimeError("DATABASE_URL is required")
    source_name = str(payload.get("source_name", "")).strip()
    endpoint = str(payload.get("endpoint", "")).strip()
    source_type = str(payload.get("source_type", "json_api")).strip().lower()
    if not source_name:
        raise RuntimeError("source_name is required")
    if source_type == "json_api":
        source = JsonApiSource(endpoint=endpoint)
    elif source_type == "cppp_feed":
        source = CpppFeedSource(endpoint=endpoint)
    elif source_type == "state_portal_table":
        source = StatePortalTableSource(endpoint=endpoint)
    else:
        raise RuntimeError(f"unsupported source_type: {source_type}")
    items = source.fetch()
    inserted = 0
    with get_db_session() as db:
        for row in items[:500]:
            ext_id = str(row.get("id") or row.get("external_id") or row.get("tender_id") or "")
            if not ext_id:
                ext_id = f"{source_name}-{uuid.uuid4().hex[:12]}"
            exists = (
                db.query(IngestedTender)
                .filter(IngestedTender.source_name == source_name, IngestedTender.external_id == ext_id)
                .first()
            )
            if exists:
                continue
            title = str(row.get("title") or row.get("brief") or row.get("tender_name") or "")
            org = str(row.get("org_name") or row.get("organization") or "")
            deadline = str(row.get("deadline") or row.get("bid_submission_date") or "")
            ref = str(row.get("ref_no") or row.get("reference_no") or "")
            raw_text = str(row.get("raw_text") or row.get("description") or "")
            db.add(
                IngestedTender(
                    source_name=source_name,
                    external_id=ext_id,
                    title=title[:1000],
                    org_name=org[:255],
                    deadline=deadline[:120],
                    reference_no=ref[:255],
                    payload_json=json.dumps(row, default=str),
                    raw_text=raw_text[:20000],
                )
            )
            inserted += 1
    return {"inserted": inserted, "fetched": len(items), "source_name": source_name, "source_type": source_type}


def _handle_clause_index(payload: dict) -> dict:
    if not settings.database_enabled:
        raise RuntimeError("DATABASE_URL is required")
    t247_id = str(payload.get("t247_id", "")).strip()
    source_record_id = int(payload.get("source_record_id") or 0)
    if not t247_id and not source_record_id:
        raise RuntimeError("t247_id or source_record_id is required")

    if t247_id:
        tender = get_tender(t247_id)
        text_blob = " ".join([
            str(tender.get("eligibility", "")),
            str(tender.get("checklist", "")),
            str(tender.get("raw_text", "")),
            str(tender.get("brief", "")),
        ])
    else:
        with get_db_session() as db:
            row = db.query(IngestedTender).filter(IngestedTender.id == source_record_id).first()
            if not row:
                raise RuntimeError("source record not found")
            text_blob = " ".join([row.title or "", row.raw_text or "", row.payload_json or ""])

    clauses = _extract_clause_candidates(text_blob)
    with get_db_session() as db:
        if t247_id:
            db.query(ClauseEvidence).filter(ClauseEvidence.t247_id == t247_id).delete()
        elif source_record_id:
            db.query(ClauseEvidence).filter(ClauseEvidence.source_record_id == source_record_id).delete()
        for c in clauses:
            db.add(
                ClauseEvidence(
                    t247_id=t247_id,
                    source_record_id=source_record_id,
                    clause_type=c["clause_type"],
                    clause_text=c["clause_text"],
                    evidence_text=c["evidence_text"],
                    confidence=c["confidence"],
                )
            )
    return {"indexed": len(clauses), "t247_id": t247_id, "source_record_id": source_record_id}

# ── FIX 4: Lifespan replaces deprecated @app.on_event("startup") ────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Clean up old job files on startup
    import time as _startup_time
    try:
        if JOBS_DIR.exists():
            cutoff = _startup_time.time() - 7200  # 2 hours
            for jf in JOBS_DIR.glob("*.json"):
                try:
                    if jf.stat().st_mtime < cutoff:
                        jf.unlink(missing_ok=True)
                        b64f = jf.with_suffix(".b64")
                        if b64f.exists(): b64f.unlink(missing_ok=True)
                except Exception:
                    pass
    except Exception:
        pass
    import time
    print("Starting Bid/No-Bid System v6.2...")
    OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
    TEMP_DIR.mkdir(exist_ok=True, parents=True)

    # Cleanup stale temp dirs from previous crashed runs
    for stale in TEMP_DIR.glob("tender_*"):
        try:
            if stale.is_dir():
                age = datetime.now().timestamp() - stale.stat().st_mtime
                if age > 3600:
                    shutil.rmtree(stale, ignore_errors=True)
        except Exception:
            pass

    drive_ok = init_drive()
    print(f"Google Drive: {'Connected' if drive_ok else 'Not configured'}")
    if not drive_ok:
        print("WARNING: Google Drive not configured — tenders_db.json is ephemeral on Render!")

    if drive_ok:
        for attempt in range(3):
            try:
                if load_from_drive(DB_FILE):
                    print(f"Loaded {len(load_db().get('tenders', {}))} tenders from Drive")
                    break
                time.sleep(2)
            except Exception as e:
                print(f"Drive load attempt {attempt+1} failed: {e}")
                time.sleep(2)
    else:
        count = len(load_db().get("tenders", {})) if DB_FILE.exists() else 0
        print(f"Local DB: {count} tenders")

    print(f"BOQ: {'loaded' if BOQ_AVAILABLE else 'missing boq_engine.py'}")
    print(f"Admin token: {'set' if ADMIN_TOKEN else 'not set (admin routes open)'}")
    print("Server ready — v6.2")
    if settings.database_enabled and engine is not None:
        try:
            Base.metadata.create_all(bind=engine)
            work_queue.register("tender_scoring", _handle_tender_scoring)
            work_queue.register("ingestion_sync", _handle_ingestion_sync)
            work_queue.register("clause_index", _handle_clause_index)
            work_queue.start()
            with get_db_session() as db:
                if not db.query(TenderSource).first():
                    db.add(TenderSource(name="manual", source_type="manual", base_url="", is_active=True, config_json="{}"))
            print("Postgres connected and worker started")
        except Exception as e:
            print(f"Database bootstrap failed: {e}")
    else:
        print("DATABASE_URL not set — running in file-db mode")
    # Restore profile from Drive (always, to get latest edits)
    try:
        if drive_available():
            _prof_local = BASE_DIR / "nascent_profile.json"
            _prof_drive = OUTPUT_DIR / "nascent_profile.json"
            if load_from_drive(_prof_drive, filename="nascent_profile.json"):
                _prof_bytes = _prof_drive.read_bytes()
                if len(_prof_bytes) > 50:
                    _prof_local.write_bytes(_prof_bytes)
                    print("✅ Profile restored from Drive")
    except Exception:
        pass

    # Restore token usage log from Drive (persists across restarts)
    try:
        if drive_available():
            _tok_drive = OUTPUT_DIR / "token_usage.json"
            if load_from_drive(_tok_drive, filename="token_usage.json"):
                _tok_bytes = _tok_drive.read_bytes()
                if len(_tok_bytes) > 10:
                    _TOKEN_LOG_FILE.write_bytes(_tok_bytes)
                    print("✅ Token usage log restored from Drive")
    except Exception:
        pass

    # Background Drive sync every 5 minutes — safety net if a save_db call was skipped
    def _periodic_drive_sync():
        _time.sleep(90)  # wait for startup to settle
        while True:
            try:
                if drive_available() and DB_FILE.exists():
                    raw = DB_FILE.read_bytes()
                    if len(raw) > 10 and json.loads(raw).get("tenders"):
                        save_to_drive(DB_FILE)
            except Exception as _pe:
                print(f"⚠️ Periodic Drive sync error: {_pe}")
            _time.sleep(300)

    threading.Thread(target=_periodic_drive_sync, daemon=True, name="drive-sync").start()
    threading.Thread(target=_run_t247_sync_scheduler, daemon=True, name="t247-sync").start()
    threading.Thread(target=_run_daily_digest_scheduler, daemon=True, name="daily-digest").start()

    yield
    # Shutdown: nothing needed
    try:
        work_queue.stop()
    except Exception:
        pass
    try:
        _t247_sync_stop.set()
    except Exception:
        pass

app = FastAPI(title=settings.app_name, version=settings.app_version, lifespan=lifespan)

# ── FIX 5: CORS locked to env var, not wildcard ──────────────────────────────
_allowed_origin = os.environ.get("ALLOWED_ORIGIN", "*")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[_allowed_origin] if _allowed_origin != "*" else ["*"],
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# ── DB helpers ───────────────────────────────────────────────────────────────
def load_db() -> dict:
    global _last_drive_restore
    # Fast path: local disk has data
    with _db_lock:
        if DB_FILE.exists():
            try:
                raw = DB_FILE.read_text(encoding="utf-8")
                parsed = json.loads(raw)
                if parsed.get("tenders"):
                    return parsed
            except Exception:
                pass
    # Slow path: Drive restore — rate-limited to once per 60s to stop log spam
    now = _time.time()
    if drive_available() and (now - _last_drive_restore) > 60.0:
        _last_drive_restore = now
        try:
            DB_FILE.parent.mkdir(exist_ok=True, parents=True)
            if load_from_drive(DB_FILE):
                data = json.loads(DB_FILE.read_text(encoding="utf-8"))
                if data.get("tenders"):
                    print(f"🔄 load_db: restored {len(data['tenders'])} tenders from Drive")
                    return data
        except Exception:
            pass
    return {"tenders": {}}

def save_db(db: dict):
    with _db_lock:
        DB_FILE.parent.mkdir(exist_ok=True, parents=True)
        DB_FILE.write_text(json.dumps(db, indent=2, default=str), encoding="utf-8")
    try:
        ok = save_to_drive(DB_FILE)
        if not ok and drive_available():
            print(f"⚠️ Drive sync skipped/failed — {len(db.get('tenders', {}))} tenders on disk only")
    except Exception as _e:
        print(f"⚠️ Drive sync exception: {_e}")

def get_tender(t247_id: str) -> dict:
    return load_db()["tenders"].get(str(t247_id), {})

def save_tender(t247_id: str, data: dict):
    db = load_db()
    db["tenders"][str(t247_id)] = data
    save_db(db)

def _safe_doc_type(doc_type: str) -> str:
    return re.sub(r"[^\w\-]", "_", (doc_type or "analysis").strip().lower())[:40] or "analysis"

def _draft_path(t247_id: str, doc_type: str = "analysis") -> Path:
    safe_tid = re.sub(r"[^\w\-]", "_", str(t247_id))[:40] or "unknown"
    return DRAFTS_DIR / f"{safe_tid}_{_safe_doc_type(doc_type)}.md"

def _build_tender_draft(tender: dict, doc_type: str = "analysis") -> str:
    title = tender.get("tender_name") or tender.get("brief") or f"Tender {tender.get('t247_id', '')}"
    verdict = tender.get("verdict", "REVIEW")
    sections = [
        f"# {doc_type.replace('_', ' ').title()} Draft",
        "",
        "## Tender",
        f"- T247 ID: {tender.get('t247_id','')}",
        f"- Reference: {tender.get('ref_no','')}",
        f"- Name: {title}",
        f"- Organization: {tender.get('org_name','')}",
        f"- Deadline: {tender.get('deadline','')}",
        f"- Verdict: {verdict}",
        "",
        "## Executive Summary",
        str(tender.get("reason", "Summary pending.")),
        "",
        "## Key Eligibility",
        str(tender.get("eligibility", "Not available")),
        "",
        "## Compliance & Risk Notes",
        str(tender.get("compliance_notes", tender.get("risk_flags", "Add compliance comments here."))),
        "",
        "## Pre-Bid Queries",
    ]
    queries = tender.get("prebid_queries", []) or []
    if queries:
        for i, q in enumerate(queries, 1):
            qtxt = q.get("query") if isinstance(q, dict) else str(q)
            sections.append(f"{i}. {qtxt}")
    else:
        sections.append("1. No pending pre-bid queries.")
    sections.extend(["", "## Draft Document Body", "Write or edit final content here before download."])
    return "\n".join(sections)

def extract_all_zips(folder: Path):
    for nested in list(folder.rglob("*.zip")):
        try:
            out = nested.parent / (nested.stem + "_inner")
            out.mkdir(exist_ok=True)
            with zipfile.ZipFile(nested, "r") as zf:
                zf.extractall(out)
            extract_all_zips(out)
        except Exception:
            pass

def days_left(deadline_str: str) -> int:
    if not deadline_str:
        return 999
    for fmt in ["%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%d %b %Y"]:
        try:
            d = datetime.strptime(str(deadline_str).split()[0], fmt).date()
            return (d - date.today()).days
        except Exception:
            continue
    return 999

def prebid_passed(date_str: str) -> bool:
    return days_left(date_str) < 0

# ══════════════════════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/", response_class=HTMLResponse)
async def root():
    index = BASE_DIR / "index.html"
    if index.exists():
        return HTMLResponse(content=index.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>Bid/No-Bid v6.2</h1>")

@app.get("/healthz")
async def healthz():
    return {"status": "ok"}


@app.get("/health/deep")
async def health_deep():
    db_state = "disabled"
    if settings.database_enabled and engine is not None:
        try:
            with get_db_session() as db:
                db.execute(text("SELECT 1"))
            db_state = "ok"
        except Exception:
            db_state = "error"
    return {
        "status": "ok" if db_state != "error" else "degraded",
        "database": db_state,
        "drive": "ok" if drive_available() else "disabled",
        "worker": "running" if settings.database_enabled else "disabled",
        "sources": ingestion_registry.list_sources(),
    }

@app.get("/health")
async def health():
    config = load_config()
    db = load_db()
    ai_keys = get_all_api_keys()
    return {
        "status": "ok",
        "version": "6.2",
        "ai_configured": bool(ai_keys),
        "ai_keys_count": len(ai_keys),
        "drive_sync": drive_available(),
        "tenders_loaded": len(db.get("tenders", {})),
        "boq_available": BOQ_AVAILABLE,
        "drive_warning": not drive_available(),
    }


@app.post("/auth/bootstrap-admin")
async def bootstrap_admin(data: dict = Body(...)):
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required for user auth")
    username = str(data.get("username", "")).strip()
    password = str(data.get("password", "")).strip()
    if len(username) < 3 or len(password) < 8:
        raise HTTPException(400, "Username/password too short")
    with get_db_session() as db:
        if db.query(User).filter(User.username == username).first():
            return {"status": "exists"}
        user = User(username=username, password_hash=hash_password(password), role="admin")
        db.add(user)
    return {"status": "created", "username": username}


@app.post("/auth/login")
async def login(data: dict = Body(...)):
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required for user auth")
    username = str(data.get("username", "")).strip()
    password = str(data.get("password", "")).strip()
    with get_db_session() as db:
        user = db.query(User).filter(User.username == username, User.is_active == True).first()
        if not user or not verify_password(password, user.password_hash):
            raise HTTPException(401, "Invalid credentials")
        token = create_access_token(subject=user.username, role=user.role)
    return {"access_token": token, "token_type": "bearer", "role": user.role}


@app.get("/auth/me")
async def auth_me(credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme)):
    user = _current_user(credentials)
    if not user:
        raise HTTPException(401, "Not authenticated")
    return {"username": user.get("sub"), "role": user.get("role")}


@app.post("/work-items")
async def create_work_item(
    data: dict = Body(...),
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required for work queue")
    work_type = str(data.get("work_type", "")).strip()
    payload = data.get("payload", {})
    if not work_type:
        raise HTTPException(400, "work_type is required")
    with get_db_session() as db:
        item = WorkItem(work_type=work_type, payload_json=json.dumps(payload))
        db.add(item)
        db.flush()
        item_id = item.id
    return {"status": "queued", "id": item_id}


@app.get("/work-items/{item_id}")
async def get_work_item(
    item_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst", "viewer"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required for work queue")
    with get_db_session() as db:
        item = db.query(WorkItem).filter(WorkItem.id == item_id).first()
        if not item:
            raise HTTPException(404, "Work item not found")
        return {
            "id": item.id,
            "work_type": item.work_type,
            "status": item.status,
            "attempts": item.attempts,
            "payload": json.loads(item.payload_json or "{}"),
            "result": json.loads(item.result_json or "{}"),
            "error": item.error_text,
        }


@app.post("/platform/sync-json-to-postgres")
async def sync_json_to_postgres(credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme)):
    _require_role(credentials, {"admin"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    db_data = load_db()
    tenders = db_data.get("tenders", {})
    synced = 0
    with get_db_session() as db:
        for t247_id, t in tenders.items():
            row = db.query(TenderRecord).filter(TenderRecord.t247_id == str(t247_id)).first()
            if not row:
                row = TenderRecord(t247_id=str(t247_id))
                db.add(row)
            row.verdict = str(t.get("verdict", ""))
            row.org_name = str(t.get("org_name", ""))
            row.tender_name = str(t.get("tender_name", t.get("brief", "")))
            row.estimated_cost = str(t.get("estimated_cost", t.get("estimated_cost_cr", "")))
            row.status = str(t.get("status", "Identified"))
            row.payload_json = json.dumps(t, default=str)
            synced += 1
    return {"status": "ok", "synced": synced}


@app.get("/platform/tenders")
async def list_platform_tenders(
    limit: int = 50,
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst", "viewer"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    limit = max(1, min(limit, 500))
    with get_db_session() as db:
        rows = (
            db.query(TenderRecord)
            .order_by(TenderRecord.updated_at.desc())
            .limit(limit)
            .all()
        )
        return {
            "items": [
                {
                    "t247_id": r.t247_id,
                    "verdict": r.verdict,
                    "org_name": r.org_name,
                    "tender_name": r.tender_name,
                    "estimated_cost": r.estimated_cost,
                    "win_probability": r.win_probability,
                    "risk_score": r.risk_score,
                    "status": r.status,
                    "updated_at": r.updated_at.isoformat() if r.updated_at else "",
                }
                for r in rows
            ]
        }


@app.get("/platform/sources")
async def list_platform_sources(credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme)):
    _require_role(credentials, {"admin", "analyst", "viewer"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    with get_db_session() as db:
        rows = db.query(TenderSource).order_by(TenderSource.name.asc()).all()
        return {
            "items": [
                {
                    "id": r.id,
                    "name": r.name,
                    "source_type": r.source_type,
                    "base_url": r.base_url,
                    "is_active": r.is_active,
                    "config": json.loads(r.config_json or "{}"),
                }
                for r in rows
            ]
        }


@app.post("/platform/sources")
async def upsert_platform_source(
    data: dict = Body(...),
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    name = str(data.get("name", "")).strip()
    source_type = str(data.get("source_type", "json_api")).strip()
    base_url = str(data.get("base_url", "")).strip()
    config = data.get("config", {}) or {}
    if not name:
        raise HTTPException(400, "name is required")
    with get_db_session() as db:
        row = db.query(TenderSource).filter(TenderSource.name == name).first()
        if not row:
            row = TenderSource(name=name, source_type=source_type)
            db.add(row)
        row.source_type = source_type
        row.base_url = base_url
        row.is_active = bool(data.get("is_active", True))
        row.config_json = json.dumps(config, default=str)
    return {"status": "saved", "name": name}


@app.post("/platform/ingestion/run")
async def run_ingestion(
    data: dict = Body(...),
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    source_name = str(data.get("source_name", "")).strip()
    if not source_name:
        raise HTTPException(400, "source_name is required")
    with get_db_session() as db:
        src = db.query(TenderSource).filter(TenderSource.name == source_name).first()
        if not src:
            raise HTTPException(404, "source not found")
        payload = {
            "source_name": src.name,
            "source_type": src.source_type,
            "endpoint": src.base_url,
            "config": json.loads(src.config_json or "{}"),
        }
        item = WorkItem(work_type="ingestion_sync", payload_json=json.dumps(payload))
        db.add(item)
        db.flush()
        work_id = item.id
    return {"status": "queued", "work_item_id": work_id}


@app.post("/platform/ingestion/preview")
async def preview_ingestion(
    data: dict = Body(...),
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst"})
    source_type = str(data.get("source_type", "json_api")).strip().lower()
    endpoint = str(data.get("endpoint", "")).strip()
    if not endpoint:
        raise HTTPException(400, "endpoint is required")

    if source_type == "json_api":
        source = JsonApiSource(endpoint=endpoint)
    elif source_type == "cppp_feed":
        source = CpppFeedSource(endpoint=endpoint)
    elif source_type == "state_portal_table":
        source = StatePortalTableSource(endpoint=endpoint)
    else:
        raise HTTPException(400, f"Unsupported source_type: {source_type}")

    rows = source.fetch()
    return {
        "status": "ok",
        "source_type": source_type,
        "count": len(rows),
        "sample": rows[:5],
    }


@app.get("/platform/ingested")
async def list_ingested(
    limit: int = 100,
    source_name: str = "",
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst", "viewer"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    limit = max(1, min(limit, 1000))
    with get_db_session() as db:
        q = db.query(IngestedTender)
        if source_name:
            q = q.filter(IngestedTender.source_name == source_name)
        rows = q.order_by(IngestedTender.updated_at.desc()).limit(limit).all()
        return {
            "items": [
                {
                    "id": r.id,
                    "source_name": r.source_name,
                    "external_id": r.external_id,
                    "title": r.title,
                    "org_name": r.org_name,
                    "deadline": r.deadline,
                    "reference_no": r.reference_no,
                    "updated_at": r.updated_at.isoformat() if r.updated_at else "",
                }
                for r in rows
            ]
        }


@app.post("/platform/clauses/index")
async def enqueue_clause_index(
    data: dict = Body(...),
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    payload = {
        "t247_id": str(data.get("t247_id", "")).strip(),
        "source_record_id": int(data.get("source_record_id") or 0),
    }
    with get_db_session() as db:
        item = WorkItem(work_type="clause_index", payload_json=json.dumps(payload))
        db.add(item)
        db.flush()
        work_id = item.id
    return {"status": "queued", "work_item_id": work_id}


@app.get("/platform/clauses")
async def list_clause_evidence(
    t247_id: str = "",
    source_record_id: int = 0,
    clause_type: str = "",
    limit: int = 200,
    credentials: HTTPAuthorizationCredentials | None = Depends(auth_scheme),
):
    _require_role(credentials, {"admin", "analyst", "viewer"})
    if not settings.database_enabled:
        raise HTTPException(400, "DATABASE_URL is required")
    limit = max(1, min(limit, 2000))
    with get_db_session() as db:
        q = db.query(ClauseEvidence)
        if t247_id:
            q = q.filter(ClauseEvidence.t247_id == t247_id)
        if source_record_id:
            q = q.filter(ClauseEvidence.source_record_id == source_record_id)
        if clause_type:
            q = q.filter(ClauseEvidence.clause_type == clause_type)
        rows = q.order_by(ClauseEvidence.created_at.desc()).limit(limit).all()
        return {
            "items": [
                {
                    "id": r.id,
                    "t247_id": r.t247_id,
                    "source_record_id": r.source_record_id,
                    "clause_type": r.clause_type,
                    "clause_text": r.clause_text,
                    "evidence_text": r.evidence_text,
                    "confidence": r.confidence,
                }
                for r in rows
            ]
        }

@app.get("/api-quota-status")
async def api_quota_status():
    """Lightweight status endpoint used by dashboard UI."""
    keys = get_all_api_keys()
    return {
        "status": "ok" if keys else "no_key",
        "daily_limit": "unknown",
        "today_used": "unknown",
        "keys_count": len(keys),
    }

_TOKEN_LOG_FILE = OUTPUT_DIR / "token_usage.json"
_token_log_lock = threading.Lock()

def _load_token_log() -> dict:
    if _TOKEN_LOG_FILE.exists():
        try:
            return json.loads(_TOKEN_LOG_FILE.read_text())
        except Exception:
            pass
    return {"today": {}, "total": {"calls": 0, "input_tokens": 0, "output_tokens": 0}}

def _save_token_log(data: dict):
    """Write token log locally + backup to Drive so it survives restarts."""
    try:
        _TOKEN_LOG_FILE.write_text(json.dumps(data, indent=2))
    except Exception:
        pass
    # Drive backup — async so it doesn't slow down the analysis path
    def _drive_save():
        try:
            if drive_available():
                save_to_drive(_TOKEN_LOG_FILE, filename="token_usage.json")
        except Exception:
            pass
    import threading as _thr
    _thr.Thread(target=_drive_save, daemon=True).start()


def _load_token_log() -> dict:
    """Load token log — prefer Drive-restored file if present."""
    if _TOKEN_LOG_FILE.exists():
        try:
            return json.loads(_TOKEN_LOG_FILE.read_text())
        except Exception:
            pass
    return {"today": {}, "history": {}, "total": {"calls": 0, "input_tokens": 0, "output_tokens": 0}}


def record_token_usage(input_tokens: int, output_tokens: int, key_masked: str = ""):
    """Call after each Gemini API call — persists to Drive, survives restarts."""
    today = datetime.now().strftime("%Y-%m-%d")
    with _token_log_lock:
        data = _load_token_log()
        # Today summary
        d = data.setdefault("today", {})
        if d.get("date") != today:
            # Roll yesterday into history
            if d.get("date"):
                hist = data.setdefault("history", {})
                hist[d["date"]] = {"calls": d.get("calls", 0),
                                   "input_tokens": d.get("input_tokens", 0),
                                   "output_tokens": d.get("output_tokens", 0)}
                # Keep only last 30 days
                for old in sorted(hist.keys())[:-30]:
                    del hist[old]
            data["today"] = {"date": today, "calls": 0, "input_tokens": 0,
                             "output_tokens": 0, "keys": {}}
            d = data["today"]
        d["calls"] = d.get("calls", 0) + 1
        d["input_tokens"] = d.get("input_tokens", 0) + input_tokens
        d["output_tokens"] = d.get("output_tokens", 0) + output_tokens
        # Per-key tracking
        if key_masked:
            kd = d.setdefault("keys", {}).setdefault(key_masked, {"calls": 0, "tokens": 0})
            kd["calls"] += 1
            kd["tokens"] += input_tokens + output_tokens
        # All-time total
        t = data.setdefault("total", {"calls": 0, "input_tokens": 0, "output_tokens": 0})
        t["calls"] = t.get("calls", 0) + 1
        t["input_tokens"] = t.get("input_tokens", 0) + input_tokens
        t["output_tokens"] = t.get("output_tokens", 0) + output_tokens
        _save_token_log(data)


@app.get("/token-usage")
async def token_usage():
    """Return today's and all-time Gemini token/call usage — persisted on Drive."""
    keys = get_all_api_keys()
    data = _load_token_log()
    today_str = datetime.now().strftime("%Y-%m-%d")
    today = data.get("today", {})
    if today.get("date") != today_str:
        today = {"date": today_str, "calls": 0, "input_tokens": 0, "output_tokens": 0}
    total = data.get("total", {"calls": 0, "input_tokens": 0, "output_tokens": 0})
    history = data.get("history", {})

    # Per free tier: 1500 requests/day per key, 1M tokens/day per key
    RPD_PER_KEY = 1500
    keys_count = len(keys)
    total_rpd   = RPD_PER_KEY * max(keys_count, 1)
    calls_today = today.get("calls", 0)
    rpd_used_pct = round(calls_today / max(total_rpd, 1) * 100, 1)
    rpd_remaining = max(0, total_rpd - calls_today)

    TOKENS_PER_KEY = 1_000_000
    total_tok_limit = TOKENS_PER_KEY * max(keys_count, 1)
    used_tok = today.get("input_tokens", 0) + today.get("output_tokens", 0)
    tok_remaining = max(0, total_tok_limit - used_tok)
    tok_pct = round(used_tok / max(total_tok_limit, 1) * 100, 1)

    # Last 7 days history for sparkline
    last7 = []
    for i in range(6, -1, -1):
        from datetime import timedelta
        day = (datetime.now() - timedelta(days=i)).strftime("%Y-%m-%d")
        if day == today_str:
            last7.append({"date": day, "calls": calls_today})
        else:
            h = history.get(day, {})
            last7.append({"date": day, "calls": h.get("calls", 0)})

    return {
        "status": "ok" if keys else "no_key",
        "keys_count": keys_count,
        "today_calls": calls_today,
        "today_tokens_used": used_tok,
        "rpd_limit": total_rpd,
        "rpd_remaining": rpd_remaining,
        "rpd_pct_used": rpd_used_pct,
        "token_limit": total_tok_limit,
        "token_remaining": tok_remaining,
        "token_pct_used": tok_pct,
        "total_all_time": total,
        "per_key_today": today.get("keys", {}),
        "last7_days": last7,
        "reset_time": "Midnight IST (Google quota resets daily)",
    }

# ══ SELF-DIAGNOSE ════════════════════════════════════════════════════════════
@app.get("/diagnose")
async def diagnose():
    results = []
    overall = "OK"

    def chk(name, fn):
        nonlocal overall
        try:
            status, detail, fix = fn()
            if status == "ERROR": overall = "ERROR"
            elif status == "WARN" and overall == "OK": overall = "WARN"
            results.append({"name": name, "status": status, "detail": detail, "fix": fix})
        except Exception as e:
            overall = "ERROR"
            results.append({"name": name, "status": "ERROR", "detail": str(e), "fix": "Check server logs"})

    chk("Database", lambda: ("WARN","Database empty","Import Excel or upload tenders_db.json") if len(load_db().get("tenders",{}))==0 else ("OK",f"{len(load_db().get('tenders',{}))} tenders in database",""))
    chk("Google Drive", lambda: ("OK","Google Drive connected","") if drive_available() else (("ERROR","GDRIVE_CREDENTIALS not set — data will be LOST on restart","Add env var in Render dashboard") if not os.environ.get("GDRIVE_CREDENTIALS") else ("WARN","Drive credentials set but connection failed","Check JSON format")))
    chk("Gemini AI", lambda: ("OK",f"Key configured ({load_config().get('gemini_api_key','')[:8]}...)","") if load_config().get("gemini_api_key") else ("ERROR","No Gemini API key","Settings → add key from aistudio.google.com"))
    chk("BOQ Engine", lambda: ("OK","BOQ engine loaded","") if BOQ_AVAILABLE else ("ERROR","boq_engine.py missing","Add boq_engine.py to GitHub repo"))
    chk("Required Files", lambda: (("OK","All required files present","") if not [f for f in ["extractor.py","doc_generator.py","nascent_checker.py","ai_analyzer.py","excel_processor.py","prebid_generator.py","chatbot.py","gdrive_sync.py","tracker.py","nascent_profile.json"] if not (BASE_DIR/f).exists()] else ("ERROR",f"Missing: {', '.join([f for f in ['extractor.py','doc_generator.py','nascent_checker.py','ai_analyzer.py','excel_processor.py','prebid_generator.py','chatbot.py','gdrive_sync.py','tracker.py','nascent_profile.json'] if not (BASE_DIR/f).exists()])}", "Add to GitHub")))
    chk("Company Profile", lambda: ("OK","Profile complete","") if (BASE_DIR/"nascent_profile.json").exists() and all(k in json.loads((BASE_DIR/"nascent_profile.json").read_text()) for k in ["company","finance","certifications","employees","projects","bid_rules"]) else ("WARN","Profile incomplete or missing","Company Profile → fill all sections"))
    def _check_data_dir():
        try:
            OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
            probe = OUTPUT_DIR / "_test.tmp"
            probe.write_text("ok", encoding="utf-8")
            probe.unlink(missing_ok=True)
            return ("OK", "Data directory writable", "")
        except Exception:
            return ("ERROR", "Cannot write data", "Check Render disk")
    chk("Data Directory", _check_data_dir)
    chk("Admin Security", lambda: ("OK","Admin token configured","") if ADMIN_TOKEN else ("WARN","ADMIN_TOKEN env var not set — config endpoints are open","Set ADMIN_TOKEN in Render environment"))

    return {"overall": overall, "timestamp": datetime.now().isoformat(), "checks": results,
            "summary": {"ok": sum(1 for r in results if r["status"]=="OK"),
                        "warn": sum(1 for r in results if r["status"]=="WARN"),
                        "error": sum(1 for r in results if r["status"]=="ERROR")}}

@app.post("/diagnose/ai")
async def diagnose_with_ai(data: dict = Body(...)):
    error_text = data.get("error", "").strip()
    if not error_text:
        raise HTTPException(400, "No error text provided")
    config = load_config()
    api_key = config.get("gemini_api_key", "")
    if not api_key:
        return {"root_cause": "No Gemini API key — go to Settings first.", "fixes": [], "severity": "WARN", "affected_file": "config"}
    prompt = f"""You are a senior Python/FastAPI developer diagnosing an error in the NIT Bid/No-Bid tender management system.
Stack: Python 3.11, FastAPI, Uvicorn, Gemini AI, Google Drive API, Render free tier.
Files: main.py, ai_analyzer.py, doc_generator.py, nascent_checker.py, boq_engine.py, gdrive_sync.py, excel_processor.py, extractor.py, tracker.py, chatbot.py, index.html

ERROR:
{error_text[:3000]}

Return ONLY valid JSON, no markdown:
{{"root_cause":"one sentence what went wrong","affected_file":"filename or deployment","severity":"CRASH|WARN|MINOR","fixes":[{{"step":1,"action":"exact action","where":"GitHub|Render|Settings|Code"}}],"can_auto_fix":false}}"""
    try:
        from ai_analyzer import call_gemini, clean_json
        return clean_json(call_gemini(prompt, api_key))
    except Exception as e:
        return {"root_cause": f"AI diagnosis unavailable: {e}", "fixes": [{"step": 1, "action": "Check API key in Settings", "where": "Settings"}], "severity": "WARN", "affected_file": "unknown", "can_auto_fix": False}

# ══ EXCEL IMPORT ════════════════════════════════════════════════════════════
@app.post("/import-excel")
async def import_excel(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(400, "Please upload an Excel file")
    tmp = Path(tempfile.mktemp(suffix=".xlsx", dir=str(TEMP_DIR)))
    try:
        content = await file.read()
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(413, "File too large (max 50MB)")
        tmp.write_bytes(content)
        LATEST_EXCEL_FILE.write_bytes(content)
        tenders = process_excel(str(tmp))
        db = load_db()
        added = updated = 0
        for t in tenders:
            tid = str(t.get("t247_id", ""))
            if not tid:
                continue
            existing = db["tenders"].get(tid, {})
            if existing:
                for field in ["ref_no","brief","org_name","location","estimated_cost_raw","estimated_cost_cr","deadline","days_left","deadline_status","doc_fee","emd","msme_exemption","eligibility","checklist","is_gem"]:
                    if t.get(field) is not None:
                        existing[field] = t[field]
                if not existing.get("bid_no_bid_done"):
                    existing["verdict"] = t.get("verdict")
                    existing["verdict_color"] = t.get("verdict_color")
                    existing["reason"] = t.get("reason")
                db["tenders"][tid] = existing
                updated += 1
            else:
                db["tenders"][tid] = t
                added += 1
        save_db(db)
        return {"status": "success", "total": len(tenders), "added": added, "updated": updated,
                "imported": len(tenders),
                "bid": sum(1 for t in tenders if t.get("verdict") == "BID"),
                "no_bid": sum(1 for t in tenders if t.get("verdict") == "NO-BID"),
                "tenders": tenders}
    finally:
        tmp.unlink(missing_ok=True)

@app.post("/sync-excel-latest")
async def sync_excel_latest():
    if not LATEST_EXCEL_FILE.exists():
        raise HTTPException(404, "No previously imported Excel found. Import once first.")
    tenders = process_excel(str(LATEST_EXCEL_FILE))
    db = load_db()
    added = updated = 0
    for t in tenders:
        tid = str(t.get("t247_id", ""))
        if not tid:
            continue
        if tid in db["tenders"]:
            db["tenders"][tid].update(t)
            updated += 1
        else:
            db["tenders"][tid] = t
            added += 1
    save_db(db)
    return {"status": "success", "source": str(LATEST_EXCEL_FILE), "total": len(tenders), "added": added, "updated": updated}

# ══ DASHBOARD ═══════════════════════════════════════════════════════════════
@app.get("/dashboard")
async def dashboard():
    db = load_db()
    tenders = list(db["tenders"].values())
    return {"stats": {
        "total": len(tenders),
        "bid": sum(1 for t in tenders if t.get("verdict") == "BID"),
        "no_bid": sum(1 for t in tenders if t.get("verdict") == "NO-BID"),
        "conditional": sum(1 for t in tenders if t.get("verdict") == "CONDITIONAL"),
        "review": sum(1 for t in tenders if t.get("verdict") == "REVIEW"),
        "analysed": sum(1 for t in tenders if t.get("bid_no_bid_done")),
        "deadline_today": sum(1 for t in tenders if days_left(t.get("deadline", "")) == 0),
        "deadline_3days": sum(1 for t in tenders if 0 < days_left(t.get("deadline", "")) <= 3),
        "has_boq": sum(1 for t in tenders if t.get("boq")),
    }, "tenders": sorted(tenders, key=lambda t: days_left(t.get("deadline", "999")))}

@app.get("/ops/daily-report")
async def ops_daily_report():
    summary = _compute_ops_summary()
    return {"status": "success", "summary": summary}

def _compute_ops_summary() -> dict:
    db = load_db()
    tenders = list(db.get("tenders", {}).values())
    now = datetime.now()
    day_ago = now - timedelta(days=1)
    today_new = 0
    urgent = 0
    no_docs = 0
    pending_analysis = 0
    corrigendum = 0
    for t in tenders:
        created = str(t.get("created_at", "") or t.get("updated_at", "") or "")
        try:
            if created:
                ts = datetime.fromisoformat(created.replace("Z", "+00:00").replace("+00:00", ""))
                if ts >= day_ago:
                    today_new += 1
        except Exception:
            pass
        dl = days_left(t.get("deadline", ""))
        if 0 <= dl <= 3:
            urgent += 1
        if not t.get("docs_available") and not t.get("t247_doc_hash"):
            no_docs += 1
        if not t.get("bid_no_bid_done"):
            pending_analysis += 1
        if t.get("has_corrigendum"):
            corrigendum += 1
    top_urgent = sorted(
        [t for t in tenders if 0 <= days_left(t.get("deadline", "")) <= 3],
        key=lambda t: days_left(t.get("deadline", "")),
    )[:5]
    return {
        "total_tenders": len(tenders),
        "new_last_24h": today_new,
        "urgent_deadlines_3d": urgent,
        "no_docs_yet": no_docs,
        "pending_analysis": pending_analysis,
        "corrigendum_flagged": corrigendum,
        "top_urgent": [
            {
                "t247_id": str(t.get("t247_id", "")),
                "brief": str(t.get("brief", ""))[:120],
                "deadline": t.get("deadline", ""),
                "days_left": days_left(t.get("deadline", "")),
                "verdict": t.get("verdict", "REVIEW"),
            }
            for t in top_urgent
        ],
    }

def _build_daily_digest() -> dict:
    s = _compute_ops_summary()
    today = datetime.now().strftime("%d %b %Y")
    subject = f"Daily Tender Digest | {today} | Total {s['total_tenders']} | Urgent {s['urgent_deadlines_3d']}"
    urgent_lines = []
    for t in s.get("top_urgent", []):
        urgent_lines.append(
            f"- {t.get('t247_id','')} | {t.get('brief','')} | D-{t.get('days_left','?')} | {t.get('verdict','REVIEW')}"
        )
    if not urgent_lines:
        urgent_lines = ["- No urgent tenders in next 3 days"]
    email_body = "\n".join([
        f"Daily Tender Digest - {today}",
        "",
        f"Total tenders: {s['total_tenders']}",
        f"New in last 24h: {s['new_last_24h']}",
        f"Urgent deadlines (<=3d): {s['urgent_deadlines_3d']}",
        f"No docs yet: {s['no_docs_yet']}",
        f"Pending analysis: {s['pending_analysis']}",
        f"Corrigendum flagged: {s['corrigendum_flagged']}",
        "",
        "Top urgent tenders:",
        *urgent_lines,
        "",
        "Action plan:",
        "1) Download missing docs",
        "2) Complete bid/no-bid review for pending tenders",
        "3) Prioritize D-0 and D-1 submissions",
    ])
    whatsapp_text = (
        f"Daily Digest ({today})\n"
        f"Total:{s['total_tenders']} | New24h:{s['new_last_24h']} | Urgent<=3d:{s['urgent_deadlines_3d']}\n"
        f"NoDocs:{s['no_docs_yet']} | PendingAnalysis:{s['pending_analysis']} | Corrigendum:{s['corrigendum_flagged']}\n"
        + "\n".join([f"* {u}" for u in urgent_lines[:3]])
    )
    digest = {
        "generated_at": datetime.now().isoformat(),
        "subject": subject,
        "email_body": email_body,
        "whatsapp_text": whatsapp_text,
        "summary": s,
    }
    with _digest_lock:
        _digest_state["last_generated_at"] = digest["generated_at"]
        _digest_state["last_generated_date"] = datetime.now().strftime("%Y-%m-%d")
        _digest_state["status"] = "success"
        _digest_state["error"] = ""
    try:
        (OUTPUT_DIR / "daily_digest_latest.json").write_text(json.dumps(digest, indent=2), encoding="utf-8")
        (OUTPUT_DIR / "daily_digest_latest.txt").write_text(
            f"SUBJECT: {subject}\n\nEMAIL:\n{email_body}\n\nWHATSAPP:\n{whatsapp_text}\n",
            encoding="utf-8",
        )
    except Exception:
        pass
    return digest

@app.get("/tenders")
async def get_all_tenders():
    return {"tenders": list(load_db()["tenders"].values())}

# ══ TENDER OPS ══════════════════════════════════════════════════════════════
@app.post("/prebid-queries")
async def get_prebid_queries_post(data: dict = Body(...)):
    return {"queries": generate_prebid_queries(data)}

@app.get("/prebid-queries/{t247_id}")
async def get_saved_prebid_queries(t247_id: str):
    return {"queries": get_tender(t247_id).get("prebid_queries", [])}

@app.post("/tender/{t247_id}/generate-prebid-letter")
async def generate_prebid_letter(t247_id: str):
    """
    Generate pre-bid query letter as Word doc (.docx) with proper NIT letterhead.
    Returns base64-encoded doc for direct browser download.
    """
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found. Analyse tender first.")

    queries = tender.get("prebid_queries", [])
    if not queries:
        raise HTTPException(400, "No pre-bid queries found. Analyse this tender first.")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import base64, io

        doc = DocxDocument()

        # Page setup
        sec = doc.sections[0]
        sec.page_width = Cm(21); sec.page_height = Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin = sec.bottom_margin = Cm(2)

        # Header — NIT details
        def add_para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p

        add_para("NASCENT INFO TECHNOLOGIES PVT. LTD.", bold=True, size=14,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 70, 127), space_after=2)
        add_para("A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad – 380015",
                 size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_para("📞 +91-79-40200400 | ✉ nascent.tender@nascentinfo.com | 🌐 www.nascentinfo.com",
                 size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_para("CIN: U72200GJ2006PTC048723 | PAN: AACCN3670J | GSTIN: 24AACCN3670J1ZG | MSME: UDYAM-GJ-01-0007420",
                 size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100,100,100), space_after=4)

        # Horizontal line
        hr = doc.add_paragraph()
        hr.paragraph_format.space_after = Pt(8)
        hr_run = hr.add_run("─" * 90)
        hr_run.font.size = Pt(8); hr_run.font.color.rgb = RGBColor(0, 70, 127)

        # Date + Ref
        today_str = datetime.now().strftime("%d %B %Y")
        add_para(f"Date: {today_str}", size=10, space_after=10)

        # To block
        org = str(tender.get("org_name") or "The Tender Inviting Authority")
        tender_no = str(tender.get("tender_no") or t247_id)
        tender_name = str(tender.get("tender_name") or tender.get("brief") or "")
        contact = str(tender.get("contact") or "")

        add_para("To,", size=11, bold=True, space_after=0)
        add_para(f"The Tender Inviting Authority", size=11, space_after=0)
        add_para(org, size=11, space_after=2)
        if contact and contact not in ("—", "Not specified"):
            add_para(f"Contact: {contact}", size=10, color=(80,80,80), space_after=2)

        # Subject
        subj_p = doc.add_paragraph()
        subj_p.paragraph_format.space_before = Pt(10)
        subj_p.paragraph_format.space_after = Pt(10)
        subj_run = subj_p.add_run(f"Subject: Pre-Bid Queries — Ref: {tender_no}")
        subj_run.bold = True; subj_run.font.size = Pt(11)
        subj_run.font.color.rgb = RGBColor(0, 70, 127)

        if tender_name and tender_name != "—":
            add_para(f"Tender: {tender_name[:120]}", size=10, color=(80,80,80), space_after=6)

        add_para("Respected Sir/Madam,", size=11, space_after=6)
        add_para(
            "With reference to the above-mentioned tender, Nascent Info Technologies Pvt. Ltd. "
            "requests clarification on the following points. We request your kind consideration "
            "and responses at the earliest, preferably before the pre-bid meeting:",
            size=11, space_after=10
        )

        # Query table
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for cell, text, w in zip(hdr, ["Sr.", "Clause Ref", "RFP Extract", "Our Query"], [1.2, 2.5, 6, 6]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.width = Cm(w)

        for i, q in enumerate(queries, 1):
            if isinstance(q, dict):
                clause = str(q.get("clause_ref") or q.get("clause") or "—")
                rfp_text = str(q.get("rfp_text") or "—")
                qtxt = str(q.get("query") or q.get("text") or str(q))
            else:
                clause = "—"; rfp_text = "—"; qtxt = str(q)

            row = tbl.add_row().cells
            row[0].text = str(i)
            row[1].text = clause
            row[2].text = rfp_text[:200]
            row[3].text = qtxt[:300]
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        # Closing
        doc.add_paragraph()
        add_para("We hope for your prompt response to ensure clarity for a competitive bid.",
                 size=11, space_after=14)
        add_para("Yours sincerely,", size=11, space_after=20)
        add_para("Authorised Signatory", size=11, bold=True, space_after=2)
        add_para("Nascent Info Technologies Pvt. Ltd.", size=11, space_after=2)
        add_para("MSME | CMMI L3 V2.0 | ISO 9001 | ISO 27001 | ISO 20000", size=9, color=(100,100,100))

        # Save to bytes → base64
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()
        doc_b64 = base64.b64encode(doc_bytes).decode("utf-8")

        fname = f"PreBid_{re.sub(r'[^\w\-]', '_', tender_no)[:40]}.docx"

        # Also save to disk for /download/ fallback
        try:
            (OUTPUT_DIR / fname).write_bytes(doc_bytes)
        except Exception:
            pass

        return {
            "status": "success",
            "filename": fname,
            "query_count": len(queries),
            "download_url": f"/download/{fname}",
            "doc_b64": doc_b64,
        }

    except Exception as e:
        import traceback
        raise HTTPException(500, f"Pre-bid letter generation failed: {str(e)}\n{traceback.format_exc()[:500]}")


@app.post("/tender/{t247_id}/analyze-compliance")
async def analyze_compliance(t247_id: str):
    """
    Basic compliance scanner (rule-based) to avoid UI failures.
    """
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")

    text = " ".join([
        str(tender.get("tender_name", "")),
        str(tender.get("reason", "")),
        str(tender.get("emd_exemption", "")),
        str(tender.get("jv_allowed", "")),
    ]).lower()
    issues = []
    if "not allowed" in text and ("consortium" in text or "joint venture" in text):
        issues.append({
            "clause_no": "JV/Consortium",
            "issue_type": "Restriction",
            "severity": "MEDIUM",
            "what_law_says": "Verify consortium restrictions and bid as single entity if required.",
        })
    if "emd" in text and "exempt" not in text and "msme" in text:
        issues.append({
            "clause_no": "EMD",
            "issue_type": "MSME exemption unclear",
            "severity": "LOW",
            "what_law_says": "Seek written clarification on MSME EMD exemption applicability.",
        })

    return {
        "status": "success",
        "violations_found": len(issues),
        "analysis": {"clause_violations": issues},
    }

@app.post("/tender/{t247_id}/status")
async def update_status(t247_id: str, data: dict = Body(...)):
    t = get_tender(t247_id)
    t.update(data)
    save_tender(t247_id, t)
    return {"status": "saved"}

@app.get("/tender/{t247_id}")
async def get_tender_detail(t247_id: str):
    t = get_tender(t247_id)
    if not t:
        raise HTTPException(404, f"Tender {t247_id} not found")
    return t

@app.post("/tender/{t247_id}/reanalyse")
async def reanalyse_tender(t247_id: str):
    """Re-run AI analysis using saved raw text from DB."""
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")

    saved_text = tender.get("raw_text", "")
    if not saved_text or len(str(saved_text).strip()) < 100:
        raise HTTPException(400, "No saved document text found. Upload and analyse again.")

    cfg = load_config()
    api_key = cfg.get("gemini_api_key", "")
    if not api_key:
        raise HTTPException(400, "Gemini API key not configured. Go to Settings.")

    prebid_flag = bool(tender.get("prebid_passed", False))
    ai_result = analyze_with_gemini(saved_text, prebid_flag)
    if "error" in ai_result:
        raise HTTPException(502, ai_result.get("error", "AI reanalysis failed"))

    merged = merge_results(tender, ai_result, prebid_flag)
    merged["bid_no_bid_done"] = True
    merged["analysed_at"] = datetime.now().isoformat()
    merged["raw_text"] = saved_text
    save_tender(t247_id, merged)

    return {
        "status": "success",
        "t247_id": t247_id,
        "verdict": merged.get("verdict", merged.get("overall_verdict", {}).get("verdict", "REVIEW")),
        "tender_data": merged,
    }

@app.get("/tender-quickview/{t247_id}")
async def tender_quickview(t247_id: str):
    t = get_tender(t247_id)
    if not t:
        raise HTTPException(404, "Not found")
    return t

@app.post("/tender/{t247_id}/skip")
async def skip_tender(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    t.update({"status": "Not Interested", "skip_reason": data.get("reason", "Not interested"), "skipped_at": datetime.now().isoformat()})
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "skipped"}

@app.post("/tender/{t247_id}/restore")
async def restore_tender(t247_id: str):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    t["status"] = "Identified"
    t.pop("skip_reason", None)
    t.pop("skipped_at", None)
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "restored"}

@app.post("/tender/{t247_id}/reclassify")
async def reclassify_tender(t247_id: str):
    from excel_processor import classify_tender
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    if not t:
        raise HTTPException(404, "Not found")
    if t.get("manual_verdict"):
        return {"status": "skipped", "verdict": t["verdict"], "reason": "Manual verdict locked"}
    r = classify_tender(t.get("brief", ""), t.get("estimated_cost_raw", 0), t.get("eligibility", ""), t.get("checklist", ""))
    t.update({"verdict": r["verdict"], "verdict_color": r["verdict_color"], "reason": r["reason"]})
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "reclassified", "verdict": r["verdict"], "reason": r["reason"]}

@app.post("/tender/{t247_id}/set-verdict")
async def set_verdict_manual(t247_id: str, payload: dict):
    """Human override — lock verdict so AI reclassify cannot change it."""
    verdict = (payload.get("verdict") or "").upper().strip()
    if verdict not in ("BID", "NO-BID", "CONDITIONAL", "REVIEW"):
        raise HTTPException(400, "verdict must be BID / NO-BID / CONDITIONAL / REVIEW")
    db = load_db()
    t = db["tenders"].get(t247_id)
    if not t:
        raise HTTPException(404, "Tender not found")
    t["verdict"] = verdict
    if payload.get("_clear_manual"):
        t.pop("manual_verdict", None)
        t.pop("manual_verdict_note", None)
    else:
        t["manual_verdict"] = True
        t["manual_verdict_note"] = (payload.get("note") or "").strip()[:200]
    db["tenders"][t247_id] = t
    save_db(db)
    asyncio.create_task(sync_to_drive_async())
    return {"status": "ok", "verdict": verdict, "manual_verdict": not payload.get("_clear_manual")}

@app.post("/reclassify-all")
async def reclassify_all():
    from excel_processor import classify_tender
    db = load_db()
    counts = {}
    for tid, t in db["tenders"].items():
        if t.get("bid_no_bid_done") or t.get("manual_verdict"):
            continue
        r = classify_tender(t.get("brief", ""), t.get("estimated_cost_raw", 0), t.get("eligibility", ""), t.get("checklist", ""))
        t.update({"verdict": r["verdict"], "verdict_color": r["verdict_color"], "reason": r["reason"]})
        counts[r["verdict"]] = counts.get(r["verdict"], 0) + 1
    save_db(db)
    return {"status": "done", "reclassified": sum(counts.values()), "breakdown": counts}

# ══ PROCESS FILES ════════════════════════════════════════════════════════════
@app.post("/process")
async def process_zip(file: UploadFile = File(...), t247_id: str = ""):
    return await process_files(files=[file], t247_id=t247_id)

@app.post("/process-files")
async def process_files(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...), t247_id: str = ""):
    if not files:
        raise HTTPException(400, "No files uploaded")

    # Read files into memory immediately (before background task)
    file_contents = []
    for upload in files:
        content = await upload.read()
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(413, f"File {upload.filename} too large (max 50MB)")
        file_contents.append((upload.filename or "upload", content))

    job_id = str(uuid.uuid4())[:12]
    import time
    _set_job(job_id, status="running", progress="Starting…", result=None, error=None, t247_id=t247_id, started_at=time.time())
    background_tasks.add_task(_run_analysis_job, job_id, file_contents, t247_id)
    return {"job_id": job_id, "status": "running"}


@app.get("/analyse-status/{job_id}")
async def analyse_status(job_id: str):
    import time as _time
    job = _get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    # Auto-fail jobs stuck > 12 minutes (Gemini timeout / server restart)
    if job.get("status") == "running" and job.get("started_at"):
        elapsed = _time.time() - job["started_at"]
        if elapsed > 720:
            _set_job(job_id, status="error", error=f"Analysis timed out after {int(elapsed//60)}min. The server may have restarted. Please re-upload and try again.")
            return _get_job(job_id)
    return job


def _run_analysis_job(job_id: str, file_contents: list, t247_id: str):
    """Runs in background thread — full analysis pipeline.
    Acquires a JobSlot so Render stays within RAM ceiling while
    multiple tenders are analyzed concurrently."""
    slot_held = False
    if API_POOL_AVAILABLE:
        try:
            slots = get_slots()
            _set_job(job_id, progress=f"Queued (active={slots.snapshot()['active']}/{slots.max})…")
            if slots.acquire(timeout=900.0):
                slot_held = True
            else:
                _set_job(job_id, status="error", error="Analyst queue busy — try again in a minute.")
                return
        except Exception:
            slot_held = False
    tmp_dir = tempfile.mkdtemp(prefix="tender_", dir=str(TEMP_DIR))
    job_ok = False
    try:
        extract_dir = Path(tmp_dir) / "extracted"
        extract_dir.mkdir()

        _set_job(job_id, progress="Extracting documents…", step=0)
        for fname, content in file_contents:
            dest = Path(tmp_dir) / fname
            dest.write_bytes(content)
            if dest.suffix.lower() == ".zip":
                with zipfile.ZipFile(dest, "r") as zf:
                    zf.extractall(extract_dir)
                extract_all_zips(extract_dir)
            else:
                shutil.copy2(dest, extract_dir / fname)

        doc_files = []
        for ext in ["*.pdf", "*.docx", "*.doc", "*.txt", "*.html", "*.htm", "*.xlsx", "*.xls"]:
            doc_files.extend(extract_dir.rglob(ext))
        seen, unique = set(), []
        for f in doc_files:
            if f.name not in seen:
                seen.add(f.name)
                unique.append(f)
        doc_files = unique

        if not doc_files:
            _set_job(job_id, status="error", error="No readable documents found in uploaded files.")
            return

        corr = [f for f in doc_files if any(k in f.name.lower() for k in ["corrigendum","addendum","amendment","corr_","addend","revised","rectification"])]
        main_files = [f for f in doc_files if f not in corr]

        _set_job(job_id, progress="Reading documents…", step=1)
        extractor = TenderExtractor()
        tender_data = extractor.process_documents(main_files if main_files else doc_files)

        if corr:
            cd = TenderExtractor().process_documents(corr)
            for field in ["bid_submission_date","bid_opening_date","bid_start_date","prebid_query_date","estimated_cost","emd","tender_fee"]:
                val = cd.get(field, "")
                if val and val not in ["—","Refer document","Not specified",""]:
                    tender_data[field] = val
            tender_data["has_corrigendum"] = True
            tender_data["corrigendum_files"] = [f.name for f in corr]

        _set_job(job_id, progress="Reading full text…", step=1)
        all_text = ""
        MAX_CORPUS = 350_000  # ~350KB — keeps RAM under 512MB on Render free tier
        for f in sorted(doc_files, key=lambda x: (0 if any(k in x.name.lower() for k in ["rfp","nit","tender","bid"]) else 1 if any(k in x.name.lower() for k in ["corrigendum","addendum"]) else 2)):
            if len(all_text) >= MAX_CORPUS:
                break
            t = read_document(f)
            if t and t.strip():
                remaining = MAX_CORPUS - len(all_text)
                all_text += f"\n\n=== FILE: {f.name} ===\n{t[:remaining]}"
        # Free doc file handles from memory before AI pipeline
        del doc_files

        config = load_config()
        api_key = config.get("gemini_api_key", "")
        ai_used = False
        passed = prebid_passed(tender_data.get("prebid_query_date", ""))

        if api_key and all_text.strip():
            import time as _time_seg
            _seg_start_times = {}
            _set_job(job_id, progress="AI pipeline starting (parallel 9 segments)…", step=2)
            _STAGE_STEP = {
                "snapshot": 1, "corrig": 1,
                "scope": 2, "pq": 3, "tq": 4, "workshed": 4,
                "payment": 5, "assessment": 6, "prebid": 7, "checklist": 7,
            }
            _STAGE_LABEL = {
                "snapshot": "Snapshot — dates, EMD, cost",
                "corrig": "Corrigendum check",
                "scope": "Scope of Work",
                "pq": "PQ / Eligibility Criteria",
                "tq": "TQ / Marking Criteria",
                "workshed": "Work Schedule",
                "payment": "Payment Schedule",
                "assessment": "Bid/No-Bid Assessment",
                "prebid": "Pre-bid Queries",
                "checklist": "Submission Checklist",
            }
            def _seg_progress(stage, done, total, seg_output=None):
                try:
                    s = _STAGE_STEP.get(stage, 2)
                    elapsed = round(_time_seg.time() - _seg_start_times.get(stage, _time_seg.time()), 1)
                    label = _STAGE_LABEL.get(stage, stage)
                    count_hint = ""
                    seg_save = {}
                    if seg_output:
                        pq = seg_output.get("pq_criteria", [])
                        tq = seg_output.get("tq_criteria", [])
                        scope = seg_output.get("scope_sections", seg_output.get("scope_items", []))
                        if pq: count_hint = f" — {len(pq)} criteria"
                        elif tq: count_hint = f" — {len(tq)} criteria"
                        elif scope: count_hint = f" — {len(scope)} sections"
                        seg_save = {stage: seg_output}
                    _set_job(
                        job_id,
                        progress=f"✓ {label}{count_hint} ({elapsed}s)",
                        step=s,
                        segments=seg_save,
                        seg_log={stage: {"label": label, "elapsed": elapsed,
                                         "count": count_hint.strip(" —"), "done": True}},
                    )
                except Exception:
                    pass
            ai_result = {"error": "timeout"}

            # Ticker thread — updates progress every 8s so frontend shows elapsed time
            import threading as _threading
            _ai_done_evt = _threading.Event()
            _ai_start_t = _time_seg.time()
            def _ai_ticker():
                _models = ["gemini-1.5-pro","gemini-2.0-flash","gemini-1.5-flash"]
                _mi = 0
                while not _ai_done_evt.wait(timeout=8):
                    _el = int(_time_seg.time() - _ai_start_t)
                    _mod = _models[_mi % len(_models)]
                    _set_job(job_id,
                             progress=f"🤖 AI reading tender… ({_el}s) — {_mod}",
                             step=3)
                    _mi += 1
            _ticker = _threading.Thread(target=_ai_ticker, daemon=True)
            _ticker.start()

            # Inner progress callback — fires once with full result to populate stream panels
            def _ai_done_cb(stage, done, total, seg_output=None):
                if stage == "done" and seg_output:
                    _ai_done_evt.set()  # stop ticker before setting higher steps
                    import time as _t2
                    _now = _t2.time()
                    # Fire segment callbacks so live panels populate
                    pq = seg_output.get("pq_criteria", [])
                    tq = seg_output.get("tq_criteria", [])
                    scope = seg_output.get("scope_items", [])
                    payment = seg_output.get("payment_terms", [])
                    snap = {k: seg_output.get(k) for k in [
                        "tender_no","org_name","estimated_cost","emd",
                        "bid_submission_date","bid_opening_date","prebid_meeting",
                        "tender_fee","contract_period","location"
                    ] if seg_output.get(k)}
                    elapsed = round(_t2.time() - _ai_start_t, 1)
                    if snap:
                        _set_job(job_id,
                                 progress=f"✓ Snapshot — dates, EMD, cost ({elapsed}s)",
                                 step=3,
                                 segments={"snapshot": snap},
                                 seg_log={"snapshot": {"label":"Snapshot — dates, EMD, cost",
                                                        "elapsed":elapsed,"count":"","done":True}})
                    if pq:
                        _set_job(job_id,
                                 progress=f"✓ PQ / Eligibility Criteria — {len(pq)} criteria ({elapsed}s)",
                                 step=4,
                                 segments={"pq": {"pq_criteria": pq}},
                                 seg_log={"pq": {"label":"PQ / Eligibility Criteria",
                                                  "elapsed":elapsed,"count":f"{len(pq)} criteria","done":True}})
                    if tq:
                        _set_job(job_id,
                                 progress=f"✓ TQ / Marking Criteria — {len(tq)} criteria ({elapsed}s)",
                                 step=5,
                                 segments={"tq": {"tq_criteria": tq}},
                                 seg_log={"tq": {"label":"TQ / Marking Criteria",
                                                  "elapsed":elapsed,"count":f"{len(tq)} criteria","done":True}})
                    if scope:
                        _set_job(job_id,
                                 progress=f"✓ Scope of Work — {len(scope)} items ({elapsed}s)",
                                 step=5,
                                 segments={"scope": {"scope_items": scope}},
                                 seg_log={"scope": {"label":"Scope of Work",
                                                     "elapsed":elapsed,"count":f"{len(scope)} items","done":True}})
                    if payment:
                        _set_job(job_id,
                                 progress=f"✓ Payment Schedule — {len(payment)} milestones ({elapsed}s)",
                                 step=5,
                                 segments={"payment": {"payment_terms": payment}},
                                 seg_log={"payment": {"label":"Payment Schedule",
                                                       "elapsed":elapsed,"count":f"{len(payment)} milestones","done":True}})

            try:
                import concurrent.futures as _cf
                with _cf.ThreadPoolExecutor(max_workers=1) as _ex:
                    if PARALLEL_ANALYST_AVAILABLE:
                        _fut = _ex.submit(analyze_with_gemini_parallel, all_text, passed, _seg_progress)
                    else:
                        _fut = _ex.submit(analyze_with_gemini, all_text, passed, _ai_done_cb)
                    ai_result = _fut.result(timeout=420)  # 7 min hard cap (Render kills at 10min)
            except _cf.TimeoutError:
                ai_result = {"error": "AI analysis timed out after 7 minutes — quota may be exhausted"}
                _set_job(job_id, progress="AI timed out — returning basic extraction…")
            except Exception as _ai_exc:
                ai_result = {"error": str(_ai_exc)}
            finally:
                _ai_done_evt.set()  # stop ticker thread

            try:
                in_tok = len(all_text) // 4
                record_token_usage(in_tok, 2000)
            except Exception:
                pass
            if "error" not in ai_result:
                tender_data = merge_results(tender_data, ai_result, passed)
                ai_used = True
            else:
                err_msg = ai_result.get("error", "AI error")
                import logging as _log
                _log.getLogger(__name__).error(f"AI FAILED job={job_id}: {err_msg}")
                if "429" in err_msg or "quota" in err_msg.lower():
                    warn = ("Gemini quota exhausted. All API keys hit their daily limit. "
                            "Add a fresh key from aistudio.google.com or wait until tomorrow.")
                elif "timeout" in err_msg.lower():
                    warn = "AI analysis timed out (7 min). Try again — large PDFs can be slow."
                elif "invalid JSON" in err_msg or "JSONDecodeError" in err_msg:
                    warn = ("AI returned malformed JSON — response may have been cut off. "
                            "This usually means the tender is very large. Try again.")
                elif "No Gemini API key" in err_msg or not api_key:
                    warn = "Gemini API key not configured. Go to Settings to add your key from aistudio.google.com."
                else:
                    warn = f"AI error: {err_msg[:200]}"
                tender_data["ai_warning"] = warn
                _set_job(job_id, progress=f"AI unavailable: {err_msg[:80]}")
        elif not api_key:
            tender_data["ai_warning"] = "Gemini API key not configured. Go to Settings → add key from aistudio.google.com (free)."

        raw_text_preview = all_text[:20000]
        del all_text  # free corpus memory before eligibility check
        _set_job(job_id, progress="Checking eligibility…", step=6)
        checker = NascentChecker()
        if not tender_data.get("overall_verdict"):
            tender_data["pq_criteria"] = checker.check_all(tender_data.get("pq_criteria", []))
            tender_data["tq_criteria"] = checker.check_all(tender_data.get("tq_criteria", []))
            tender_data["overall_verdict"] = checker.get_overall_verdict(tender_data["pq_criteria"] + tender_data["tq_criteria"])

        _set_job(job_id, progress="Generating Word report…", step=8)
        output_filename = ""
        doc_b64 = ""
        try:
            import base64
            generator = BidDocGenerator()
            safe_no = re.sub(r'[^\w\-]', '_', tender_data.get("tender_no", "Report"))[:50]
            output_filename = f"BidNoBid_{safe_no}.docx"
            out_path = str(OUTPUT_DIR / output_filename)
            generator.generate(tender_data, out_path)
            # Read back as base64 so frontend can download directly (Render disk is ephemeral)
            with open(out_path, "rb") as docf:
                doc_b64 = base64.b64encode(docf.read()).decode("utf-8")
        except Exception as doc_err:
            tender_data["doc_warning"] = f"Word report failed: {str(doc_err)[:100]}"

        if t247_id:
            _set_job(job_id, progress="Saving to database…")
            db_record = get_tender(t247_id)
            db_record.update({
                "t247_id": t247_id,
                "tender_no": tender_data.get("tender_no"),
                "org_name": tender_data.get("org_name"),
                "tender_name": tender_data.get("tender_name"),
                "bid_submission_date": tender_data.get("bid_submission_date"),
                "emd": tender_data.get("emd"),
                "estimated_cost": tender_data.get("estimated_cost"),
                "verdict": (tender_data.get("overall_verdict") or {}).get("verdict", ""),
                "verdict_color": (tender_data.get("overall_verdict") or {}).get("color", ""),
                "bid_no_bid_done": True,
                "report_file": output_filename,
                "analysed_at": datetime.now().isoformat(),
                "has_corrigendum": tender_data.get("has_corrigendum", False),
                "ai_used": ai_used,
                "scope_items": tender_data.get("scope_items", []),
                "contract_period": tender_data.get("contract_period", ""),
                "pq_criteria": tender_data.get("pq_criteria", []),
                "tq_criteria": tender_data.get("tq_criteria", []),
                "payment_terms": tender_data.get("payment_terms", []),
                "notes": tender_data.get("notes", []),
                "overall_verdict": tender_data.get("overall_verdict", {}),
                "prebid_queries": tender_data.get("prebid_queries", []),
                "raw_text": raw_text_preview,
                "prebid_passed": passed,
                "scope_sections": tender_data.get("scope_sections", []),
                "tq_total_marks": tender_data.get("tq_total_marks"),
                "tq_nascent_estimated_total": tender_data.get("tq_nascent_estimated_total"),
                "submission_checklist": tender_data.get("submission_checklist", []),
            })
            save_tender(t247_id, db_record)

        _set_job(job_id,
            status="done",
            progress="Complete",
            result={
                "status": "success",
                "ai_used": ai_used,
                "has_corrigendum": tender_data.get("has_corrigendum", False),
                "files_processed": [fc[0] for fc in file_contents],
                "tender_data": tender_data,
                "download_file": output_filename,
                "doc_b64": doc_b64,
                "doc_filename": output_filename,
            }
        )
        job_ok = True

    except Exception as e:
        import traceback
        _set_job(job_id, status="error", error=str(e), traceback=traceback.format_exc())
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        if slot_held and API_POOL_AVAILABLE:
            try:
                get_slots().release(job_ok)
            except Exception:
                pass



# ══ GENERATE DOCS ════════════════════════════════════════════════════════════
@app.post("/generate-docs/{t247_id}")
async def generate_docs(t247_id: str):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found. Analyse the tender first.")
    try:
        import base64, io, zipfile as _zf
        files_b64 = []   # [{name, filename, b64}]
        zip_b64 = ""
        zip_filename = ""

        if SUBMISSION_GEN_AVAILABLE:
            pkg = generate_submission_package(tender, OUTPUT_DIR)
            if "error" not in pkg and pkg.get("files"):
                pkg_dir = Path(pkg.get("pkg_dir", ""))
                # Encode each docx as base64
                for item in pkg.get("files", []):
                    fpath = pkg_dir / item["filename"]
                    if fpath.exists():
                        b64 = base64.b64encode(fpath.read_bytes()).decode()
                        files_b64.append({
                            "name": item["name"],
                            "filename": item["filename"],
                            "b64": b64,
                            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        })
                # Encode ZIP
                if pkg.get("zip_file"):
                    zip_path = OUTPUT_DIR / pkg["zip_file"]
                    if zip_path.exists():
                        zip_b64 = base64.b64encode(zip_path.read_bytes()).decode()
                        zip_filename = pkg["zip_file"]

                return {
                    "status": "success",
                    "files": files_b64,
                    "zip_b64": zip_b64,
                    "zip_filename": zip_filename,
                    "doc_count": len(files_b64),
                }

        # Fallback — just bid/no-bid report
        generator = BidDocGenerator()
        safe_no = re.sub(r'[^\w\-]', '_', tender.get("tender_no") or tender.get("brief", "Report"))[:50]
        output_filename = f"BidNoBid_{safe_no}.docx"
        tender_data = dict(tender)
        if not tender_data.get("overall_verdict") and tender.get("verdict"):
            tender_data["overall_verdict"] = {
                "verdict": tender.get("verdict", ""),
                "color": tender.get("verdict_color", "BLUE"),
                "reason": tender.get("reason", ""),
                "green": 0, "amber": 0, "red": 0
            }
        out_path = OUTPUT_DIR / output_filename
        generator.generate(tender_data, str(out_path))
        b64 = base64.b64encode(out_path.read_bytes()).decode()
        return {
            "status": "success",
            "files": [{"name": "Bid/No-Bid Report", "filename": output_filename, "b64": b64,
                       "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}],
            "zip_b64": "",
            "zip_filename": "",
            "doc_count": 1,
        }
    except Exception as e:
        import traceback
        raise HTTPException(500, f"Document generation failed: {str(e)}\n{traceback.format_exc()[:400]}")


@app.post("/generate-technical-proposal/{t247_id}")
async def generate_technical_proposal(t247_id: str):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    if not TECH_PROPOSAL_AVAILABLE:
        raise HTTPException(500, "technical_proposal_generator not available")
    safe_no = re.sub(r'[^\w\-]', '_', str(tender.get("tender_no", t247_id)))[:30]
    out_path = str(OUTPUT_DIR / f"TechProposal_{safe_no}.docx")
    result = _gen_tech_proposal(tender, out_path)
    if result.get("status") == "error":
        raise HTTPException(500, result.get("message", "Generation failed"))
    filename = Path(out_path).name
    # also store b64 for ephemeral-safe download
    try:
        with open(out_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
    except Exception:
        b64 = None
    return {"status": "success", "filename": filename, "doc_b64": b64, "pages": result.get("pages")}

# alias used by frontend
@app.post("/generate-tech-proposal/{t247_id}")
async def generate_tech_proposal_short(t247_id: str):
    return await generate_technical_proposal(t247_id)

@app.post("/tender/{t247_id}/technical-proposal")
async def generate_technical_proposal_alias(t247_id: str):
    return await generate_technical_proposal(t247_id)

@app.post("/merge-submission-pdf/{t247_id}")
async def merge_submission_pdf(t247_id: str):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    if not PDF_MERGE_AVAILABLE:
        raise HTTPException(500, "pdf_merger not available")

    safe_no = re.sub(r'[^\w\-]', '_', str(tender.get("tender_no", t247_id)))[:30]
    pkg_dir = OUTPUT_DIR / f"SubmissionPackage_{safe_no}"
    source_dirs = [pkg_dir, OUTPUT_DIR]
    source_dirs = [d for d in source_dirs if d.exists()]
    if not source_dirs:
        raise HTTPException(400, "No generated documents found. Generate docs first.")

    merged = merge_submission_package(
        t247_id=t247_id,
        tender_data=tender,
        source_dirs=source_dirs,
        output_dir=OUTPUT_DIR,
        include_cover=True,
    )
    if merged.get("status") != "success":
        errs = merged.get("errors", ["Merge failed"])
        # Return a non-5xx response so UI can show a friendly message instead of hard failure.
        return {
            "status": "unavailable",
            "message": "Merged PDF could not be generated in this environment.",
            "errors": errs,
        }

    fname = merged.get("filename")
    return {
        "status": "success",
        "filename": fname,
        "download_url": f"/download/{fname}" if fname else "",
        "page_count": merged.get("page_count", 0),
        "file_count": merged.get("file_count", 0),
    }

@app.post("/tender/{t247_id}/merge-pdf")
async def merge_submission_pdf_alias(t247_id: str):
    return await merge_submission_pdf(t247_id)

@app.post("/tender/{t247_id}/draft/generate")
async def generate_tender_draft(t247_id: str, data: dict = Body(default={})):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    doc_type = _safe_doc_type(data.get("doc_type", "analysis"))
    content = _build_tender_draft(tender, doc_type)
    path = _draft_path(t247_id, doc_type)
    path.write_text(content, encoding="utf-8")
    return {"status": "success", "t247_id": t247_id, "doc_type": doc_type, "filename": path.name, "content": content}

@app.get("/tender/{t247_id}/draft")
async def get_tender_draft(t247_id: str, doc_type: str = "analysis"):
    doc_type = _safe_doc_type(doc_type)
    path = _draft_path(t247_id, doc_type)
    if not path.exists():
        tender = get_tender(t247_id)
        if not tender:
            raise HTTPException(404, "Tender not found")
        content = _build_tender_draft(tender, doc_type)
        path.write_text(content, encoding="utf-8")
    return {"status": "success", "t247_id": t247_id, "doc_type": doc_type, "filename": path.name, "content": path.read_text(encoding="utf-8")}

@app.post("/tender/{t247_id}/draft/save")
async def save_tender_draft(t247_id: str, data: dict = Body(...)):
    content = str(data.get("content", ""))
    doc_type = _safe_doc_type(data.get("doc_type", "analysis"))
    if not content.strip():
        raise HTTPException(400, "Draft content is empty")
    path = _draft_path(t247_id, doc_type)
    path.write_text(content, encoding="utf-8")
    return {"status": "saved", "filename": path.name, "size_kb": round(path.stat().st_size / 1024, 1)}

@app.post("/tender/{t247_id}/draft/chat-edit")
async def chat_edit_tender_draft(t247_id: str, data: dict = Body(...)):
    instruction = str(data.get("instruction", "")).strip()
    doc_type = _safe_doc_type(data.get("doc_type", "analysis"))
    content = str(data.get("content", "")).strip()
    if not instruction:
        raise HTTPException(400, "Edit instruction is required")
    if not content:
        draft = await get_tender_draft(t247_id, doc_type)
        content = draft.get("content", "")

    cfg = load_config()
    keys = get_all_api_keys()
    if not keys:
        raise HTTPException(400, "No AI API key configured in Settings")

    prompt = (
        "You are an expert tender document editor.\n"
        "Rewrite the draft as per the user's instruction.\n"
        "Keep factual details intact unless user asks to change.\n"
        "Return only updated draft text.\n\n"
        f"USER INSTRUCTION:\n{instruction}\n\n"
        f"CURRENT DRAFT:\n{content}"
    )

    edited = None
    for key in keys:
        try:
            out = call_gemini(prompt, key)
            if out and len(out.strip()) > 10:
                edited = out.strip()
                break
        except Exception:
            continue
    if not edited:
        raise HTTPException(503, "AI edit unavailable right now (quota or API issue)")

    path = _draft_path(t247_id, doc_type)
    path.write_text(edited, encoding="utf-8")
    return {"status": "success", "t247_id": t247_id, "doc_type": doc_type, "filename": path.name, "content": edited}

@app.get("/tender/{t247_id}/draft/download")
async def download_tender_draft(t247_id: str, doc_type: str = "analysis"):
    doc_type = _safe_doc_type(doc_type)
    path = _draft_path(t247_id, doc_type)
    if not path.exists():
        raise HTTPException(404, "Draft not found")
    return FileResponse(str(path), filename=path.name, media_type="text/markdown")

# ══ DOWNLOAD ─────────────────────────────────────────────────────────────────
@app.get("/download/{filename}")
async def download_file(filename: str):
    # FIX: Path traversal guard — resolve and assert inside OUTPUT_DIR
    fp = (OUTPUT_DIR / Path(filename).name).resolve()
    output_resolved = OUTPUT_DIR.resolve()
    if not str(fp).startswith(str(output_resolved)):
        raise HTTPException(400, "Invalid filename")
    if not fp.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(
        path=str(fp),
        filename=fp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/reports")
async def list_reports():
    return [{"filename": f.name, "size_kb": round(f.stat().st_size / 1024, 1),
             "created": datetime.fromtimestamp(f.stat().st_mtime).strftime("%d %b %Y %H:%M")}
            for f in sorted(OUTPUT_DIR.glob("BidNoBid_*.docx"), key=lambda f: f.stat().st_mtime, reverse=True)[:100]]

@app.get("/reports-list")
async def reports_list():
    try:
        db = load_db()
        reports = []
        for fname in sorted(OUTPUT_DIR.glob("BidNoBid_*.docx"), reverse=True):
            tender = next((t for tid, t in db["tenders"].items() if tid in fname.stem or (t.get("tender_no", "") and t.get("tender_no", "").replace("/", "_") in fname.stem)), None)
            reports.append({
                "filename": fname.name,
                "created": datetime.fromtimestamp(fname.stat().st_mtime).strftime("%d-%b-%Y %H:%M"),
                "size_kb": round(fname.stat().st_size / 1024, 1),
                "t247_id": tender.get("t247_id", "—") if tender else "—",
                "tender_name": tender.get("brief", "")[:60] if tender else fname.stem[:60],
                "org": tender.get("org_name", "—") if tender else "—",
                "verdict": tender.get("verdict", "—") if tender else "—",
            })
        return {"reports": reports}
    except Exception as e:
        return {"reports": [], "error": str(e)}

# ══ CONFIG (admin-guarded) ════════════════════════════════════════════════════
@app.get("/config-full")
@app.get("/config")
async def get_config_route(request: Request):
    check_admin(request)
    config = load_config()
    key = config.get("gemini_api_key", "")
    keys = config.get("gemini_api_keys", [])
    # Ensure primary key is first in list
    if key and key not in keys:
        keys = [key] + keys
    groq_key = config.get("groq_api_key", "")
    return {
        "gemini_api_key_set": bool(key),
        "gemini_api_key": key,
        "gemini_api_key_preview": (key[:8] + "..." + key[-4:]) if key else "",
        "gemini_api_keys": keys,
        "gemini_api_key_2": keys[1] if len(keys) > 1 else "",
        "gemini_api_key_3": keys[2] if len(keys) > 2 else "",
        "gemini_api_key_4": keys[3] if len(keys) > 3 else "",
        "groq_api_key": groq_key,
        "t247_username": config.get("t247_username", ""),
        "t247_auto_sync_enabled": bool(config.get("t247_auto_sync_enabled", True)),
        "t247_auto_sync_minutes": int(config.get("t247_auto_sync_minutes", 180) or 180),
        "daily_digest_enabled": bool(config.get("daily_digest_enabled", True)),
        "daily_digest_hour": int(config.get("daily_digest_hour", 9) or 9),
        "daily_digest_minute": int(config.get("daily_digest_minute", 0) or 0),
    }

@app.post("/config")
async def update_config_route(request: Request, data: dict = Body(...)):
    check_admin(request)
    config = load_config()
    if data.get("gemini_api_key"):
        config["gemini_api_key"] = data["gemini_api_key"]
    if data.get("gemini_api_keys"):
        keys = [k.strip() for k in data["gemini_api_keys"] if k and k.strip()]
        config["gemini_api_keys"] = keys
        if keys:
            config["gemini_api_key"] = keys[0]
    # Also accept individual key fields from UI (key1/key2/key3/key4)
    ui_keys = []
    for field in ["gemini_api_key", "gemini_api_key_2", "gemini_api_key_3", "gemini_api_key_4"]:
        v = str(data.get(field, "") or "").strip()
        if v and len(v) > 20:
            ui_keys.append(v)
    if ui_keys and not data.get("gemini_api_keys"):
        config["gemini_api_key"] = ui_keys[0]
        config["gemini_api_keys"] = ui_keys
    if data.get("groq_api_key"):
        config["groq_api_key"] = data["groq_api_key"].strip()
    for src, dst in [("t247_username","t247_username"),("t247_user","t247_username"),
                     ("t247_password","t247_password"),("t247_pass","t247_password")]:
        if data.get(src):
            config[dst] = data[src]
    if "t247_auto_sync_enabled" in data:
        config["t247_auto_sync_enabled"] = bool(data.get("t247_auto_sync_enabled"))
    if "t247_auto_sync_minutes" in data:
        try:
            mins = int(data.get("t247_auto_sync_minutes") or 180)
            config["t247_auto_sync_minutes"] = max(15, min(720, mins))
        except Exception:
            pass
    if "daily_digest_enabled" in data:
        config["daily_digest_enabled"] = bool(data.get("daily_digest_enabled"))
    if "daily_digest_hour" in data:
        try:
            config["daily_digest_hour"] = max(0, min(23, int(data.get("daily_digest_hour") or 9)))
        except Exception:
            pass
    if "daily_digest_minute" in data:
        try:
            config["daily_digest_minute"] = max(0, min(59, int(data.get("daily_digest_minute") or 0)))
        except Exception:
            pass
    save_config(config)
    return {"status": "saved"}

# ══ PROFILE ══════════════════════════════════════════════════════════════════
@app.get("/profile")
async def get_profile():
    from nascent_checker import load_profile
    profile = load_profile()
    projects = profile.get("projects", []) or []
    if not profile.get("project_tabs"):
        grouped = {}
        for p in projects:
            tab = str((p or {}).get("tab", "General") or "General").strip() or "General"
            grouped.setdefault(tab, []).append(p)
        profile["project_tabs"] = [{"name": name, "projects": items} for name, items in grouped.items()] or [{"name": "General", "projects": []}]
    return profile

@app.post("/profile")
async def update_profile(data: dict = Body(...)):
    project_tabs = data.get("project_tabs", []) or []
    if project_tabs:
        flat_projects = []
        for tab in project_tabs:
            tab_name = str((tab or {}).get("name", "General")).strip() or "General"
            for p in (tab or {}).get("projects", []) or []:
                if isinstance(p, dict):
                    item = dict(p)
                    item["tab"] = tab_name
                    flat_projects.append(item)
        data["projects"] = flat_projects

    payload = json.dumps(data, indent=2, ensure_ascii=False)
    payload_bytes = payload.encode("utf-8")

    # Write to all locations
    repo_path = BASE_DIR / "nascent_profile.json"
    try:
        repo_path.write_text(payload, encoding="utf-8")
    except Exception as e:
        print(f"⚠️ repo profile write failed: {e}")

    runtime_dir = Path(os.environ.get("BIDNOBID_RUNTIME_DIR", "/tmp/bid-nobid"))
    runtime_dir.mkdir(parents=True, exist_ok=True)
    runtime_path = runtime_dir / "nascent_profile.json"
    try:
        runtime_path.write_text(payload, encoding="utf-8")
    except Exception as e:
        print(f"⚠️ runtime profile write failed: {e}")

    # Drive backup — this is the ONLY one that survives redeploy
    drive_saved = False
    try:
        if drive_available():
            prof_tmp = OUTPUT_DIR / "nascent_profile.json"
            OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
            prof_tmp.write_bytes(payload_bytes)
            drive_saved = save_to_drive(prof_tmp, filename="nascent_profile.json")
            if drive_saved:
                print("✅ Profile saved to Drive")
            else:
                print("⚠️ Profile Drive save failed")
    except Exception as e:
        print(f"⚠️ Profile Drive backup failed: {e}")

    try:
        from excel_processor import invalidate_rules_cache
        invalidate_rules_cache()
    except (ImportError, AttributeError):
        pass

    return {"status": "saved", "drive_synced": drive_saved}


@app.post("/profile/import-excel")
async def import_profile_excel(file: UploadFile = File(...)):
    """Parse an Excel file and extract company profile data (projects, company info)."""
    if not file.filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(400, "Only .xlsx / .xls files accepted")
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(500, "openpyxl not installed — cannot read Excel")

    content = await file.read()
    tmp = TEMP_DIR / f"profile_import_{uuid.uuid4().hex}.xlsx"
    try:
        tmp.write_bytes(content)
        wb = openpyxl.load_workbook(tmp, data_only=True)
    except Exception as e:
        raise HTTPException(400, f"Cannot open Excel: {e}")
    finally:
        try: tmp.unlink()
        except Exception: pass

    projects = []
    company_updates = {}

    # Look for project data in any sheet
    PROJECT_COLS = {
        "title": ["title","project","project name","name","work","description","tender name"],
        "client": ["client","department","dept","owner","organisation","org","authority","ministry"],
        "sector": ["sector","domain","type","category","technology","area"],
        "value_cr": ["value","amount","cost","cr","crore","contract value","₹ cr","value (cr)"],
        "year": ["year","fy","financial year","completion","awarded"],
    }
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows: continue

        # Find header row (first row with >2 non-empty cells)
        header_row_idx = None
        for idx, row in enumerate(rows[:10]):
            non_empty = [c for c in row if c is not None and str(c).strip()]
            if len(non_empty) >= 2:
                header_row_idx = idx
                break
        if header_row_idx is None: continue

        headers = [str(h).strip().lower() if h else "" for h in rows[header_row_idx]]

        # Map columns
        col_map = {}
        for field, aliases in PROJECT_COLS.items():
            for i, h in enumerate(headers):
                if any(a in h for a in aliases):
                    col_map[field] = i
                    break

        # If we found at least title + client or title + value, treat as project sheet
        if "title" not in col_map:
            continue

        for row in rows[header_row_idx + 1:]:
            if not any(c for c in row if c is not None):
                continue
            def cell(field):
                idx = col_map.get(field)
                return str(row[idx]).strip() if idx is not None and row[idx] is not None else ""

            title = cell("title")
            if not title or title.lower() in ("none","nan","—","-",""):
                continue
            val_raw = cell("value_cr")
            try: val_cr = float(re.sub(r"[^\d.]", "", val_raw)) if val_raw else 0.0
            except Exception: val_cr = 0.0

            projects.append({
                "title": title,
                "client": cell("client"),
                "sector": cell("sector"),
                "value_cr": val_cr,
                "year": cell("year"),
            })

    # Deduplicate by title
    seen_titles = set()
    unique_projects = []
    for p in projects:
        key = p["title"].lower().strip()
        if key not in seen_titles:
            seen_titles.add(key)
            unique_projects.append(p)

    return {
        "status": "ok",
        "sheets_scanned": len(wb.sheetnames),
        "projects": unique_projects,
        "company": company_updates if company_updates else None,
        "message": f"Found {len(unique_projects)} projects across {len(wb.sheetnames)} sheets",
    }


# ══ BOQ ══════════════════════════════════════════════════════════════════════
@app.get("/boq/constants")
async def boq_constants():
    return get_boq_constants()

@app.get("/boq/{t247_id}")
async def get_boq(t247_id: str):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    boq = tender.get("boq")
    name = tender.get("tender_name") or tender.get("brief", "")
    if boq:
        return {"t247_id": t247_id, "tender_name": name, "boq": boq, "source": "saved"}
    items = extract_boq_from_scope(tender)
    return {"t247_id": t247_id, "tender_name": name, "boq": {"items": items, "margin_pct": 15.0, "gst_pct": 18.0}, "source": "auto"}

@app.post("/boq/{t247_id}")
async def save_boq(t247_id: str, data: dict = Body(...)):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    result = calculate_boq_totals(data.get("items", []), float(data.get("margin_pct", 15)), float(data.get("gst_pct", 18)))
    tender["boq"] = {**result, "margin_pct": data.get("margin_pct", 15), "gst_pct": data.get("gst_pct", 18), "saved_at": datetime.now().isoformat()}
    save_tender(t247_id, tender)
    return {"status": "saved", "totals": result}

@app.post("/boq/{t247_id}/regenerate")
async def regenerate_boq(t247_id: str):
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    return {"t247_id": t247_id, "boq": {"items": extract_boq_from_scope(tender), "margin_pct": 15.0, "gst_pct": 18.0}, "source": "regenerated"}

# ══ CHECKLIST ════════════════════════════════════════════════════════════════
@app.get("/checklist/{t247_id}")
async def get_checklist(t247_id: str):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    if "doc_checklist" in t:
        return {"checklist": t["doc_checklist"], "t247_id": t247_id}
    return {"checklist": generate_doc_checklist(t), "t247_id": t247_id}

@app.post("/checklist/{t247_id}")
async def save_checklist(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    t["doc_checklist"] = data.get("checklist", [])
    pct = round(sum(1 for d in t["doc_checklist"] if d.get("done")) / max(len(t["doc_checklist"]), 1) * 100)
    t["checklist_pct"] = pct
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved", "completion_pct": pct}

@app.post("/checklist/{t247_id}/item")
async def toggle_checklist_item(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    cl = t.get("doc_checklist", [])
    for item in cl:
        if str(item.get("id")) == str(data.get("id")):
            item["done"] = data.get("done", False)
            break
    t["doc_checklist"] = cl
    pct = round(sum(1 for d in cl if d.get("done")) / max(len(cl), 1) * 100)
    t["checklist_pct"] = pct
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved", "completion_pct": pct}

# ══ PIPELINE / ANALYTICS ═════════════════════════════════════════════════════
@app.get("/alerts")
async def get_alerts():
    return {"alerts": get_deadline_alerts()}

@app.get("/pipeline")
async def get_pipeline():
    return {"stages": get_pipeline_stats(), "stage_list": PIPELINE_STAGES, "stage_colors": STAGE_COLORS}

@app.get("/win-loss")
async def get_win_loss():
    return get_win_loss_stats()

@app.get("/analytics/win-loss")
async def get_win_loss_alias():
    return get_win_loss_stats()

@app.get("/analytics")
async def get_analytics_alias():
    return get_win_loss_stats()

@app.post("/tender/{t247_id}/stage")
async def update_stage(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    if "status" in data: t["status"] = data["status"]
    if "notes" in data: t["notes_internal"] = data["notes"]
    if "outcome_value" in data: t["outcome_value"] = data["outcome_value"]
    if "outcome_notes" in data: t["outcome_notes"] = data["outcome_notes"]
    t["status_updated_at"] = datetime.now().isoformat()
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved", "new_stage": t.get("status")}

@app.post("/bid-result/{t247_id}")
@app.post("/tender/{t247_id}/bid-result")
async def save_bid_result(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    t.update({"outcome": data.get("outcome", ""), "outcome_value": data.get("value", ""),
              "outcome_competitor": data.get("competitor", ""), "outcome_notes": data.get("notes", ""),
              "outcome_date": datetime.now().isoformat()})
    if data.get("outcome") == "Won": t["status"] = "Won"
    elif data.get("outcome") == "Lost": t["status"] = "Lost"
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved"}

@app.post("/bid-result")
async def save_bid_result_without_path(data: dict = Body(...)):
    t247_id = str(data.get("t247_id", "")).strip()
    if not t247_id:
        raise HTTPException(400, "t247_id is required")
    return await save_bid_result(t247_id, data)

# ══ PRE-BID ══════════════════════════════════════════════════════════════════
@app.post("/prebid-sent/{t247_id}")
async def mark_prebid_sent(t247_id: str, data: dict = Body(...)):
    db = load_db()
    t = db["tenders"].get(t247_id, {})
    t.update({"prebid_sent": True, "prebid_sent_at": datetime.now().isoformat(),
              "prebid_sent_to": data.get("email", ""), "status": "Pre-bid Sent",
              "status_updated_at": datetime.now().isoformat()})
    db["tenders"][t247_id] = t
    save_db(db)
    return {"status": "saved"}

# ══ CHAT ════════════════════════════════════════════════════════════════════
@app.post("/chat")
async def chat(data: dict = Body(...)):
    message = data.get("message", "").strip()
    if not message:
        raise HTTPException(400, "Empty message")
    return process_message(message, load_history())

@app.get("/chat/history")
async def get_chat_history():
    return {"history": load_history()}

@app.delete("/chat/history")
async def clear_chat_history():
    h = OUTPUT_DIR / "chat_history.json"
    if h.exists():
        h.unlink()
    return {"status": "cleared"}

# ══ SKIPPED ═════════════════════════════════════════════════════════════════
@app.get("/skipped")
async def get_skipped():
    db = load_db()
    return {"skipped": [t for t in db["tenders"].values() if t.get("status") == "Not Interested"]}

@app.get("/skipped-tenders")
async def get_skipped_alias():
    db = load_db()
    return {"tenders": [t for t in db["tenders"].values() if t.get("status") == "Not Interested"]}

# ══ DRIVE (admin-guarded) ════════════════════════════════════════════════════
@app.post("/sync-drive")
async def sync_drive(request: Request):
    check_admin(request)
    if not drive_available():
        return JSONResponse({"status": "error", "message": "Google Drive not connected"}, status_code=400)
    try:
        db = load_db()
        ok = save_to_drive(DB_FILE)
        if ok:
            return {"status": "ok", "message": f"Synced {len(db.get('tenders', {}))} tenders to Drive"}
        return JSONResponse({"status": "error", "message": "Sync failed"}, status_code=500)
    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)

@app.post("/sync-sheets")
async def sync_sheets(request: Request):
    return await sync_drive(request)

@app.get("/drive-status")
async def drive_status():
    db = load_db()
    return {"drive_connected": drive_available(),
            "tenders_in_memory": len(db.get("tenders", {})),
            "db_file_exists": DB_FILE.exists(),
            "db_size_kb": round(DB_FILE.stat().st_size / 1024) if DB_FILE.exists() else 0}

@app.get("/letterhead/status")
async def letterhead_status():
    lh_docx = OUTPUT_DIR / "letterhead.docx"
    return {
        "has_letterhead": lh_docx.exists(),
        "filename": lh_docx.name if lh_docx.exists() else "",
        "size_kb": round(lh_docx.stat().st_size / 1024, 1) if lh_docx.exists() else 0,
    }

@app.get("/letterhead-status")
async def letterhead_status_alias():
    return await letterhead_status()

@app.post("/letterhead/upload")
async def upload_letterhead(file: UploadFile = File(...)):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(400, "Upload a .docx file for letterhead template")
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 50MB)")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    dest = OUTPUT_DIR / "letterhead.docx"
    dest.write_bytes(content)
    return {"status": "ok", "filename": dest.name, "size_kb": round(dest.stat().st_size / 1024, 1)}

@app.post("/upload-letterhead")
async def upload_letterhead_alias(file: UploadFile = File(...)):
    return await upload_letterhead(file)

@app.post("/upload-db")
async def upload_db(request: Request, file: UploadFile = File(...)):
    check_admin(request)
    try:
        content = await file.read()
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(413, "File too large")
        data = json.loads(content)
        count = len(data.get("tenders", {}))
        if count == 0:
            raise HTTPException(400, "File has 0 tenders")
        DB_FILE.write_bytes(content)
        drive_ok = save_to_drive(DB_FILE) if drive_available() else False
        return {"status": "ok", "tenders": count, "drive_saved": drive_ok}
    except json.JSONDecodeError:
        raise HTTPException(400, "Invalid JSON file")

# ══ EXPORT ══════════════════════════════════════════════════════════════════
@app.get("/export-tenders")
async def export_tenders(verdict: str = "", search: str = ""):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        db = load_db()
        tenders = list(db["tenders"].values())
        if verdict and verdict != "ALL":
            tenders = [t for t in tenders if t.get("verdict") == verdict]
        if search:
            s = search.lower()
            tenders = [t for t in tenders if any(s in str(t.get(f, "")).lower() for f in ["t247_id","ref_no","brief","org_name","location","verdict"])]
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Tenders"
        headers = ["Sr.","T247 ID","Reference No.","Brief","Organization","Location","Cost (Cr)","EMD","Doc Fee","MSME Exempt","Deadline","Days Left","Verdict","Stage","Analysed","BOQ","Checklist %","Reason"]
        col_widths = [5,12,25,45,30,20,10,12,10,12,14,10,14,18,10,8,12,35]
        hdr_fill = PatternFill("solid", fgColor="1E2A3B")
        hdr_font = Font(bold=True, color="FFFFFF", size=11)
        for ci, (hdr, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.column_dimensions[cell.column_letter].width = w
        ws.row_dimensions[1].height = 30
        vc = {"BID": "E2EFDA", "CONDITIONAL": "FFF2CC", "NO-BID": "FCE4D6", "REVIEW": "DEEAF1"}
        def dl(t):
            try:
                s = t.get("deadline", "")
                for fmt in ["%d-%m-%Y", "%d/%m/%Y"]:
                    try:
                        return (datetime.strptime(s.split()[0], fmt).date() - date.today()).days
                    except Exception:
                        continue
            except Exception:
                pass
            return 999
        for ri, t in enumerate(sorted(tenders, key=dl), 2):
            days = dl(t)
            v = t.get("verdict", "")
            rf = PatternFill("solid", fgColor=vc.get(v, "FFFFFF"))
            vals = [ri-1, t.get("t247_id",""), t.get("ref_no",""), t.get("brief",""), t.get("org_name",""), t.get("location",""),
                    t.get("estimated_cost_cr",""), t.get("emd",""), t.get("doc_fee",""), t.get("msme_exemption",""),
                    t.get("deadline",""), days if days < 999 else "—", v, t.get("status","Identified"),
                    "Yes" if t.get("bid_no_bid_done") else "No", "Yes" if t.get("boq") else "No",
                    str(t.get("checklist_pct","0")) + "%", t.get("reason","")[:100]]
            for ci, val in enumerate(vals, 1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.fill = rf
                cell.alignment = Alignment(vertical="center", wrap_text=True)
        ws.freeze_panes = "A2"
        fname = f"Tenders_Export_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx"
        fpath = OUTPUT_DIR / fname
        wb.save(str(fpath))
        return FileResponse(str(fpath), filename=fname, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        raise HTTPException(500, f"Export failed: {str(e)}")

# ══ TEST ════════════════════════════════════════════════════════════════════
@app.get("/test-ai")
async def test_ai():
    """Quick single-key test (legacy)."""
    from ai_analyzer import get_api_key, call_gemini
    key = get_api_key()
    if not key:
        return {"status": "error", "message": "No API key configured"}
    try:
        result = call_gemini('Return this exact JSON: {"status":"ok"}', key)
        return {"status": "success", "gemini_response": result[:80]}
    except Exception as e:
        return {"status": "error", "error": str(e)[:200]}


@app.get("/test-ai-keys")
async def test_ai_keys():
    """Ping every configured key with a tiny 5-token call.
    Returns live status per key — working / rate_limited / invalid / no_key."""
    import urllib.request, urllib.error, json as _json, time as _t
    from ai_analyzer import get_all_api_keys, GEMINI_MODELS

    PING_PROMPT = 'Reply with exactly: {"ok":true}'
    TEST_MODEL  = "gemini-2.0-flash"   # fastest, cheapest ping

    keys = get_all_api_keys()
    if not keys:
        return {"keys": [], "summary": "No keys configured"}

    results = []
    for idx, key in enumerate(keys):
        masked = key[:8] + "…" + key[-4:]
        url = (
            "https://generativelanguage.googleapis.com/v1beta/models/"
            f"{TEST_MODEL}:generateContent?key={key}"
        )
        payload = _json.dumps({
            "contents": [{"parts": [{"text": PING_PROMPT}]}],
            "generationConfig": {"temperature": 0, "maxOutputTokens": 20}
        }).encode()
        req = urllib.request.Request(
            url, data=payload,
            headers={"Content-Type": "application/json"}, method="POST"
        )
        t0 = _t.time()
        try:
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = _json.loads(resp.read())
                ms = int((_t.time() - t0) * 1000)
                results.append({
                    "key_no": idx + 1,
                    "masked": masked,
                    "status": "working",
                    "model": TEST_MODEL,
                    "latency_ms": ms,
                    "color": "green",
                })
        except urllib.error.HTTPError as e:
            ms = int((_t.time() - t0) * 1000)
            body = e.read().decode("utf-8", errors="ignore")[:120]
            if e.code == 429:
                status, color = "rate_limited", "amber"
            elif e.code in [400, 403]:
                status, color = "invalid_key", "red"
            elif e.code == 404:
                status, color = "model_unavailable", "amber"
            else:
                status, color = f"error_{e.code}", "red"
            results.append({
                "key_no": idx + 1,
                "masked": masked,
                "status": status,
                "error": body,
                "latency_ms": ms,
                "color": color,
            })
        except Exception as e:
            results.append({
                "key_no": idx + 1,
                "masked": masked,
                "status": "unreachable",
                "error": str(e)[:120],
                "color": "red",
            })

    working = sum(1 for r in results if r["status"] == "working")
    limited = sum(1 for r in results if r["status"] == "rate_limited")
    return {
        "keys": results,
        "total": len(results),
        "working": working,
        "rate_limited": limited,
        "summary": f"{working}/{len(results)} keys working" + (f", {limited} rate-limited" if limited else ""),
        "note": "Each ping uses ~5 Gemini tokens total"
    }

def _t247_decode_jwt(token: str) -> dict:
    """Decode JWT payload (no signature check — we trust T247's server)."""
    try:
        import base64 as _b64
        payload_b64 = token.split(".")[1]
        payload_b64 += "=" * (4 - len(payload_b64) % 4)
        return json.loads(_b64.b64decode(payload_b64).decode("utf-8"))
    except Exception:
        return {}

def _t247_api_headers(token: str) -> dict:
    return {
        "accept": "*/*",
        "authorization": f"Bearer {token}",
        "content-type": "application/json",
        "origin": "https://www.tender247.com",
        "referer": "https://www.tender247.com/",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/147.0.0.0 Safari/537.36",
    }

# Confirmed from browser DevTools (F12): exact service name and path
_T247_LOGIN_URL = "https://t247_api.tender247.com/apigateway/T247ApiTender/api/auth/login"
_T247_USER_QUERY_URL = "https://t247_api.tender247.com/apigateway/T247ApiTender/api/auth/user-login-query"
_T247_LOGIN_ENDPOINTS = [_T247_LOGIN_URL]  # kept for probe endpoint compatibility

def _t247_auto_login(cfg: dict) -> str:
    """Login to T247 with stored email+password. Two-step: login → user-login-query."""
    import requests as _req
    email = str(cfg.get("t247_email", "") or cfg.get("t247_username", "") or "").strip()
    password = str(cfg.get("t247_password", "") or "").strip()
    if not email or not password:
        raise ValueError("T247 credentials not saved. Go to Settings → T247 Connection and enter your email + password.")

    login_headers = {
        "accept": "application/json, text/plain, */*",
        "content-type": "application/json",
        "origin": "https://www.tender247.com",
        "referer": "https://www.tender247.com/auth",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/147.0.0.0 Safari/537.36",
    }

    # Step 1 — authenticate, get initial JWT
    # verify=False required: t247_api.tender247.com has underscore — SSL cert hostname mismatch
    try:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        resp = _req.post(
            _T247_LOGIN_URL,
            json={"email_id": email, "password": password},
            headers=login_headers,
            timeout=20,
            verify=False,
        )
    except _req.exceptions.ConnectionError as e:
        raise ValueError(f"Cannot reach T247 login server: {e}")
    except _req.exceptions.Timeout:
        raise ValueError("T247 login request timed out.")

    if resp.status_code == 401:
        raise ValueError("T247 login failed — wrong email or password. Check credentials in Settings.")
    if resp.status_code != 200:
        raise ValueError(f"T247 login failed (HTTP {resp.status_code}): {resp.text[:200]}")

    try:
        data = resp.json()
    except Exception:
        raise ValueError(f"T247 login returned non-JSON: {resp.text[:200]}")

    # Extract token — T247 returns it in data.token or top-level token
    token = (
        (data.get("data") or {}).get("token")
        or (data.get("data") or {}).get("accessToken")
        or data.get("token")
        or data.get("accessToken")
        or data.get("access_token")
        or data.get("jwtToken")
        or ""
    )
    if not token and isinstance(data.get("data"), str) and len(data["data"]) > 20:
        token = data["data"]
    token = str(token).strip()
    if token.lower().startswith("bearer "):
        token = token[7:].strip()
    if not token or len(token) < 20:
        raise ValueError(f"T247 login succeeded but no token in response: {str(data)[:300]}")

    jwt_payload = _t247_decode_jwt(token)
    user_id = jwt_payload.get("UserId") or jwt_payload.get("user_id")
    company_service_ids = jwt_payload.get("CompanyServiceids") or []
    company_service_id = company_service_ids[0] if company_service_ids else 1

    # Step 2 — user-login-query (gets subscription/query data, refreshes session)
    query_id = None
    try:
        q_headers = {**login_headers, "authorization": f"Bearer {token}"}
        q_resp = _req.post(
            _T247_USER_QUERY_URL,
            json={"user_id": user_id, "company_service_id": company_service_id, "is_grace": False},
            headers=q_headers,
            timeout=20,
            verify=False,
        )
        if q_resp.status_code == 200:
            q_data = q_resp.json()
            query_id = (
                (q_data.get("data") or {}).get("query_id")
                or (q_data.get("data") or {}).get("queryId")
                or q_data.get("query_id")
                or q_data.get("queryId")
            )
            print(f"T247 user-login-query OK — query_id={query_id}")
        else:
            print(f"T247 user-login-query non-200: {q_resp.status_code} {q_resp.text[:120]}")
    except Exception as e:
        print(f"T247 user-login-query failed (non-fatal): {e}")

    # Persist everything
    fresh_cfg = load_config()
    fresh_cfg["t247_bearer_token"] = token
    fresh_cfg["t247_email"] = email
    fresh_cfg["t247_password"] = password
    if user_id:
        fresh_cfg["t247_user_id"] = user_id
    if query_id:
        fresh_cfg["t247_query_id"] = query_id
    save_config(fresh_cfg)

    exp_str = datetime.fromtimestamp(jwt_payload["exp"]).strftime("%d %b %H:%M") if jwt_payload.get("exp") else "?"
    print(f"✅ T247 auto-login OK — user {user_id} exp {exp_str}")
    return token

    raise ValueError(f"T247 login failed. Check credentials in Settings. Last error: {last_err}")


def _t247_get_token(cfg: dict) -> str:
    """Return valid T247 JWT — auto-login if missing or expired."""
    import time as _t
    token = str(cfg.get("t247_bearer_token", "") or "").strip()
    if token:
        payload = _t247_decode_jwt(token)
        exp = payload.get("exp", 0)
        # If token valid for >5 more minutes, use it
        if not exp or _t.time() < (exp - 300):
            return token
        # Token expiring soon or expired — try auto-refresh if credentials present
        has_creds = bool(
            (cfg.get("t247_email") or cfg.get("t247_username"))
            and cfg.get("t247_password")
        )
        if not has_creds:
            exp_str = datetime.fromtimestamp(exp).strftime('%d %b %Y %H:%M') if exp else "unknown"
            raise ValueError(f"T247 token expired at {exp_str}. Enter your email + password in Settings → T247 Connection to enable auto-refresh.")
    else:
        has_creds = bool(
            (cfg.get("t247_email") or cfg.get("t247_username"))
            and cfg.get("t247_password")
        )
        if not has_creds:
            raise ValueError("T247 not connected. Go to Settings → T247 Connection and enter your email + password.")
    # Auto-login with stored credentials
    return _t247_auto_login(cfg)

def _t247_merge_tenders(tenders: list) -> dict:
    """Merge tender list into DB. Returns {added, updated}."""
    db = load_db()
    added = updated = 0
    added_ids = []
    updated_ids = []
    for t in tenders:
        tid = str(t.get("t247_id", "")).strip()
        if not tid:
            continue
        if tid in db["tenders"]:
            existing = db["tenders"][tid]
            for field in ["ref_no","brief","org_name","location","estimated_cost_raw","estimated_cost_cr",
                          "deadline","days_left","deadline_status","doc_fee","emd","msme_exemption",
                          "eligibility","checklist","is_gem","tender_name"]:
                if t.get(field) is not None:
                    existing[field] = t[field]
            db["tenders"][tid] = existing
            updated += 1
            updated_ids.append(tid)
        else:
            db["tenders"][tid] = t
            added += 1
            added_ids.append(tid)
    save_db(db)
    return {"added": added, "updated": updated, "added_ids": added_ids, "updated_ids": updated_ids}

def _auto_score_tender_v1(tender: dict) -> dict:
    """Rule-based bid/no-bid scoring from Nascent profile."""
    try:
        from nascent_checker import load_profile
        profile = load_profile() or {}
    except Exception:
        profile = {}

    rules = profile.get("bid_rules", {}) or {}
    preferred = [str(x).lower() for x in (rules.get("preferred_sectors", []) or []) if str(x).strip()]
    do_not_bid = [str(x).lower() for x in (rules.get("do_not_bid", []) or []) if str(x).strip()]
    conditional = [str(x).lower() for x in (rules.get("conditional", []) or []) if str(x).strip()]
    min_v = float(rules.get("min_project_value_cr", 0) or 0)
    max_v = float(rules.get("max_project_value_cr", 0) or 0)

    text_blob = " ".join([
        str(tender.get("brief", "") or ""),
        str(tender.get("org_name", "") or ""),
        str(tender.get("eligibility", "") or ""),
    ]).lower()
    value = float(tender.get("estimated_cost_cr", 0) or 0)
    score = 50
    reasons = []

    pref_hits = [k for k in preferred if k and k in text_blob]
    if pref_hits:
        score += min(25, len(pref_hits) * 8)
        reasons.append(f"sector match: {', '.join(pref_hits[:3])}")

    dnb_hits = [k for k in do_not_bid if k and k in text_blob]
    if dnb_hits:
        score -= min(40, len(dnb_hits) * 15)
        reasons.append(f"do-not-bid hit: {', '.join(dnb_hits[:3])}")

    cond_hits = [k for k in conditional if k and k in text_blob]
    if cond_hits:
        score -= min(15, len(cond_hits) * 6)
        reasons.append(f"conditional hit: {', '.join(cond_hits[:3])}")

    if value:
        if min_v and value < min_v:
            score -= 12
            reasons.append(f"value below min ({value} < {min_v} Cr)")
        if max_v and value > max_v:
            score -= 10
            reasons.append(f"value above max ({value} > {max_v} Cr)")

    score = max(0, min(100, int(round(score))))
    if score >= 65:
        verdict = "BID"
    elif score >= 45:
        verdict = "CONDITIONAL"
    else:
        verdict = "NO-BID"

    tender["win_probability"] = score
    tender["verdict"] = verdict
    tender["auto_scored"] = True
    tender["auto_scored_at"] = datetime.now().isoformat()
    tender["reason"] = "; ".join(reasons[:4]) or "Rule-based profile fit scoring"
    return tender

def _try_enrich_tender_from_t247(tender: dict):
    """Best-effort enrichment: doc count + corrigendum hint from document list API."""
    tid = str(tender.get("t247_id", "")).strip()
    if not tid:
        return
    import requests as _req
    t_id_num = tid if tid.isdigit() else re.sub(r"\D", "", tid)
    if not t_id_num:
        return
    url = f"https://t247_api.tender247.com/apigateway/T247Tender/api/tender/tender-document-list/{t_id_num}"
    headers = {
        "accept": "application/json",
        "content-length": "0",
        "origin": "https://www.tender247.com",
        "referer": "https://www.tender247.com/",
        "user-agent": "Mozilla/5.0",
    }
    try:
        resp = _req.post(url, headers=headers, timeout=20, verify=False)
        if resp.status_code != 200:
            return
        data = resp.json()
        items = data if isinstance(data, list) else (data.get("data") or data.get("result") or data.get("documents") or [])
        if isinstance(items, dict):
            items = [items]
        doc_count = len(items or [])
        tender["t247_doc_count"] = doc_count
        tender["docs_available"] = doc_count > 0
        corr = 0
        for it in (items or []):
            if not isinstance(it, dict):
                continue
            blob = json.dumps(it).lower()
            if "corrig" in blob:
                corr += 1
        tender["corrigendum_count"] = corr
        tender["has_corrigendum"] = corr > 0
    except Exception:
        return

def _run_t247_sync_once() -> dict:
    """Run one Tender247 sync and return merge summary."""
    cfg = load_config()
    token = _t247_get_token(cfg)
    payload = _t247_decode_jwt(token)
    user_id = int(payload.get("UserId") or cfg.get("t247_user_id") or 0)
    query_id = int(cfg.get("t247_query_id") or 328890)

    import requests as _req
    headers = _t247_api_headers(token)
    request_body = {
        "tab_id": 2, "tender_id": 0, "tender_number": "", "search_text": "",
        "refine_search_text": "", "tender_value_operator": 0,
        "tender_value_from": 0, "tender_value_to": 0,
        "publication_date_from": "", "publication_date_to": "",
        "closing_date_from": "", "closing_date_to": "",
        "search_by_location": False, "statezone_ids": "", "city_ids": "",
        "state_ids": "", "organization_ids": "", "organization_name": "",
        "sort_by": 1, "sort_type": 2, "page_no": 1, "record_per_page": 500,
        "keyword_id": "", "mfa": "", "nameof_website": "",
        "tender_typeid": 0, "is_tender_doc_uploaded": False,
        "user_id": user_id, "user_email_service_query_id": query_id,
        "exact_search": False, "exact_search_text": False,
        "search_by_split_word": False, "product_id": "",
        "organization_type_id": "", "sub_industry_id": "",
        "search_by": 0, "guest_user_id": 0, "quantity": "",
        "quantity_operator": 0, "msme_exemption": 0,
        "startup_exemption": 0, "gem": 0, "mail_date": "",
        "tab_status": 0, "is_ai_summary": False, "boq": 0,
        "is_grace": False, "surety_bond": False, "limited_tender": False,
    }

    r = _req.post(
        "https://t247_api.tender247.com/apigateway/T247Tender/api/tender/auth/tender-excel-download",
        headers=headers, json=request_body, timeout=120, verify=False,
    )
    if r.status_code == 401:
        raise ValueError("T247 token expired. Refresh token in Settings.")
    if r.status_code == 403:
        raise ValueError("T247 access denied. Check subscription/account access.")
    if r.status_code != 200:
        raise RuntimeError(f"T247 API returned HTTP {r.status_code}: {r.text[:300]}")

    ct = r.headers.get("Content-Type", "")
    xl_bytes = r.content
    if "json" in ct:
        resp_json = r.json()
        dl_url = resp_json.get("url") or resp_json.get("file_url") or resp_json.get("download_url")
        if not dl_url:
            raise RuntimeError(f"T247 returned JSON instead of Excel: {str(resp_json)[:300]}")
        r2 = _req.get(dl_url, headers=headers, timeout=120)
        xl_bytes = r2.content

    if len(xl_bytes) < 100:
        raise RuntimeError(f"T247 returned empty file ({len(xl_bytes)} bytes)")

    LATEST_EXCEL_FILE.write_bytes(xl_bytes)
    tmp = Path(tempfile.mktemp(suffix=".xlsx", dir=str(TEMP_DIR)))
    tmp.write_bytes(xl_bytes)
    try:
        tenders = process_excel(str(tmp))
    finally:
        try:
            tmp.unlink()
        except Exception:
            pass

    result = _t247_merge_tenders(tenders)
    touched_ids = (result.get("added_ids", []) + result.get("updated_ids", []))[:120]
    if touched_ids:
        db = load_db()
        changed = 0
        for tid in touched_ids:
            t = db.get("tenders", {}).get(tid, {})
            if not t:
                continue
            _try_enrich_tender_from_t247(t)
            if not t.get("bid_no_bid_done"):
                t = _auto_score_tender_v1(t)
            db["tenders"][tid] = t
            changed += 1
        if changed:
            save_db(db)
    return {
        "status": "success",
        "total": len(tenders),
        "added": result["added"],
        "updated": result["updated"],
        "enriched": len(touched_ids),
        "file_size_kb": len(xl_bytes) // 1024,
        "source": "T247 API (real-time)",
    }

def _update_t247_sync_state(status: str, message: str = "", total: int = 0, added: int = 0, updated: int = 0):
    with _t247_sync_lock:
        _t247_sync_state["last_run_at"] = datetime.now().isoformat()
        _t247_sync_state["last_status"] = status
        _t247_sync_state["last_message"] = message
        _t247_sync_state["last_total"] = int(total or 0)
        _t247_sync_state["last_added"] = int(added or 0)
        _t247_sync_state["last_updated"] = int(updated or 0)

def _run_t247_sync_scheduler():
    """Periodic Tender247 sync loop (manual trigger still available)."""
    _time.sleep(60)
    while not _t247_sync_stop.is_set():
        try:
            cfg = load_config()
            enabled = bool(cfg.get("t247_auto_sync_enabled", True))
            minutes = int(cfg.get("t247_auto_sync_minutes", 180) or 180)
            minutes = max(15, min(720, minutes))
            if enabled:
                result = _run_t247_sync_once()
                _update_t247_sync_state(
                    "success",
                    "auto sync completed",
                    total=result.get("total", 0),
                    added=result.get("added", 0),
                    updated=result.get("updated", 0),
                )
            sleep_for = minutes * 60
        except Exception as e:
            _update_t247_sync_state("error", f"auto sync failed: {e}")
            sleep_for = 600
        _t247_sync_stop.wait(sleep_for)

def _run_daily_digest_scheduler():
    """Generate email/whatsapp style digest daily at configured time."""
    _time.sleep(45)
    while not _t247_sync_stop.is_set():
        try:
            cfg = load_config()
            enabled = bool(cfg.get("daily_digest_enabled", True))
            hour = int(cfg.get("daily_digest_hour", 9) or 9)
            minute = int(cfg.get("daily_digest_minute", 0) or 0)
            hour = max(0, min(23, hour))
            minute = max(0, min(59, minute))
            now = datetime.now()
            today_key = now.strftime("%Y-%m-%d")
            should_run = enabled and now.hour == hour and now.minute == minute
            already_done = _digest_state.get("last_generated_date") == today_key
            if should_run and not already_done:
                _build_daily_digest()
        except Exception as e:
            with _digest_lock:
                _digest_state["status"] = "error"
                _digest_state["error"] = str(e)
        _t247_sync_stop.wait(55)

@app.post("/t247/connect")
async def t247_connect(data: dict = Body(...)):
    """Save T247 email+password and immediately login to verify + store token."""
    email = str(data.get("email", "") or "").strip()
    password = str(data.get("password", "") or "").strip()
    if not email or not password:
        raise HTTPException(400, "Email and password are required.")
    cfg = load_config()
    cfg["t247_email"] = email
    cfg["t247_password"] = password
    cfg["t247_bearer_token"] = ""  # force fresh login
    save_config(cfg)
    try:
        token = _t247_auto_login(cfg)
        payload = _t247_decode_jwt(token)
        import time as _t
        exp = payload.get("exp", 0)
        remaining_hrs = round((exp - _t.time()) / 3600, 1) if exp else 0
        return {
            "status": "connected",
            "message": f"Connected — token valid for {remaining_hrs:.0f}h",
            "user_id": payload.get("UserId"),
            "bidder_name": payload.get("bidder_name", ""),
            "email": email,
            "expires_in_hours": remaining_hrs,
            "expires_at": datetime.fromtimestamp(exp).strftime("%d %b %Y %H:%M") if exp else "",
        }
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(502, f"T247 login error: {e}")


@app.get("/t247/connection-status")
async def t247_connection_status():
    """Return current T247 connection state without triggering a login."""
    import time as _t
    cfg = load_config()
    email = str(cfg.get("t247_email", "") or cfg.get("t247_username", "") or "").strip()
    has_creds = bool(email and cfg.get("t247_password"))
    token = str(cfg.get("t247_bearer_token", "") or "").strip()
    if not token:
        return {"status": "disconnected", "has_credentials": has_creds, "email": email,
                "message": "Not connected." + (" Enter credentials in Settings." if not has_creds else " Click Connect.")}
    payload = _t247_decode_jwt(token)
    exp = payload.get("exp", 0)
    if exp and _t.time() > exp:
        remaining = round((_t.time() - exp) / 3600, 1)
        return {"status": "expired", "has_credentials": has_creds, "email": email,
                "message": f"Token expired {remaining:.1f}h ago. {'Auto-refresh on next sync.' if has_creds else 'Re-enter credentials.'}",
                "expired_at": datetime.fromtimestamp(exp).strftime("%d %b %Y %H:%M")}
    remaining_hrs = round((exp - _t.time()) / 3600, 1) if exp else 0
    return {
        "status": "connected",
        "has_credentials": has_creds,
        "email": email,
        "user_id": payload.get("UserId"),
        "bidder_name": payload.get("bidder_name", ""),
        "expires_in_hours": remaining_hrs,
        "expires_at": datetime.fromtimestamp(exp).strftime("%d %b %Y %H:%M") if exp else "",
        "message": f"Connected — {remaining_hrs:.0f}h remaining",
    }


@app.get("/t247/probe-login")
async def t247_probe_login():
    """Probe candidate login URLs — find real endpoint by status code."""
    import requests as _req
    results = []
    probe_headers = {
        "accept": "application/json, text/plain, */*",
        "content-type": "application/json",
        "origin": "https://www.tender247.com",
        "referer": "https://www.tender247.com/auth",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/147.0.0.0 Safari/537.36",
    }
    probe_body = {"email": "probe@test.com", "password": "probe"}
    found = []
    for url in _T247_LOGIN_ENDPOINTS:
        try:
            r = _req.post(url, json=probe_body, headers=probe_headers, timeout=10)
            is_json = "json" in r.headers.get("content-type", "")
            preview = r.text[:300] if is_json else f"[HTML response, {len(r.text)} chars]"
            entry = {"url": url, "status": r.status_code, "is_json": is_json, "response_preview": preview}
            results.append(entry)
            if r.status_code in (200, 401, 400) and is_json:
                found.append(url)
        except _req.exceptions.ConnectionError:
            results.append({"url": url, "status": "ConnectionError", "is_json": False, "response_preview": "path not routed"})
        except _req.exceptions.Timeout:
            results.append({"url": url, "status": "Timeout", "is_json": False, "response_preview": ""})
        except Exception as e:
            results.append({"url": url, "status": "Error", "is_json": False, "response_preview": str(e)})
    return {
        "likely_login_endpoints": found,
        "probe_results": results,
        "hint": "likely_login_endpoints shows URLs that returned JSON with 200/401/400. Share this full result.",
    }


@app.get("/test-t247")
async def test_t247():
    return await t247_connection_status()

@app.get("/t247-token-status")
async def t247_token_status():
    return await t247_connection_status()

@app.post("/fetch-t247-excel")
async def fetch_t247_excel(background_tasks: BackgroundTasks):
    """Manual trigger for Tender247 sync."""
    try:
        result = _run_t247_sync_once()
        _update_t247_sync_state(
            "success",
            "manual sync completed",
            total=result.get("total", 0),
            added=result.get("added", 0),
            updated=result.get("updated", 0),
        )
        return result
    except ValueError as e:
        _update_t247_sync_state("error", str(e))
        raise HTTPException(400, str(e))
    except HTTPException:
        raise
    except Exception as e:
        _update_t247_sync_state("error", str(e))
        raise HTTPException(502, f"T247 sync failed: {e}")

@app.post("/fetch-t247-excel/retry")
async def retry_t247_sync():
    return await fetch_t247_excel(BackgroundTasks())

@app.post("/ops/daily-digest/generate")
async def generate_daily_digest():
    try:
        digest = _build_daily_digest()
        return {"status": "success", **digest}
    except Exception as e:
        raise HTTPException(500, f"Digest generation failed: {e}")

@app.get("/ops/daily-digest")
async def get_daily_digest():
    p = OUTPUT_DIR / "daily_digest_latest.json"
    if p.exists():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            return {"status": "success", **data, "scheduler": dict(_digest_state)}
        except Exception:
            pass
    digest = _build_daily_digest()
    return {"status": "success", **digest, "scheduler": dict(_digest_state)}

@app.get("/ops/daily-digest.txt")
async def download_daily_digest_text():
    p = OUTPUT_DIR / "daily_digest_latest.txt"
    if not p.exists():
        _build_daily_digest()
    return FileResponse(
        path=str(p),
        filename=f"daily_digest_{datetime.now().strftime('%Y%m%d')}.txt",
        media_type="text/plain",
    )

@app.get("/t247-sync-status")
async def t247_sync_status():
    cfg = load_config()
    return {
        "status": "success",
        "auto_sync_enabled": bool(cfg.get("t247_auto_sync_enabled", True)),
        "auto_sync_minutes": int(cfg.get("t247_auto_sync_minutes", 180) or 180),
        **_t247_sync_state,
    }

def _t247_doc_download_headers(token: str = "") -> dict:
    """Headers for documents.tender247.com — include Bearer if available."""
    h = {
        "accept": "application/json, text/plain, */*",
        "origin": "https://www.tender247.com",
        "referer": "https://www.tender247.com/",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/147.0.0.0 Safari/537.36",
    }
    if token:
        h["authorization"] = f"Bearer {token}"
    return h

def _save_tender_doc_to_vault(t247_id: str, filename: str, content: bytes, mime: str, doc_hash: str = "") -> dict:
    """Persist downloaded Tender247 document bundle in vault and link to tender."""
    checksum = hashlib.sha256(content).hexdigest()
    vault = _load_vault()
    for item in vault:
        if (
            str(item.get("source")) == "t247"
            and str(item.get("t247_id")) == str(t247_id)
            and str(item.get("checksum")) == checksum
        ):
            return {k: v for k, v in item.items() if k != "b64"}

    entry = {
        "id": str(uuid.uuid4()),
        "name": filename,
        "category": "tender247",
        "source": "t247",
        "t247_id": str(t247_id),
        "doc_hash": str(doc_hash or ""),
        "checksum": checksum,
        "size": len(content),
        "uploaded_at": datetime.now().isoformat(),
        "b64": base64.b64encode(content).decode(),
        "mime": mime or "application/octet-stream",
    }
    vault.append(entry)
    _save_vault(vault)

    db = load_db()
    tender = db.get("tenders", {}).get(str(t247_id), {})
    attachments = tender.get("attachments", [])
    attachments = [a for a in attachments if a.get("vault_id") != entry["id"]]
    attachments.append({
        "vault_id": entry["id"],
        "name": entry["name"],
        "source": "t247",
        "uploaded_at": entry["uploaded_at"],
        "size": entry["size"],
        "mime": entry["mime"],
    })
    tender["attachments"] = attachments[-20:]
    if doc_hash:
        tender["t247_doc_hash"] = doc_hash
    db.setdefault("tenders", {})[str(t247_id)] = tender
    save_db(db)
    return {k: v for k, v in entry.items() if k != "b64"}

@app.get("/tender/{t247_id}/doc-download")
async def download_tender_docs(t247_id: str):
    """Download tender documents from T247 using auto-login token."""
    import requests as _req, io as _io

    safe_tid = re.sub(r"[^\w\-]", "_", str(t247_id))[:40]

    # Get valid auth token (auto-login if needed)
    cfg = load_config()
    try:
        token = _t247_get_token(cfg)
    except ValueError as e:
        raise HTTPException(401, str(e))

    t_id_num = t247_id if str(t247_id).isdigit() else re.sub(r"\D", "", str(t247_id))

    # Step 1: find stored doc hash for this tender
    db = load_db()
    tender = db.get("tenders", {}).get(t247_id, {})
    doc_hash = (tender.get("t247_doc_hash") or tender.get("doc_hash")
                or tender.get("document_hash") or "").strip()

    # Step 2: fetch document list with auth token to get hash
    if not doc_hash:
        doc_list_url = f"https://t247_api.tender247.com/apigateway/T247Tender/api/tender/tender-document-list/{t_id_num}"
        try:
            r_list = _req.post(
                doc_list_url,
                headers=_t247_api_headers(token),
                timeout=30,
                verify=False,
            )
            if r_list.status_code == 401:
                # Token rejected — force re-login once
                try:
                    token = _t247_auto_login(cfg)
                    r_list = _req.post(doc_list_url, headers=_t247_api_headers(token), timeout=30, verify=False)
                except Exception:
                    pass
            if r_list.status_code == 200:
                d = r_list.json()
                items = d if isinstance(d, list) else (
                    d.get("data") or d.get("result") or d.get("documents") or d.get("list") or []
                )
                if isinstance(items, dict):
                    items = [items]
                for item in (items or []):
                    if not isinstance(item, dict):
                        continue
                    h = (str(
                        item.get("document_hash","") or item.get("doc_hash","")
                        or item.get("download_hash","") or item.get("hash","")
                        or item.get("documentHash","") or item.get("downloadHash","")
                    ).strip())
                    if h and len(h) >= 20:
                        doc_hash = h
                        tender["t247_doc_hash"] = doc_hash
                        db["tenders"][t247_id] = tender
                        save_db(db)
                        break
        except Exception as ex:
            print(f"⚠️ Doc list fetch error for {t247_id}: {ex}")

    if not doc_hash:
        raise HTTPException(
            404,
            f"No documents found for tender {t247_id}. "
            "The tender may not have uploaded documents on T247 yet."
        )

    # Step 3: download ZIP from documents.tender247.com using hash + auth
    download_url = f"https://documents.tender247.com/tender/download-document-all/{doc_hash}"
    try:
        r = _req.get(download_url, headers=_t247_doc_download_headers(token), timeout=120)
        if r.status_code == 401:
            # Retry once with fresh login
            try:
                token = _t247_auto_login(cfg)
                r = _req.get(download_url, headers=_t247_doc_download_headers(token), timeout=120)
            except Exception:
                pass
    except Exception as e:
        raise HTTPException(502, f"Document server unreachable: {e}")

    if r.status_code == 304:
        # Not Modified — shouldn't happen on fresh request but handle it
        raise HTTPException(502, "T247 documents server returned 304 — try again")
    if r.status_code != 200:
        raise HTTPException(r.status_code, f"Document server returned HTTP {r.status_code}")

    content = r.content
    if not content or len(content) < 200:
        raise HTTPException(404, f"Document file empty ({len(content)} bytes) — hash may be invalid")

    from fastapi.responses import Response
    ct_out = "application/zip"
    fname = f"T247_{safe_tid}_docs.zip"
    try:
        if not zipfile.is_zipfile(_io.BytesIO(content)):
            ct_out = r.headers.get("Content-Type", "application/octet-stream")
            ext = ".zip" if "zip" in ct_out else ".pdf" if "pdf" in ct_out else ".bin"
            fname = f"T247_{safe_tid}_docs{ext}"
    except Exception:
        pass

    try:
        _save_tender_doc_to_vault(
            t247_id=t247_id,
            filename=fname,
            content=content,
            mime=ct_out,
            doc_hash=doc_hash,
        )
    except Exception as e:
        print(f"⚠️ Vault save skipped for {t247_id}: {e}")

    return Response(
        content=content, media_type=ct_out,
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )

@app.post("/tender/{t247_id}/store-doc-hash")
async def store_doc_hash(t247_id: str, data: dict = Body(...)):
    """Store a document hash for a tender (called from frontend when hash is known)."""
    doc_hash = str(data.get("doc_hash", "")).strip()
    if not doc_hash or len(doc_hash) < 32:
        raise HTTPException(400, "Invalid doc_hash")
    db = load_db()
    if t247_id not in db.get("tenders", {}):
        raise HTTPException(404, f"Tender {t247_id} not in database")
    db["tenders"][t247_id]["t247_doc_hash"] = doc_hash
    save_db(db)
    return {"status": "ok", "t247_id": t247_id, "doc_hash": doc_hash}


@app.post("/auto-download/{t247_id}")
async def auto_download_tender(t247_id: str):
    return await download_tender_docs(t247_id)

@app.post("/tender/{t247_id}/auto-download")
async def auto_download_tender_alias(t247_id: str):
    return await download_tender_docs(t247_id)


# ── TENDER UPDATE (alias for status) ─────────────────────────
@app.post("/tender/{t247_id}/update")
async def update_tender_fields(t247_id: str, body: dict = Body(...)):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        allowed = {"verdict","reason","status","notes","notes_internal","pipeline_stage",
                   "outcome","outcome_value","outcome_competitor","outcome_notes"}
        for k, v in body.items():
            if k in allowed:
                t[k] = v
        t["updated_at"] = datetime.now().isoformat()
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok", "t247_id": t247_id}


# ── VAULT (doc storage — disk-backed so it survives restarts) ────
_VAULT_FILE = OUTPUT_DIR / "vault_index.json"
_vault_lock = threading.Lock()

def _load_vault() -> list:
    with _vault_lock:
        try:
            if _VAULT_FILE.exists():
                return json.loads(_VAULT_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
        return []

def _save_vault(vault: list):
    with _vault_lock:
        _VAULT_FILE.parent.mkdir(exist_ok=True, parents=True)
        _VAULT_FILE.write_text(json.dumps(vault, indent=2), encoding="utf-8")

@app.get("/vault/list")
async def vault_list():
    vault = _load_vault()
    return {"files": [{k: v for k, v in f.items() if k != "b64"} for f in vault]}

@app.post("/vault/upload")
async def vault_upload(file: UploadFile = File(...), category: str = "general"):
    import base64
    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "File too large (max 50MB)")
    entry = {
        "id": str(uuid.uuid4()),
        "name": file.filename,
        "category": category,
        "size": len(data),
        "uploaded_at": datetime.now().isoformat(),
        "b64": base64.b64encode(data).decode(),
        "mime": file.content_type or "application/octet-stream",
    }
    vault = _load_vault()
    vault.append(entry)
    _save_vault(vault)
    return {"status": "ok", "file": {k: v for k, v in entry.items() if k != "b64"}}

@app.delete("/vault/delete/{file_id}")
async def vault_delete(file_id: str):
    vault = _load_vault()
    before = len(vault)
    vault = [f for f in vault if f.get("id") != file_id]
    _save_vault(vault)
    return {"status": "ok", "deleted": before - len(vault)}

@app.get("/vault/download/{file_id}")
async def vault_download(file_id: str):
    import base64
    from fastapi.responses import Response
    vault = _load_vault()
    f = next((x for x in vault if x.get("id") == file_id), None)
    if not f:
        raise HTTPException(404, "File not found")
    return Response(content=base64.b64decode(f["b64"]),
                    media_type=f.get("mime","application/octet-stream"),
                    headers={"Content-Disposition": f'attachment; filename="{f["name"]}"'})


# ── POST-AWARD MILESTONES ─────────────────────────────────────
@app.get("/post-award/{t247_id}/milestones")
async def get_milestones(t247_id: str):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id), {})
    return {"milestones": t.get("milestones", []), "invoices": t.get("invoices", [])}

@app.get("/milestones/{t247_id}")
async def get_milestones_alias(t247_id: str):
    return await get_milestones(t247_id)

@app.post("/post-award/{t247_id}/milestones")
async def add_milestone(t247_id: str, body: dict = Body(...)):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        ms = t.get("milestones", [])
        new_ms = {
            "id": str(uuid.uuid4())[:8],
            "name": body.get("name", "Milestone"),
            "due_date": body.get("due_date", ""),
            "value_pct": body.get("value_pct", 0),
            "status": body.get("status", "Pending"),
            "notes": body.get("notes", ""),
            "created_at": datetime.now().isoformat(),
        }
        ms.append(new_ms)
        t["milestones"] = ms
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok", "milestone": new_ms}

@app.post("/post-award/{t247_id}/milestones/setup")
async def setup_milestones(t247_id: str, body: dict = Body(...)):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        t["milestones"] = body.get("milestones", [])
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok", "count": len(t["milestones"])}

@app.post("/milestones/{t247_id}/setup")
async def setup_milestones_alias(t247_id: str, body: dict = Body(default={})):
    return await setup_milestones(t247_id, {"milestones": body.get("milestones", [])})

@app.patch("/post-award/{t247_id}/milestones/{mid}")
async def update_milestone(t247_id: str, mid: str, body: dict = Body(...)):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        for ms in t.get("milestones", []):
            if ms.get("id") == mid:
                ms.update({k: v for k, v in body.items()})
                ms["updated_at"] = datetime.now().isoformat()
                break
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok"}

@app.post("/milestones/{t247_id}/{mid}/done")
async def mark_milestone_done_alias(t247_id: str, mid: str):
    return await update_milestone(t247_id, mid, {"status": "Done"})

@app.post("/post-award/{t247_id}/invoice")
async def add_invoice(t247_id: str, body: dict = Body(...)):
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        invoices = t.get("invoices", [])
        inv = {
            "id": str(uuid.uuid4())[:8],
            "invoice_no": body.get("invoice_no", ""),
            "amount": body.get("amount", 0),
            "date": body.get("date", ""),
            "status": body.get("status", "Raised"),
            "milestone_id": body.get("milestone_id", ""),
            "created_at": datetime.now().isoformat(),
        }
        invoices.append(inv)
        t["invoices"] = invoices
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok", "invoice": inv}

@app.post("/tender/{t247_id}/invoice")
async def add_invoice_alias(t247_id: str, body: dict = Body(...)):
    return await add_invoice(t247_id, body)

@app.post("/post-award/{t247_id}/{doc_type}")
async def post_award_doc(t247_id: str, doc_type: str, body: dict = Body(...)):
    """Generic post-award doc save (amc, extension, etc.)"""
    with db_lock:
        db = load_db()
        t = db["tenders"].get(str(t247_id))
        if not t:
            raise HTTPException(404, "Tender not found")
        t[f"pa_{doc_type}"] = body
        t[f"pa_{doc_type}_updated"] = datetime.now().isoformat()
        db["tenders"][str(t247_id)] = t
        save_db(db)
    return {"status": "ok", "type": doc_type}

@app.post("/tender/{t247_id}/letter/{doc_type}")
async def post_award_doc_alias(t247_id: str, doc_type: str, body: dict = Body(default={})):
    if not body:
        body = {"generated_at": datetime.now().isoformat()}
    return await post_award_doc(t247_id, doc_type, body)


# ══════════════════════════════════════════════════════════════════════════════
# v8.1 — DOC PREVIEW + AI-EDIT + LOCAL DOWNLOAD + VERSIONS + RISK + COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════

def _locate_tender_docx(t247_id: str) -> Path:
    """Find latest BidNoBid_*.docx for given tender id; fallback to report_file field."""
    tender = get_tender(t247_id) or {}
    rf = tender.get("report_file")
    if rf:
        p = (OUTPUT_DIR / Path(rf).name)
        if p.exists():
            return p
    candidates = list(OUTPUT_DIR.glob(f"BidNoBid_*{t247_id}*.docx"))
    if not candidates:
        candidates = sorted(OUTPUT_DIR.glob("BidNoBid_*.docx"),
                            key=lambda f: f.stat().st_mtime, reverse=True)
        if tender.get("tender_no"):
            tn = re.sub(r'[^\w\-]', '_', str(tender.get("tender_no")))[:50]
            matched = [p for p in candidates if tn in p.stem]
            if matched:
                return matched[0]
    if candidates:
        return candidates[0]
    raise HTTPException(404, f"No report found for {t247_id}. Run Analyse first.")


@app.get("/tender/{t247_id}/doc-html")
async def tender_doc_html(t247_id: str):
    """Render current report docx as HTML for in-browser preview."""
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    docx_path = _locate_tender_docx(t247_id)
    html = doc_editor.docx_to_html(docx_path)
    return {"status": "ok", "html": html, "filename": docx_path.name,
            "modified": datetime.fromtimestamp(docx_path.stat().st_mtime).isoformat(),
            "size_kb": round(docx_path.stat().st_size / 1024, 1)}


@app.post("/tender/{t247_id}/ai-edit")
async def tender_ai_edit(t247_id: str, body: dict = Body(...)):
    """Apply chatbot instruction to current doc HTML. Returns edited HTML (not saved)."""
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    html = (body.get("html") or "").strip()
    instruction = (body.get("instruction") or "").strip()
    if not instruction:
        raise HTTPException(400, "instruction required")
    if not html:
        docx_path = _locate_tender_docx(t247_id)
        html = doc_editor.docx_to_html(docx_path)
    result = doc_editor.ai_edit_html(html, instruction)
    if "error" in result:
        raise HTTPException(502, result["error"])
    return {"status": "ok", "html": result["html"], "chars": result["chars"]}


@app.post("/tender/{t247_id}/doc-save")
async def tender_doc_save(t247_id: str, body: dict = Body(...)):
    """Save HTML back to docx. Creates a version snapshot first."""
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    html = body.get("html") or ""
    note = body.get("note") or "manual edit"
    if len(html) < 20:
        raise HTTPException(400, "HTML content empty")
    docx_path = _locate_tender_docx(t247_id)
    try:
        doc_editor.snapshot_version(docx_path, note="auto-snapshot before save")
    except Exception as e:
        print(f"[doc-save] snapshot skipped: {e}")
    try:
        doc_editor.html_to_docx(html, docx_path)
    except Exception as e:
        raise HTTPException(500, f"Save failed: {e}")
    import base64
    b64 = base64.b64encode(docx_path.read_bytes()).decode()
    return {"status": "ok", "filename": docx_path.name,
            "size_kb": round(docx_path.stat().st_size / 1024, 1),
            "doc_b64": b64, "note": note,
            "saved_at": datetime.now().isoformat()}


@app.get("/tender/{t247_id}/analysis-doc-download")
async def tender_doc_download_local(t247_id: str):
    """Explicit local-download button endpoint — streams docx with correct filename."""
    docx_path = _locate_tender_docx(t247_id)
    return FileResponse(
        path=str(docx_path),
        filename=docx_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/tender/{t247_id}/doc-versions")
async def tender_doc_versions(t247_id: str):
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    docx_path = _locate_tender_docx(t247_id)
    return {"status": "ok", "filename": docx_path.name,
            "versions": doc_editor.list_versions(docx_path)}


@app.post("/tender/{t247_id}/doc-restore/{version_id}")
async def tender_doc_restore(t247_id: str, version_id: str):
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    docx_path = _locate_tender_docx(t247_id)
    result = doc_editor.restore_version(docx_path, version_id)
    if "error" in result:
        raise HTTPException(404, result["error"])
    return {"status": "ok", **result}


@app.get("/tender/{t247_id}/risk-score")
async def tender_risk_score(t247_id: str):
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    return {"status": "ok", **doc_editor.compute_risk_score(tender)}


@app.get("/tender/{t247_id}/compliance-matrix")
async def tender_compliance_matrix(t247_id: str):
    if not DOC_EDITOR_AVAILABLE:
        raise HTTPException(500, "doc_editor not available")
    tender = get_tender(t247_id)
    if not tender:
        raise HTTPException(404, "Tender not found")
    matrix = doc_editor.build_compliance_matrix(tender)
    covered = sum(1 for m in matrix if m.get("color") == "GREEN")
    partial = sum(1 for m in matrix if m.get("color") == "AMBER")
    gap     = sum(1 for m in matrix if m.get("color") == "RED")
    return {"status": "ok", "total": len(matrix), "covered": covered,
            "partial": partial, "gap": gap, "matrix": matrix}


@app.get("/platform/api-pool-stats")
async def api_pool_stats():
    """Observability for key-pool + analyst slots."""
    if not API_POOL_AVAILABLE:
        return {"status": "unavailable"}
    try:
        refresh_pool()
        return {
            "status": "ok",
            "slots": get_slots().snapshot(),
            "keys": get_pool().stats(),
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.get("/platform/analyst-capacity")
async def analyst_capacity():
    """Quick sanity check — how many analysts can run today?"""
    if not API_POOL_AVAILABLE:
        return {"status": "unavailable", "hint": "core.api_pool import failed"}
    refresh_pool()
    stats = get_pool().stats()
    slots = get_slots().snapshot()
    total_rpd = sum(max(0, 1400 - s.get("rpd_used", 0)) for s in stats)  # 1400 cap per key
    per_tender_calls = 9  # 9 segments
    max_today = max(0, int(total_rpd / per_tender_calls))
    return {
        "status": "ok",
        "api_keys_configured": len(stats),
        "concurrent_slots": slots.get("max"),
        "active_now": slots.get("active"),
        "completed_today": slots.get("completed_today"),
        "theoretical_daily_capacity": max_today,
        "recommendation": "Add more GEMINI_API_KEY_2..5 env vars to increase throughput"
            if len(stats) < 3 else "Capacity is sufficient for 10+ tenders/day.",
    }


# ── URL compatibility aliases ────────────────────────────────────────────────

@app.post("/merge-pdf/{t247_id}")
async def merge_pdf_url_alias(t247_id: str):
    """Frontend calls /merge-pdf/{id} — route to /tender/{id}/merge-pdf"""
    return await merge_submission_pdf(t247_id)

@app.get("/risk-scan/{t247_id}")
async def risk_scan_url_alias(t247_id: str):
    """Frontend calls /risk-scan/{id} — route to /tender/{id}/risk-score"""
    return await tender_risk_score(t247_id)

@app.post("/setup-milestones/{t247_id}")
async def setup_milestones_url_alias(t247_id: str, body: dict = Body(default={})):
    """Frontend calls /setup-milestones/{id} — route to /milestones/{id}/setup"""
    return await setup_milestones_alias(t247_id, body)

@app.post("/restore-tender/{t247_id}")
async def restore_tender_url_alias(t247_id: str):
    """Frontend calls /restore-tender/{id} — route to /tender/{id}/restore"""
    return await restore_tender(t247_id)

@app.post("/generate-letter/{doc_type}")
async def generate_letter_url_alias(doc_type: str, t247_id: str = ""):
    """Frontend calls /generate-letter/{type}?t247_id={id} — route to /tender/{id}/letter/{type}"""
    if not t247_id:
        raise HTTPException(400, "t247_id required")
    body = {"generated_at": datetime.now().isoformat()}
    return await post_award_doc(t247_id, doc_type, body)

@app.post("/generate-invoice")
async def generate_invoice_wrapper(body: dict = Body(...)):
    """Frontend calls /generate-invoice with {t247_id, description, amount} in body"""
    t247_id = body.get("t247_id", "")
    if not t247_id:
        raise HTTPException(400, "t247_id required")
    return await add_invoice(t247_id, body)

@app.get("/boq-search")
async def boq_search(q: str = ""):
    """Search BOQ items across all tenders by keyword"""
    q_lower = q.lower().strip()
    if len(q_lower) < 2:
        return {"results": []}
    db = load_db()
    results = []
    seen = set()
    for tid, tender in db.get("tenders", {}).items():
        boq = tender.get("boq", {})
        for item in boq.get("items", []):
            name = str(item.get("item", item.get("name", ""))).lower()
            if q_lower in name:
                key = name
                if key not in seen:
                    seen.add(key)
                    results.append({
                        "item": item.get("item", item.get("name", "")),
                        "category": item.get("category", ""),
                        "unit": item.get("unit", ""),
                        "min_price": item.get("unit_cost", item.get("min_price", "")),
                        "max_price": item.get("max_price", ""),
                        "avg_price": item.get("unit_cost", item.get("avg_price", "")),
                        "last_updated": tender.get("last_analysed", ""),
                    })
                    if len(results) >= 50:
                        break
        if len(results) >= 50:
            break
    return {"results": results}
