"""
Post-Award Module v1.0 — Nascent Info Technologies Bid/No-Bid System

For tenders that Nascent has WON — manages:
1. Agreement / Contract documents (LoA acceptance, performance security letter, contract)
2. Milestone tracker (deliverable deadlines, payment schedule)
3. Running Account (RA) Bill / Invoice generator
4. Project completion certificate request letter
"""

import json
import re
from pathlib import Path
from datetime import datetime, date, timedelta
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

BASE_DIR   = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "data"
DOCS_DIR   = BASE_DIR / "docs"
DB_FILE    = OUTPUT_DIR / "tenders_db.json"

NAVY  = RGBColor(31, 56, 100)  if DOCX_OK else None
BLUE  = RGBColor(46, 117, 182) if DOCX_OK else None
WHITE = RGBColor(255, 255, 255) if DOCX_OK else None

N = {
    "name":     "Nascent Info Technologies Pvt. Ltd.",
    "address":  "A-805, Shapath IV, SG Highway, Prahlad Nagar, Ahmedabad – 380015, Gujarat",
    "pan":      "AACCN3670J",
    "gstin":    "24AACCN3670J1ZG",
    "phone":    "+91-79-40200400",
    "email":    "nascent.tender@nascentinfo.com",
    "sign":     "Hitesh Patel",
    "desig":    "Chief Administrative Officer",
}


def _load_db() -> dict:
    if DB_FILE.exists():
        try: return json.loads(DB_FILE.read_text(encoding="utf-8"))
        except: pass
    return {"tenders": {}}


def _save_db(db: dict):
    DB_FILE.write_text(json.dumps(db, indent=2, default=str), encoding="utf-8")


# ─────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell(cell, text: str, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.paragraphs[0].alignment = align
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top','bottom','left','right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '80')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _para(doc, text: str, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p   = doc.add_paragraph()
    run = p.add_run(text or "")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    return p


def _heading(doc, text: str, level=1):
    p   = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)


def _letterhead(doc, tender: dict):
    """Add Nascent letterhead to top of document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(N["name"].upper())
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY; r.font.name = "Arial"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(N["address"])
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(80,80,80); r2.font.name = "Arial"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Ph: {N['phone']}  |  {N['email']}")
    r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(80,80,80); r3.font.name = "Arial"

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"PAN: {N['pan']}  |  GSTIN: {N['gstin']}")
    r4.font.size = Pt(9); r4.font.color.rgb = BLUE; r4.font.name = "Arial"

    # Horizontal line
    phr = doc.add_paragraph()
    phr.paragraph_format.space_after = Pt(4)
    from docx.oxml import OxmlElement
    pPr = phr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _signature(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    _para(doc, "Authorised Signatory")
    _para(doc, N["sign"], bold=True)
    _para(doc, N["desig"])
    _para(doc, N["name"], bold=True)
    _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    _para(doc, "Company Seal:", bold=True)


# ─────────────────────────────────────────────────────────────────
# 1. LOA ACCEPTANCE LETTER
# ─────────────────────────────────────────────────────────────────

def generate_loa_acceptance(tender_data: dict, out_path: str) -> dict:
    """Generate Letter of Award (LoA) acceptance letter."""
    if not DOCX_OK:
        return {"status": "error", "message": "python-docx not installed"}
    try:
        doc = Document()
        section = doc.sections[0]
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

        _letterhead(doc, tender_data)
        doc.add_paragraph()

        _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
        doc.add_paragraph()
        _para(doc, "To,")
        _para(doc, tender_data.get("client_contact", "The Authorised Officer,"))
        _para(doc, tender_data.get("org_name", "[Organisation Name]"))
        _para(doc, tender_data.get("org_address", "[Address]"))
        doc.add_paragraph()

        ref = tender_data.get("loa_ref", "[LoA Reference Number]")
        tender_no = tender_data.get("tender_no", tender_data.get("ref_no", ""))
        _para(doc, f"Sub: Acceptance of Letter of Award (LoA) for {tender_data.get('brief', 'the above-referenced work')}", bold=True)
        _para(doc, f"Ref: (1) Tender No.: {tender_no}  (2) Your LoA No.: {ref}", bold=True)
        doc.add_paragraph()
        _para(doc, "Dear Sir / Madam,")
        doc.add_paragraph()
        _para(doc, f"We, {N['name']}, do hereby acknowledge receipt of your Letter of Award No. {ref} and are pleased to convey our unconditional acceptance of the same.")
        doc.add_paragraph()
        _para(doc, "We confirm the following:", bold=True)
        for item in [
            f"1. We accept the award for: {tender_data.get('brief', 'the above work')}",
            f"2. Contract value: Rs. {tender_data.get('contract_value', '[Amount]')} Cr (as per LoA)",
            f"3. We shall commence work as per the timelines stipulated in the LoA / Agreement",
            f"4. We shall submit the Performance Security / Security Deposit within the stipulated period",
            f"5. We shall execute the formal Agreement as per the schedule notified by your office",
            f"6. All terms and conditions of the LoA and original tender document are unconditionally accepted",
        ]:
            _para(doc, item)
        doc.add_paragraph()
        _para(doc, "We look forward to successfully executing this project and building a long-term partnership with your esteemed organisation.")
        doc.add_paragraph()
        _para(doc, "Thanking you,")
        _para(doc, "Yours faithfully,")
        _signature(doc)
        doc.save(out_path)
        return {"status": "success", "path": out_path}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ─────────────────────────────────────────────────────────────────
# 2. PERFORMANCE SECURITY LETTER
# ─────────────────────────────────────────────────────────────────

def generate_performance_security_letter(tender_data: dict, out_path: str) -> dict:
    """Generate covering letter for Performance Security / Bank Guarantee submission."""
    if not DOCX_OK:
        return {"status": "error", "message": "python-docx not installed"}
    try:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2)

        _letterhead(doc, tender_data)
        doc.add_paragraph()
        _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
        doc.add_paragraph()
        _para(doc, "To,")
        _para(doc, tender_data.get("client_contact", "The Authorised Officer,"))
        _para(doc, tender_data.get("org_name", ""))
        doc.add_paragraph()
        _para(doc, f"Sub: Submission of Performance Security — {tender_data.get('brief','')[:80]}", bold=True)
        _para(doc, f"Ref: LoA No.: {tender_data.get('loa_ref','[LoA Ref]')} dated {tender_data.get('loa_date','[LoA Date]')}", bold=True)
        doc.add_paragraph()
        _para(doc, "Dear Sir / Madam,")
        doc.add_paragraph()
        _para(doc, "With reference to the above LoA, we hereby submit the Performance Security as under:")
        doc.add_paragraph()

        # Performance security table
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = 'Table Grid'
        rows_data = [
            ("Performance Security Amount", f"Rs. {tender_data.get('perf_security_amount','[Amount]')}"),
            ("Form of Security", tender_data.get('perf_security_form','Bank Guarantee / DD')),
            ("Bank / Instrument No.", tender_data.get('perf_security_ref','[BG/DD Number]')),
            ("Issuing Bank", tender_data.get('perf_security_bank','State Bank of India, Ahmedabad')),
            ("Validity", tender_data.get('perf_security_validity','[Validity period]')),
        ]
        for i, (label, val) in enumerate(rows_data):
            _set_cell_bg(tbl.cell(i,0), "EEF2F7")
            _cell(tbl.cell(i,0), label, bold=True, color=NAVY)
            _cell(tbl.cell(i,1), val)

        doc.add_paragraph()
        _para(doc, "We request you to verify the instrument and confirm acceptance at the earliest so that we may proceed with project commencement activities.")
        doc.add_paragraph()
        _para(doc, "Thanking you,")
        _para(doc, "Yours faithfully,")
        _signature(doc)
        doc.save(out_path)
        return {"status": "success", "path": out_path}
    except Exception as e:
        return {"status": "error", "message": str(e)}


# ─────────────────────────────────────────────────────────────────
# 3. MILESTONE TRACKER SETUP
# ─────────────────────────────────────────────────────────────────

STANDARD_MILESTONES = [
    {"name": "Project Kickoff",           "pct_payment": 0,   "days_from_start": 0},
    {"name": "SRS / Requirement Sign-off","pct_payment": 10,  "days_from_start": 45},
    {"name": "UI/UX Design Approval",     "pct_payment": 5,   "days_from_start": 75},
    {"name": "Development (50%)",          "pct_payment": 15,  "days_from_start": 120},
    {"name": "Development Complete",       "pct_payment": 20,  "days_from_start": 150},
    {"name": "Testing & UAT Sign-off",     "pct_payment": 15,  "days_from_start": 165},
    {"name": "Go-Live",                    "pct_payment": 20,  "days_from_start": 180},
    {"name": "Documentation Handover",     "pct_payment": 10,  "days_from_start": 195},
    {"name": "Warranty / Defect Period",   "pct_payment": 5,   "days_from_start": 365},
]

AMC_MILESTONES = [
    {"name": "AMC Q1 Report",  "pct_payment": 25, "days_from_start": 90},
    {"name": "AMC Q2 Report",  "pct_payment": 25, "days_from_start": 180},
    {"name": "AMC Q3 Report",  "pct_payment": 25, "days_from_start": 270},
    {"name": "AMC Q4 Report",  "pct_payment": 25, "days_from_start": 365},
]


def setup_milestones(t247_id: str, start_date: str,
                     contract_value_cr: float,
                     milestones: list = None,
                     is_amc: bool = False) -> dict:
    """
    Set up milestone tracker for a won project.
    start_date: "DD-MM-YYYY"
    """
    db     = _load_db()
    tender = db["tenders"].get(t247_id, {})
    if not tender:
        return {"success": False, "error": "Tender not found"}

    template = AMC_MILESTONES if is_amc else STANDARD_MILESTONES
    ms_list  = milestones if milestones else template

    # Parse start date
    try:
        for fmt in ["%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d"]:
            try:
                start = datetime.strptime(start_date, fmt).date()
                break
            except Exception:
                continue
    except Exception:
        start = date.today()

    built_ms = []
    for i, ms in enumerate(ms_list):
        due   = start + timedelta(days=ms.get("days_from_start", 0))
        pct   = ms.get("pct_payment", 0)
        value = round(contract_value_cr * pct / 100, 4)
        built_ms.append({
            "id":            i + 1,
            "name":          ms.get("name", f"Milestone {i+1}"),
            "due_date":      due.strftime("%d-%m-%Y"),
            "pct_payment":   pct,
            "payment_value": value,
            "status":        "Pending",  # Pending / In Progress / Completed / Delayed
            "completed_on":  "",
            "invoice_raised": False,
            "invoice_ref":   "",
            "notes":         ms.get("notes", ""),
        })

    tender["milestones"]       = built_ms
    tender["project_start"]    = start_date
    tender["contract_value"]   = contract_value_cr
    tender["milestone_setup"]  = datetime.now().isoformat()
    db["tenders"][t247_id]     = tender
    _save_db(db)

    return {
        "success":    True,
        "milestones": built_ms,
        "total_milestones": len(built_ms),
        "contract_value_cr": contract_value_cr,
    }


def update_milestone(t247_id: str, milestone_id: int,
                     status: str, completed_on: str = "",
                     invoice_raised: bool = False,
                     invoice_ref: str = "", notes: str = "") -> dict:
    """Update a milestone status."""
    db     = _load_db()
    tender = db["tenders"].get(t247_id, {})
    if not tender:
        return {"success": False, "error": "Tender not found"}

    for ms in tender.get("milestones", []):
        if ms["id"] == milestone_id:
            ms["status"]        = status
            ms["completed_on"]  = completed_on or (
                date.today().strftime("%d-%m-%Y") if status == "Completed" else "")
            ms["invoice_raised"] = invoice_raised
            ms["invoice_ref"]   = invoice_ref
            if notes:
                ms["notes"] = notes
            break

    db["tenders"][t247_id] = tender
    _save_db(db)
    return {"success": True, "milestone_id": milestone_id, "status": status}


def get_milestone_summary(t247_id: str) -> dict:
    """Get milestone completion summary for a project."""
    db     = _load_db()
    tender = db["tenders"].get(t247_id, {})
    ms     = tender.get("milestones", [])
    if not ms:
        return {"has_milestones": False}

    total      = len(ms)
    completed  = sum(1 for m in ms if m["status"] == "Completed")
    delayed    = sum(1 for m in ms if m["status"] == "Delayed" or (
        m["status"] == "Pending" and m["due_date"] and
        datetime.strptime(m["due_date"], "%d-%m-%Y").date() < date.today()
    ))
    pending    = total - completed - delayed
    paid_cr    = sum(m.get("payment_value",0) for m in ms if m["status"]=="Completed")
    pending_cr = sum(m.get("payment_value",0) for m in ms if m["status"]!="Completed")

    # Upcoming milestones (next 30 days)
    upcoming = []
    for m in ms:
        if m["status"] != "Pending" or not m.get("due_date"):
            continue
        try:
            due = datetime.strptime(m["due_date"], "%d-%m-%Y").date()
            days = (due - date.today()).days
            if 0 <= days <= 30:
                upcoming.append({**m, "days_left": days})
        except Exception:
            pass

    return {
        "has_milestones":   True,
        "total":            total,
        "completed":        completed,
        "pending":          pending,
        "delayed":          delayed,
        "completion_pct":   round(completed / total * 100) if total else 0,
        "paid_cr":          round(paid_cr, 3),
        "pending_cr":       round(pending_cr, 3),
        "upcoming":         upcoming,
        "milestones":       ms,
    }


# ─────────────────────────────────────────────────────────────────
# 4. RUNNING ACCOUNT BILL / INVOICE GENERATOR
# ─────────────────────────────────────────────────────────────────

def generate_ra_bill(
    t247_id:         str,
    milestone_id:    int,
    invoice_no:      str,
    invoice_date:    str,
    work_done_desc:  str,
    amount_before_tax: float,
    gst_rate_pct:    float = 18.0,
    tds_pct:         float = 2.0,
    out_path:        str = "",
) -> dict:
    """Generate a Running Account (RA) Bill / Tax Invoice."""
    if not DOCX_OK:
        return {"status": "error", "message": "python-docx not installed"}

    db     = _load_db()
    tender = db["tenders"].get(t247_id, {})
    if not tender:
        return {"status": "error", "message": "Tender not found"}

    # Find milestone
    ms = next((m for m in tender.get("milestones", [])
               if m["id"] == milestone_id), {})

    # Compute amounts
    cgst  = round(amount_before_tax * gst_rate_pct / 2 / 100, 2)
    sgst  = cgst
    igst  = 0.0  # Interstate: use IGST instead
    total_with_gst = round(amount_before_tax + cgst + sgst + igst, 2)
    tds   = round(total_with_gst * tds_pct / 100, 2)
    net_payable = round(total_with_gst - tds, 2)

    if not out_path:
        safe_id = re.sub(r'[^\w\-]', '_', t247_id)[:30]
        out_path = str(OUTPUT_DIR / f"Invoice_{safe_id}_M{milestone_id}.docx")

    try:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin  = section.bottom_margin = Cm(1.5)

        _letterhead(doc, tender)
        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("TAX INVOICE / RUNNING ACCOUNT BILL")
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY; r.font.name = "Arial"
        doc.add_paragraph()

        # Invoice meta table
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        left  = tbl.cell(0,0)
        right = tbl.cell(0,1)
        left_lines = [
            f"Invoice No.: {invoice_no}",
            f"Invoice Date: {invoice_date or date.today().strftime('%d-%m-%Y')}",
            f"Milestone: {ms.get('name','Milestone '+str(milestone_id))}",
        ]
        right_lines = [
            f"Project: {tender.get('brief','')[:50]}",
            f"Client: {tender.get('org_name','')}",
            f"Tender No.: {tender.get('tender_no', tender.get('ref_no',''))}",
        ]
        for lines, cell in [(left_lines, left), (right_lines, right)]:
            cell.paragraphs[0].clear()
            for line in lines:
                p = cell.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(10); r.font.name = "Arial"

        doc.add_paragraph()

        # Work done description
        _heading(doc, "Description of Work Done", 2)
        _para(doc, work_done_desc)
        doc.add_paragraph()

        # Amount table
        _heading(doc, "Invoice Summary", 2)
        tbl2 = doc.add_table(rows=7, cols=2)
        tbl2.style = 'Table Grid'
        amount_rows = [
            ("Work Value (before tax)",    f"Rs. {amount_before_tax:,.2f}", False, "FFFFFF"),
            (f"CGST @ {gst_rate_pct/2}%", f"Rs. {cgst:,.2f}",             False, "FFFFFF"),
            (f"SGST @ {gst_rate_pct/2}%", f"Rs. {sgst:,.2f}",             False, "FFFFFF"),
            ("Total with GST",             f"Rs. {total_with_gst:,.2f}",   True,  "EBF3FB"),
            (f"Less: TDS @ {tds_pct}%",   f"Rs. {tds:,.2f}",              False, "FFF2CC"),
            ("",                           "",                              False, "FFFFFF"),
            ("NET AMOUNT PAYABLE",         f"Rs. {net_payable:,.2f}",      True,  "1F3864"),
        ]
        for i, (label, value, bold, bg) in enumerate(amount_rows):
            _set_cell_bg(tbl2.cell(i,0), bg)
            _set_cell_bg(tbl2.cell(i,1), bg)
            color = WHITE if bg == "1F3864" else NAVY
            _cell(tbl2.cell(i,0), label, bold=bold, color=color, size=11)
            _cell(tbl2.cell(i,1), value, bold=bold, color=color,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

        doc.add_paragraph()
        # GST details
        _para(doc, f"Our GSTIN: {N['gstin']}  |  HSN/SAC Code: 998314 (IT services)")
        _para(doc, "Bank Details: State Bank of India, SG Highway Branch, Ahmedabad")
        _para(doc, "Account Name: Nascent Info Technologies Pvt. Ltd.")
        _para(doc, "Account No.: [ACCOUNT NUMBER]  |  IFSC: [IFSC CODE]")
        doc.add_paragraph()
        _para(doc, "Please process payment within 30 days of invoice receipt as per contract terms.")
        doc.add_paragraph()
        _para(doc, "Thanking you,")
        _para(doc, "Yours faithfully,")
        _signature(doc)
        doc.save(out_path)

        # Update milestone with invoice ref
        update_milestone(t247_id, milestone_id,
                         status="Completed",
                         invoice_raised=True,
                         invoice_ref=invoice_no)

        return {
            "status":        "success",
            "path":          out_path,
            "invoice_no":    invoice_no,
            "gross_amount":  total_with_gst,
            "tds":           tds,
            "net_payable":   net_payable,
        }

    except Exception as e:
        return {"status": "error", "message": str(e)}


# ─────────────────────────────────────────────────────────────────
# 5. COMPLETION CERTIFICATE REQUEST
# ─────────────────────────────────────────────────────────────────

def generate_completion_cert_request(tender_data: dict, out_path: str) -> dict:
    """Generate a letter requesting project completion certificate."""
    if not DOCX_OK:
        return {"status": "error", "message": "python-docx not installed"}
    try:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2)

        _letterhead(doc, tender_data)
        doc.add_paragraph()
        _para(doc, f"Date: {date.today().strftime('%d %B %Y')}")
        doc.add_paragraph()
        _para(doc, "To,")
        _para(doc, tender_data.get("client_contact", "The Authorised Officer,"))
        _para(doc, tender_data.get("org_name", ""))
        doc.add_paragraph()
        _para(doc, f"Sub: Request for Project Completion Certificate — {tender_data.get('brief','')[:80]}", bold=True)
        _para(doc, f"Ref: Contract/LoA No.: {tender_data.get('loa_ref','[Ref]')}  |  Tender No.: {tender_data.get('tender_no','')}", bold=True)
        doc.add_paragraph()
        _para(doc, "Dear Sir / Madam,")
        doc.add_paragraph()
        _para(doc, f"We, {N['name']}, have successfully completed the work assigned to us under the above-referenced contract. All deliverables have been handed over, UAT has been completed and signed off, and the system has been live and operational as per the Go-Live schedule.")
        doc.add_paragraph()
        _para(doc, "We request your good office to kindly issue a Project Completion Certificate at the earliest, specifying:", bold=True)
        for item in [
            "1. Scope of work completed",
            "2. Contract value",
            f"3. Period: {tender_data.get('project_start','')} to {date.today().strftime('%d-%m-%Y')}",
            "4. Quality and performance of work",
            "5. Any other relevant details",
        ]:
            _para(doc, item)
        doc.add_paragraph()
        _para(doc, "The completion certificate is required for our records and for future tender eligibility. We assure you of our continued support during the AMC / warranty period.")
        doc.add_paragraph()
        _para(doc, "Thanking you,")
        _para(doc, "Yours faithfully,")
        _signature(doc)
        doc.save(out_path)
        return {"status": "success", "path": out_path}
    except Exception as e:
        return {"status": "error", "message": str(e)}
