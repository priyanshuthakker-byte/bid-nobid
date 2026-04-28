"""
TenderExtractor v3
Priority order for text extraction:
  1. HTML file in ZIP  → parsed for portal/snapshot fields (most reliable)
  2. Searchable PDF    → pdfplumber / PyPDF2
  3. Scanned PDF       → OCR via ocr_engine.py (pytesseract)
  4. DOCX              → python-docx
  5. XLSX              → openpyxl (for BOQ / schedule data)
  6. TXT / HTM         → plain read

Snapshot fields extracted from HTML portal data first,
then supplemented/overridden by RFP document content.
"""

import re, logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ── FILE READERS ───────────────────────────────────────────────

def read_html(path: Path) -> tuple[str, dict]:
    """
    Read HTML file. Returns (plain_text, portal_fields_dict).
    portal_fields_dict contains structured snapshot data parsed from
    the T247/eproc portal HTML (labels + values).
    """
    try:
        from bs4 import BeautifulSoup
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "lxml")

        # Remove scripts and styles
        for tag in soup(["script", "style"]):
            tag.decompose()

        # Extract plain text
        plain_text = soup.get_text(separator="\n", strip=True)

        # Parse portal field pairs (label → value) from table cells
        portal_fields = {}
        tables = soup.find_all("table")
        for table in tables:
            rows = table.find_all("tr")
            for row in rows:
                cells = row.find_all(["td", "th"])
                if len(cells) >= 2:
                    for i in range(0, len(cells) - 1, 2):
                        label = cells[i].get_text(strip=True)
                        value = cells[i + 1].get_text(strip=True)
                        if label and value and len(label) < 120:
                            portal_fields[label.lower()] = value

        # Also parse label+value divs (Bootstrap/Angular portal style)
        label_els = soup.find_all(class_=re.compile(r'label|font-bold|field-label', re.I))
        for el in label_els:
            label = el.get_text(strip=True)
            sibling = el.find_next_sibling()
            if sibling:
                value = sibling.get_text(strip=True)
                if label and value and len(label) < 120:
                    portal_fields[label.lower()] = value

        return plain_text, portal_fields

    except Exception as e:
        logger.warning(f"HTML read failed for {path.name}: {e}")
        try:
            return path.read_text(encoding="utf-8", errors="ignore"), {}
        except Exception:
            return "", {}


def read_pdf(path: Path, max_chars: int = 120_000) -> str:
    """Read PDF — memory-safe: caps at max_chars, PyPDF2 primary (lighter than pdfplumber)."""
    text = ""

    # PyPDF2 primary — lighter memory footprint than pdfplumber
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            if sum(len(p) for p in parts) >= max_chars:
                break
            t = page.extract_text()
            if t:
                parts.append(t)
        text = "\n\n".join(parts)
        del reader, parts
    except Exception as e:
        logger.debug(f"PyPDF2 failed for {path.name}: {e}")

    # pdfplumber fallback only if PyPDF2 gave nothing
    if not text or len(text.strip()) < 100:
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    if sum(len(p) for p in parts) >= max_chars:
                        break
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            text = "\n\n".join(parts)
            del parts
        except Exception as e:
            logger.debug(f"pdfplumber failed for {path.name}: {e}")

    return text[:max_chars]




def read_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"DOCX read failed for {path.name}: {e}")
        return ""


def read_xlsx(path: Path) -> str:
    ext = path.suffix.lower()
    # Old .xls format — use xlrd
    if ext == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(str(path))
            parts = []
            for sheet in wb.sheets():
                parts.append(f"=== Sheet: {sheet.name} ===")
                for rx in range(sheet.nrows):
                    row_text = " | ".join(str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols) if str(sheet.cell_value(rx, cx)).strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            msg = str(e).lower()
            if "encrypted" in msg or "password" in msg:
                logger.info(f"XLS skipped (encrypted/password-protected): {path.name}")
            else:
                logger.warning(f"XLS read failed for {path.name}: {e}")
            return ""
    # Modern .xlsx format — use openpyxl
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(v) for v in row if v is not None and str(v).strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        msg = str(e).lower()
        if "encrypted" in msg or "password" in msg:
            logger.info(f"XLSX skipped (encrypted/password-protected): {path.name}")
        else:
            logger.warning(f"XLSX read failed for {path.name}: {e}")
        return ""


def read_document(path: Path) -> str:
    """Read any supported document type. Returns plain text."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return read_pdf(path)
    elif ext in (".docx", ".doc"):
        return read_docx(path)
    elif ext in (".xlsx", ".xls"):
        return read_xlsx(path)
    elif ext in (".html", ".htm"):
        text, _ = read_html(path)
        return text
    elif ext == ".txt":
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    return ""


# ── PORTAL HTML FIELD MAPPER ───────────────────────────────────

def extract_portal_snapshot(portal_fields: dict) -> dict:
    """
    Map T247/eproc portal HTML label→value pairs to snapshot fields.
    These are the fields typically in the portal header / critical data sheet.
    The field names below are normalised (lowercase) label fragments.
    """
    snap = {}

    def find(keywords: list) -> str:
        for kw in keywords:
            for label, value in portal_fields.items():
                if kw in label and value and value.strip() not in ("—", "", "N/A", "null"):
                    return value.strip()
        return ""

    snap["tender_no"]          = find(["tender reference no", "nit no", "tender no", "ref no"])
    snap["tender_id"]          = find(["system tender no", "tender id", "bid no"])
    snap["org_name"]           = find(["organization hierarchy", "organization", "organisation", "dept name"])
    snap["tender_name"]        = find(["detailed description", "nit", "title", "description"])
    snap["tender_type"]        = find(["category", "bid type", "tender type", "form of contract"])
    snap["portal"]             = find(["portal", "website", "url", "submission url", "bid submit"])
    snap["bid_start_date"]     = find(["bid submission start date", "start date", "publish date", "publishing date"])
    snap["bid_submission_date"]= find(["bid submission due date", "bid submission end date", "last date", "closing date", "end date"])
    snap["bid_opening_date"]   = find(["bid open date", "technical bid opening", "opening date"])
    snap["prebid_meeting"]     = find(["pre-bid meeting", "pre bid", "prebid", "pre bid discussion"])
    snap["prebid_query_date"]  = find(["last date.*clarification", "written queries", "query deadline", "clarification"])
    snap["estimated_cost"]     = find(["estimated value", "estimated cost", "project value", "nit value"])
    snap["tender_fee"]         = find(["non-refundable tender fee", "tender fee", "document fee", "bid fee"])
    snap["processing_fee"]     = find(["processing fee"])
    snap["emd"]                = find(["earnest money deposit", "emd", "bid security"])
    snap["offer_validity"]     = find(["offer validity", "bid validity", "validity"])
    snap["mode_of_selection"]  = find(["evaluation method", "mode of selection", "method of evaluation", "qcbs"])
    snap["no_of_covers"]       = find(["bid parts", "covers", "envelopes", "no of covers"])
    snap["contact"]            = find(["contact person", "tender issuing authority", "contact"])
    snap["location"]           = find(["location", "place", "state"])

    # Remove empty fields
    return {k: v for k, v in snap.items() if v}


# ── REGEX-BASED FIELD EXTRACTION FROM RFP TEXT ─────────────────

def _find(text: str, patterns: list, flags=re.IGNORECASE) -> str:
    for pat in patterns:
        try:
            m = re.search(pat, text, flags)
            if m:
                # Try group(1) first, fall back to group(0) if no capture group
                try:
                    val = m.group(1).strip()
                except IndexError:
                    val = m.group(0).strip()
                if val and val not in ("—", "") and len(val) < 500:
                    return val
        except re.error:
            continue
    return ""


def extract_snapshot_from_text(text: str) -> dict:
    """Extract snapshot fields from RFP text using regex patterns."""
    snap = {}

    snap["tender_no"] = _find(text, [
        r'(?:Tender|NIT)\s+(?:No|Number|Ref)\.?\s*[:\-]?\s*([A-Z0-9/_\-\.]{4,40})',
        r'(?:Reference\s+No\.?|Ref\.?\s*No\.?)\s*[:\-]\s*([A-Z0-9/_\-\.]{4,40})',
        r'(?:Tender\s+Notice\s+No\.?)\s*[:\-]?\s*([A-Z0-9/_\-\.]{4,40})',
    ])

    snap["org_name"] = _find(text, [
        r'(?:Government of|Govt\.? of)\s+([A-Za-z\s]{5,60})\n',
        r'(?:Department|Ministry|Authority|Corporation|Board|Commission)\s+of\s+([A-Za-z\s,]{5,80})',
        r'(?:Issued by|Issuing Authority)\s*[:\-]\s*([A-Za-z\s,\.]{10,100})',
    ])

    snap["tender_name"] = _find(text, [
        r'(?:Request for Proposal|RFP)\s*(?:for|:)?\s*(.{20,200}?)(?:\n|$)',
        r'(?:Tender for|NIT for|Bid for)\s*:?\s*(.{20,200}?)(?:\n|$)',
        r'(?:Supply|Implementation|Development|Design)\s+(?:of|for)\s+(.{20,200}?)(?:\n|$)',
    ])

    snap["bid_submission_date"] = _find(text, [
        r'(?:Bid\s+Submission\s+(?:End\s+)?(?:Date|Deadline)|Last\s+Date\s+(?:of|for)\s+(?:Submission|Bid))\s*[:\-]?\s*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4}(?:\s*[\,\s]+\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)?)',
        r'(?:Submission\s+Due\s+Date)\s*[:\-]?\s*([0-9\-/\.:\s]+(?:AM|PM)?)',
    ])

    snap["bid_opening_date"] = _find(text, [
        r'(?:Technical\s+Bid\s+Opening|Bid\s+Opening\s+Date)\s*[:\-]?\s*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4}(?:\s*[\,\s]+\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)?)',
    ])

    snap["bid_start_date"] = _find(text, [
        r'(?:Bid\s+Submission\s+Start|Bid\s+Availability|Publishing\s+Date)\s*[:\-]?\s*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4}(?:\s*[\,\s]+\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)?)',
    ])

    snap["prebid_query_date"] = _find(text, [
        r'(?:Last\s+Date\s+(?:and\s+Time\s+)?(?:for|of)\s+(?:Submission\s+of\s+)?(?:written\s+)?(?:queries?|clarification|pre.bid))\s*[:\-]?\s*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4}(?:\s*[\,\s]+\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)?)',
        r'(?:Pre.bid\s+[Qq]uery\s+[Dd]eadline|Query\s+Submission\s+Date)\s*[:\-]?\s*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4})',
    ])

    snap["prebid_meeting"] = _find(text, [
        r'(?:Pre.Bid\s+(?:Meeting|Conference|Discussion))\s*[:\-]?\s*(.{20,200}?)(?:\n|$)',
    ])

    snap["estimated_cost"] = _find(text, [
        r'(?:Estimated\s+(?:Cost|Value)|NIT\s+Value|Project\s+(?:Cost|Value))\s*[:\-]?\s*((?:Rs\.?|INR|₹)\s*[\d,\.]+(?:\s*(?:Crore|Lakh|Lac|CR|L))?)',
        r'(?:Estimated\s+Cost)\s*[:\-]\s*([\d,\.]+)',
    ])

    snap["tender_fee"] = _find(text, [
        r'(?:Non.refundable\s+(?:Tender|Bid|Document)\s+Fee|Tender\s+Fee|Document\s+Fee|Bid\s+Fee)\s*[:\-]?\s*((?:Rs\.?|INR|₹)?\s*[\d,\.]+(?:[/\-]\s*)?(?:\s*\+\s*GST)?(?:\s*only)?)',
        r'(?:Tender\s+Document\s+(?:Fee|Cost))\s*[:\-]?\s*((?:Rs\.?|INR)?\s*[\d,]+)',
    ])

    snap["processing_fee"] = _find(text, [
        r'(?:Processing\s+Fee)\s*[:\-]?\s*((?:Rs\.?|INR|₹)?\s*[\d,\.]+)',
    ])

    snap["emd"] = _find(text, [
        r'(?:Earnest\s+Money\s+Deposit|EMD|Bid\s+Security|Bid\s+Guarantee)\s*[:\-]?\s*((?:Rs\.?|INR|₹)?\s*[\d,\.]+(?:\s*(?:Crore|Lakh|Lac|CR|L|/-|only))?)',
        r'(?:EMD\s+Amount)\s*[:\-]?\s*((?:Rs\.?|INR)?\s*[\d,\.]+)',
    ])

    snap["emd_exemption"] = _find(text, [
        r'(?:EMD\s+Exemption|Exemption\s+from\s+EMD)\s*[:\-]?\s*(.{20,300}?)(?:\n|$)',
        r'(?:MSME|Startup)\s*(?:are|is)?\s*(?:exempt|exempted)\s+from\s+(EMD[^\n]{0,100})',
    ])

    snap["performance_security"] = _find(text, [
        r'(?:Performance\s+(?:Bank\s+)?(?:Guarantee|Security|Bond)|PBG)\s*[:\-]?\s*(.{10,200}?)(?:\n|$)',
        r'(\d+(?:\.\d+)?%?\s+of\s+(?:contract|awarded|project)\s+value)',
    ])

    snap["contract_period"] = _find(text, [
        r'(?:Period\s+of\s+(?:Work|Contract|Agreement)|Contract\s+Duration|Project\s+Duration)\s*[:\-]?\s*(.{10,200}?)(?:\n|$)',
        r'(?:Completion\s+Period|Duration)\s*[:\-]\s*(.{5,100}?)(?:\n|$)',
    ])

    snap["bid_validity"] = _find(text, [
        r'(?:Bid\s+Validity|Offer\s+Validity|Valid(?:ity)?\s+(?:Period|for\s+acceptance))\s*[:\-]?\s*(.{5,100}?)(?:\n|$)',
        r'(\d+\s+days?\s+(?:from|after|of))',
    ])

    snap["contact"] = _find(text, [
        r'(?:Contact\s+(?:Person|Officer|for\s+queries)|For\s+(?:any|technical)\s+queries?\s+(?:contact|please\s+contact))\s*[:\-]?\s*(.{20,300}?)(?:\n\n|$)',
    ])

    snap["portal"] = _find(text, [
        r'(https?://(?:eproc|eprocure|etender|gem|portal|tender)\S+)',
        r'(?:bids?\s+must\s+be\s+submitted\s+(?:through|via|on|at))\s+(https?://\S+)',
        r'(https?://[a-zA-Z0-9\-\.]+\.(?:gov|nic|in|org)\S*)',
    ])

    snap["jv_allowed"] = _find(text, [
        r'(?:JV|Joint\s+Venture|Consortium)\s*(?:is|are)?\s*((?:not\s+)?(?:allowed|permitted|acceptable))[^\n]{0,100}',
        r'(?:joint\s+venture|consortium)\s*[:\-]?\s*(.{20,200}?)(?:\n|$)',
    ])

    snap["mode_of_selection"] = _find(text, [
        r'(?:Method\s+of\s+(?:Tender\s+)?Evaluation|Mode\s+of\s+Selection|Evaluation\s+Method)\s*[:\-]?\s*((?:L1|QCBS|LCS|QBS|FBS|CQS|Open|Single)[^\n]{0,60})',
    ])

    snap["post_implementation"] = _find(text, [
        r'(?:O&M|AMC|CAMC|Operation\s+and\s+Maintenance|Annual\s+Maintenance)\s+(?:period|for)\s*[:\-]?\s*(.{5,100}?)(?:\n|$)',
    ])

    snap["location"] = _find(text, [
        r'(?:Project\s+Location|Place\s+of\s+Work|Location)\s*[:\-]\s*([A-Za-z\s,]+(?:\n|$))',
    ])

    return {k: v for k, v in snap.items() if v and len(v.strip()) > 2}


# ── MAIN EXTRACTOR CLASS ───────────────────────────────────────

class TenderExtractor:

    def process_documents(self, doc_files: list) -> dict:
        """
        Process a list of document files and return merged snapshot data.
        HTML is parsed first for portal fields. PDF/DOCX supplement the rest.
        """
        result = {
            "tender_no": "—", "tender_id": "—", "org_name": "—",
            "tender_name": "—", "portal": "—", "tender_type": "—",
            "bid_start_date": "—", "bid_submission_date": "—",
            "bid_opening_date": "—", "commercial_opening_date": "—",
            "prebid_meeting": "Not specified", "prebid_query_date": "Not specified",
            "estimated_cost": "Not specified", "tender_fee": "—",
            "processing_fee": "—", "emd": "—", "emd_exemption": "—",
            "performance_security": "As per tender",
            "contract_period": "—", "bid_validity": "—",
            "post_implementation": "—", "mode_of_selection": "—",
            "no_of_covers": "—", "jv_allowed": "Not specified",
            "technology_mandatory": "—", "location": "—", "contact": "—",
            "files_processed": [],
            "pq_criteria": [], "tq_criteria": [], "scope_items": [],
            "payment_terms": [], "notes": [],
        }

        # Separate HTML files from document files
        html_files = [f for f in doc_files if Path(f).suffix.lower() in (".html", ".htm")]
        non_html = [f for f in doc_files if Path(f).suffix.lower() not in (".html", ".htm")]

        portal_snapshot = {}

        # ── STEP 1: Parse HTML for portal/snapshot fields ──────
        for hf in html_files:
            hf = Path(hf)
            logger.info(f"[Extractor] Parsing HTML portal data: {hf.name}")
            try:
                _, portal_fields = read_html(hf)
                if portal_fields:
                    snap = extract_portal_snapshot(portal_fields)
                    portal_snapshot.update(snap)
                    result["files_processed"].append(hf.name)
                    logger.info(f"[Extractor] HTML gave {len(snap)} portal fields")
            except Exception as e:
                logger.warning(f"[Extractor] HTML parse failed: {e}")

        # Apply portal snapshot to result (highest priority for these fields)
        PORTAL_PRIORITY_FIELDS = [
            "tender_no", "tender_id", "org_name", "tender_name", "portal",
            "tender_type", "bid_start_date", "bid_submission_date",
            "bid_opening_date", "prebid_meeting", "prebid_query_date",
            "estimated_cost", "tender_fee", "processing_fee", "emd",
            "offer_validity", "mode_of_selection", "no_of_covers",
            "contact", "location",
        ]
        for field in PORTAL_PRIORITY_FIELDS:
            if portal_snapshot.get(field):
                result[field] = portal_snapshot[field]

        # ── STEP 2: Extract text from PDFs / DOCX ─────────────
        for doc in non_html:
            doc = Path(doc)
            try:
                text = read_document(doc)
                if text and text.strip():
                    result["files_processed"].append(doc.name)
                    # Extract snapshot fields from RFP text
                    rfp_snap = extract_snapshot_from_text(text)
                    # Only fill in fields not already populated by portal HTML
                    EMPTY = {"—", "Not specified", "", "As per tender", None}
                    for field, value in rfp_snap.items():
                        if str(result.get(field, "")).strip() in EMPTY and value:
                            result[field] = value
                    logger.info(f"[Extractor] {doc.name}: {len(text)} chars, {len(rfp_snap)} fields extracted")
            except Exception as e:
                logger.warning(f"[Extractor] Failed to read {doc.name}: {e}")

        return result
